# Extracted from WMH_Analysis.ipynb, code cell 13.
# Notebook heading: # Interaction OLS of HSN-WMH (log1p) with BH-FDR on interaction p (Peri/Deep only)
# Run this file from the repository root unless a local CONFIG section is edited.

# Interaction OLS of HSN-WMH (log1p) with BH-FDR on interaction p (Peri/Deep only)
# ----------------------------------------------------------------------------------
# Standalone script aligned with user's stratified style.
#
# Key logic:
#   • For each stratification variable, define the interaction analytic subset
#     as matched sets where the Study arm belongs to any eligible level.
#   • Fit: outcome ~ C(group) * C(strat_col) + prespecified covariates
#   • SE: cluster-robust by match_id when feasible; otherwise HC3
#   • Extract GLOBAL interaction p (joint test of all group×strat terms)
#   • Apply BH-FDR within each cohort × stratification to {Periventricular, Deep}
#     only; Total excluded (q = NaN)
#
# Cohorts:
#   - primary_cohort_stratified.csv
#   - sensitivity_cohort_stratified.csv
#
# Outputs:
#   - Interaction/<Cohort>/Interaction_<StratKey>.csv
#   - Interaction/Interaction_AllCohorts_Combined.csv
#   - Interaction/Interaction_AllCohorts_Supplement.docx
# ----------------------------------------------------------------------------------

import os, re
import numpy as np
import pandas as pd
import statsmodels.api as sm
from patsy import dmatrices
from statsmodels.stats.multitest import multipletests

# ---------- Optional Word export ----------
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception:
    Document = None

# ---------------- Paths & cohorts ----------------
COHORTS = {
    "Primary":     "primary_cohort_stratified.csv",
    "Sensitivity": "sensitivity_cohort_stratified.csv",
}
OUTDIR = "Interaction"
os.makedirs(OUTDIR, exist_ok=True)

# ---------------- Columns & settings ----------------
GROUP_COL, MATCH_ID = "group", "match_id"
SCALE_COL = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
RAW_DEEP  = "Total_volume_of_deep_white_matter_hyperintensities_Instance_2"
RAW_PV    = "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"
RAW_TOT   = ("Total_volume_of_white_matter_hyperintensities_from_T1_"
             "and_T2_FLAIR_images_Instance_2")

OUTCOME_BUILD = [
    ("Log_HSNorm_Total_WMH",           RAW_TOT,  "Total WMH"),
    ("Log_HSNorm_PeriVentricular_WMH", RAW_PV,   "Periventricular WMH"),
    ("Log_HSNorm_Deep_WMH",            RAW_DEEP, "Deep WMH"),
]

OUTCOME_ORDER = ["Total WMH", "Periventricular WMH", "Deep WMH"]

# Prespecified covariates (NO CMC here)
ADJUST_ALL = [
    "Sex","Age_at_Instance_2","Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0","Genetic_ethnic_grouping",
    "Smoking_Ever","Alcohol_intake_frequency_ordinal",
]
CATEGORICAL_ALL = {"Sex","Genetic_ethnic_grouping","Smoking_Ever"}

# ---------------- Helpers ----------------
def safe_name(s: str) -> str:
    return re.sub(r"[^0-9A-Za-z._-]+", "_", str(s))

def clean_cols(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns=lambda c: pd.Series(c).str.replace(r"[^0-9a-zA-Z]+","_", regex=True).iloc[0])

def ensure_outcomes(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    if GROUP_COL in df.columns:
        df[GROUP_COL] = pd.Categorical(df[GROUP_COL], ["Control","Study"])
    scale = pd.to_numeric(df.get(SCALE_COL, 1.0), errors="coerce").fillna(1.0)
    for out_col, src, _lab in OUTCOME_BUILD:
        v = pd.to_numeric(df.get(src), errors="coerce")
        v = v * scale
        df[out_col] = np.log1p(v)
    return df

def fit_cluster(formula: str, data: pd.DataFrame):
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    base = sm.OLS(y, X).fit(cov_type="HC3")
    if MATCH_ID not in data.columns:
        return base
    groups = data.loc[y.index, MATCH_ID].dropna()
    if groups.nunique() <= 1:
        return base
    return sm.OLS(y.loc[groups.index], X.loc[groups.index]).fit(
        cov_type="cluster", cov_kwds={"groups": groups}
    )

# ---------------- Stratification spec (same logic as user's code) ----------------
def build_strata(df_example: pd.DataFrame):
    strata = []

    # Sex (adaptive to "Male/Female" or 1/0)
    if "Sex" in df_example.columns:
        uniq = set(str(x) for x in pd.Series(df_example["Sex"]).dropna().unique())
        if {"Male","Female"} & uniq:
            levels = ["Male","Female"]; lmap = {"Male":"Male","Female":"Female"}
        else:
            levels = [1,0];            lmap = {1:"Male", 0:"Female"}
        strata.append(("Sex", levels, lmap, "Sex", "Sex"))

    if "Sleepy_SA_Stratum" in df_example.columns:
        strata.append(("Sleepy_SA_Stratum", [1,0], {1:"Sleepy", 0:"Not Sleepy"},
                       "Daytime Sleepiness", "Sleepiness"))

    if "Snoring_Stratum" in df_example.columns:
        strata.append(("Snoring_Stratum", [1,0], {1:"Snoring", 0:"No Snore"},
                       "Snoring Status", "Snoring"))

    if "SA_MRI_interval_stratum" in df_example.columns:
        strata.append(("SA_MRI_interval_stratum",
                       ["LE10_years","GT10_years"],
                       {"LE10_years":"≤10 y", "GT10_years":">10 y"},
                       "Diagnosis–MRI Interval", "Dx_MRI_Interval"))

    # SA diagnosis method
    if "SA_ascertain_group" in df_example.columns:
        strata.append(("SA_ascertain_group",
                       ["Self_Report_Only","Hospital_Secondary","Hospital_Primary"],
                       {"Self_Report_Only":"Self-Report",
                        "Hospital_Secondary":"Hosp-Secondary",
                        "Hospital_Primary":"Hosp-Primary"},
                       "SA Diagnosis Method", "SA_Diagnosis"))

    return strata

# ---------------- Interaction-specific helpers ----------------
def build_interaction_formula(outcome: str, strat_col: str, data: pd.DataFrame) -> str:
    rhs = [f"C({GROUP_COL}) * C({strat_col})"]
    for v in ADJUST_ALL:
        if v in data.columns:
            rhs.append(f"C({v})" if v in CATEGORICAL_ALL else v)
    return f"{outcome} ~ " + " + ".join(rhs)

def get_interaction_subset(df: pd.DataFrame, strat_col: str, levels: list) -> pd.DataFrame:
    # Matched sets where Study belongs to ANY eligible level
    if strat_col not in df.columns:
        return pd.DataFrame()
    study_ids = df[(df[GROUP_COL] == "Study") & (df[strat_col].isin(levels))][MATCH_ID].dropna().unique()
    if study_ids.size == 0:
        return pd.DataFrame()
    return df[df[MATCH_ID].isin(study_ids)].copy()

def extract_global_interaction_p(model, strat_col: str) -> float:
    # Identify all interaction terms comparing Study vs reference level(s)
    # Pattern like: C(group)[T.Study]:C(strat_col)[T.xxx]
    pat = f"C\\({GROUP_COL}\\)\\[T\\.Study\\]:C\\({strat_col}\\)\\[T\\."
    inter_params = [p for p in model.params.index if re.search(pat, p)]
    if not inter_params:
        return np.nan

    pnames = list(model.params.index)
    k = len(pnames)
    rows = []
    for ip in inter_params:
        r = np.zeros(k)
        r[pnames.index(ip)] = 1.0
        rows.append(r)
    R = np.vstack(rows)

    try:
        w = model.wald_test(R)
        return float(w.pvalue)
    except Exception:
        return np.nan

def impute_covariates_like_user(sub: pd.DataFrame) -> pd.DataFrame:
    sub = sub.copy()
    for v in ADJUST_ALL:
        if v not in sub.columns:
            continue
        if pd.api.types.is_numeric_dtype(sub[v]):
            sub[v] = pd.to_numeric(sub[v], errors="coerce").fillna(sub[v].median())
        else:
            m = sub[v].mode(dropna=True)
            sub[v] = sub[v].fillna(m.iloc[0] if not m.empty else sub[v])
    return sub

def run_one_interaction_block(df: pd.DataFrame,
                              strat_col: str,
                              levels: list,
                              label_map: dict,
                              title: str,
                              cohort_label: str) -> pd.DataFrame:
    """
    For a given stratification family:
      - Build interaction subset (Study-defined matched sets)
      - Fit global interaction tests for three WMH outcomes
      - Apply BH-FDR within this block for {Periventricular, Deep}
    """
    sub = get_interaction_subset(df, strat_col, levels)
    if sub.empty:
        return pd.DataFrame()

    sub = impute_covariates_like_user(sub)

    # Helpful counts for table
    n_sets = int(sub[MATCH_ID].nunique()) if MATCH_ID in sub.columns else np.nan
    n_total = int(sub.shape[0])

    outcome_map = {
        "Log_HSNorm_Total_WMH": "Total WMH",
        "Log_HSNorm_PeriVentricular_WMH": "Periventricular WMH",
        "Log_HSNorm_Deep_WMH": "Deep WMH",
    }

    rows = []
    for out_col, out_lab in outcome_map.items():
        if out_col not in sub.columns:
            continue
        formula = build_interaction_formula(out_col, strat_col, sub)
        model = fit_cluster(formula, sub)

        p_int = extract_global_interaction_p(model, strat_col)

        rows.append({
            "Analytic Cohort": cohort_label,
            "Stratification": title,
            "Strat_Column": strat_col,
            "Outcome": out_lab,
            "p_interaction_global": p_int,
            "q_interaction_primary": np.nan,
            "q_interaction_sensitivity": np.nan,
            "N_sets": n_sets,
            "N_total": n_total,
        })

    if not rows:
        return pd.DataFrame()

    # BH-FDR within this stratification block for Peri/Deep only
    idx_fdr = [i for i, r in enumerate(rows)
               if r["Outcome"] in ("Periventricular WMH", "Deep WMH")
               and pd.notna(r["p_interaction_global"])]

    if idx_fdr:
        pvals = np.array([rows[i]["p_interaction_global"] for i in idx_fdr], dtype=float)
        qvals = multipletests(pvals, method="fdr_bh")[1]
        for i_row, qv in zip(idx_fdr, qvals):
            if cohort_label == "Primary":
                rows[i_row]["q_interaction_primary"] = float(qv)
            elif cohort_label == "Sensitivity":
                rows[i_row]["q_interaction_sensitivity"] = float(qv)

    return pd.DataFrame(rows)

# ---------------- Word table helpers (aligned with user's style) ----------------
def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge); tag = OxmlElement(f'w:{edge}')
            for key in ("val","sz","color","space"):
                if key in edge_data: tag.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(tag)

def _apply_three_line_table(table, header_row_idx=0):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, left={"val":"nil"}, right={"val":"nil"},
                                   top={"val":"nil"},  bottom={"val":"nil"})
    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell, top={"val":"single","sz":8,"color":"000000"},
                              bottom={"val":"single","sz":8,"color":"000000"})
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom={"val":"single","sz":8,"color":"000000"})

def _right_align_numeric(table, numeric_col_idx):
    for row in table.rows[1:]:
        for j in numeric_col_idx:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def export_interaction_word(df_all: pd.DataFrame, outdir: str,
                            title: str = "eTable. Interaction analyses of WMH by SA group",
                            legend: str | None = None,
                            filename: str = "Interaction_AllCohorts_Supplement.docx"):
    if Document is None or df_all is None or df_all.empty:
        print("Word export skipped."); return None

    df = df_all.copy()

    # Cohort display names
    df["Analytic Cohort"] = df["Analytic Cohort"].replace({
        "Primary": "Primary PSM cohort",
        "Sensitivity": "PSM–sensitivity cohort"
    })

    def _fmt_p(x):
        if pd.isna(x): return "—"
        x = float(x)
        return "<0.001" if x < 1e-3 else f"{x:.3f}"

    def _pick_q(row):
        c = str(row.get("Analytic Cohort","")).lower()
        if c.startswith("primary"):
            return row.get("q_interaction_primary", np.nan)
        if "sensitivity" in c:
            return row.get("q_interaction_sensitivity", np.nan)
        qa = row.get("q_interaction_primary", np.nan)
        qb = row.get("q_interaction_sensitivity", np.nan)
        return qa if pd.notna(qa) else qb

    df["q_interaction"] = df.apply(_pick_q, axis=1)

    # Order outcomes
    # ---- Force ordering: Primary first; Sex first ----
    cohort_order = ["Primary PSM cohort", "PSM–sensitivity cohort"]
    strat_order = [
        "Sex",
        "Daytime Sleepiness",
        "Snoring Status",
        "Diagnosis–MRI Interval",
        "SA Diagnosis Method",
    ]

    # Cohort order
    df["Analytic Cohort"] = pd.Categorical(
        df["Analytic Cohort"],
        categories=cohort_order,
        ordered=True
    )

    # Stratification order (keep only those actually present)
    present_strata = [s for s in strat_order if s in set(df["Stratification"].dropna().unique())]
    # If any unexpected strata exist, append them to the end to avoid dropping
    extra_strata = [s for s in df["Stratification"].dropna().unique() if s not in strat_order]
    final_strata_order = present_strata + extra_strata

    df["Stratification"] = pd.Categorical(
        df["Stratification"],
        categories=final_strata_order,
        ordered=True
    )

    # Outcome order
    df["Outcome"] = pd.Categorical(df["Outcome"], categories=OUTCOME_ORDER, ordered=True)

    # Sort for docx
    df = df.sort_values(["Analytic Cohort", "Stratification", "Outcome"]).reset_index(drop=True)


    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run(title).bold = True

    headers = ["Analytic Cohort","Stratification","Outcome",
               "Matched sets (N)","Participants (N)",
               "Global interaction p","q (BH-FDR*)"]

    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, r in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(r.get("Analytic Cohort",""))
        cells[1].text = str(r.get("Stratification",""))
        cells[2].text = str(r.get("Outcome",""))
        cells[3].text = "" if pd.isna(r.get("N_sets", np.nan)) else str(int(r["N_sets"]))
        cells[4].text = "" if pd.isna(r.get("N_total", np.nan)) else str(int(r["N_total"]))
        cells[5].text = _fmt_p(r.get("p_interaction_global", np.nan))
        cells[6].text = _fmt_p(r.get("q_interaction", np.nan)) if pd.notna(r.get("q_interaction", np.nan)) else "—"

    _right_align_numeric(t, numeric_col_idx=[3,4,5,6])
    _apply_three_line_table(t, header_row_idx=0)

    if legend is None:
        legend = (
            "Legend: Global interaction p values test whether the adjusted Study vs Control difference "
            "in log1p-transformed, head-size–normalized WMH varies across levels of each stratification "
            "variable, within matched sets where the Study arm belongs to the eligible strata for that "
            "stratification family. Models adjust for prespecified covariates (no CMC). Standard errors "
            "are cluster-robust by matched set (match_id) when feasible; otherwise HC3. "
            "*Benjamini–Hochberg FDR is applied within each cohort and stratification family to the "
            "interaction p values of {Periventricular, Deep}; Total WMH is excluded."
        )
    doc.add_paragraph().add_run(legend).italic = True

    path = os.path.join(outdir, filename)
    doc.save(path)
    print(f"Word table saved: {path}")
    return path

# ---------------- One cohort end-to-end ----------------
def run_for_cohort(label: str, csv_path: str) -> pd.DataFrame:
    print(f"\n===== {label} ({csv_path}) =====")
    subdir = os.path.join(OUTDIR, label)
    os.makedirs(subdir, exist_ok=True)

    df_raw = pd.read_csv(csv_path)
    df = ensure_outcomes(clean_cols(df_raw))

    STRATA = build_strata(df)
    all_blocks = []

    for strat_col, levels, lmap, panel_title, safe_key in STRATA:
        block_df = run_one_interaction_block(df, strat_col, levels, lmap, panel_title, label)
        if block_df is None or block_df.empty:
            continue

        all_blocks.append(block_df)

        block_csv = os.path.join(subdir, f"Interaction_{safe_key}.csv")
        block_df.to_csv(block_csv, index=False, encoding="utf-8-sig")

        # Console preview
        show = block_df.copy()
        show["Outcome"] = pd.Categorical(show["Outcome"], OUTCOME_ORDER, ordered=True)
        show = show.sort_values(["Outcome"])
        cols = ["Stratification","Outcome","p_interaction_global",
                "q_interaction_primary","q_interaction_sensitivity","N_sets","N_total"]
        print("\n--", panel_title, "--")
        print(show[cols].to_string(index=False))

    if all_blocks:
        res_all = pd.concat(all_blocks, ignore_index=True)
        res_csv = os.path.join(subdir, "Interaction_AllResults.csv")
        res_all.to_csv(res_csv, index=False, encoding="utf-8-sig")
    else:
        res_all = pd.DataFrame()

    # README
    note = [
        "Interaction models: outcome ~ C(group) * C(strat_col) + covariates.",
        "Head-size–normalized WMH with log1p transform.",
        "Covariates: " + ", ".join(ADJUST_ALL) + " (no CMC).",
        "Cluster-robust SEs by match_id where feasible; otherwise HC3.",
        "Global interaction p derived from joint Wald tests of all group×stratum terms.",
        "Within each cohort × stratification family: BH-FDR to {Periventricular, Deep} interaction p; Total excluded."
    ]
    with open(os.path.join(subdir, "README_Interaction.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(note))

    return res_all

# ---------------- Run both cohorts & export combined table ----------------
all_res = []
for lab, fp in COHORTS.items():
    if not os.path.exists(fp):
        print(f"[Warning] File not found: {fp} → skip {lab}")
        continue
    res = run_for_cohort(lab, fp)
    if res is not None and not res.empty:
        all_res.append(res)

if all_res:
    comb = pd.concat(all_res, ignore_index=True)

    # Standardize cohort display names for CSV
    comb["Analytic Cohort"] = comb["Analytic Cohort"].replace({
        "Primary": "Primary PSM cohort",
        "Sensitivity": "PSM–sensitivity cohort",
        "primary": "Primary PSM cohort",
        "sensitivity": "PSM–sensitivity cohort",
    })

    out_csv = os.path.join(OUTDIR, "Interaction_AllCohorts_Combined.csv")
    comb.to_csv(out_csv, index=False, encoding="utf-8-sig")
    print(f"\nCombined CSV saved: {out_csv}")

    export_interaction_word(comb, OUTDIR,
                            filename="Interaction_AllCohorts_Supplement.docx")

print(f"\nAll interaction outputs saved under: {os.path.abspath(OUTDIR)}")
