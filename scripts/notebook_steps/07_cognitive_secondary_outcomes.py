# Extracted from WMH_Analysis.ipynb, code cell 6.
# Notebook heading: # Cognitive secondary outcomes analysis
# Run this file from the repository root unless a local CONFIG section is edited.

# Cognitive secondary outcomes analysis
# =============================================================================
# Design:
#   - Two cohorts: Primary & Sensitivity (same CSVs as main WMH analysis).
#   - Outcomes (all standardized to SD units):
#       1) Reaction time (RT):
#            Mean_time_to_correctly_identify_matches_Instance_2
#            -> drop NaN
#            -> log transform (natural log)
#            -> z-score
#            -> multiply by -1 (higher = better)
#       2) Trail making test-B (TMT-B):
#            Duration_to_complete_alphanumeric_path_trail_2_Instance_2
#            -> drop NaN and 0
#            -> log transform
#            -> z-score
#            -> multiply by -1 (higher = better)
#       3) Fluid intelligence (FI):
#            Fluid_intelligence_score_Instance_2
#            -> treat -1 as missing (drop)
#            -> z-score
#       4) Digit span (Memory):
#            Maximum_digits_remembered_correctly_Instance_2
#            -> treat -1 as missing (drop)
#            -> z-score
#
#   - For each endpoint, missing (after these rules) are excluded *before* modeling.
#     The script prints N removed / N retained per cohort & endpoint.
#
#   - Covariates (Base model):
#       Sex, Age_at_Instance_2, Townsend_deprivation_index_at_recruitment,
#       Body_mass_index_BMI_Instance_0, Genetic_ethnic_grouping,
#       Smoking_Ever, Alcohol_intake_frequency_ordinal,
#       has_degree, e4_count
#     Categorical: Sex, Genetic_ethnic_grouping, Smoking_Ever, has_degree.
#
#   - WMH covariates (for sensitivity models):
#       Head-size normalized + log1p only (NO z-score):
#         Log_HSNorm_Total_WMH_T1_T2
#         Log_HSNorm_PeriVentricular_WMH
#         Log_HSNorm_Deep_WMH
#
#   - CMC covariate:
#       Prefer CMC_score_cat (categorical), else CMC_score_raw (continuous).
#
#   - Models (per endpoint, per WMH family):
#       Model 1: Base
#       Model 2: Base + corresponding WMH (Total / PWMH / DWMH)
#       Model 3: Base + CMC
#       Model 4: Base + corresponding WMH + CMC
#
#   - Outcomes are standardized cognitive scores, so effect size β is in SD units.
#
#   - Outputs:
#       For each cohort:
#         /Cognitive/<Cohort>/Cognitive_Results.csv
#         /Cognitive/<Cohort>/Cognitive_Results.docx  (publication-ready table)
#         /Cognitive/<Cohort>/Cognitive_Forest_3Panel.png/.pdf/.tiff
#       Root:
#         /Cognitive/Cognitive_Results_AllCohorts_Combined.csv
#         /Cognitive/Cognitive_Results_AllCohorts.xlsx (if xlsxwriter available)
#         /Cognitive/Cognitive_Results_AllCohorts.docx (combined table)
#         /Cognitive/Cognitive_Figure_Legend.docx + .txt
#
#   - Style:
#       Match main WMH analysis:
#         Arial for figures, Times New Roman 10pt for tables,
#         deep blue markers, dashed vertical zero-line,
#         top/right spines removed, 600 dpi, TIFF LZW.
# =============================================================================

import os
import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import statsmodels.api as sm
from patsy import dmatrices
from statsmodels.stats.multitest import multipletests
from datetime import datetime

# Optional Word export
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception:
    Document = None

# ======================== CONFIG ========================

ROOT_DIR = "Cognitive"
os.makedirs(ROOT_DIR, exist_ok=True)

COHORT_FILES = {
    "Primary PSM cohort":       "primary_cohort.csv",
    "PSM–sensitivity cohort":   "sensitivity_cohort.csv",
}

GROUP_COL = "group"
CLUSTER_VAR = "match_id"
GROUP_LEVELS = ["Control", "Study"]

SCALE_COL = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
RAW_TOT = "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"
RAW_PV  = "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"
RAW_DW  = "Total_volume_of_deep_white_matter_hyperintensities_Instance_2"

# Cognitive raw variables
RT_COL   = "Mean_time_to_correctly_identify_matches_Instance_2"
TM_COL   = "Duration_to_complete_alphanumeric_path_trail_2_Instance_2"
FI_COL   = "Fluid_intelligence_score_Instance_2"
MEM_COL  = "Maximum_digits_remembered_correctly_Instance_2"

# Labels
COG_ENDPOINTS = [
    ("RT_z",   "Reaction time"),
    ("TM_z",   "Trail making test-B"),
    ("FI_z",   "Fluid intelligence"),
    ("MEM_z",  "Digit span"),
]

WMH_PANELS = [
    ("Log_HSNorm_Total_WMH_T1_T2",     "Total WMH",          "A"),
    ("Log_HSNorm_PeriVentricular_WMH", "PWMH",               "B"),
    ("Log_HSNorm_Deep_WMH",            "DWMH",               "C"),
]

MODEL_SPECS = [
    ("Model 1 (Base)",              1),
    ("Model 2 (+WMH)",              2),
    ("Model 3 (+CMC)",              3),
    ("Model 4 (+WMH+CMC)",          4),
]

# Base covariates
BASE_COVARIATES = [
    "Sex",
    "Age_at_Instance_2",
    "Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "Alcohol_intake_frequency_ordinal",
    "has_degree",
    "e4_count",
]
CATEGORICAL_VARS = {
    "Sex",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "has_degree",
}

# Export style
SAVEFIG_DPI = 600
PNG_KW  = {"dpi": SAVEFIG_DPI, "bbox_inches": "tight", "facecolor": "white"}
PDF_KW  = {"bbox_inches": "tight"}
TIFF_KW = {
    "dpi": SAVEFIG_DPI,
    "bbox_inches": "tight",
    "facecolor": "white",
    "format": "tiff",
    "pil_kwargs": {"compression": "tiff_lzw"},
}

plt.rcParams.update({
    "font.family": "Arial",
    "font.size": 12,
    "axes.labelsize": 12,
    "xtick.labelsize": 11,
    "ytick.labelsize": 11,
    "savefig.dpi": SAVEFIG_DPI,
})

# ======================== HELPERS ========================

def clean_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Normalize column names to underscores."""
    df = df.copy()
    df.columns = df.columns.str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
    return df

def ensure_group(df: pd.DataFrame) -> pd.DataFrame:
    """Ensure group column exists and is a categorical with fixed levels."""
    if GROUP_COL not in df.columns:
        raise ValueError("Missing 'group' column.")
    df[GROUP_COL] = pd.Categorical(df[GROUP_COL], GROUP_LEVELS)
    return df

def build_wmh_covariates(df: pd.DataFrame) -> pd.DataFrame:
    """Create head-size normalized log1p WMH covariates (no z-score)."""
    df = df.copy()
    if SCALE_COL not in df.columns:
        raise ValueError("Missing scaling factor column for WMH.")
    scale = pd.to_numeric(df[SCALE_COL], errors="coerce")

    df["HSNorm_Total_WMH_T1_T2"] = pd.to_numeric(df[RAW_TOT], errors="coerce") * scale
    df["HSNorm_PeriVentricular_WMH"] = pd.to_numeric(df[RAW_PV], errors="coerce") * scale
    df["HSNorm_Deep_WMH"] = pd.to_numeric(df[RAW_DW], errors="coerce") * scale

    df["Log_HSNorm_Total_WMH_T1_T2"] = np.log1p(df["HSNorm_Total_WMH_T1_T2"])
    df["Log_HSNorm_PeriVentricular_WMH"] = np.log1p(df["HSNorm_PeriVentricular_WMH"])
    df["Log_HSNorm_Deep_WMH"] = np.log1p(df["HSNorm_Deep_WMH"])

    return df

def choose_cmc(df: pd.DataFrame):
    """Prefer categorical CMC, else continuous; return (varname, is_categorical)."""
    if "CMC_score_cat" in df.columns:
        return "CMC_score_cat", True
    if "CMC_score_raw" in df.columns:
        return "CMC_score_raw", False
    return None, None

def impute_covariates(df: pd.DataFrame,
                      covariates: list,
                      categoricals: set) -> pd.DataFrame:
    """Simple median/mode imputation for covariates only (not outcomes)."""
    df = df.copy()
    for v in covariates:
        if v not in df.columns:
            continue
        if v in categoricals:
            mode_vals = df[v].mode(dropna=True)
            if not mode_vals.empty:
                df[v] = df[v].fillna(mode_vals.iloc[0])
        else:
            df[v] = pd.to_numeric(df[v], errors="coerce")
            df[v] = df[v].fillna(df[v].median())
    return df

def standardize(series: pd.Series) -> pd.Series:
    """Z-score with ddof=1; returns NaN vector if sd==0."""
    s = pd.to_numeric(series, errors="coerce")
    mu = s.mean()
    sd = s.std(ddof=1)
    return (s - mu) / sd if sd > 0 else pd.Series(np.nan, index=s.index)

def prepare_cognitive_scores(df: pd.DataFrame,
                             cohort_label: str) -> pd.DataFrame:
    """Create cleaned & z-scored cognitive endpoints; print N removed/retained."""
    df = df.copy()

    # Reaction time
    rt = pd.to_numeric(df[RT_COL], errors="coerce")
    n_total = rt.notna().sum()
    rt_z = standardize(np.log(rt))
    rt_z = -rt_z   # higher = better
    n_used = rt_z.notna().sum()
    print(f"[{cohort_label}] RT: removed {n_total - n_used}, retained {n_used}")
    df["RT_z"] = rt_z

    # Trail making test-B
    tm = pd.to_numeric(df[TM_COL], errors="coerce")
    invalid = tm.isna() | (tm == 0)
    n_total = (~tm.isna()).sum()
    tm_clean = tm.mask(invalid)
    tm_z = standardize(np.log(tm_clean))
    tm_z = -tm_z
    n_used = tm_z.notna().sum()
    print(f"[{cohort_label}] TMT-B: removed {n_total - n_used} (incl. zeros), retained {n_used}")
    df["TM_z"] = tm_z

    # Fluid intelligence
    fi = pd.to_numeric(df[FI_COL], errors="coerce")
    invalid = fi.isna() | (fi == -1)
    n_total = (~fi.isna()).sum()
    fi_clean = fi.mask(invalid)
    fi_z = standardize(fi_clean)
    n_used = fi_z.notna().sum()
    print(f"[{cohort_label}] FI: removed {n_total - n_used} (incl. -1), retained {n_used}")
    df["FI_z"] = fi_z

    # Digit span
    mem = pd.to_numeric(df[MEM_COL], errors="coerce")
    invalid = mem.isna() | (mem == -1)
    n_total = (~mem.isna()).sum()
    mem_clean = mem.mask(invalid)
    mem_z = standardize(mem_clean)
    n_used = mem_z.notna().sum()
    print(f"[{cohort_label}] Digit span: removed {n_total - n_used} (incl. -1), retained {n_used}")
    df["MEM_z"] = mem_z

    return df

def build_formula(outcome, covariates, categoricals, group_var="group"):
    """Construct Patsy formula with categorical expansion for specified vars."""
    terms = [f"C({group_var})"]
    for v in covariates:
        if v == group_var:
            continue
        if v in categoricals:
            terms.append(f"C({v})")
        else:
            terms.append(v)
    return f"{outcome} ~ " + " + ".join(terms)

def find_group_coef(params_index, group_var="group", level="Study"):
    """Find the coefficient name for Study vs Control (may vary by backend)."""
    prefix = f"C({group_var})[T."
    for name in params_index:
        if name.startswith(prefix) and name.endswith(f"{level}]"):
            return name
    # fallback: first group contrast if naming differs
    for name in params_index:
        if name.startswith(prefix):
            return name
    return None

def fit_cluster_ols(formula, data, cluster_var):
    """Fit OLS with cluster-robust SE by cluster_var where available; else HC3."""
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    if y.empty:
        return None, None, 0, 0
    if cluster_var not in data.columns:
        model = sm.OLS(y, X).fit(cov_type="HC3")
        return model, "HC3 (no cluster var)", 0, len(y)
    groups = data.loc[y.index, cluster_var]
    if groups.notna().sum() < 2 or groups.nunique() < 2:
        model = sm.OLS(y, X).fit(cov_type="HC3")
        return model, "HC3 (<=1 cluster)", int(groups.nunique()), len(y)
    ok = groups.notna()
    y2, X2, g2 = y.loc[ok], X.loc[ok], groups.loc[ok]
    model = sm.OLS(y2, X2).fit(cov_type="cluster", cov_kwds={"groups": g2})
    return model, f"cluster (n_clusters={g2.nunique()})", int(g2.nunique()), len(y2)

def fit_cluster_logit(formula, data, cluster_var):
    """
    Fit logistic regression (Binomial GLM) with robust covariance.

    Outcome in `formula` must be 0/1.

    Returns
    -------
    res : GLMResults
        Fitted model with robust covariance.
    se_info : str
        Description of SE type.
    n_clusters : int
        Number of clusters used (0 if none).
    n_obs : int
        Number of observations used.
    """
    # Build design matrices
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    if y.empty:
        return None, None, 0, 0

    n_obs = int(y.shape[0])

    # Default: HC3 robust SE
    def _fit_hc3():
        model = sm.GLM(y, X, family=sm.families.Binomial())
        try:
            res = model.fit(cov_type="HC3")
        except TypeError:
            # older statsmodels without cov_type in fit
            res = model.fit()
        return res, "HC3 (no/insufficient clusters)", 0

    # If we have a cluster var, try cluster-robust
    if cluster_var in data.columns:
        groups = data.loc[y.index, cluster_var]

        if groups.notna().sum() >= 2 and groups.nunique() >= 2:
            model = sm.GLM(y, X, family=sm.families.Binomial())
            try:
                res = model.fit(
                    cov_type="cluster",
                    cov_kwds={"groups": groups}
                )
                se_info = f"cluster (n_clusters={groups.nunique()})"
                n_clusters = int(groups.nunique())
                return res, se_info, n_clusters, n_obs
            except TypeError:
                # if this statsmodels doesn't support cov_type here, fall back
                pass

        # if too few clusters or cov_type not supported → HC3 fallback
        res, se_info, n_clusters = _fit_hc3()
        return res, se_info, n_clusters, n_obs

    # No cluster var → HC3
    res, se_info, n_clusters = _fit_hc3()
    return res, se_info, n_clusters, n_obs



# ========= Word table helpers =========

def _set_cell_border(cell, **kwargs):
    """Apply thin borders selectively to a table cell (used for three-line style)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = OxmlElement(f'w:{edge}')
            for key in ("val","sz","color","space"):
                if key in edge_data:
                    tag.set(qn(f"w:{key}"), str(edge_data[key]))
            tcBorders.append(tag)

def _apply_three_line_table(table, header_row_idx=0):
    """Make a three-line academic table: header top+bottom lines, final bottom line."""
    # clear all
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                             left={"val":"nil"}, right={"val":"nil"},
                             top={"val":"nil"}, bottom={"val":"nil"})
    # header top+bottom
    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell,
                         top={"val":"single","sz":8,"color":"000000"},
                         bottom={"val":"single","sz":8,"color":"000000"})
    # bottom line
    for cell in table.rows[-1].cells:
        _set_cell_border(cell,
                         bottom={"val":"single","sz":8,"color":"000000"})

def _right_align_numeric(table, idx_list):
    """Right-align numeric-looking columns by index list (0-based)."""
    for row in table.rows[1:]:
        for j in idx_list:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def export_cog_table_word(df_all: pd.DataFrame,
                          out_path: str,
                          title: str,
                          legend: str):
    """
    Word table (no vertical merges), columns:
      Analytic Cohort | Outcome | N | Models | β (95% CI) | p

    Row layout per (Cohort, Outcome):
      - Model 1 (Base)
      - Model 2 (+CMC)                      <-- source = "Model 3 (+CMC)"
      - Model 3 (+Total WMH)                <-- source = "Model 2 (+WMH)" w/ WMH_family=Total WMH
      - Model 3 (+PWMH)                     <-- source = "Model 2 (+WMH)" w/ WMH_family=PWMH
      - Model 3 (+DWMH)                     <-- source = "Model 2 (+WMH)" w/ WMH_family=DWMH
      - Model 4 (+Total WMH + CMC)          <-- source = "Model 4 (+WMH+CMC)" per family
      - Model 4 (+PWMH + CMC)
      - Model 4 (+DWMH + CMC)

    Notes
    -----
    - Expects df_all columns: Cohort, WMH_family, Endpoint, N, Model, Beta, CI_lower, CI_upper, p
    - N 已是 "SA/Control" 形式（脚本前面已生成）
    - 不输出 q
    """
    if Document is None or df_all.empty:
        return None

    needed = {"Cohort","WMH_family","Endpoint","N","Model",
              "Beta","CI_lower","CI_upper","p"}
    miss = [c for c in needed if c not in df_all.columns]
    if miss:
        raise ValueError(f"Missing columns for export: {miss}")

    # 规范顺序
    df = df_all.copy()
    wmh_order = ["Total WMH", "PWMH", "DWMH"]

    # 小工具：按条件取一行，取不到则返回 None
    def _get_row(cohort, endpoint, model, wmh=None):
        q = (df["Cohort"]==cohort) & (df["Endpoint"]==endpoint) & (df["Model"]==model)
        if wmh is not None:
            q &= (df["WMH_family"]==wmh)
        sub = df[q]
        if sub.empty:
            return None
        return sub.iloc[0]

    def _fmt_beta_ci(beta, lo, hi):
        if any(pd.isna([beta, lo, hi])):
            return "—"
        return f"{beta:.3f} ({lo:.3f}, {hi:.3f})"

    def _fmt_p(pv):
        if pd.isna(pv): return "—"
        pv = float(pv)
        return "<0.001" if pv < 1e-3 else f"{pv:.3f}"

    # 文档与表头
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph(); r = p.add_run(title); r.bold = True

    headers = ["Analytic Cohort", "Outcome", "N",
               "Models", "β (95% CI)", "p"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    for j,h in enumerate(headers):
        tbl.rows[0].cells[j].text = h

    # 以 (Cohort, Endpoint) 为块输出
    for (cohort, endpoint), blk in df.groupby(["Cohort","Endpoint"], sort=False):
        # 这个 endpoint 的 N（任意行的 N 都相同）
        n_str = str(blk["N"].dropna().iloc[0]) if not blk["N"].dropna().empty else ""

        # ---- Model 1 (Base)
        r1 = _get_row(cohort, endpoint, "Model 1 (Base)")
        row = tbl.add_row().cells
        row[0].text = str(cohort)
        row[1].text = str(endpoint)
        row[2].text = n_str
        row[3].text = "Model 1 (Base)"
        if r1 is not None:
            row[4].text = _fmt_beta_ci(r1["Beta"], r1["CI_lower"], r1["CI_upper"])
            row[5].text = _fmt_p(r1["p"])
        else:
            row[4].text = "—"; row[5].text = "—"

        # ---- Model 2 (+CMC)  ← 源自你结果里的 “Model 3 (+CMC)”
        r2 = _get_row(cohort, endpoint, "Model 3 (+CMC)")
        row = tbl.add_row().cells
        row[0].text = str(cohort)
        row[1].text = str(endpoint)
        row[2].text = n_str
        row[3].text = "Model 2 (+CMC)"
        if r2 is not None:
            row[4].text = _fmt_beta_ci(r2["Beta"], r2["CI_lower"], r2["CI_upper"])
            row[5].text = _fmt_p(r2["p"])
        else:
            row[4].text = "—"; row[5].text = "—"

        # ---- Model 3：分别 + 三种 WMH（来自 “Model 2 (+WMH)”）
        for fam in wmh_order:
            r3 = _get_row(cohort, endpoint, "Model 2 (+WMH)", fam)
            row = tbl.add_row().cells
            row[0].text = str(cohort)
            row[1].text = str(endpoint)
            row[2].text = n_str
            row[3].text = f"Model 3 (+{fam})"
            if r3 is not None:
                row[4].text = _fmt_beta_ci(r3["Beta"], r3["CI_lower"], r3["CI_upper"])
                row[5].text = _fmt_p(r3["p"])
            else:
                row[4].text = "—"; row[5].text = "—"

        # ---- Model 4：分别 + 三种 WMH + CMC（来自 “Model 4 (+WMH+CMC)”）
        for fam in wmh_order:
            r4 = _get_row(cohort, endpoint, "Model 4 (+WMH+CMC)", fam)
            row = tbl.add_row().cells
            row[0].text = str(cohort)
            row[1].text = str(endpoint)
            row[2].text = n_str
            row[3].text = f"Model 4 (+{fam} + CMC)"
            if r4 is not None:
                row[4].text = _fmt_beta_ci(r4["Beta"], r4["CI_lower"], r4["CI_upper"])
                row[5].text = _fmt_p(r4["p"])
            else:
                row[4].text = "—"; row[5].text = "—"

    # 右对齐数值列
    _right_align_numeric(tbl, [2,4,5])
    # 三线样式（表头上下 + 全表底线）
    _apply_three_line_table(tbl, header_row_idx=0)

    # 备注
    leg = doc.add_paragraph()
    leg.add_run(legend).italic = True

    doc.save(out_path)
    return out_path




def analyze_tmtb_noncompletion(df: pd.DataFrame,
                               cohort_label: str,
                               subdir: str,
                               cmc_var: str,
                               cmc_is_cat: bool):
    """
    For a given cohort dataframe, run logistic models for TMT-B non-completion.

    Outcome:
        TMTB_noncomp = 1 if Duration_to_complete_alphanumeric_path_trail_2 == 0
                       0 if > 0
        Rows with missing TMT-B duration are excluded.

    Models per WMH family:
        Model 1 (Base)              : group + base covariates
        Model 2 (+WMH)              : Model 1 + that WMH family
        Model 3 (+CMC)              : Model 1 + CMC
        Model 4 (+WMH+CMC)          : Model 1 + that WMH family + CMC

    Effect:
        Study vs Control association with non-completion:
        log(OR), OR, 95% CI, p-value.
    """
    if TM_COL not in df.columns:
        print(f"[{cohort_label}] TMT-B column missing; skip non-completion analysis.")
        return None

    d = df.copy()
    tm_raw = pd.to_numeric(d[TM_COL], errors="coerce")

    # Define binary outcome: 1 = not completed (0 sec), 0 = completed (>0)
    d["TMTB_noncomp"] = np.where(tm_raw == 0, 1,
                          np.where(tm_raw > 0, 0, np.nan))

    # Keep rows with defined outcome and group
    d = d[~d["TMTB_noncomp"].isna()]
    d = d.dropna(subset=[GROUP_COL])
    if d.empty:
        print(f"[{cohort_label}] No TMT-B non-completion data.")
        return None

    rows = []

    for wmh_var, wmh_label, _tag in WMH_PANELS:
        if wmh_var not in d.columns:
            print(f"[{cohort_label}] Missing {wmh_var}; skip {wmh_label} in TMT-B non-completion.")
            continue

        for model_label, model_id in MODEL_SPECS:
            # Start from base covariates
            covs = [v for v in BASE_COVARIATES if v in d.columns]
            cats = set(CATEGORICAL_VARS) & set(covs)

            # +WMH terms for models 2 & 4
            if model_id in (2, 4):
                covs = covs + [wmh_var]

            # +CMC terms for models 3 & 4
            if model_id in (3, 4) and cmc_var is not None:
                covs = covs + [cmc_var]
                if cmc_is_cat:
                    cats.add(cmc_var)

            # Impute covariates (outcome and group are not imputed here)
            d_model = impute_covariates(d, covs, cats)

            # Build logistic formula
            formula = build_formula("TMTB_noncomp", covs, cats)

            # Fit clustered logistic
            res, se_info, n_clu, n_obs = fit_cluster_logit(formula, d_model, CLUSTER_VAR)
            if res is None:
                continue

            # Extract Study vs Control effect
            coef_name = find_group_coef(res.params.index, GROUP_COL, "Study")
            if coef_name is None or coef_name not in res.params.index:
                continue

            beta = float(res.params[coef_name])  # log(OR)
            ci_l, ci_u = [float(x) for x in res.conf_int().loc[coef_name]]
            pval = float(res.pvalues[coef_name])

            rows.append({
                "Cohort": cohort_label,
                "WMH_family": wmh_label,
                "Model": model_label,
                "log_OR": beta,
                "log_OR_CI_lower": ci_l,
                "log_OR_CI_upper": ci_u,
                "OR": np.exp(beta),
                "OR_CI_lower": np.exp(ci_l),
                "OR_CI_upper": np.exp(ci_u),
                "p": pval,
                "SE_type": se_info,
                "n_obs": n_obs,
                "n_clusters": n_clu,
            })

    res_df = pd.DataFrame(rows)
    if res_df.empty:
        print(f"[{cohort_label}] No TMT-B non-completion logistic results.")
        return None

    out_path = os.path.join(subdir, "TMTB_NonCompletion_Logistic_Results.csv")
    res_df.to_csv(out_path, index=False, encoding="utf-8-sig")
    print(f"[{cohort_label}] TMT-B non-completion logistic results saved: {out_path}")

    return res_df

# ========= Forest plot (3 panels; A/B/C) =========

def plot_cog_three_panel(df: pd.DataFrame,
                         cohort_label: str,
                         outdir: str):
    """
    Three-panel forest plot for cognitive outcomes by WMH family.

    Panels:
        A = Total WMH
        B = PWMH
        C = DWMH

    Endpoint display order (top → bottom), restricted to those present:
        Reaction time, Trail making test-B, Fluid intelligence, Digit span

    Legend is placed at the bottom center.
    Panel letters are fixed at the top-left corner of each subplot.
    """
    if df.empty:
        return

    # Restrict to the selected cohort
    df = df[df["Cohort"] == cohort_label].copy()
    if df.empty:
        return

    # Full desired endpoint order (labels)
    endpoint_order_full = [label for _, label in COG_ENDPOINTS]

    # Keep only endpoints that actually appear in the results for this cohort
    endpoint_order = [
        ep for ep in endpoint_order_full
        if (df["Endpoint"] == ep).any()
    ]
    if not endpoint_order:
        return

    model_order = [m for m, _ in MODEL_SPECS]
    wmh_order = [fam for _, fam, _ in WMH_PANELS]
    panel_tag = {fam: tag for _, fam, tag in WMH_PANELS}

    # --------- Build vertical positions (y) ----------
    # Each endpoint forms a block of len(model_order) rows, with a small gap between blocks.
    n_models = len(model_order)
    gap = 1.0  # vertical gap between endpoint blocks

    n_blocks = len(endpoint_order)
    rows_per_block = n_models
    total_rows = n_blocks * rows_per_block + (n_blocks - 1) * gap

    y_map = {}      # (endpoint, model) -> y
    tick_pos = {}   # endpoint -> y center

    current = total_rows
    for ep_label in endpoint_order:
        ys = []
        for model_label in model_order:
            current -= 1.0
            y_map[(ep_label, model_label)] = current
            ys.append(current)
        tick_pos[ep_label] = float(np.mean(ys))
        current -= gap

    # --------- X-axis limits (from all CIs) ----------
    xmin = float(df["CI_lower"].min())
    xmax = float(df["CI_upper"].max())
    xmin = min(xmin, 0.0)
    xmax = max(xmax, 0.0)
    span = xmax - xmin if xmax > xmin else 0.2
    pad = 0.12 * span
    xmin -= pad
    xmax += pad

    # Markers for models
    markers = {
        "Model 1 (Base)": "o",
        "Model 2 (+WMH)": "s",
        "Model 3 (+CMC)": "^",
        "Model 4 (+WMH+CMC)": "D",
    }

    fig, axes = plt.subplots(1, 3, figsize=(11, 4.0), sharey=True)

    # --------- Draw each panel ----------
    for ax, fam in zip(axes, wmh_order):
        sub = df[df["WMH_family"] == fam].copy()
        if sub.empty:
            ax.axis("off")
            continue

        # Draw in fixed endpoint/model order to guarantee alignment
        for ep_label in endpoint_order:
            for model_label in model_order:
                row = sub[
                    (sub["Endpoint"] == ep_label) &
                    (sub["Model"] == model_label)
                ]
                if row.empty:
                    continue

                r = row.iloc[0]
                beta = r["Beta"]
                lo = r["CI_lower"]
                hi = r["CI_upper"]
                y = y_map[(ep_label, model_label)]
                m = markers[model_label]

                ax.errorbar(
                    beta, y,
                    xerr=[[beta - lo], [hi - beta]],
                    fmt=m,
                    mfc="#1f3b4d",
                    mec="#1f3b4d",
                    ecolor="#1f3b4d",
                    elinewidth=1.6,
                    capsize=4,
                    markersize=5.8,
                    linestyle="none",
                    zorder=3,
                )

        # Zero line
        ax.axvline(0, color="#9e9e9e", linestyle="--", linewidth=1.0)

        # Axes style
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(0, total_rows + 0.5)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(False)

        # WMH family title (respecting global rcParams sizes)
        ax.set_title(fam, pad=8)

        # Panel letter (A/B/C) at top-left (axes coordinates)
        tag = panel_tag.get(fam, "")
        ax.text(
            0.02, 0.98,
            tag,
            transform=ax.transAxes,
            fontsize=12,
            fontweight="bold",
            va="top",
            ha="left",
        )

    # --------- Y-axis ticks & labels ----------
    # Only on the leftmost panel; positions match block centers.
    axes[0].set_yticks([tick_pos[ep] for ep in endpoint_order])
    axes[0].set_yticklabels(endpoint_order)
    axes[0].set_ylabel("Cognitive endpoints", labelpad=8)

    # Hide y tick labels on other panels, but keep shared scale
    for ax in axes[1:]:
        ax.tick_params(axis="y", labelleft=False)

    # --------- Global X label & Legend (bottom) ----------
    try:
        fig.supxlabel("β (SD units)  Study vs Control", y=0.10)
    except Exception:
        fig.text(
            0.5, 0.10,
            "β (SD units)  Study vs Control",
            ha="center",
        )

    handles = [
        plt.Line2D(
            [0], [0],
            marker=markers[m],
            linestyle="none",
            markerfacecolor="#1f3b4d",
            markeredgecolor="#1f3b4d",
            markersize=6,
            label=m,
        )
        for m in model_order
    ]

    fig.legend(
        handles=handles,
        loc="lower center",
        bbox_to_anchor=(0.5, 0.03),
        ncol=4,
        frameon=False,
    )

    # Layout: leave space for left labels and bottom legend
    plt.subplots_adjust(
        left=0.18,
        right=0.98,
        top=0.82,
        bottom=0.22,
        wspace=0.22,
    )

    # Save figure
    safe = re.sub(r"[^0-9a-zA-Z]+", "_", cohort_label).strip("_")
    base = os.path.join(outdir, f"{safe}_Cognitive_Forest_3Panel")
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)

def plot_cog_four_panel_by_endpoint(df: pd.DataFrame,
                                    cohort_label: str,
                                    outdir: str):
    """
    Four-panel forest plot by cognitive endpoint, single cohort.

    For each endpoint (one panel):
        Y-axis rows (top -> bottom):
            Base          : Model 1 (Base)
            Base + CMC    : Model 3 (+CMC)
            Total WMH     : Model 2 (+WMH) and Model 4 (+WMH+CMC) for Total WMH
            PWMH          : Model 2 (+WMH) and Model 4 (+WMH+CMC) for PWMH
            DWMH          : Model 2 (+WMH) and Model 4 (+WMH+CMC) for DWMH
    """

    if df.empty:
        return

    # Restrict to this cohort
    df = df[df["Cohort"] == cohort_label].copy()
    if df.empty:
        return

    # Endpoint order (labels from config)
    endpoint_order = [label for _, label in COG_ENDPOINTS]
    endpoint_order = [ep for ep in endpoint_order if (df["Endpoint"] == ep).any()]
    if not endpoint_order:
        return

    # WMH families
    wmh_order = [fam for _, fam, _ in WMH_PANELS]

    # X-axis limits across all models for this cohort
    xmin = float(df["CI_lower"].min())
    xmax = float(df["CI_upper"].max())
    xmin = min(xmin, 0.0)
    xmax = max(xmax, 0.0)
    span = xmax - xmin if xmax > xmin else 0.2
    pad = 0.12 * span
    xmin -= pad
    xmax += pad

    # Marker shapes by underlying model label
    markers = {
        "Model 1 (Base)": "o",
        "Model 3 (+CMC)": "^",  # displayed as Model 2 (Base + CMC)
        "Model 2 (+WMH)": "s",  # displayed as Model 3 (Base + WMH)
        "Model 4 (+WMH+CMC)": "D",
    }

    # Vertical offsets for WMH rows so two models don't overlap
    offsets = {
        "Model 1 (Base)": 0.0,
        "Model 3 (+CMC)": 0.0,
        "Model 2 (+WMH)": +0.10,
        "Model 4 (+WMH+CMC)": -0.10,
    }

    # Fixed row layout
    row_labels = ["Base", "Base + CMC", "Total WMH", "PWMH", "DWMH"]
    y_positions = {
        "Base": 5.0,
        "Base + CMC": 4.0,
        "Total WMH": 3.0,
        "PWMH": 2.0,
        "DWMH": 1.0,
    }

    # Figure and axes (2x2 panels)
    fig, axes = plt.subplots(2, 2, figsize=(11, 7), sharex=True)
    axes = axes.flatten()
    panel_letters = ["A", "B", "C", "D"]

    for ax, ep_label, panel_letter in zip(axes, endpoint_order, panel_letters):
        sub_ep = df[df["Endpoint"] == ep_label].copy()
        if sub_ep.empty:
            ax.axis("off")
            continue

        # ----- Base row: Model 1 (Base) -----
        m1 = sub_ep[sub_ep["Model"] == "Model 1 (Base)"]
        if not m1.empty:
            r = m1.iloc[0]
            beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
            y = y_positions["Base"] + offsets["Model 1 (Base)"]
            ax.errorbar(
                beta, y,
                xerr=[[beta - lo], [hi - beta]],
                fmt=markers["Model 1 (Base)"],
                mfc="#1f3b4d", mec="#1f3b4d",
                ecolor="#1f3b4d",
                elinewidth=1.5, capsize=4,
                markersize=5.5, linestyle="none", zorder=3,
            )

        # ----- Base + CMC row: underlying "Model 3 (+CMC)" -----
        m3 = sub_ep[sub_ep["Model"] == "Model 3 (+CMC)"]
        if not m3.empty:
            r = m3.iloc[0]
            beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
            y = y_positions["Base + CMC"] + offsets["Model 3 (+CMC)"]
            ax.errorbar(
                beta, y,
                xerr=[[beta - lo], [hi - beta]],
                fmt=markers["Model 3 (+CMC)"],
                mfc="#1f3b4d", mec="#1f3b4d",
                ecolor="#1f3b4d",
                elinewidth=1.5, capsize=4,
                markersize=5.5, linestyle="none", zorder=3,
            )


        # ----- WMH rows: underlying Model 2 (+WMH) and Model 4 (+WMH+CMC) -----
        for wmh_label in wmh_order:
            base_y = y_positions[wmh_label]

            # +WMH
            m2 = sub_ep[
                (sub_ep["WMH_family"] == wmh_label) &
                (sub_ep["Model"] == "Model 2 (+WMH)")
            ]
            if not m2.empty:
                r = m2.iloc[0]
                beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
                y = base_y + offsets["Model 2 (+WMH)"]
                ax.errorbar(
                    beta, y,
                    xerr=[[beta - lo], [hi - beta]],
                    fmt=markers["Model 2 (+WMH)"],
                    mfc="#1f3b4d", mec="#1f3b4d",
                    ecolor="#1f3b4d",
                    elinewidth=1.5, capsize=4,
                    markersize=5.5, linestyle="none", zorder=3,
                )

            # +WMH+CMC
            m4 = sub_ep[
                (sub_ep["WMH_family"] == wmh_label) &
                (sub_ep["Model"] == "Model 4 (+WMH+CMC)")
            ]
            if not m4.empty:
                r = m4.iloc[0]
                beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
                y = base_y + offsets["Model 4 (+WMH+CMC)"]
                ax.errorbar(
                    beta, y,
                    xerr=[[beta - lo], [hi - beta]],
                    fmt=markers["Model 4 (+WMH+CMC)"],
                    mfc="#1f3b4d", mec="#1f3b4d",
                    ecolor="#1f3b4d",
                    elinewidth=1.5, capsize=4,
                    markersize=5.5, linestyle="none", zorder=3,
                )

        # Zero line
        ax.axvline(0, color="#9e9e9e", linestyle="--", linewidth=1.0, zorder=1)

        # Y-axis labels
        ax.set_yticks([y_positions[lbl] for lbl in row_labels])
        ax.set_yticklabels(row_labels)

        # Axes style
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(0.5, 5.5)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(False)

        # Panel title and letter
        ax.set_title(ep_label, pad=6, fontweight="bold")
        ax.text(
            -0.15, 0.99,
            panel_letter,
            transform=ax.transAxes,
            fontsize=14,
            fontweight="bold",
            va="top",
            ha="right",
            clip_on=False,
        )

    # Hide unused axes if <4 endpoints
    if len(endpoint_order) < 4:
        for j in range(len(endpoint_order), 4):
            axes[j].axis("off")

    # Bottom x-label
    try:
        fig.supxlabel("β (SD units)  SA vs Control", y= -0.045)
    except Exception:
        fig.text(0.5, -0.045, "β (SD units)  SA vs Control", ha="center")

    # Legend mapping (display only)
    display_label_map = {
        "Model 1 (Base)": "Model 1 (Base)",
        "Model 3 (+CMC)": "Model 2 (Base + CMC)",
        "Model 2 (+WMH)": "Model 3 (Base + WMH)",
        "Model 4 (+WMH+CMC)": "Model 4 (Base + WMH + CMC)",
    }

    legend_order = [
        "Model 1 (Base)",
        "Model 3 (+CMC)",
        "Model 2 (+WMH)",
        "Model 4 (+WMH+CMC)",
    ]

    handles = []
    for key in legend_order:
        handles.append(
            plt.Line2D(
                [0], [0],
                marker=markers[key],
                linestyle="none",
                markerfacecolor="#1f3b4d",
                markeredgecolor="#1f3b4d",
                markersize=5.5,
                label=display_label_map[key],
            )
        )

    fig.legend(
        handles=handles,
        loc="lower center",
        bbox_to_anchor=(0.5, -0.01),
        ncol=4,
        frameon=False,
    )

    plt.subplots_adjust(
        left=0.10,
        right=0.98,
        top=0.90,
        bottom=0.24,
        wspace=0.30,
        hspace=0.30,
    )

    safe = re.sub(r"[^0-9a-zA-Z]+", "_", cohort_label).strip("_")
    base = os.path.join(outdir, f"{safe}_Cognitive_Forest_4Panel_ByEndpoint")
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)


def plot_cog_four_panel_by_endpoint_combined(df: pd.DataFrame,
                                             primary_cohort: str,
                                             sensitivity_cohort: str,
                                             outdir: str):
    """
    Four-panel forest plot by cognitive endpoint, combining
    primary and sensitivity cohorts.

    - One panel per cognitive endpoint.
    - Y-axis rows (top -> bottom):
        Base
        Base + CMC
        Total WMH
        PWMH
        DWMH
    - Primary cohort: deep blue markers.
    - Sensitivity cohort: grey markers.
    - Same marker shapes for the same model.
    - Legend layout:
        Row 1 (blue markers): model definitions.
        Row 2 (colored squares): cohort definitions.
    """

    if df.empty:
        return

    df_primary = df[df["Cohort"] == primary_cohort].copy()
    df_sens = df[df["Cohort"] == sensitivity_cohort].copy()

    if df_primary.empty or df_sens.empty:
        print("[Combined plot] Missing one of the cohorts; skip combined figure.")
        return

    # Endpoint order
    endpoint_order = [label for _, label in COG_ENDPOINTS]
    endpoint_order = [ep for ep in endpoint_order if (df["Endpoint"] == ep).any()]
    if not endpoint_order:
        return

    # WMH families
    wmh_order = [fam for _, fam, _ in WMH_PANELS]

    # X-axis limits across both cohorts
    xmin = float(df["CI_lower"].min())
    xmax = float(df["CI_upper"].max())
    xmin = min(xmin, 0.0)
    xmax = max(xmax, 0.0)
    span = xmax - xmin if xmax > xmin else 0.2
    pad = 0.12 * span
    xmin -= pad
    xmax += pad

    # Colors
    primary_color = "#1f3b4d"
    sens_color = "#b3b3b3"

    # Marker shapes (by underlying model label)
    markers = {
        "Model 1 (Base)": "o",
        "Model 3 (+CMC)": "^",
        "Model 2 (+WMH)": "s",
        "Model 4 (+WMH+CMC)": "D",
    }

    # Vertical offsets to separate +WMH and +WMH+CMC within each WMH row
    offsets = {
        "Model 1 (Base)": 0.0,
        "Model 3 (+CMC)": 0.0,
        "Model 2 (+WMH)": +0.10,
        "Model 4 (+WMH+CMC)": -0.10,
    }

    # Y rows
    row_labels = ["Base", "Base + CMC", "Total WMH", "PWMH", "DWMH"]
    y_positions = {
        "Base": 5.0,
        "Base + CMC": 4.0,
        "Total WMH": 3.0,
        "PWMH": 2.0,
        "DWMH": 1.0,
    }

    def get_row(sub_df, model_label, endpoint_label, wmh_label=None):
        q = (sub_df["Endpoint"] == endpoint_label) & (sub_df["Model"] == model_label)
        if wmh_label is not None:
            q &= (sub_df["WMH_family"] == wmh_label)
        rows = sub_df[q]
        if rows.empty:
            return None
        return rows.iloc[0]

    fig, axes = plt.subplots(2, 2, figsize=(11, 7), sharex=True)
    axes = axes.flatten()
    panel_letters = ["A", "B", "C", "D"]

    for ax, ep_label, panel_letter in zip(axes, endpoint_order, panel_letters):
        sub_p = df_primary[df_primary["Endpoint"] == ep_label].copy()
        sub_s = df_sens[df_sens["Endpoint"] == ep_label].copy()

        if sub_p.empty and sub_s.empty:
            ax.axis("off")
            continue

        # ---- Base (Model 1) ----
        for sub_df, color in ((df_sens, sens_color), (df_primary, primary_color)):
            r = get_row(sub_df, "Model 1 (Base)", ep_label)
            if r is not None:
                beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
                y = y_positions["Base"] + offsets["Model 1 (Base)"]
                ax.errorbar(
                    beta, y,
                    xerr=[[beta - lo], [hi - beta]],
                    fmt=markers["Model 1 (Base)"],
                    mfc=color, mec=color,
                    ecolor=color,
                    elinewidth=1.5, capsize=4,
                    markersize=5.5, linestyle="none", zorder=3,
                )

        # ---- Base + CMC (Model 3 (+CMC)) ----
        for sub_df, color in ((df_sens, sens_color), (df_primary, primary_color)):
            r = get_row(sub_df, "Model 3 (+CMC)", ep_label)
            if r is not None:
                beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
                y = y_positions["Base + CMC"] + offsets["Model 3 (+CMC)"]
                ax.errorbar(
                    beta, y,
                    xerr=[[beta - lo], [hi - beta]],
                    fmt=markers["Model 3 (+CMC)"],
                    mfc=color, mec=color,
                    ecolor=color,
                    elinewidth=1.5, capsize=4,
                    markersize=5.5, linestyle="none", zorder=3,
                )

        # ---- WMH rows: +WMH & +WMH+CMC ----
        for wmh_label in wmh_order:
            base_y = y_positions[wmh_label]

            # +WMH (Model 2)
            for sub_df, color in ((df_sens, sens_color), (df_primary, primary_color)):
                r = get_row(sub_df, "Model 2 (+WMH)", ep_label, wmh_label)
                if r is not None:
                    beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
                    y = base_y + offsets["Model 2 (+WMH)"]
                    ax.errorbar(
                        beta, y,
                        xerr=[[beta - lo], [hi - beta]],
                        fmt=markers["Model 2 (+WMH)"],
                        mfc=color, mec=color,
                        ecolor=color,
                        elinewidth=1.5, capsize=4,
                        markersize=5.5, linestyle="none", zorder=3,
                    )

            # +WMH + CMC (Model 4)
            for sub_df, color in ((df_sens, sens_color), (df_primary, primary_color)):
                r = get_row(sub_df, "Model 4 (+WMH+CMC)", ep_label, wmh_label)
                if r is not None:
                    beta, lo, hi = r["Beta"], r["CI_lower"], r["CI_upper"]
                    y = base_y + offsets["Model 4 (+WMH+CMC)"]
                    ax.errorbar(
                        beta, y,
                        xerr=[[beta - lo], [hi - beta]],
                        fmt=markers["Model 4 (+WMH+CMC)"],
                        mfc=color, mec=color,
                        ecolor=color,
                        elinewidth=1.5, capsize=4,
                        markersize=5.5, linestyle="none", zorder=3,
                    )

        # Axis style
        ax.axvline(0, color="#9e9e9e", linestyle="--", linewidth=1.0, zorder=1)
        ax.set_yticks([y_positions[lbl] for lbl in row_labels])
        ax.set_yticklabels(row_labels)
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(0.5, 5.5)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(False)

        ax.set_title(ep_label, pad=6, fontweight="bold")
        ax.text(
            -0.15, 0.99,
            panel_letter,
            transform=ax.transAxes,
            fontsize=14,
            fontweight="bold",
            va="top",
            ha="right",
            clip_on=False,
        )

    # Hide unused axes if <4 endpoints
    if len(endpoint_order) < 4:
        for j in range(len(endpoint_order), 4):
            axes[j].axis("off")

    # ----- Global x-label + legends (fixed vertical layout) -----

    plt.subplots_adjust(
        left=0.10,
        right=0.98,
        top=0.90,
        bottom=0.25,
        wspace=0.30,
        hspace=0.30,
    )

    # β label
    fig.text(
        0.5, -0.005,
        "β (SD units)  SA vs Control",
        ha="center",
        va="center",
        fontsize=14
    )

    # First legend: models (blue markers)
    model_handles = [
        plt.Line2D([0], [0], marker=markers["Model 1 (Base)"], linestyle="none",
                   markerfacecolor=primary_color, markeredgecolor=primary_color,
                   markersize=5.5, label="Model 1 (Base)"),
        plt.Line2D([0], [0], marker=markers["Model 3 (+CMC)"], linestyle="none",
                   markerfacecolor=primary_color, markeredgecolor=primary_color,
                   markersize=5.5, label="Model 2 (+ CMC)"),
        plt.Line2D([0], [0], marker=markers["Model 2 (+WMH)"], linestyle="none",
                   markerfacecolor=primary_color, markeredgecolor=primary_color,
                   markersize=5.5, label="Model 3 (+ WMH)"),
        plt.Line2D([0], [0], marker=markers["Model 4 (+WMH+CMC)"], linestyle="none",
                   markerfacecolor=primary_color, markeredgecolor=primary_color,
                   markersize=5.5, label="Model 4 (+ WMH + CMC)"),
    ]
    leg1 = fig.legend(
        handles=model_handles,
        loc="lower center",
        bbox_to_anchor=(0.5, -0.075),
        ncol=4,
        frameon=False,
    )
    fig.add_artist(leg1)

    # Second legend: cohorts (color blocks)
    color_handles = [
        plt.Line2D([0], [0], marker="o", linestyle="none",
                   markerfacecolor=primary_color, markeredgecolor=primary_color,
                   markersize=5.5, label="Primary PSM cohort"),
        plt.Line2D([0], [0], marker="o", linestyle="none",
                   markerfacecolor="#b3b3b3", markeredgecolor="#b3b3b3",
                   markersize=5.5, label="PSM–sensitivity cohort"),
    ]
    fig.legend(
        handles=color_handles,
        loc="lower center",
        bbox_to_anchor=(0.5, -0.11),
        ncol=2,
        frameon=False,
    )

    base = os.path.join(
        outdir,
        "Primary_Sensitivity_Combined_Cognitive_Forest_4Panel_ByEndpoint"
    )
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)



# ======================== MAIN PIPELINE ========================

all_results = []

for cohort_label, fp in COHORT_FILES.items():
    if not os.path.exists(fp):
        print(f"[Warning] File not found: {fp} (skip {cohort_label})")
        continue

    print(f"\n=== {cohort_label} ===")
    subdir = os.path.join(ROOT_DIR, re.sub(r'[^0-9A-Za-z]+', '_', cohort_label).strip("_"))
    os.makedirs(subdir, exist_ok=True)

    df = pd.read_csv(fp)
    df = clean_columns(df)
    df = ensure_group(df)
    df = build_wmh_covariates(df)
    df = prepare_cognitive_scores(df, cohort_label)

    cmc_var, cmc_is_cat = choose_cmc(df)

    # Logistic analysis for TMT-B non-completion (0 sec vs >0 sec)
    analyze_tmtb_noncompletion(df, cohort_label, subdir, cmc_var, cmc_is_cat)

    # Loop over WMH panels & models & endpoints
    rows = []
    for wmh_var, wmh_label, _tag in WMH_PANELS:
        if wmh_var not in df.columns:
            print(f"[{cohort_label}] Missing {wmh_var}; skip {wmh_label} panel.")
            continue

        for ep_col, ep_label in COG_ENDPOINTS:
            # Subset to non-missing endpoint and group
            d = df.copy()
            d = d[~d[ep_col].isna()]
            d = d.dropna(subset=[GROUP_COL])
            if d.empty:
                continue

            # ---------- NEW: N as "SA/Control" (Study/Control) per endpoint ----------
            # Counts are taken *after* endpoint filtering and group availability.
            n_sa = int((d[GROUP_COL] == "Study").sum())
            n_ctrl = int((d[GROUP_COL] == "Control").sum())
            n_str = f"{n_sa}/{n_ctrl}"

            for model_label, model_id in MODEL_SPECS:
                covs = [v for v in BASE_COVARIATES if v in d.columns]
                cats = set(CATEGORICAL_VARS) & set(covs)

                if model_id in (2, 4):
                    covs = covs + [wmh_var]
                if model_id in (3, 4) and cmc_var is not None:
                    covs = covs + [cmc_var]
                    if cmc_is_cat:
                        cats.add(cmc_var)

                d_model = impute_covariates(d, covs, cats)

                formula = build_formula(ep_col, covs, cats)
                model, se_info, n_clu, n_obs = fit_cluster_ols(formula, d_model, CLUSTER_VAR)
                if model is None:
                    continue

                coef_name = find_group_coef(model.params.index, GROUP_COL, "Study")
                if coef_name is None or coef_name not in model.params.index:
                    continue

                beta = float(model.params[coef_name])
                ci_l, ci_u = [float(x) for x in model.conf_int().loc[coef_name]]
                pval = float(model.pvalues[coef_name])

                rows.append({
                    "Cohort": cohort_label,
                    "WMH_family": wmh_label,
                    "Endpoint": ep_label,
                    "Model": model_label,
                    "N": n_str,                 # <-- keep SA/Control here
                    "Beta": beta,
                    "CI_lower": ci_l,
                    "CI_upper": ci_u,
                    "p": pval,
                    "SE_type": se_info,
                    "n_obs": n_obs,
                    "n_clusters": n_clu,
                })

    # ---- Build per-cohort results and apply FDR (PWMH/DWMH only) ----
# After collecting results for one cohort:
    res_df = pd.DataFrame(rows)
    if res_df.empty:
        print(f"[{cohort_label}] No cognitive results.")
        continue

    # ---- Save per-cohort results ----
    res_path = os.path.join(subdir, "Cognitive_Results.csv")
    res_df.to_csv(res_path, index=False, encoding="utf-8-sig")
    print(f"[{cohort_label}] CSV saved: {res_path}")

    # ---- Word table ----
    legend_text = (
        "Legend: β coefficients represent standardized mean differences in cognitive performance "
        "(Study vs Control) in SD units. Reaction time and Trail making test-B were log-transformed "
        "and multiplied by −1 so that higher values indicate better performance. "
        "Fluid intelligence and Digit span were standardized after excluding invalid codes "
        "(-1 for Fluid intelligence and Digit span; 0 for Trail making). "
        "WMH covariates are head-size–normalized and log1p-transformed (no z-score). "
        "CMC covariate uses the categorical comorbidity score when available, otherwise the continuous score. "
        "Models use cluster-robust standard errors by matched set (match_id) where feasible. "
        "N is reported as SA/Control counts after endpoint-specific exclusions."
    )
    word_path = os.path.join(subdir, "Cognitive_Results.docx")
    export_cog_table_word(res_df, word_path,
                        title=f"Table. Cognitive outcomes — {cohort_label}",
                        legend=legend_text)
    print(f"[{cohort_label}] Word table saved: {word_path}")

    # ---- Figures ----
    plot_cog_three_panel(res_df, cohort_label, subdir)
    plot_cog_four_panel_by_endpoint(res_df, cohort_label, subdir)


    all_results.append(res_df)

# ===== Combined outputs =====
if all_results:
    comb = pd.concat(all_results, ignore_index=True)

    comb_csv = os.path.join(ROOT_DIR, "Cognitive_Results_AllCohorts_Combined.csv")
    comb.to_csv(comb_csv, index=False, encoding="utf-8-sig")
    print(f"Combined CSV saved: {comb_csv}")

    # Combined Word
    comb_word = os.path.join(ROOT_DIR, "Cognitive_Results_AllCohorts_Combined.docx")
    export_cog_table_word(
        comb,
        comb_word,
        title="Table. Cognitive outcomes — Primary PSM and PSM–sensitivity cohorts",
        legend=legend_text
    )
    print(f"Combined Word table saved: {comb_word}")

    # Combined Figure
    plot_cog_four_panel_by_endpoint_combined(
        comb,
        primary_cohort="Primary PSM cohort",
        sensitivity_cohort="PSM–sensitivity cohort",
        outdir=ROOT_DIR,
    )
    print("Combined 4-panel cognitive figure saved.")


print("\nCognitive secondary endpoint analysis completed.")


# ===== Cognitive Figure 3 legend (4-panel by endpoint, combined cohorts) =====

cog_fig3_legend = (
    "Figure 3. Association of SA with cognitive outcomes in propensity-matched cohorts after additional adjustment for WMH burden and cardiovascular–metabolic conditions\n\n"
    "Forest plots show adjusted differences in standardized cognitive outcomes (β, SD units; SA vs control) for four cognitive measures in the Primary PSM cohort (dark blue) "
    "and the PSM–sensitivity cohort (light grey). Panels A–D display reaction time, Trail making test-B, fluid intelligence, and digit span, respectively. Within each panel, "
    "rows correspond to the base model (Base), the base model additionally adjusted for cardiovascular and metabolic conditions (Base + CMC), and models including total WMH, "
    "periventricular WMH (PWMH), or deep WMH (DWMH), with and without additional CMC adjustment. Circles, triangles, squares, and diamonds represent Model 1 (Base), "
    "Model 2 (+ CMC), Model 3 (+ WMH), and Model 4 (+ WMH + CMC), respectively. All cognitive outcomes were standardized using z-scores; reaction time and Trail making test-B "
    "were log-transformed before standardization, and all measures were coded so that higher values indicate better cognitive performance. The Base model was adjusted for "
    "age at Instance 2, sex, Townsend deprivation index at recruitment, body mass index at baseline, genetic ethnic grouping, smoking status, alcohol intake frequency, "
    "educational attainment (degree vs no degree), and APOE ε4 allele count. Positive β values therefore indicate better cognitive performance in participants with SA relative "
    "to matched controls. Models including CMC further adjusted for cardiovascular and metabolic conditions, and WMH models additionally included the corresponding WMH measures. "
    "All models used cluster-robust standard errors by matched set."
)

# TXT version
cog_fig3_txt_path = os.path.join(ROOT_DIR, "Cognitive_Figure3_Legend.txt")
with open(cog_fig3_txt_path, "w", encoding="utf-8") as f:
    f.write(cog_fig3_legend)
print(f"Cognitive Figure 3 legend (TXT) saved: {cog_fig3_txt_path}")

# Word version (if python-docx is available)
if Document is not None:
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    # Title line in bold
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(
        "Figure 3. Association of SA with cognitive outcomes in propensity-matched cohorts after additional adjustment for WMH burden and cardiovascular–metabolic conditions"
    )
    r_title.bold = True

    # Main legend text
    doc.add_paragraph(
        "Forest plots show adjusted differences in standardized cognitive outcomes (β, SD units; SA vs control) for four cognitive measures in the Primary PSM cohort (dark blue) "
        "and the PSM–sensitivity cohort (light grey). Panels A–D display reaction time, Trail making test-B, fluid intelligence, and digit span, respectively. Within each panel, "
        "rows correspond to the base model (Base), the base model additionally adjusted for cardiovascular and metabolic conditions (Base + CMC), and models including total WMH, "
        "periventricular WMH (PWMH), or deep WMH (DWMH), with and without additional CMC adjustment. Circles, triangles, squares, and diamonds represent Model 1 (Base), "
        "Model 2 (+ CMC), Model 3 (+ WMH), and Model 4 (+ CMC + WMH), respectively. All cognitive outcomes were standardized using z-scores; reaction time and Trail making test-B "
        "were log-transformed before standardization, and all measures were coded so that higher values indicate better cognitive performance. The Base model was adjusted for "
        "age at Instance 2, sex, Townsend deprivation index at recruitment, body mass index at baseline, genetic ethnic grouping, smoking status, alcohol intake frequency, "
        "educational attainment (degree vs no degree), and APOE ε4 allele count. Positive β values therefore indicate better cognitive performance in participants with SA relative "
        "to matched controls. Models including CMC further adjusted for cardiovascular and metabolic conditions, and WMH models additionally included the corresponding WMH measures. "
        "All models used cluster-robust standard errors by matched set."
    )

    # Abbreviations line
    doc.add_paragraph(
        "SA = sleep apnea; WMH = white matter hyperintensities; PWMH = periventricular WMH; "
        "DWMH = deep WMH; PSM = propensity score matching; CMC = cardiovascular and metabolic conditions."
    )

    cog_fig3_docx_path = os.path.join(ROOT_DIR, "Cognitive_Figure3_Legend.docx")
    doc.save(cog_fig3_docx_path)
    print(f"Cognitive Figure 3 legend (Word) saved: {cog_fig3_docx_path}")
else:
    print("Cognitive Figure 3 legend (Word) not saved: python-docx not available.")
