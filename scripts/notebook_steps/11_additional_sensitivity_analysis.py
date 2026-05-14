# Extracted from WMH_Analysis.ipynb, code cell 10.
# Notebook heading: # Additional Sensitivity Analysis Pipeline
# Run this file from the repository root unless a local CONFIG section is edited.

# Additional Sensitivity Analysis Pipeline
"""
This pipeline performs additional sensitivity analyses of WMH outcomes
using three alternative cohorts:
1. ATO Symmetric exclusion (Instance 2 MRI-anchored neuro exclusion)
2. ATO No exclusion (full sample with ATO weighting, no neuro exclusion)
3. Full-sample adjusted (no weighting, covariate-adjusted regression)

Outputs
-------
- Cohort-specific OLS/WLS results (CSV)
- Forest plots (single cohort + overlay)
- Manuscript-ready Word tables (three-line format)
- Figure legends (TXT + Word)
- Supplementary combined table (3 cohorts together, with q)
"""
import os, re
import pandas as pd
import numpy as np
import statsmodels.api as sm
import matplotlib.pyplot as plt
from patsy import dmatrices
from datetime import datetime
from statsmodels.stats.multitest import multipletests

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except:
    Document=None

# ---------------- CONFIG ----------------
OUT_DIR = "Additional_Sensitivity_Analysis"
os.makedirs(OUT_DIR, exist_ok=True)

file_sets = {
    "ATO Symmetric exclusion": "ato_sensitivity_sym.csv",
    "ATO No exclusion":        "ato_sensitivity_noexclusion.csv",
    "Full-sample adjusted":    "data_processed.csv"
}

SHORT = {
    "Log_HSNorm_Total_WMH_T1_T2":"Total WMH",
    "Log_HSNorm_PeriVentricular_WMH":"Periventricular WMH",
    "Log_HSNorm_Deep_WMH":"Deep WMH"
}
PLOT_ORDER=["Total WMH","PWMH","DWMH"]

BASE_ADJ=[
    "Sex","Age_at_Instance_2","Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0","Genetic_ethnic_grouping","Smoking_Ever",
    "Alcohol_intake_frequency_ordinal"
]
CATEGORICAL={"Sex","Smoking_Ever","Genetic_ethnic_grouping"}

# === Publication style & multi-format export (match main analysis) ===
# Fonts & sizes consistent with the main analysis
PUB_RC = {
    "font.family": "Arial",
    "font.size": 12,
    "axes.labelsize": 12,
    "xtick.labelsize": 11,
    "ytick.labelsize": 11,
    "savefig.dpi": 600,
    "figure.autolayout": True,
}

# --- High-quality TIFF export defaults ---
SAVEFIG_DPI = 600
PNG_KW  = {"dpi": SAVEFIG_DPI, "bbox_inches": "tight", "facecolor": "white"}
PDF_KW  = {"bbox_inches": "tight"}
SVG_KW  = {"bbox_inches": "tight"}
TIFF_KW = {
    "dpi": SAVEFIG_DPI,
    "bbox_inches": "tight",
    "facecolor": "white",
    "format": "tiff",
    "pil_kwargs": {"compression": "tiff_lzw"}  
}

# -Cohort Name----
alias_to_final = {
    "ATO Symmetric exclusion": "ATO–restricted cohort",
    "ATO No exclusion":        "ATO–full cohort",
    "Full-sample adjusted":    "Regression-adjusted",
    "Primary 1:10 (NoNeuro)":  "Primary 1:10 (NoNeuro)",
    "Sensitivity 1:10 (WithNeuro)": "Sensitivity 1:10 (WithNeuro)",
    # passthroughs (if upstream already standardized)
    "Regression-adjusted":     "Regression-adjusted",
}

# ---------------- UTILITIES ----------------
def build_formula(outcome, adjust_vars, categorical_vars, group_var="group"):
    terms=[f"C({group_var})"]
    for v in adjust_vars:
        if v==group_var: continue
        terms.append(f"C({v})" if v in categorical_vars else v)
    return f"{outcome} ~ " + " + ".join(terms)

def find_group_term(params_index, group_var="group", target_level="Study"):
    cands=[p for p in params_index if p.startswith(f"C({group_var})[T.")]
    for p in cands:
        if p.endswith(f"[T.{target_level}]"): return p
    return None

def fit_hc3(formula, data, weight_col=None):
    """
    Fit OLS (HC3) or WLS (HC3 if weights provided).
    """
    y,X=dmatrices(formula,data,return_type="dataframe",NA_action="drop")
    if weight_col and weight_col in data.columns:
        w = data.loc[y.index, weight_col]
        model=sm.WLS(y,X,weights=w).fit(cov_type="HC3")
    else:
        model=sm.OLS(y,X).fit(cov_type="HC3")
    return model,y,X

def plot_overlay_forest(res_dict, outdir, basename):
    """
    Overlay forest plot across three cohorts with publication-ready styling.

    Changes in this version:
    - Unify display names to the finalized style:
        "ATO–restricted cohort", "ATO–full cohort", "Regression-adjusted"
      (An alias map keeps backward-compatibility for old keys like
       "ATO Symmetric exclusion", "ATO No exclusion", "Full-sample adjusted".)
    - Remove top/right spines; legend on the right center without frame.
    - Apply global PUB_RC for consistent typography.
    """
    if not res_dict:
        return

    # Apply global publication style (Arial, sizes, dpi, etc.)
    plt.rcParams.update(PUB_RC)

    # ---- (1) Normalize incoming keys to the finalized display names ----
    # Backward-compatible alias mapping from OLD -> NEW (display) names
    alias_to_final = {
        "ATO Symmetric exclusion": "ATO–restricted cohort",
        "ATO No exclusion":        "ATO–full cohort",
        "Full-sample adjusted":    "Regression-adjusted",
        # passthroughs if already standardized:
        "ATO–restricted cohort":   "ATO–restricted cohort",
        "ATO–full cohort":         "ATO–full cohort",
        "Regression-adjusted":     "Regression-adjusted",
    }
    # Build a standardized copy so downstream uses only the NEW names
    std_res = {}
    for k, df in res_dict.items():
        std_name = alias_to_final.get(k, k)  # default: passthrough if unknown
        std_res[std_name] = df

    # Finalized order for plotting and legend
    labels_order = ["ATO–restricted cohort", "ATO–full cohort", "Regression-adjusted"]

    # ---- (2) Visual mappings (colors/markers) keyed by NEW names ----
    colors = {
        "ATO–restricted cohort": "#1f3b4d",  # deep blue (primary)
        "ATO–full cohort":       "#6e6e6e",  # neutral gray
        "Regression-adjusted":   "#d95f02"   # accent orange
    }
    markers = {
        "ATO–restricted cohort": "o",
        "ATO–full cohort":       "s",
        "Regression-adjusted":   "^"
    }

    # ---- (3) Figure + axis ----
    fig, ax = plt.subplots(figsize=(7, 4))

    # Map outcomes to positions using your global PLOT_ORDER
    order_map = {v: i for i, v in enumerate(PLOT_ORDER)}
    y = np.arange(len(PLOT_ORDER))[::-1]
    pad_y = 0.45
    jitter = 0.18  # horizontal stacking offset

    # ---- (4) Draw each cohort as a jittered series ----
    for i, lbl in enumerate(labels_order):
        if lbl not in std_res:
            continue
        res = std_res[lbl].sort_values("Outcome", key=lambda s: s.map(order_map))
        # Map raw region names to display labels for y-axis
        region_map = {
            "Total": "Total WMH",
            "Periventricular": "PWMH",
            "Deep": "DWMH"
        }
        res["Outcome"] = res["Outcome"].replace(region_map)
        y_shift = y - jitter * (i - 1)

        ax.errorbar(
            res["% Change"], y_shift,
            xerr=[res["% Change"] - res["% CI Lower"],
                  res["% CI Upper"] - res["% Change"]],
            fmt=markers[lbl], ms=5.0,
            # Fill only the primary series; others hollow for contrast
            mfc=("white" if lbl != "ATO–restricted cohort" else colors[lbl]),
            mec=colors[lbl], ecolor=colors[lbl],
            elinewidth=1.4, capsize=4, label=lbl
        )

    # Reference line at zero
    ax.axvline(0, color="grey", ls="--", lw=1)

    # Y axis: categories and padding
    ax.set_yticks(y, PLOT_ORDER)
    ax.set_ylim(y.min() - pad_y, y.max() + pad_y)

    # X axis label (kept as previously used text)
    ax.set_xlabel("% change (SA - Control)")

    # Cosmetics: remove top/right spines for a modern look
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # ---- (5) Legend on the right, vertically centered, no frame ----
    handles, labels = ax.get_legend_handles_labels()
    by_label = dict(zip(labels, handles))
    ordered = [by_label[l] for l in labels_order if l in by_label]

    ax.legend(
        ordered, labels_order,
        frameon=False,
        loc="center left",
        bbox_to_anchor=(1.02, 0.5),
        fontsize=8,
        handlelength=1.5
    )

    # Tight layout + safe filename
    plt.tight_layout()
    safe = re.sub(r"[^0-9a-zA-Z]+", "_", basename).strip("_")
    base = os.path.join(outdir, safe)

    # Save high-resolution outputs
    plt.savefig(base + ".png",  **PNG_KW)
    plt.savefig(base + ".pdf",  **PDF_KW)
    plt.savefig(base + ".svg",  **SVG_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)   # 新增高品质 TIFF (600 dpi, LZW)
    plt.close(fig)

def export_table1_word(df: pd.DataFrame, analysis_label: str, outdir: str, timestamp: str = None):
    """
    Export a 'Table 1' style baseline table to Word (.docx).

    Parameters
    ----------
    df : pandas.DataFrame
        Aggregated baseline table produced by make_table1(...).
        Typical columns: ["Analysis","Characteristic","Study","Control", ... optionally p_value / SMD]
    analysis_label : str
        Label used in the table title and file name.
    outdir : str
        Output directory.
    timestamp : str, optional
        Timestamp for the file name. If None, a timestamp will be generated.

    Returns
    -------
    str or None
        Path to the generated .docx file, or None if python-docx is unavailable.
    """
    if Document is None:
        print("[Warn] python-docx not available; skipping Word export.")
        return None

    os.makedirs(outdir, exist_ok=True)
    if timestamp is None:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")

    safe_label = re.sub(r"[^0-9a-zA-Z]+", "_", analysis_label).strip("_")
    word_path = os.path.join(outdir, f"Table1_Baseline_{safe_label}_{timestamp}.docx")

    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"Table Sx. Baseline characteristics — {analysis_label}")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Table header
    cols = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = str(c)

    # Table rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(cols):
            val = row[c] if c in row else ""
            if isinstance(val, (int, float, np.floating)):
                cl = str(c).lower()
                if cl in {"p", "p_value", "p-value"}:
                    cells[j].text = "<0.001" if float(val) < 0.001 else f"{float(val):.3f}"
                elif cl in {"smd"}:
                    cells[j].text = f"{float(val):.3f}"
                else:
                    cells[j].text = f"{float(val):.2f}"
            else:
                cells[j].text = "" if pd.isna(val) else str(val)

    # Right-align numeric columns
    numeric_idx = [j for j, c in enumerate(cols) if pd.api.types.is_numeric_dtype(df[c])]
    for row in table.rows[1:]:
        for j in numeric_idx:
            for p in row.cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Optional styling hooks (safe if helpers are absent)
    try:
        apply_three_line_table(table, header_row_idx=0)
    except Exception:
        pass
    try:
        set_table_font(table, font_name="Times New Roman", font_size=10)
    except Exception:
        pass

    # Legend
    legend = doc.add_paragraph()
    note = legend.add_run(
        "Legend: Values are mean (SD), median [IQR], or n (%). "
        "p-values from Welch’s t test (continuous), Mann–Whitney U (alcohol intake), "
        "and χ² or Fisher’s exact test (categorical). "
        "Study = sleep apnea; Control = matched controls."
    )
    note.italic = True

    doc.save(word_path)
    return word_path


# ---------------- MAIN ----------------
timestamp=datetime.now().strftime("%Y%m%d-%H%M%S")
res_by_label={}

for label,file in file_sets.items():
    print(f"\n=== {label} ===")
    df=pd.read_csv(file)
    df.columns=df.columns.str.replace(r"[^0-9a-zA-Z]+","_",regex=True)

    if "group" not in df.columns:
        if "treatment_var" in df.columns:
            df["group"]=df["treatment_var"].map({0:"Control",1:"Study"})
        else:
            continue

    scale=pd.to_numeric(df.get("Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"),errors="coerce")
    df["HSNorm_Deep_WMH"]=pd.to_numeric(df.get("Total_volume_of_deep_white_matter_hyperintensities_Instance_2"),errors="coerce")*scale
    df["HSNorm_PeriVentricular_WMH"]=pd.to_numeric(df.get("Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"),errors="coerce")*scale
    df["HSNorm_Total_WMH_T1_T2"]=pd.to_numeric(df.get("Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"),errors="coerce")*scale
    df["Log_HSNorm_Deep_WMH"]=np.log1p(df["HSNorm_Deep_WMH"])
    df["Log_HSNorm_PeriVentricular_WMH"]=np.log1p(df["HSNorm_PeriVentricular_WMH"])
    df["Log_HSNorm_Total_WMH_T1_T2"]=np.log1p(df["HSNorm_Total_WMH_T1_T2"])

    present_adj=[v for v in BASE_ADJ if v in df.columns]
    categorical_in=set([v for v in present_adj if v in CATEGORICAL])

    dmod=df.copy()
    for v in present_adj:
        if pd.api.types.is_numeric_dtype(dmod[v]):
            dmod[v]=pd.to_numeric(dmod[v],errors="coerce").fillna(dmod[v].median())
        else:
            mode=dmod[v].mode(dropna=True)
            dmod[v]=dmod[v].fillna(mode.iloc[0] if not mode.empty else dmod[v])

    outcomes=["Log_HSNorm_Total_WMH_T1_T2","Log_HSNorm_PeriVentricular_WMH","Log_HSNorm_Deep_WMH"]

    rows=[]
    for oc in outcomes:
        if oc not in dmod.columns: continue
        fml=build_formula(oc,BASE_ADJ,categorical_in)

        # Use ATO weights if available
        weight_col=None
        if label in ["ATO Symmetric exclusion","ATO No exclusion"]:
            weight_col="ato_weight_norm"

        model,y,X=fit_hc3(fml,dmod,weight_col=weight_col)
        key=find_group_term(model.params.index,"group","Study")
        if key is None: continue
        beta=float(model.params[key])
        ci_l,ci_u=[float(x) for x in model.conf_int().loc[key]]
        pval=float(model.pvalues[key])
        pct,lo,hi=100*(np.exp(beta)-1),100*(np.exp(ci_l)-1),100*(np.exp(ci_u)-1)
        rows.append({"Outcome":SHORT[oc],"% Change":pct,"% CI Lower":lo,"% CI Upper":hi,"p_value":pval})

    res=pd.DataFrame(rows)
    if not res.empty:
        mask=res["Outcome"].isin(["Periventricular WMH","Deep WMH"])
        if mask.any():
            rej,qv,_,_=multipletests(res.loc[mask,"p_value"],method="fdr_bh")
            res.loc[mask,"q_value"]=qv
        res.loc[res["Outcome"]=="Total WMH",["q_value"]]=np.nan
    new_label = alias_to_final.get(label, label)
    res_by_label[new_label] = res
    plot_single_forest(res,label,OUT_DIR)
    safe=re.sub(r"[^0-9a-zA-Z]+","_",label)
    res.to_csv(os.path.join(OUT_DIR,f"{safe}_Results.csv"),index=False,encoding="utf-8-sig")

# Overlay figure
plot_overlay_forest(res_by_label,OUT_DIR,f"Overlay_Forest_Sensitivity")

# -------- Combine 5 cohorts into one publication-ready table --------
def make_five_cohort_table(all_results: dict, outdir: str) -> pd.DataFrame:
    cohort_order = [
        "Primary 1:10 (NoNeuro)",
        "Sensitivity 1:10 (WithNeuro)",
        "ATO–restricted cohort",
        "ATO–full cohort",
        "Regression-adjusted"
    ]

    outcome_order = ["Total WMH", "Periventricular WMH", "Deep WMH"]

    rows = []
    for outcome in outcome_order:
        for c in cohort_order:
            if c not in all_results or all_results[c].empty:
                continue
            df = all_results[c]

            row = df[df["Outcome"] == outcome]
            if row.empty:
                continue

            r = row.iloc[0]
            rows.append({
                "Outcome": outcome,
                "Cohort": c,
                "%Δ (95% CI)": f"{r['% Change']:.1f}% ({r['% CI Lower']:.1f}, {r['% CI Upper']:.1f})",
                "p": "<0.001" if r["p_value"] < 0.001 else f"{r['p_value']:.3f}",
                "q": "—" if pd.isna(r.get("q_value")) else (
                    "<0.001" if r["q_value"] < 0.001 else f"{r['q_value']:.3f}"
                )
            })

    df_out = pd.DataFrame(rows)

    # ---- 排序 ----
    if not df_out.empty:
        df_out["Outcome"] = pd.Categorical(df_out["Outcome"], categories=outcome_order, ordered=True)
        df_out["Cohort"] = pd.Categorical(df_out["Cohort"], categories=cohort_order, ordered=True)
        df_out = df_out.sort_values(["Outcome", "Cohort"]).reset_index(drop=True)

    # Save CSV
    csv_path = os.path.join(outdir, "Combined_FiveCohort_Table.csv")
    df_out.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"[OK] Five-cohort CSV saved: {csv_path}")

    # Save Word
    if Document is not None and not df_out.empty:
        doc = Document()
        title = doc.add_paragraph()
        r = title.add_run("Table Sx. Combined sensitivity analyses of WMH (five cohorts)")
        r.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        cols = ["Outcome", "Cohort", "%Δ (95% CI)", "p", "q"]
        t = doc.add_table(rows=1, cols=len(cols))
        hdr = t.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = c

        for _, row in df_out.iterrows():
            cells = t.add_row().cells
            for j, c in enumerate(cols):
                cells[j].text = str(row[c])

        # numeric right align
        for row in t.rows[1:]:
            for j in [2, 3, 4]:
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        apply_three_line_table(t, header_row_idx=0)

        legend = doc.add_paragraph()
        legend.add_run(
            "Legend: Adjusted between-group differences in log1p WMH expressed as percent change "
            "(Study vs Control) with 95% confidence intervals. Cohorts include Primary (main analysis), "
            "Matched Sensitivity (with neuro), ATO Symmetric exclusion, ATO No exclusion, and Regression-adjusted. "
            "Models adjusted for prespecified covariates; cluster-robust or HC3 SEs applied. "
            "FDR (q) is applied only within secondary outcomes (Periventricular, Deep); "
            "Total WMH is primary outcome and not FDR-adjusted."
        ).italic = True

        word_path = os.path.join(outdir, "Combined_FiveCohort_Table.docx")
        doc.save(word_path)
        print(f"[OK] Five-cohort Word saved: {word_path}")

    return df_out


# ---------------- RUN COMBINED 5-COHORT TABLE ----------------
try:
    # === Load already computed Primary & Sensitivity results ===
    primary_file = "Primary_1_10_NoNeuro_LOG1P_results_gatekept.csv"
    sensitivity_file = "Sensitivity_1_10_WithNeuro_LOG1P_results_gatekept.csv"

    if os.path.exists(primary_file):
        df_p = pd.read_csv(primary_file)
        res_by_label["Primary 1:10 (NoNeuro)"] = df_p

    if os.path.exists(sensitivity_file):
        df_s = pd.read_csv(sensitivity_file)
        res_by_label["Sensitivity 1:10 (WithNeuro)"] = df_s

    # Run five-cohort table
    combined5 = make_five_cohort_table(res_by_label, OUT_DIR)
    if not combined5.empty:
        print(combined5.head(10))  
except Exception as e:
    print(f"Five-cohort combined table failed: {e}")

# === Baseline Table Export for Supplementary Sensitivity Cohorts ===
"""
This block generates baseline descriptive tables for the three supplementary
sensitivity cohorts (ATO–restricted, ATO–full, Regression-adjusted).
Output style matches the main Table 1 (three-line table, right-aligned numeric).
Exports: CSV + Word (Table Sx).
"""

SUPP_OUT_DIR = os.path.join(OUT_DIR, "Supplementary_Baseline")
os.makedirs(SUPP_OUT_DIR, exist_ok=True)

supp_file_sets = {
    "ATO–restricted cohort": "ato_sensitivity_sym.csv",
    "ATO–full cohort":       "ato_sensitivity_noexclusion.csv",
    "Regression-adjusted":   "data_processed.csv"
}

all_supp_tbls = []

for label, file in supp_file_sets.items():
    if not os.path.exists(file):
        print(f"[Skip] File not found: {file}")
        continue

    print(f"\n=== Baseline table for {label} ===")
    df = pd.read_csv(file)
    df.columns = df.columns.str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)

    # Ensure group variable
    if "group" not in df.columns:
        if "treatment_var" in df.columns:
            df["group"] = df["treatment_var"].map({0: "Control", 1: "Study"})
        else:
            print(f"[Skip] Missing group variable in {file}")
            continue

    # Impute covariates (same rules as main)
    present_adj = [v for v in BASE_ADJ if v in df.columns]
    dmod = df.copy()
    for v in present_adj:
        if pd.api.types.is_numeric_dtype(dmod[v]):
            dmod[v] = pd.to_numeric(dmod[v], errors="coerce").fillna(dmod[v].median())
        else:
            mode_vals = dmod[v].mode(dropna=True)
            dmod[v] = dmod[v].fillna(mode_vals.iloc[0] if not mode_vals.empty else dmod[v])

    dmod = dmod.dropna(subset=["group"])

    # --- Build Table 1 using the same helper ---
    tbl = make_table1(dmod, analysis_label=label, id_col="Participant_ID")

    # Save CSV
    safe = re.sub(r"[^0-9a-zA-Z]+", "_", label).strip("_")
    csv_path = os.path.join(SUPP_OUT_DIR, f"{safe}_Table1_Baseline.csv")
    tbl.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"[OK] Saved CSV: {csv_path}")

    # Save Word
    docx_path = export_table1_word(tbl, analysis_label=label,outdir=SUPP_OUT_DIR)
    if docx_path:
        print(f"[OK] Saved Word: {docx_path}")

    all_supp_tbls.append(tbl)

# --- Combine all three supplementary cohorts into one Word table ---
if all_supp_tbls and Document is not None:
    df_combined = pd.concat(all_supp_tbls, ignore_index=True)

    doc = Document()
    r = doc.add_paragraph().add_run("Table Sx. Baseline characteristics — Supplementary sensitivity cohorts")
    r.bold = True
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    cols = ["Analysis", "Characteristic", "Study", "Control"]
    if TABLE1_USE_PVALUE: cols.append("p_value")
    if TABLE1_INCLUDE_SMD: cols.append("SMD")

    t = doc.add_table(rows=1, cols=len(cols))
    hdr = t.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c

    for _, row in df_combined[cols].iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(cols):
            val = row[c] if c in row else ""
            if c == "p_value" and isinstance(val, float):
                cells[j].text = "<0.001" if val < 0.001 else f"{val:.3f}"
            else:
                cells[j].text = "" if pd.isna(val) else str(val)

    # Right-align numeric
    num_idx = list(range(2, len(cols)))
    _right_align_numeric(t, num_idx)

    # Apply three-line + font
    if THREELINE_TABLES:
        apply_three_line_table(t)
    set_table_font(t, font_name="Times New Roman", font_size=10)

    legend = doc.add_paragraph()
    legend.add_run(
        "Legend: Values are mean (SD), median [IQR], or n (%). "
        "p-values from Welch’s t test (continuous), Mann–Whitney U (alcohol intake), "
        "and χ² or Fisher’s exact test (categorical). "
        "Study = sleep apnea; Control = matched controls."
    ).italic = True

    combined_docx = os.path.join(SUPP_OUT_DIR, "TableSx_Baseline_Supplementary.docx")
    doc.save(combined_docx)
    print(f"[OK] Combined Word table saved: {combined_docx}")
