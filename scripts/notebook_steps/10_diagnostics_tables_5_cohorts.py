# Extracted from WMH_Analysis.ipynb, code cell 9.
# Notebook heading: # Diagnostics Tables for 5 Cohorts (with ATO weights support)
# Run this file from the repository root unless a local CONFIG section is edited.

# Diagnostics Tables for 5 Cohorts (with ATO weights support)
# Output: eTable 1 (fit + tests), eTable 2 (VIF)
# Format: landscape, three-line table, Times New Roman 10pt
# ============================================================

import os, re
import numpy as np
import pandas as pd
import statsmodels.api as sm
from patsy import dmatrices
from pathlib import Path

# ---- diagnostics ----
from statsmodels.stats.diagnostic import het_breuschpagan, het_white
from statsmodels.stats.stattools import jarque_bera
from statsmodels.stats.outliers_influence import OLSInfluence, variance_inflation_factor

# ---- Word export ----
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

# ============================================================
# Paths & Settings
# ============================================================

OUT_ROOT = Path("Diagnostics_Summary")
SUPP_DIR = OUT_ROOT / "Supplement"
SUPP_DIR.mkdir(parents=True, exist_ok=True)

COHORTS = {
    "Primary": "primary_cohort.csv",
    "PSM–sensitivity": "sensitivity_cohort.csv",
    "ATO–restricted": "ato_sensitivity_sym.csv",
    "ATO–full": "ato_sensitivity_noexclusion.csv",
    "Fully adjusted": "data_processed.csv"
}

COHORT_ORDER = list(COHORTS.keys())

GROUP_COL = "group"
MATCH_ID  = "match_id"

SCALE_I2 = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
RAW_TOT  = "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"
RAW_PV   = "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"
RAW_DEEP = "Total_volume_of_deep_white_matter_hyperintensities_Instance_2"

OUTCOME_BUILD = [
    ("Log_HSNorm_Total_WMH_T1_T2",     RAW_TOT,  "Total WMH"),
    ("Log_HSNorm_PeriVentricular_WMH", RAW_PV,   "PWMH"),
    ("Log_HSNorm_Deep_WMH",            RAW_DEEP, "DWMH"),
]
OUTCOME_ORDER = ["Total WMH","PWMH","DWMH"]

BASE_COVARS = [
    "Sex","Age_at_Instance_2","Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0","Genetic_ethnic_grouping",
    "Smoking_Ever","Alcohol_intake_frequency_ordinal"
]
CATEGORICAL = {"Sex","Genetic_ethnic_grouping","Smoking_Ever"}

# ============================================================
# Helpers
# ============================================================

def clean_cols(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns=lambda c: re.sub(r"_{2,}", "_",
           re.sub(r"[^0-9a-zA-Z]+","_", str(c))).strip("_"))

def ensure_outcomes(df: pd.DataFrame) -> pd.DataFrame:
    d = df.copy()
    if GROUP_COL in d.columns:
        d[GROUP_COL] = pd.Categorical(d[GROUP_COL], ["Control","Study"])
    scale = pd.to_numeric(d.get(SCALE_I2), errors="coerce").fillna(1.0)
    for y, src, _lab in OUTCOME_BUILD:
        v = pd.to_numeric(d.get(src), errors="coerce")
        d[y] = np.log1p(v * scale)
    return d

def build_formula(y_col: str, df: pd.DataFrame) -> str:
    rhs = ["C(group)"]
    for v in BASE_COVARS:
        if v not in df.columns: continue
        rhs.append(f"C({v})" if v in CATEGORICAL else v)
    return f"{y_col} ~ " + " + ".join(rhs)

def fit_cluster_or_hc3(formula: str, data: pd.DataFrame):
    """
    Fit OLS with cluster-robust SE or WLS if weight column is present.
    """
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    weights = data.loc[y.index, "ato_weight"] if "ato_weight" in data.columns else None
    if weights is not None:
        base = sm.WLS(y, X, weights=weights).fit(cov_type="HC3")
    else:
        base = sm.OLS(y, X).fit(cov_type="HC3")

    if MATCH_ID not in data.columns: 
        return base, y, X, 0
    g = data.loc[y.index, MATCH_ID].dropna()
    if g.nunique() <= 1: 
        return base, y, X, int(g.nunique())

    if weights is not None:
        fit = sm.WLS(y.loc[g.index], X.loc[g.index], weights=weights.loc[g.index]).fit(
            cov_type="cluster", cov_kwds={"groups": g})
    else:
        fit = sm.OLS(y.loc[g.index], X.loc[g.index]).fit(
            cov_type="cluster", cov_kwds={"groups": g})
    return fit, y.loc[g.index], X.loc[g.index], int(g.nunique())

def tests_bp_white(fit):
    resid, exog = fit.resid, fit.model.exog
    try: bp_stat, bp_p, _, _ = het_breuschpagan(resid, exog)
    except: bp_stat, bp_p = np.nan, np.nan
    try: wh_stat, wh_p, _, _ = het_white(resid, exog)
    except: wh_stat, wh_p = np.nan, np.nan
    return bp_stat, bp_p, wh_stat, wh_p

def test_jb(fit):
    jb_stat, jb_p, _, _ = jarque_bera(fit.resid)
    return jb_stat, jb_p

def compute_vif_df(X: pd.DataFrame) -> pd.DataFrame:
    cols = [c for c in X.columns if c.lower() != "intercept"]
    Xv = X[cols].copy()
    out=[]
    for i,c in enumerate(Xv.columns):
        try: v = variance_inflation_factor(Xv.values, i)
        except: v = np.nan
        out.append({"Term": c, "VIF": float(v) if np.isfinite(v) else np.nan})
    return pd.DataFrame(out)

def condition_number(X: pd.DataFrame) -> float:
    cols = [c for c in X.columns if c.lower() != "intercept"]
    Xm = X[cols].to_numpy(dtype=float)
    try:
        s = np.linalg.svd(Xm, compute_uv=False)
        s = s[s>0]
        return float((s.max() / s.min())) if s.size>0 else np.nan
    except: return np.nan

# ============================================================
# Word export (three-line, landscape)
# ============================================================

def _apply_normal_style(doc):
    style = doc.styles["Normal"]; style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)

def _set_cell_border(cell, **kwargs):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        tag = 'w:' + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); tcBorders.append(el)
        for k in ["val","sz","space","color"]:
            if k in attrs: el.set(qn('w:'+k), str(attrs[k]))

def apply_three_line_table(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top={"val":"none"}, bottom={"val":"none"},
                             left={"val":"none"}, right={"val":"none"})
    for cell in table.rows[0].cells:
        _set_cell_border(cell, top={"val":"single","sz":"20","color":"000000"},
                               bottom={"val":"single","sz":"15","color":"000000"})
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom={"val":"single","sz":"20","color":"000000"})

def write_docx_three_line(path_docx, title, df, legend_text):
    doc = Document(); _apply_normal_style(doc)
    # landscape
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = new_width, new_height

    p = doc.add_paragraph(); run=p.add_run(title); run.bold=True
    table = doc.add_table(rows=1, cols=len(df.columns))
    for j,c in enumerate(df.columns): 
        header = "Analytic cohort" if c=="Cohort" else str(c)
        table.rows[0].cells[j].text = header
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j,c in enumerate(df.columns):
            val = row[c]
            if isinstance(val, float) and not np.isnan(val):
                cells[j].text = f"{val:.4g}"
            else:
                cells[j].text = "" if pd.isna(val) else str(val)
    apply_three_line_table(table)
    doc.add_paragraph(legend_text).italic = True
    doc.save(path_docx)

# ============================================================
# Main
# ============================================================

all_fit_rows = []
all_vif_rows = []

for label, fp in COHORTS.items():
    if not os.path.exists(fp):
        print(f"[Skip] {label}: not found")
        continue
    df = ensure_outcomes(clean_cols(pd.read_csv(fp)))

    if GROUP_COL not in df.columns:
        if "treatment_var" in df.columns:
            df[GROUP_COL] = df["treatment_var"].map({0:"Control",1:"Study"})
        elif "SA" in df.columns:
            df[GROUP_COL] = df["SA"].map({0:"Control",1:"Study"})
        else:
            raise ValueError(f"{label}: no group col")

    for v in BASE_COVARS:
        if v not in df.columns: continue
        if pd.api.types.is_numeric_dtype(df[v]):
            df[v] = pd.to_numeric(df[v], errors="coerce")
            if df[v].notna().any():
                df[v] = df[v].fillna(df[v].median())
        else:
            m = df[v].mode(dropna=True)
            if not m.empty: df[v] = df[v].fillna(m.iloc[0])

    for y_col, _src, lab in OUTCOME_BUILD:
        if y_col not in df.columns: continue
        formula = build_formula(y_col, df)
        fit, y, X, n_clusters = fit_cluster_or_hc3(formula, df)

        jb_stat, jb_p = test_jb(fit)
        bp_stat, bp_p, wh_stat, wh_p = tests_bp_white(fit)
        cooks = OLSInfluence(fit).cooks_distance[0]

        # Table 1
        all_fit_rows.append({
            "Cohort": label, "Outcome": lab,
            "N": int(fit.nobs), "Clusters": n_clusters,
            "R^2": float(fit.rsquared), "Adj R^2": float(fit.rsquared_adj),
            "AIC": float(fit.aic), "BIC": float(fit.bic),
            "SE type": fit.cov_type + (f" (clusters={n_clusters})" if fit.cov_type=="cluster" else ""),
            "JB χ²": jb_stat, "JB p": jb_p,
            "BP χ²": bp_stat, "BP p": bp_p,
            "White χ²": wh_stat, "White p": wh_p,
            "κ (condition number)": condition_number(X),
            "Max Cook’s D": float(np.nanmax(cooks))
        })

        # Table 2
        vif_df = compute_vif_df(X)
        if label == "Primary":
            vif_df.insert(0, "Cohort", label)
            vif_df.insert(1, "Outcome", lab)
            vif_df["κ (condition number)"] = np.where(
                vif_df.index==0, condition_number(X), np.nan)
            all_vif_rows.append(vif_df)
        else:
            all_vif_rows.append(pd.DataFrame([{
                "Cohort": label, "Outcome": lab,
                "Term": "Summary", 
                "VIF": vif_df["VIF"].max(),
                "κ (condition number)": condition_number(X)
            }]))

# ============================================================
# Export
# ============================================================

df_fit = pd.DataFrame(all_fit_rows)
df_fit["Outcome"] = pd.Categorical(df_fit["Outcome"], categories=OUTCOME_ORDER, ordered=True)
df_fit["Cohort"] = pd.Categorical(df_fit["Cohort"], categories=COHORT_ORDER, ordered=True)
df_fit = df_fit.sort_values(["Cohort","Outcome"])
write_docx_three_line(
    SUPP_DIR / "eTable_1_ModelFit_AssumptionTests.docx",
    "eTable 1. Model fit and assumption tests across analytic cohorts",
    df_fit,
    ("Reported metrics include N, clusters, R², adjusted R², AIC, BIC, "
     "standard error type, Jarque–Bera normality test, Breusch–Pagan and White heteroscedasticity tests, "
     "condition number κ, and maximum Cook’s distance.")
)

df_vif = pd.concat(all_vif_rows, ignore_index=True)
df_vif["Outcome"] = pd.Categorical(df_vif["Outcome"], categories=OUTCOME_ORDER, ordered=True)
df_vif["Cohort"] = pd.Categorical(df_vif["Cohort"], categories=COHORT_ORDER, ordered=True)
df_vif = df_vif.sort_values(["Cohort","Outcome"])
write_docx_three_line(
    SUPP_DIR / "eTable_2_VIF_and_Kappa.docx",
    "eTable 2. Multicollinearity diagnostics (VIF and condition number κ)",
    df_vif,
    ("For the Primary cohort, VIFs are reported for each covariate; for other cohorts, "
     "only the maximum VIF and condition number κ are shown. Values >5–10 suggest potential collinearity, "
     "and κ>30 indicates serious collinearity.")
)

print("Diagnostics tables saved under", SUPP_DIR.resolve())
