# Extracted from WMH_Analysis.ipynb, code cell 2.
# Notebook heading: #Cleaning & Final Dataset Preparation
# Run this file from the repository root unless a local CONFIG section is edited.

#Cleaning & Final Dataset Preparation
"""
UK Biobank Cleaning & Final Dataset Preparation (Submission-Ready)
-----------------------------------------------------------------
- Reads:  data.csv
- Writes: data_processed.csv (final analysis dataset)
          Process/Derived_Variables_UKB.docx / .csv
          Process/eMethod_DataProcessing.docx
          Process/Sample_Flow.txt

Pipeline
1) Normalize column names (spaces & all special characters -> underscores).
2) Derive variables (overwrite if present):
   - Age_at_Instance_2
   - Smoking_Ever
   - Alcohol_intake_frequency_ordinal
3) Impute missing values:
   - Continuous -> median (Age_at_Instance_2, Alcohol_intake_frequency_ordinal)
   - Categorical/Binary -> mode (Smoking_Ever)
4) Add Group labels from treatment_var (Study/Control, overwrite if present).
5) Exclude participants with missing WMH outcome
   ('Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2').
6) Save final dataset and generate Word documentation via python-docx.
"""

from pathlib import Path
from textwrap import dedent
import pandas as pd
import numpy as np
import re

# --- Word generation (python-docx) ---
# pip install python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except Exception as e:
    raise SystemExit(
        "python-docx is required for Word outputs. "
        "Install with: pip install python-docx\n"
        f"Import error: {e}"
    )

# =========================
# Configuration
# =========================
INPUT_FILE   = "data.csv"
OUTPUT_FINAL = "data_processed.csv"
PROCESS_DIR  = Path("Process")
PROCESS_DIR.mkdir(parents=True, exist_ok=True)

# Raw WMH outcome column name (as it appears in the original CSV, BEFORE cleaning).
WMH_COL_RAW = "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"

# NEW: APOE calls file (after column-name cleaning we will merge it)
APOE_FILE = "APOE_calls.csv"
APOE_KEEP_COLS = ["IID", "e4_count", "e2_count", "APOE_genotype", "APOE_e4_carrier", "APOE_e2_carrier"]

# SA diagnosis date source column (as it appears in the original CSV, BEFORE cleaning).
SA_DATE_RAW = "Date_G47_first_reported_sleep_disorders"


# =========================
# Helpers: data processing
# =========================
def clean_column_names(df: pd.DataFrame) -> pd.DataFrame:
    """Normalize all column names to snake_case with underscores."""
    cols = (
        df.columns
          .str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
          .str.replace(r"_{2,}", "_", regex=True)
          .str.strip("_")
    )
    out = df.copy()
    out.columns = cols
    return out

def derive_variables(df: pd.DataFrame) -> pd.DataFrame:
    """Create/overwrite Age_at_Instance_2, Smoking_Ever, Alcohol_intake_frequency_ordinal."""
    out = df.copy()

    col_i2_date  = "Date_of_attending_assessment_centre_Instance_2"  # UKB Field 53 (I2)
    col_yob      = "Year_of_birth"                                   # UKB Field 34
    col_smoke_i2 = "Smoking_status_Instance_2"                        # UKB Field 20160 (I2)
    col_alc_i2   = "Alcohol_intake_frequency_Instance_2"             # UKB Field 1558  (I2)

    # Age_at_Instance_2 = year(Date I2) - Year of birth
    if col_i2_date in out.columns and col_yob in out.columns:
        out[col_i2_date] = pd.to_datetime(out[col_i2_date], errors="coerce")
        out["Age_at_Instance_2"] = out[col_i2_date].dt.year - out[col_yob]
    else:
        out["Age_at_Instance_2"] = np.nan

    # Smoking_Ever from coded Smoking_status_Instance_2
    # Codes: -3=Prefer not to answer (missing), 0=Never, 1=Previous, 2=Current
    if col_smoke_i2 in out.columns:
        s = out[col_smoke_i2].replace(-3, np.nan)
        out["Smoking_Ever"] = s.apply(
            lambda v: 1 if pd.notna(v) and v in [1, 2]
            else (0 if pd.notna(v) and v == 0 else np.nan)
        )
    else:
        out["Smoking_Ever"] = np.nan

    # Alcohol_intake_frequency_ordinal (ordinal mapping; -3 -> missing)
    freq_map = {6:0, 5:1, 4:2, 3:3, 2:4, 1:5}
    if col_alc_i2 in out.columns:
        a = out[col_alc_i2].replace(-3, np.nan)
        out["Alcohol_intake_frequency_ordinal"] = a.map(freq_map)
    else:
        out["Alcohol_intake_frequency_ordinal"] = np.nan

    return out

def merge_apoe_calls(df_main: pd.DataFrame,
                     apoe_path: str,
                     keep_cols: list[str]) -> tuple[pd.DataFrame, dict]:
    """
    Merge APOE_calls.csv into the main dataframe (LEFT JOIN).
    - Main key: Participant_ID  (main file)
    - APOE key: IID             (APOE file)
    - Keep only the five requested APOE columns.
    - All keys coerced to string for safe join.

    Returns:
        df_merged, stats_dict
    """
    stats = {"apoe_rows": 0, "apoe_unique_iid": 0, "merged_on": "Participant_ID ← IID",
             "matched_n": 0, "unmatched_in_main": 0}

    apoe = pd.read_csv(apoe_path)
    # Clean APOE column names to be consistent with main cleaning
    apoe.columns = (
        apoe.columns
            .str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
            .str.replace(r"_{2,}", "_", regex=True)
            .str.strip("_")
    )

    # Ensure required columns exist
    missing = [c for c in keep_cols if c not in apoe.columns]
    if missing:
        raise KeyError(f"APOE file missing columns: {missing}. Present={list(apoe.columns)}")

    apoe = apoe[keep_cols].copy()

    # Key alignment: main uses Participant_ID; APOE uses IID
    apoe["IID"] = apoe["IID"].astype(str)
    out = df_main.copy()
    if "Participant_ID" not in out.columns:
        raise KeyError("Main dataframe lacks 'Participant_ID' needed for APOE merge.")
    out["Participant_ID"] = out["Participant_ID"].astype(str)

    # Handle duplicated IID in APOE (keep first; log)
    apoe_dedup = apoe.drop_duplicates(subset=["IID"], keep="first")
    stats["apoe_rows"] = len(apoe)
    stats["apoe_unique_iid"] = apoe_dedup["IID"].nunique()
    if len(apoe_dedup) < len(apoe):
        print(f"[APOE] Duplicated IID detected: {len(apoe) - len(apoe_dedup)} rows removed (keeping first).")

    # Left join
    out = out.merge(apoe_dedup.rename(columns={"IID": "Participant_ID"}),
                    on="Participant_ID", how="left")

    # Merge stats
    matched_mask = out[["e4_count","e2_count","APOE_genotype","APOE_e4_carrier","APOE_e2_carrier"]].notna().any(axis=1)
    stats["matched_n"] = int(matched_mask.sum())
    stats["unmatched_in_main"] = int(len(out) - stats["matched_n"])

    # Write a small report
    rep = (
        f"APOE merge report\n"
        f"- Source file: {apoe_path}\n"
        f"- Join: Participant_ID (main) ← IID (APOE)\n"
        f"- APOE rows: {stats['apoe_rows']}; unique IID: {stats['apoe_unique_iid']}\n"
        f"- Main rows: {len(df_main)}\n"
        f"- Matched in main: {stats['matched_n']}\n"
        f"- Unmatched in main (APOE NA): {stats['unmatched_in_main']}\n"
        f"- Kept columns: {keep_cols[1:]}\n"
    )
    (PROCESS_DIR / "APOE_Merge_Report.txt").write_text(rep, encoding="utf-8")
    print(rep.strip())
    return out, stats


def derive_has_degree(df: pd.DataFrame) -> pd.DataFrame:
    """
    Derive binary has_degree from UK Biobank Qualifications (Data-Coding 100305, multi-select).

    Logic:
        - The qualifications field is stored as a multi-select string
          (e.g. "1", "2|3|5", "1|2|3", "-7", "-3").
        - Coding 100305:
              1  = College or University degree
              2  = A levels/AS levels or equivalent
              3  = O levels/GCSEs or equivalent
              4  = CSEs or equivalent
              5  = NVQ/HND/HNC or equivalent
              6  = Other professional qualifications
             -7  = None of the above
             -3  = Prefer not to answer

        - has_degree:
              1 if ANY code == 1
              0 if NO code == 1 AND at least one code in {2,3,4,5,6,-7}
              NaN if only -3, empty, or uninformative

    If no qualifications-like column is found, has_degree is set to NaN.
    """
    out = df.copy()

    # Try to locate the qualifications column (adjust if your export uses a fixed name)
    cand_cols = [c for c in out.columns if c.lower().startswith("qualifications")]
    if not cand_cols:
        out["has_degree"] = np.nan
        return out

    col = cand_cols[0]
    s = out[col]

    def _parse_has_degree(val):
        if pd.isna(val):
            return np.nan

        text = str(val).strip()
        if text == "":
            return np.nan

        # Split on common delimiters for multi-select encoding
        parts = re.split(r"[|,; ]+", text)
        parts = [p for p in parts if p]

        if not parts:
            return np.nan

        # Normalize tokens to integer-like strings (e.g. "1.0" -> "1")
        norm = []
        for p in parts:
            try:
                norm.append(str(int(float(p))))
            except ValueError:
                # Ignore non-numeric garbage
                continue

        if not norm:
            return np.nan

        # Any code 1 -> has degree
        if "1" in norm:
            return 1

        # No code 1, but at least one valid non-degree code -> no degree
        if any(x in {"2", "3", "4", "5", "6", "-7"} for x in norm):
            return 0

        # Only -3 (prefer not to answer) -> missing
        if all(x == "-3" for x in norm):
            return np.nan

        # Fallback: treat remaining ambiguous patterns as missing
        return np.nan

    out["has_degree"] = s.map(_parse_has_degree)
    return out






def derive_sa_flags_and_groups(df: pd.DataFrame):
    """
    Derive SA-related flags:
      - treatment_var: 1 if any ICD-10 contains 'G473' OR any self-report (20002) code == 1123; else 0
      - group        : 'Study' if treatment_var==1, else 'Control'
      - SA_ascertain_group:
            Control                      -> 'No_SA'
            Study & has_main             -> 'Hospital_Primary'
            Study & ~has_main & has_sec  -> 'Hospital_Secondary'
            Study & ~has_any             -> 'Self_Report_Only'
      - hospitalization_exposure: 1 if Diagnoses_ICD10 non-empty, else 0

    Returns:
        df_out, self_report_only_count, self_report_only_prop
    """
    out = df.copy()

    # --- ICD-10 source columns (after name cleaning) ---
    col_any_icd  = "Diagnoses_ICD10"            # pipe-separated ICD-10 codes (string)
    col_main_icd = "Diagnoses_main_ICD10"       # primary diagnosis ICD-10 (string)
    col_sec_icd  = "Diagnoses_secondary_ICD10"  # secondary diagnosis ICD-10 (string)

    # --- 1) Identify self-report SA columns (UKB Field 20002, long-form names) ---
    # Example: Non_cancer_illness_code_self_reported_Instance_0_Array_0 ... _Array_33 (Instance 0–2)
    import re
    pattern_long = re.compile(
        r"^Non_cancer_illness_code_self_reported_Instance_[0-2]_Array_([0-9]|[1-2][0-9]|3[0-3])$"
    )
    self_report_cols = [c for c in out.columns if pattern_long.match(c)]
    print(f"[SR detect] Found {len(self_report_cols)} long-name 20002 columns. "
          f"Examples: {self_report_cols[:6]}")

    # --- 2) Flag rows with self-reported SA (code 1123) robustly (numeric compare) ---
    if self_report_cols:
        block = out[self_report_cols]
        block_num = block.apply(pd.to_numeric, errors="coerce")  # "1123", "1123.0", 1123 -> 1123
        sr_has_1123 = block_num.eq(1123).any(axis=1)
        print(f"[SR detect] Rows with self-reported 1123 = {int(sr_has_1123.sum())}")
    else:
        sr_has_1123 = pd.Series(False, index=out.index)
        print("[SR detect] No 20002 columns detected; self-reported SA cannot be identified.")

    # --- 3) Flags for ICD-10 SA (G473) from any/main/secondary lists ---
    s_any  = out.get(col_any_icd,  pd.Series("", index=out.index)).fillna("").astype(str)
    s_main = out.get(col_main_icd, pd.Series("", index=out.index)).fillna("").astype(str)
    s_sec  = out.get(col_sec_icd,  pd.Series("", index=out.index)).fillna("").astype(str)

    has_any  = s_any.str.contains("G473", regex=False)
    has_main = s_main.str.contains("G473", regex=False)
    has_sec  = s_sec.str.contains("G473", regex=False)

    # --- 4) treatment_var & group ---
    out["treatment_var"] = np.where(has_any | sr_has_1123, 1, 0)
    out["group"] = np.where(out["treatment_var"].eq(1), "Study", "Control")

    # --- 5) SA ascertainment group ---
    out["SA_ascertain_group"] = pd.NA
    out.loc[out["group"] == "Control", "SA_ascertain_group"] = "No_SA"
    out.loc[(out["group"] == "Study") & has_main,
            "SA_ascertain_group"] = "Hospital_Primary"
    out.loc[(out["group"] == "Study") & has_any & (~has_main),
            "SA_ascertain_group"] = "Hospital_Secondary"
    out.loc[(out["group"] == "Study") & (~has_any),
            "SA_ascertain_group"] = "Self_Report_Only"

    # --- 6) hospitalization_exposure: Diagnoses_ICD10 non-empty -> 1, else 0
    non_empty_any_icd = s_any.str.strip().ne("")
    out["hospitalization_exposure"] = np.where(non_empty_any_icd, 1, 0)

    # --- 7) Console stats ---
    study_mask = out["group"].eq("Study")
    sro_mask   = out["SA_ascertain_group"].eq("Self_Report_Only")
    self_report_only_count = int((study_mask & sro_mask).sum())
    study_n = int(study_mask.sum())
    self_report_only_prop  = (self_report_only_count / study_n) if study_n > 0 else float("nan")

    n_hp  = int((out["SA_ascertain_group"] == "Hospital_Primary").sum())
    n_hs  = int((out["SA_ascertain_group"] == "Hospital_Secondary").sum())
    n_sro = int((out["SA_ascertain_group"] == "Self_Report_Only").sum())
    print(f"[SA ascertainment] Hospital_Primary={n_hp}, Hospital_Secondary={n_hs}, Self_Report_Only={n_sro}")
    print(f"[SA ascertainment] Self_Report_Only count = {self_report_only_count} "
          f"({self_report_only_prop:.3f} of Study)")

    return out, self_report_only_count, self_report_only_prop



def compute_cmc_variables(df: pd.DataFrame) -> pd.DataFrame:
    """
    Build CMC components (0/1), raw score (0–7), and categorical score (0/1/2),
    using diagnosis dates prior to the MRI date at Instance 2.

    Assumptions:
      - df already contains the MRI reference date column:
        "Date_of_attending_assessment_centre_Instance_2"
      - disease diagnosis columns follow the "Date_XXX_first_reported_..." naming
      - column names have been cleaned to snake_case already
    """
    out = df.copy()

    # MRI date column
    mri_col = "Date_of_attending_assessment_centre_Instance_2"
    if mri_col in out.columns:
        out[mri_col] = pd.to_datetime(out[mri_col], errors="coerce")

    # List all diagnosis-date columns we may use (exists-or-ignore)
    date_cols = [
        # MRI date (included for uniform datetime handling)
        mri_col,

        # Hypertension (I10–I15)
        "Date_I10_first_reported_essential_primary_hypertension",
        "Date_I11_first_reported_hypertensive_heart_disease",
        "Date_I12_first_reported_hypertensive_renal_disease",
        "Date_I13_first_reported_hypertensive_heart_and_renal_disease",
        "Date_I15_first_reported_secondary_hypertension",

        # Hyperlipidemia (E78)
        "Date_E78_first_reported_disorders_of_lipoprotein_metabolism_and_other_lipidaemias",

        # Arrhythmia (I47 + I48 + I49)
        "Date_I47_first_reported_paroxysmal_tachycardia",
        "Date_I48_first_reported_atrial_fibrillation_and_flutter",
        "Date_I49_first_reported_other_cardiac_arrhythmias",

        # CAD (I20–I25)
        "Date_I20_first_reported_angina_pectoris",
        "Date_I21_first_reported_acute_myocardial_infarction",
        "Date_I22_first_reported_subsequent_myocardial_infarction",
        "Date_I23_first_reported_complications_following_acute_myocardial_infarction",
        "Date_I24_first_reported_other_acute_ischaemic_heart_diseases",
        "Date_I25_first_reported_chronic_ischaemic_heart_disease",

        # Heart failure (I50)
        "Date_I50_first_reported_heart_failure",

        # Diabetes (E10–E14)
        "Date_E10_first_reported_type1_insulin_dependent_diabetes_mellitus",
        "Date_E11_first_reported_type2_non_insulin_dependent_diabetes_mellitus",
        "Date_E12_first_reported_malnutrition_related_diabetes_mellitus",
        "Date_E13_first_reported_other_specified_diabetes_mellitus",
        "Date_E14_first_reported_unspecified_diabetes_mellitus",

        # Stroke (I60–I64 only)
        "Date_I60_first_reported_subarachnoid_haemorrhage",
        "Date_I61_first_reported_intracerebral_haemorrhage",
        "Date_I62_first_reported_other_nontraumatic_intracranial_haemorrhage",
        "Date_I63_first_reported_cerebral_infarction",
        "Date_I64_first_reported_stroke_not_specified_as_haemorrhage_or_infarction",
    ]
    for c in date_cols:
        if c in out.columns:
            out[c] = pd.to_datetime(out[c], errors="coerce")

    def diagnosed_before_mri(row, candidate_cols, mri_col_name):
        """
        Return 1 if any candidate diagnosis date exists and is <= MRI date; else 0.
        If MRI date is missing, return 0 (conservative).
        """
        mri_date = row.get(mri_col_name, pd.NaT)
        if pd.isna(mri_date):
            return 0
        for cc in candidate_cols:
            if cc in row.index:
                d = row[cc]
                if (not pd.isna(d)) and (d <= mri_date):
                    return 1
        return 0

    # Component definitions
    htn_cols = [
        "Date_I10_first_reported_essential_primary_hypertension",
        "Date_I11_first_reported_hypertensive_heart_disease",
        "Date_I12_first_reported_hypertensive_renal_disease",
        "Date_I13_first_reported_hypertensive_heart_and_renal_disease",
        "Date_I15_first_reported_secondary_hypertension",
    ]
    lipid_cols = ["Date_E78_first_reported_disorders_of_lipoprotein_metabolism_and_other_lipidaemias"]
    arrhythmia_cols = [
        "Date_I47_first_reported_paroxysmal_tachycardia",
        "Date_I48_first_reported_atrial_fibrillation_and_flutter",
        "Date_I49_first_reported_other_cardiac_arrhythmias",
    ]
    cad_cols = [
        "Date_I20_first_reported_angina_pectoris",
        "Date_I21_first_reported_acute_myocardial_infarction",
        "Date_I22_first_reported_subsequent_myocardial_infarction",
        "Date_I23_first_reported_complications_following_acute_myocardial_infarction",
        "Date_I24_first_reported_other_acute_ischaemic_heart_diseases",
        "Date_I25_first_reported_chronic_ischaemic_heart_disease",
    ]
    hf_cols = ["Date_I50_first_reported_heart_failure"]
    dm_cols = [
        "Date_E10_first_reported_type1_insulin_dependent_diabetes_mellitus",
        "Date_E11_first_reported_type2_non_insulin_dependent_diabetes_mellitus",
        "Date_E12_first_reported_malnutrition_related_diabetes_mellitus",
        "Date_E13_first_reported_other_specified_diabetes_mellitus",
        "Date_E14_first_reported_unspecified_diabetes_mellitus",
    ]
    stroke_cols = [
        "Date_I60_first_reported_subarachnoid_haemorrhage",
        "Date_I61_first_reported_intracerebral_haemorrhage",
        "Date_I62_first_reported_other_nontraumatic_intracranial_haemorrhage",
        "Date_I63_first_reported_cerebral_infarction",
        "Date_I64_first_reported_stroke_not_specified_as_haemorrhage_or_infarction",
    ]

    # Compute components (0/1)
    out["CMC_hypertension"]   = out.apply(lambda r: diagnosed_before_mri(r, htn_cols,       mri_col), axis=1)
    out["CMC_hyperlipidemia"] = out.apply(lambda r: diagnosed_before_mri(r, lipid_cols,     mri_col), axis=1)
    out["CMC_arrhythmia"]     = out.apply(lambda r: diagnosed_before_mri(r, arrhythmia_cols, mri_col), axis=1)
    out["CMC_CAD"]            = out.apply(lambda r: diagnosed_before_mri(r, cad_cols,       mri_col), axis=1)
    out["CMC_heart_failure"]  = out.apply(lambda r: diagnosed_before_mri(r, hf_cols,        mri_col), axis=1)
    out["CMC_diabetes"]       = out.apply(lambda r: diagnosed_before_mri(r, dm_cols,        mri_col), axis=1)
    out["CMC_stroke"]         = out.apply(lambda r: diagnosed_before_mri(r, stroke_cols,    mri_col), axis=1)

    # Sum as raw score (0–7)
    cmc_components = [
        "CMC_hypertension","CMC_hyperlipidemia","CMC_arrhythmia",
        "CMC_CAD","CMC_heart_failure","CMC_diabetes","CMC_stroke",
    ]
    out["CMC_score_raw"] = out[cmc_components].sum(axis=1)

    # Categorical score: 0/1/2 (2 = ≥2 comorbidities)
    def _categorize_cmc(score):
        if pd.isna(score):
            return np.nan
        if score == 0:
            return 0
        elif score == 1:
            return 1
        else:
            return 2
    out["CMC_score_cat"] = out["CMC_score_raw"].apply(_categorize_cmc)

    return out

def add_sa_years_and_exclude_negatives(df: pd.DataFrame, sa_col_name: str):
    """
    Create 'Years_since_sleep_disorder' in years using a specific SA diagnosis date column:
        Years = (MRI_date - SA_diagnosis_date) / 365.25
      MRI_date: 'Date_of_attending_assessment_centre_Instance_2'
      SA_date : sa_col_name  (e.g., 'Date_G47_first_reported_sleep_disorders')

    Exclusion rule:
      - EXCLUDE only Study (treatment_var == 1) if Years < 0
      - Controls kept; if Years < 0 in controls, set Years = NaN.
    """
    out = df.copy()
    mri_col = "Date_of_attending_assessment_centre_Instance_2"

    # Ensure datetime dtype
    if mri_col in out.columns:
        out[mri_col] = pd.to_datetime(out[mri_col], errors="coerce")
    if sa_col_name in out.columns:
        out[sa_col_name] = pd.to_datetime(out[sa_col_name], errors="coerce")
    else:
        out["Years_since_sleep_disorder"] = np.nan
        print(f"[SA years] SA date column '{sa_col_name}' not present; Years_since_sleep_disorder set to all-NaN.")
        return out, 0, 0

    # Compute years only where both dates exist
    both = (~out[mri_col].isna()) & (~out[sa_col_name].isna())
    out["Years_since_sleep_disorder"] = np.nan
    if both.any():
        delta_days = (out.loc[both, mri_col] - out.loc[both, sa_col_name]).dt.days
        out.loc[both, "Years_since_sleep_disorder"] = delta_days / 365.25

    # Negative-year handling
    neg_mask     = out["Years_since_sleep_disorder"] < 0
    study_mask   = (out.get("treatment_var", pd.Series(dtype=float)) == 1)
    control_mask = (out.get("treatment_var", pd.Series(dtype=float)) == 0)

    exclude_mask = study_mask & neg_mask
    excl_sa_negative_st = int(exclude_mask.sum())
    neg_controls_kept   = int((control_mask & neg_mask).sum())

    out.loc[control_mask & neg_mask, "Years_since_sleep_disorder"] = np.nan
    out = out[~exclude_mask]

    return out, excl_sa_negative_st, neg_controls_kept



def impute_missing(df: pd.DataFrame):
    """
    Perform imputation and generate a detailed before/after missingness report.

    Returns:
        df_imputed: DataFrame after imputation
        imputation_df: DataFrame containing the imputation summary table
                       (before/after missing counts, percentages, strategy, fill value, etc.)

    Rules:
        • Continuous variables → median imputation
          (Age_at_Instance_2, Alcohol_intake_frequency_ordinal)
        • Categorical/Binary/Count variables → mode imputation
          (Smoking_Ever, has_degree, e4_count)
          - e4_count is coerced to numeric (0/1/2), then imputed by mode, kept as integer (Int64)
    """
    df_imputed = df.copy()

    # Define variables by strategy
    CONTINUOUS_MEDIAN = ["Age_at_Instance_2", "Alcohol_intake_frequency_ordinal"]
    CATEGORICAL_MODE  = ["Smoking_Ever", "has_degree", "e4_count"]  # << added e4_count

    rows = []  # rows for the imputation report

    def _miss_stat(series: pd.Series, total_n: int):
        n = int(series.isna().sum())
        pct = (n / total_n * 100.0) if total_n > 0 else 0.0
        return n, round(pct, 2)

    total_n = len(df_imputed)

    # 1) Continuous variables → median imputation
    for col in CONTINUOUS_MEDIAN:
        if col in df_imputed.columns:
            n_before, pct_before = _miss_stat(df_imputed[col], total_n)
            if df_imputed[col].notna().any():
                fill_value = float(df_imputed[col].median())
                df_imputed[col] = df_imputed[col].fillna(fill_value)
            else:
                fill_value = 0.0
                df_imputed[col] = df_imputed[col].fillna(fill_value)
            n_after, pct_after = _miss_stat(df_imputed[col], total_n)

            rows.append({
                "Variable": col,
                "Type": "Continuous",
                "Strategy": "Median",
                "Fill_value": fill_value,
                "Missing_n_before": n_before,
                "Missing_%_before": pct_before,
                "Imputed_n": n_before - n_after,
                "Missing_n_after": n_after,
                "Missing_%_after": pct_after,
            })

    # 2) Categorical/Binary/Count variables → mode imputation
    for col in CATEGORICAL_MODE:
        if col in df_imputed.columns:
            series = df_imputed[col]

            # Special handling for e4_count: coerce to numeric and keep Int64
            if col == "e4_count":
                series = pd.to_numeric(series, errors="coerce")  # coerce non-numeric → NaN
                df_imputed[col] = series  # update coerced numeric back before stats

            n_before, pct_before = _miss_stat(df_imputed[col], total_n)

            if df_imputed[col].notna().any():
                mode_series = df_imputed[col].mode(dropna=True)
                fill_value = mode_series.iloc[0] if not mode_series.empty else 0
            else:
                fill_value = 0  # all missing → fallback to 0

            df_imputed[col] = df_imputed[col].fillna(fill_value)

            # Keep e4_count as integer dtype if possible
            if col == "e4_count":
                try:
                    df_imputed[col] = df_imputed[col].astype("Int64")
                    # ensure the reported fill_value is int for consistency
                    fill_value = int(fill_value)
                except Exception:
                    pass

            n_after, pct_after = _miss_stat(df_imputed[col], total_n)

            rows.append({
                "Variable": col,
                "Type": "Categorical/Binary/Count" if col == "e4_count" else "Categorical/Binary",
                "Strategy": "Mode",
                "Fill_value": fill_value,
                "Missing_n_before": n_before,
                "Missing_%_before": pct_before,
                "Imputed_n": n_before - n_after,
                "Missing_n_after": n_after,
                "Missing_%_after": pct_after,
            })

    imputation_df = pd.DataFrame(rows, columns=[
        "Variable","Type","Strategy","Fill_value",
        "Missing_n_before","Missing_%_before",
        "Imputed_n","Missing_n_after","Missing_%_after"
    ])
    return df_imputed, imputation_df



def compute_sample_flow(df: pd.DataFrame, wmh_col_clean: str) -> dict:
    """Return sample counts and filtered df after excluding WMH-missing."""
    orig_total   = len(df)
    orig_study   = int((df.get("treatment_var", pd.Series(dtype=int)) == 1).sum())
    orig_control = int((df.get("treatment_var", pd.Series(dtype=int)) == 0).sum())

    filtered_df  = df[~df[wmh_col_clean].isna()]
    final_total  = len(filtered_df)
    final_study  = int((filtered_df.get("treatment_var", pd.Series(dtype=int)) == 1).sum())
    final_control= int((filtered_df.get("treatment_var", pd.Series(dtype=int)) == 0).sum())
    excl_total   = orig_total - final_total

    return {
        "orig_total": orig_total, "orig_study": orig_study, "orig_control": orig_control,
        "excl_total": excl_total,
        "final_total": final_total, "final_study": final_study, "final_control": final_control,
        "filtered_df": filtered_df
    }

def _group_counts(df: pd.DataFrame):
    """Return total, Study, Control counts based on treatment_var."""
    total   = len(df)
    study   = int((df.get("treatment_var", pd.Series(dtype=int)) == 1).sum())
    control = int((df.get("treatment_var", pd.Series(dtype=int)) == 0).sum())
    return total, study, control

# =========================
# Helpers: Word documents
# =========================
def _doc_apply_default_style(doc: Document):
    """Set a clean, journal-friendly default font/size."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

def export_variable_table_word_and_csv(process_dir: Path):
    """Create Derived_Variables_UKB.docx and CSV copy (now includes SA duration)."""
    rows = [
        [
            "Age_at_Instance_2",
            "Date of attending assessment centre | 53 | Instance 2; Year of birth | 34 | —",
            "Age at I2 = year(Date of attending assessment centre at Instance 2) - Year of birth.",
            "Median imputation."
        ],
        [
            "Smoking_Ever",
            "Smoking status | 20160 | Instance 2",
            "Codes: -3=Prefer not to answer (missing), 0=Never, 1=Previous, 2=Current. Smoking_Ever = 1 if code ∈ {1,2}, else 0.",
            "Mode imputation."
        ],
        [
            "Alcohol_intake_frequency_ordinal",
            "Alcohol intake frequency | 1558 | Instance 2",
            "Ordinal mapping: 6→0 (Never), 5→1, 4→2, 3→3, 2→4, 1→5; -3 treated as missing.",
            "Median imputation."
        ],
        # --- If you already added CMC rows, keep them; otherwise you can add later ---
        [
            "Years_since_sleep_disorder",
            "Date_G47_first_reported_sleep_disorders and Date of attending assessment centre | 53 | Instance 2",
            ("Years since sleep disorder diagnosis at I2 MRI: "
            "Years = (Date_of_attending_assessment_centre_Instance_2 - Date_G47_first_reported_sleep_disorders) / 365.25. "
            "Participants with negative values (<0) were excluded."),
            "No imputation (derived from dates; NaN when either date is missing)."
        ],

        [
            "has_degree",
            "Qualifications (coding 100305) | Instance 0",
            "Binary indicator for tertiary education: 1 if code = 1 (College/University degree); "
            "0 if code ∈ {2,3,4,5,6,-7}; -3 (prefer not to answer) treated as missing.",
            "Mode imputation."
        ],
    ]
    df_tbl = pd.DataFrame(rows, columns=[
        "Derived variable", "UKB source (name | Field ID | Instance)",
        "Definition / Rule", "Missing-data handling"
    ])
    df_tbl.to_csv(process_dir / "Derived_Variables_UKB.csv", index=False, encoding="utf-8")

    doc = Document()
    _doc_apply_default_style(doc)
    h = doc.add_heading("Derived Variables (UK Biobank)", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    legend = ("Legend. All derived variables were constructed from UK Biobank source fields. "
              "Missing values were imputed using the median for continuous variables and the mode for categorical variables. "
              "Date-derived variables (e.g., Years_since_sleep_disorder) were not imputed; "
              "participants with negative duration were excluded.")
    p = doc.add_paragraph(legend)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.add_table(rows=1, cols=len(df_tbl.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, col in enumerate(df_tbl.columns):
        hdr[i].text = col

    for r in rows:
        row = table.add_row().cells
        for i, val in enumerate(r):
            row[i].text = str(val)

    out_path = process_dir / "Derived_Variables_UKB.docx"
    doc.save(out_path)


def export_emethod_word(process_dir: Path, flow_summary: str, wmh_raw_name: str, imputation_df: pd.DataFrame):
    """Create eMethod_DataProcessing.docx describing the data cleaning + imputation appendix + CMC definition."""
    doc = Document()
    _doc_apply_default_style(doc)
    doc.add_heading("eMethod: Data Processing and Cleaning", level=1)

    txt = dedent(f"""
    We prepared the analysis dataset from a wide-format UK Biobank (UKB) extract as follows.

    1) Column-name normalization and Outcome completeness screen
       All column names were converted to snake_case by replacing spaces and special characters with underscores.
        Before any derivations or exclusions, participants with missing WMH outcome (“Total volume of white matter hyperintensities from T1 and T2-FLAIR images, Instance 2”) were removed to ensure a consistent analytic cohort. The number excluded at this early step was recorded.  

    2) Derived variables
       • Age_at_Instance_2 was calculated as the year of the Instance 2 assessment date (UKB Field 53) minus Year of birth (Field 34).
       • Smoking_Ever was defined from Smoking status at Instance 2 (Field 20160). Codes -3=Prefer not to answer were treated as missing; 0=Never; 1=Previous; 2=Current. Ever smokers (1 or 2) were coded as 1; never smokers as 0.
       • Alcohol_intake_frequency_ordinal was derived from Alcohol intake frequency (Field 1558) at Instance 2 using an ordinal mapping (6→0 [Never], 5→1, 4→2, 3→3, 2→4, 1→5). Code -3 was treated as missing.
       • CMC (comorbidity count) variables were constructed using first-reported diagnosis dates prior to the MRI date at Instance 2
         (Date_of_attending_assessment_centre_Instance_2). Seven components (0/1 each) were included:
           Hypertension (I10–I15), Hyperlipidemia (E78), Arrhythmia (I47/I48/I49), Coronary artery disease (I20–I25),
           Heart failure (I50), Diabetes (E10–E14), and Stroke (I60–I64).
         Raw CMC score (0–7) equals the sum of components; CMC category was 0 (none), 1 (single), or 2 (≥2).
        • Years_since_sleep_disorder was computed as
            (Date_of_attending_assessment_centre_Instance_2 - Sleep_Disorder_Diagnosis_Date) / 365.25 years.
            Participants with negative values (<0) were excluded; the number excluded is reported in the sample flow.
        • has_degree was derived from Qualifications (coding 100305): code 1 → 1; codes {2,3,4,5,6,-7} → 0; code -3 treated as missing.



    3) Missing data
       Continuous variables (Age_at_Instance_2, Alcohol_intake_frequency_ordinal) were imputed using the cohort median.
       Categorical/Binary variables (Smoking_Ever, has_degree) and the count variable e4_count (0/1/2) were imputed using the mode.
       For e4_count, values were coerced to numeric prior to imputation and kept as integer (Int64) after imputation.
       When a variable was entirely missing, a fallback of 0 was used.

    4) Group definition
       A Group column was added as Study for participants with treatment_var=1 and Control for treatment_var=0.

    5) Outcome completeness exclusion
       Participants with missing WMH outcome were excluded. The WMH outcome field used for exclusion was:
       “{wmh_raw_name}”. The same column was matched after column-name normalization.

    6) Final dataset
       The processed dataset (data_processed.csv) includes only participants with complete WMH outcomes and the derived covariates described above.

    Sample flow summary
    {flow_summary}
    """).strip()

    for para in txt.split("\n\n"):
        p = doc.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Appendix: Imputation summary (unchanged logic) ---
    doc.add_heading("Appendix: Missingness and Imputation Summary", level=2)
    if imputation_df is not None and not imputation_df.empty:
        imputation_csv = process_dir / "Imputation_Report.csv"
        imputation_df.to_csv(imputation_csv, index=False, encoding="utf-8-sig")

        cols = list(imputation_df.columns)
        table = doc.add_table(rows=1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = str(c)

        for _, r in imputation_df.iterrows():
            row = table.add_row().cells
            for i, c in enumerate(cols):
                val = r[c]
                row[i].text = str(val) if len(str(val)) <= 50 else str(val)[:47] + "..."
    else:
        doc.add_paragraph("No imputation summary available.")

    out_path = process_dir / "eMethod_DataProcessing.docx"
    doc.save(out_path)



# =========================
# Main
# =========================
def main():
    # 1) Load
    df = pd.read_csv(INPUT_FILE)

    # 2) Normalize column names
    df = clean_column_names(df)

        # --- Fix: Fill missing Genetic_ethnic_grouping as Non-White (0) ---
    # This field in UKB typically uses code 1 = White / Caucasian.
    # According to project rule, any missing value should be treated as 0 (Non-White).
    if "Genetic_ethnic_grouping" in df.columns:
        df["Genetic_ethnic_grouping"] = (
            df["Genetic_ethnic_grouping"]
            .replace("", np.nan)
            .fillna(0)
            .astype(int)
        )
        print("[Fix] Filled missing Genetic_ethnic_grouping as 0 (Non-White). "
              f"Unique values now: {df['Genetic_ethnic_grouping'].unique().tolist()}")
    else:
        print("[Warning] Column 'Genetic_ethnic_grouping' not found; skip ethnicity fix.")

    # 2.0) Merge APOE calls (after column-name cleaning; before any filtering)
    try:
        df, apoe_stats = merge_apoe_calls(df, APOE_FILE, APOE_KEEP_COLS)
        print(f"[APOE] Merge OK: matched={apoe_stats['matched_n']}, "
              f"unmatched_in_main={apoe_stats['unmatched_in_main']}")
    except FileNotFoundError:
        print(f"[APOE] WARNING: File not found: {APOE_FILE}. Skipping APOE merge.")
    except Exception as e:
        print(f"[APOE] WARNING: Merge failed: {e}. Skipping APOE merge.")


    # 2.1) Early WMH completeness filter (APPLY FIRST)
    wmh_col_clean = (
        pd.Index([WMH_COL_RAW])
          .str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
          .str.replace(r"_{2,}", "_", regex=True)
          .str.strip("_")
          .tolist()[0]
    )
    if wmh_col_clean not in df.columns:
        raise KeyError(
            f"WMH column '{WMH_COL_RAW}' not found after cleaning (expected '{wmh_col_clean}')."
        )
    n0_total = len(df)
    df = df[~df[wmh_col_clean].isna()].copy()
    excl_wmh_early = n0_total - len(df)

    # 2.2) SA flags and groups (now on WMH-complete cohort)
    df, sro_count, sro_prop = derive_sa_flags_and_groups(df)
    pre_total, pre_study, pre_control = _group_counts(df)
    print(f"[Counts | After SA flags (WMH-complete)] Study={pre_study}, Control={pre_control}, Total={pre_total}")

    # Compute the cleaned SA date column name from raw
    sa_col_clean = (
        pd.Index([SA_DATE_RAW])
        .str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
        .str.replace(r"_{2,}", "_", regex=True)
        .str.strip("_")
        .tolist()[0]
    )

    if sa_col_clean not in df.columns:
        print(f"[SA date] WARNING: '{SA_DATE_RAW}' not found after cleaning (expected '{sa_col_clean}'). "
            "Years_since_sleep_disorder will be NaN unless an alternative is supplied.")


    # 3) Derived variables
    df = derive_variables(df)
    df = derive_has_degree(df)
    df = compute_cmc_variables(df)

    # 3.3) SA years + exclude negatives (Study only)
    orig_total_pre   = len(df)
    orig_study_pre   = int((df.get("treatment_var", pd.Series(dtype=int)) == 1).sum())
    orig_control_pre = int((df.get("treatment_var", pd.Series(dtype=int)) == 0).sum())
    df, excl_sa_negative_st, neg_controls_kept = add_sa_years_and_exclude_negatives(df, sa_col_clean)
    post_sa_total, post_sa_study, post_sa_control = _group_counts(df)
    print(f"[Counts | After SA-negative exclusion] Study={post_sa_study}, Control={post_sa_control}, Total={post_sa_total}")

    # 4) Impute missing values
    df, imputation_df = impute_missing(df)

    # 5) Save final dataset (already WMH-filtered at step 2.1)
    df.to_csv(OUTPUT_FINAL, index=False, encoding="utf-8")

    # 6) Sample flow log
    flow_summary = dedent(f"""
    Early exclusion for WMH missing (applied before any SA/derivations): {excl_wmh_early}
    Counts after SA flags (on WMH-complete cohort): total={pre_total}, Study={pre_study}, Control={pre_control}
    Self-Report-Only within Study: {sro_count} ({sro_prop:.3f})
    Excluded negative SA duration in Study (Years_since_sleep_disorder < 0): {excl_sa_negative_st}
    Controls with negative SA duration kept (Years set to NaN): {neg_controls_kept}
    Final (WMH-complete + SA-negative exclusion): total={post_sa_total}, Study={post_sa_study}, Control={post_sa_control}
    """).strip()
    (PROCESS_DIR / "Sample_Flow.txt").write_text(flow_summary, encoding="utf-8")

    # 7) Documentation (Word + CSV)
    export_variable_table_word_and_csv(PROCESS_DIR)
    export_emethod_word(PROCESS_DIR, flow_summary, WMH_COL_RAW, imputation_df)

    # 8) Done
    print(f"[DONE] Final dataset saved: {OUTPUT_FINAL}")
    print("       Documentation saved in: Process/")


if __name__ == "__main__":
    main()
