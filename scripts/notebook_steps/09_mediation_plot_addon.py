# Extracted from WMH_Analysis.ipynb, code cell 8.
# Notebook heading: # Mediation Plot INCREMENTAL ADD-ON (Revised)
# Run this file from the repository root unless a local CONFIG section is edited.

# Mediation Plot INCREMENTAL ADD-ON (Revised)
# 1) Build a concise supplementary mediation table (docx) with publication-style names
# 2) Plot a 4-panel figure for indirect effects across 3 WMH mediators
#
# Revisions in this version:
#   (A) X-axis range no longer visually "forces symmetry around 0".
#       Use a negative-dominant, asymmetric padding strategy to reduce empty
#       right-side space when most effects are < 0.
#   (B) Move legend downward by ~1.5 lines and move the global x-axis label
#       downward by ~1 line distance.
# =============================================================================

import os
import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator

# Optional Word export (reuse your existing import if already loaded)
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception:
    Document = None

# ----------------------- Locate results in memory -----------------------

def _get_mediation_df_from_memory():
    """
    Try to find the combined mediation results dataframe from memory.
    Priority:
      1) comb
      2) comb_df
      3) df_all
      4) all_rows -> DataFrame
      5) fallback reading combined CSV
    """
    g = globals()
    if "comb" in g and isinstance(g["comb"], pd.DataFrame) and not g["comb"].empty:
        return g["comb"].copy()
    if "comb_df" in g and isinstance(g["comb_df"], pd.DataFrame) and not g["comb_df"].empty:
        return g["comb_df"].copy()
    if "df_all" in g and isinstance(g["df_all"], pd.DataFrame) and not g["df_all"].empty:
        return g["df_all"].copy()
    if "all_rows" in g and isinstance(g["all_rows"], list) and len(g["all_rows"]) > 0:
        return pd.DataFrame(g["all_rows"]).copy()

    # fallback to disk
    fp = os.path.join("Mediation", "Mediation_Results_AllCohorts_Combined.csv")
    if os.path.exists(fp):
        return pd.read_csv(fp)
    return pd.DataFrame()

df_med = _get_mediation_df_from_memory()
if df_med.empty:
    raise ValueError("No mediation results found in memory or on disk.")

# ----------------------- Config / orders -----------------------

OUTROOT = "Mediation"
os.makedirs(OUTROOT, exist_ok=True)

cohort_order = ["Primary PSM cohort", "PSM–sensitivity cohort"]
outcome_order = ["Reaction time", "Trail making test-B", "Fluid intelligence", "Digit span"]
mediator_order = ["Total WMH", "PWMH", "DWMH"]

# For figure aesthetics (match your style)
SAVEFIG_DPI = 600
PNG_KW  = {"dpi": SAVEFIG_DPI, "bbox_inches": "tight", "facecolor": "white"}
PDF_KW  = {"bbox_inches": "tight"}
TIFF_KW = {
    "dpi": SAVEFIG_DPI,
    "bbox_inches": "tight",
    "facecolor": "white",
    "format": "tiff",
    "pil_kwargs": {"compression": "tiff_lzw"},
}

plt.rcParams.update({
    "font.family": "Arial",
    "font.size": 12,
    "axes.labelsize": 12,
    "xtick.labelsize": 11,
    "ytick.labelsize": 11,
    "savefig.dpi": SAVEFIG_DPI,
})

# ----------------------- Choose "key" columns for Supplementary table -----------------------
# You can decide whether to label the indirect effect column as "ACME" explicitly.
USE_ACME_LABEL = False  # <- set True if you want the column heading to be "ACME"

# Formatting helpers (use same style logic as your main table)
def _fmt(x, nd=3):
    return "—" if pd.isna(x) else f"{float(x):.{nd}f}"

def _fmt_p(x):
    if pd.isna(x): return "—"
    x = float(x)
    return "<0.001" if x < 1e-3 else f"{x:.3f}"

def _fmt_ci(lo, hi, nd=3):
    if any(pd.isna([lo, hi])): return "—"
    return f"({_fmt(lo, nd)}, {_fmt(hi, nd)})"

def _fmt_beta_se(beta, se, nd=3):
    if any(pd.isna([beta, se])): return "—"
    return f"{_fmt(beta, nd)} ({_fmt(se, nd)})"

def build_supplementary_key_table(df):
    """
    Keep columns most suitable for a concise supplementary mediation table.

    Core logic:
      - Focus on exposure -> mediator -> outcome mediation results.
      - Provide a, b, indirect effect (a×b) with bootstrap CI,
        direct effect c′, total effect c, and PM (optional but useful).
      - Keep Sobel p as a traditional reference, but rely on bootstrap CI for inference.
    """
    d = df.copy()

    needed = [
        "Cohort","Outcome","Mediator","N",
        "a","a_se","a_p",
        "b","b_se","b_p",
        "ab","sobel_p",
        "ab_boot_mean","ab_ci_low","ab_ci_high",
        "cprime","cprime_se","cprime_p",
        "ctotal","ctotal_se","ctotal_p",
        "pm","pm_ci_low","pm_ci_high",
        "n_boot_ok"
    ]

    for c in needed:
        if c not in d.columns:
            d[c] = np.nan

    d["Cohort"] = pd.Categorical(d["Cohort"], categories=cohort_order, ordered=True)
    d["Outcome"] = pd.Categorical(d["Outcome"], categories=outcome_order, ordered=True)
    d["Mediator"] = pd.Categorical(d["Mediator"], categories=mediator_order, ordered=True)
    d = d.sort_values(["Cohort","Outcome","Mediator"]).reset_index(drop=True)

    indirect_name = "ACME" if USE_ACME_LABEL else "Indirect effect (a×b)"

    out = pd.DataFrame({
        "Analytic cohort": d["Cohort"],
        "Cognitive outcome": d["Outcome"],
        "WMH mediator": d["Mediator"],
        "N (SA/Control)": d["N"],

        "a: SA→WMH (β, SE)": d.apply(lambda r: _fmt_beta_se(r["a"], r["a_se"]), axis=1),
        "p_a": d["a_p"].apply(_fmt_p),

        "b: WMH→Cognition (β, SE)": d.apply(lambda r: _fmt_beta_se(r["b"], r["b_se"]), axis=1),
        "p_b": d["b_p"].apply(_fmt_p),

        f"{indirect_name} (point a×b)": d["ab"].apply(_fmt),
        "Sobel p": d["sobel_p"].apply(_fmt_p),

        f"{indirect_name} (bootstrap mean)": d["ab_boot_mean"].apply(_fmt),
        f"{indirect_name} 95% CI": d.apply(lambda r: _fmt_ci(r["ab_ci_low"], r["ab_ci_high"]), axis=1),

        "Direct effect c′ (β, SE)": d.apply(lambda r: _fmt_beta_se(r["cprime"], r["cprime_se"]), axis=1),
        "p_c′": d["cprime_p"].apply(_fmt_p),

        "Total effect c (β, SE)": d.apply(lambda r: _fmt_beta_se(r["ctotal"], r["ctotal_se"]), axis=1),
        "p_c": d["ctotal_p"].apply(_fmt_p),

        "Proportion mediated (PM)": d["pm"].apply(_fmt),
        "PM 95% CI": d.apply(lambda r: _fmt_ci(r["pm_ci_low"], r["pm_ci_high"]), axis=1),

        "Bootstrap draws used": d["n_boot_ok"].fillna("").astype(str),
    })

    return out

# ----------------------- Word three-line table helpers -----------------------

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = OxmlElement(f'w:{edge}')
            for key in ("val","sz","color","space"):
                if key in edge_data:
                    tag.set(qn(f"w:{key}"), str(edge_data[key]))
            tcBorders.append(tag)

def _apply_three_line_table(table, header_row_idx=0):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                             left={"val":"nil"}, right={"val":"nil"},
                             top={"val":"nil"}, bottom={"val":"nil"})
    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell,
                         top={"val":"single","sz":8,"color":"000000"},
                         bottom={"val":"single","sz":8,"color":"000000"})
    for cell in table.rows[-1].cells:
        _set_cell_border(cell,
                         bottom={"val":"single","sz":8,"color":"000000"})

def _right_align_numeric(table, idx_list):
    for row in table.rows[1:]:
        for j in idx_list:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def export_key_supp_word(df_key: pd.DataFrame, out_path: str, title: str, legend: str):
    if Document is None or df_key is None or df_key.empty:
        print("Supplementary Word export skipped (python-docx unavailable or empty).")
        return None

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run(title).bold = True

    headers = list(df_key.columns)
    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, r in df_key.iterrows():
        cells = t.add_row().cells
        for j, h in enumerate(headers):
            cells[j].text = str(r[h]) if pd.notna(r[h]) else "—"

    numeric_cols = []
    for j, h in enumerate(headers):
        if any(k in h for k in ["β", "SE", "p", "CI", "PM", "Bootstrap"]):
            numeric_cols.append(j)
        if h.startswith("N"):
            numeric_cols.append(j)

    numeric_cols = sorted(set(numeric_cols))
    _right_align_numeric(t, numeric_cols)
    _apply_three_line_table(t, header_row_idx=0)

    leg = doc.add_paragraph()
    leg.add_run(legend).italic = True

    doc.save(out_path)
    print(f"Supplementary Word table saved: {out_path}")
    return out_path

# ----------------------- Build and export supplementary table -----------------------

df_key = build_supplementary_key_table(df_med)

supp_legend = (
    "Legend: This supplementary table summarizes key mediation results assessing whether WMH burden "
    "mediates the association between sleep apnea (SA) and cognitive outcomes. "
    "Indirect effects were quantified as a×b with 2000 bootstrap resamples (cluster bootstrap by match_id when available). "
    "Point estimates of a and b are regression coefficients from models adjusted for age at Instance 2, sex, Townsend deprivation index, "
    "baseline BMI, genetic ethnic grouping, smoking status, alcohol intake frequency, educational attainment, and APOE ε4 allele count. "
    "Reaction time and Trail making test-B were log-transformed and multiplied by −1 prior to z-score standardization, "
    "so higher standardized outcomes indicate better performance. "
    "Abbreviations: SA = sleep apnea; WMH = white matter hyperintensities; PWMH = periventricular WMH; "
    "DWMH = deep WMH; PM = proportion mediated."
    + (" ACME denotes the average causal mediation effect (i.e., indirect effect a×b)." if USE_ACME_LABEL else "")
)

supp_path = os.path.join(OUTROOT, "Supplementary_Mediation_KeyResults.docx")
export_key_supp_word(
    df_key,
    supp_path,
    title="eTable X. Key mediation results for the association between sleep apnea and cognitive outcomes mediated by WMH burden",
    legend=supp_legend
)

# ----------------------- 4-panel figure for indirect effects -----------------------

def _compute_asymmetric_xlim(d, lo_col="ab_ci_low", hi_col="ab_ci_high"):
    """
    Compute a negative-dominant, asymmetric xlim.

    Rationale:
      - Many indirect effects are < 0 in this project.
      - Avoid visually wasting space on the positive side.
      - Still ensure 0 is included.
    """
    lo = pd.to_numeric(d[lo_col], errors="coerce").dropna()
    hi = pd.to_numeric(d[hi_col], errors="coerce").dropna()

    if lo.empty or hi.empty:
        return (-0.02, 0.01)

    xmin_data = float(lo.min())
    xmax_data = float(hi.max())

    # Ensure 0 is inside range but do NOT expand to a symmetric window.
    # Expand left more than right.
    neg_extent = abs(min(xmin_data, 0.0))
    pos_extent = max(xmax_data, 0.0)

    # Base pads
    # If positive side is tiny, give it only a small cushion.
    pad_left = 0.12 * (neg_extent if neg_extent > 0 else abs(xmin_data) if xmin_data != 0 else 1.0)
    pad_right = 0.06 * (pos_extent if pos_extent > 0 else neg_extent * 0.25 if neg_extent > 0 else 1.0)

    xmin = xmin_data - pad_left

    # Right side: keep tight, but ensure 0 is not glued to the frame edge.
    xmax_candidate = xmax_data + pad_right
    # Guarantee a small visible space beyond 0 if xmax_data <= 0
    if xmax_candidate <= 0:
        xmax_candidate = 0.0 + max(pad_right, 0.08 * neg_extent)

    xmax = xmax_candidate

    # Final safety: avoid zero-span
    if xmax <= xmin:
        xmax = xmin + 0.02

    return xmin, xmax


def plot_mediation_indirect_4panel(df: pd.DataFrame, outdir: str):
    """
    Four panels (A-D), one per cognitive outcome.
    Each panel shows indirect effect (bootstrap mean) with 95% CI
    for 3 mediators (Total WMH, PWMH, DWMH).

    Layout and style aligned with your cognitive figures.

    Revisions:
      - Asymmetric xlim (negative-dominant) to reduce empty right-side space.
      - Legend moved downward ~1.5 lines.
      - Global x-axis label moved downward ~1 line.
    """
    d = df.copy()

    for col in ["Cohort","Outcome","Mediator","ab_boot_mean","ab_ci_low","ab_ci_high"]:
        if col not in d.columns:
            raise ValueError(f"Missing column for plotting: {col}")

    d["Cohort"] = pd.Categorical(d["Cohort"], categories=cohort_order, ordered=True)
    d["Outcome"] = pd.Categorical(d["Outcome"], categories=outcome_order, ordered=True)
    d["Mediator"] = pd.Categorical(d["Mediator"], categories=mediator_order, ordered=True)
    d = d.sort_values(["Outcome","Mediator","Cohort"]).reset_index(drop=True)

    # Colors
    primary_color = "#1f3b4d"
    sens_color = "#b3b3b3"

    # Y positions
    y_labels = mediator_order
    y_pos = {m: (len(mediator_order) - i) for i, m in enumerate(mediator_order)}  # 3,2,1
    offset = {"Primary PSM cohort": +0.10, "PSM–sensitivity cohort": -0.10}

    # NEW: asymmetric x range
    xmin, xmax = _compute_asymmetric_xlim(d, "ab_ci_low", "ab_ci_high")

    fig, axes = plt.subplots(2, 2, figsize=(11, 7), sharex=True, sharey=True)
    axes = axes.flatten()
    panel_letters = ["A", "B", "C", "D"]

    for ax, outcome_label, letter in zip(axes, outcome_order, panel_letters):
        sub = d[d["Outcome"] == outcome_label].copy()
        if sub.empty:
            ax.axis("off")
            continue

        for cohort in cohort_order:
            sub_c = sub[sub["Cohort"] == cohort]
            if sub_c.empty:
                continue
            color = primary_color if cohort == "Primary PSM cohort" else sens_color

            for med in mediator_order:
                row = sub_c[sub_c["Mediator"] == med]
                if row.empty:
                    continue
                r = row.iloc[0]
                beta = r["ab_boot_mean"]
                lo = r["ab_ci_low"]
                hi = r["ab_ci_high"]

                if any(pd.isna([beta, lo, hi])):
                    continue

                y = y_pos[med] + offset.get(cohort, 0.0)

                ax.errorbar(
                    beta, y,
                    xerr=[[beta - lo], [hi - beta]],
                    fmt="o",
                    mfc=color, mec=color,
                    ecolor=color,
                    elinewidth=1.6,
                    capsize=4,
                    markersize=5.8,
                    linestyle="none",
                    zorder=3,
                )

        # zero line
        ax.axvline(0, color="#9e9e9e", linestyle="--", linewidth=1.0, zorder=1)

        # axis style
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(0.5, 3.5)

        # --- Auto reduce x-axis ticks (方案1) ---
        ax.xaxis.set_major_locator(MaxNLocator(nbins=5))
        # 如果你想更激进一点可改成 nbins=4 或 6

        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.grid(False)

        ax.set_yticks([y_pos[m] for m in y_labels])
        ax.set_yticklabels(y_labels)

        ax.set_title(outcome_label, pad=6, fontweight="bold")
        ax.text(
            -0.12, 0.99,
            letter,
            transform=ax.transAxes,
            fontsize=14,
            fontweight="bold",
            va="top",
            ha="right",
            clip_on=False,
        )

    # global x label (moved down ~1 line)
    fig.text(
        0.55, 0.11,
        "Indirect effect (a×b), Δ cognitive z-score",
        ha="center",
        fontsize=13
    )

    # legend for cohorts (moved down ~1.5 lines)
    handles = [
        plt.Line2D([0],[0], marker="o", linestyle="none",
                   markerfacecolor=primary_color, markeredgecolor=primary_color,
                   markersize=6, label="Primary PSM cohort"),
        plt.Line2D([0],[0], marker="o", linestyle="none",
                   markerfacecolor=sens_color, markeredgecolor=sens_color,
                   markersize=6, label="PSM–sensitivity cohort"),
    ]
    fig.legend(
        handles=handles,
        loc="lower center",
        bbox_to_anchor=(0.55, 0.050),
        ncol=2,
        frameon=False,
    )

    # Give a bit more bottom room to accommodate the lowered label + legend
    plt.subplots_adjust(
        left=0.18,
        right=0.98,
        top=0.92,
        bottom=0.18,
        wspace=0.25,
        hspace=0.30,
    )

    base = os.path.join(outdir, "Figure_eX_SA_WMH_Cognition_Indirect_4Panel")
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)

    print(f"4-panel indirect-effect figure saved: {base}.png/.pdf/.tiff")


plot_mediation_indirect_4panel(df_med, OUTROOT)

# ----------------------- Figure legend text (optional quick export) -----------------------

fig_legend = (
    "Figure eX. Indirect effects of sleep apnea on cognitive outcomes mediated by WMH burden.\n\n"
    "Four panels display the bootstrap-estimated indirect effects (a×b) of sleep apnea (SA) on "
    "standardized cognitive outcomes through total WMH, periventricular WMH (PWMH), and deep WMH (DWMH). "
    "Points denote bootstrap means and error bars indicate 95% bootstrap confidence intervals "
    "(2000 resamples, cluster bootstrap by match_id when available). "
    "Results are shown for the Primary PSM cohort (dark blue) and the PSM–sensitivity cohort (light grey). "
    "Cognitive outcomes were z-score standardized; reaction time and Trail making test-B were log-transformed "
    "and multiplied by −1 prior to standardization so that higher values indicate better performance. "
    "Models adjusted for age at Instance 2, sex, Townsend deprivation index at recruitment, baseline BMI, "
    "genetic ethnic grouping, smoking status, alcohol intake frequency, educational attainment, "
    "and APOE ε4 allele count.\n\n"
    "Abbreviations: SA = sleep apnea; WMH = white matter hyperintensities; PWMH = periventricular WMH; "
    "DWMH = deep WMH."
    + (" ACME = average causal mediation effect (i.e., indirect effect a×b)." if USE_ACME_LABEL else "")
)

legend_txt = os.path.join(OUTROOT, "Figure_eX_Legend.txt")
with open(legend_txt, "w", encoding="utf-8") as f:
    f.write(fig_legend)
print(f"Figure legend TXT saved: {legend_txt}")

if Document is not None:
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass
    p = doc.add_paragraph()
    r = p.add_run("Figure eX. Indirect effects of sleep apnea on cognitive outcomes mediated by WMH burden.")
    r.bold = True
    doc.add_paragraph(fig_legend.split("\n\n",1)[1] if "\n\n" in fig_legend else fig_legend)
    legend_docx = os.path.join(OUTROOT, "Figure_eX_Legend.docx")
    doc.save(legend_docx)
    print(f"Figure legend Word saved: {legend_docx}")
