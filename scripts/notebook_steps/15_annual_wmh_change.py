# Extracted from WMH_Analysis.ipynb, code cell 14.
# Notebook heading: # Annual WMH Change (I3–I2) · Primary + Sensitivity Cohorts · Final (Revised & FDR Applied)
# Run this file from the repository root unless a local CONFIG section is edited.

# Annual WMH Change (I3–I2) · Primary + Sensitivity Cohorts · Final (Revised & FDR Applied)
# Key points:
# - HSN log1p (la table): Benjamini–Hochberg FDR applied only to Periventricular & Deep; Total excluded (q=NaN).
# - Inc/Dec (proportion tests): FDR applied only to Periventricular & Deep; Total excluded. In plots, Total has no q.
# - 2-panel and single-region plots: plotting order Total → Periventricular → Deep; Total shows p only (as primary).

import os, re, warnings
warnings.filterwarnings("ignore", category=RuntimeWarning)

import numpy as np
import pandas as pd
import matplotlib as mpl
import matplotlib.pyplot as plt
from scipy import stats
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.diagnostic import het_breuschpagan
from statsmodels.stats.stattools import jarque_bera
from statsmodels.stats.multitest import multipletests
from statsmodels.stats.proportion import proportions_ztest, confint_proportions_2indep

# ---------------- Basic Configuration ----------------
COHORTS = {
    "Primary":     "primary_cohort.csv",
    "Sensitivity": "sensitivity_cohort.csv",
}
OUTDIR = "Annual_Change"
os.makedirs(OUTDIR, exist_ok=True)

mpl.rcParams.update({
    "figure.dpi": 120,
    "savefig.dpi": 600,
    "font.family": "Arial",
    "font.size": 12,
    "axes.labelsize": 12,
    "xtick.labelsize": 11,
    "ytick.labelsize": 11,
    "figure.autolayout": True,
})

# Publication style rcParams
PUB_RC = {
    "font.family": "Arial",
    "font.size": 12,
    "axes.labelsize": 12,
    "xtick.labelsize": 11,
    "ytick.labelsize": 11,
    "savefig.dpi": 600,
    "figure.autolayout": True,
}

# File export kwargs
SAVEFIG_DPI = 600
PNG_KW  = {"dpi": SAVEFIG_DPI, "bbox_inches": "tight", "facecolor": "white"}
PDF_KW  = {"bbox_inches": "tight"}
SVG_KW  = {"bbox_inches": "tight"}
TIFF_KW = {
    "dpi": SAVEFIG_DPI,
    "bbox_inches": "tight",
    "facecolor": "white",
    "format": "tiff",
    "pil_kwargs": {"compression": "tiff_lzw"},
}

# Colors / markers consistent with main analysis
COLOR_PRIMARY = "#1f3b4d"  # deep blue
ELW, CAPSIZE, MS = 1.6, 4, 5.5

# Abbreviated labels & plotting order
DISPLAY_LABEL = {"Total": "Total WMH", "Periventricular": "PWMH", "Deep": "DWMH"}
PLOT_ORDER = ["Total", "Periventricular", "Deep"]

def safe_name(s):
    return re.sub(r"[^0-9A-Za-z._-]+", "_", str(s))

def fmt_p(p):
    try:
        p = float(p)
    except Exception:
        return str(p)
    return f"{p:.1e}" if p < 1e-3 else f"{p:.3f}"

def clean_cols(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns=lambda c: pd.Series(c)
                     .str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
                     .iloc[0])

# ---------------- Variable Definitions ----------------
wmh_fields = {
    "Deep": ("Total_volume_of_deep_white_matter_hyperintensities_Instance_2",
             "Total_volume_of_deep_white_matter_hyperintensities_Instance_3"),
    "Periventricular": ("Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2",
                        "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_3"),
    "Total": ("Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2",
              "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_3"),
}

# ---------------- Annualized Change Definitions ----------------
def annual_change_raw_ratio(d, f2, f3):
    wm2 = d["Volume_of_white_matter_Instance_2"].replace(0, np.nan)
    wm3 = d["Volume_of_white_matter_Instance_3"].replace(0, np.nan)
    return ((d[f3] / wm3) - (d[f2] / wm2)) / d["Delta_years"], "WMH/WM", "raw"

def annual_change_raw_abs_hsn(d, f2, f3):
    v2 = d[f2] * d["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"]
    v3 = d[f3] * d["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_3"]
    return (v3 - v2) / d["Delta_years"], "HSN", "raw"

def annual_change_log1p_ratio(d, f2, f3):
    wm2 = d["Volume_of_white_matter_Instance_2"].replace(0, np.nan)
    wm3 = d["Volume_of_white_matter_Instance_3"].replace(0, np.nan)
    return (np.log1p(d[f3] / wm3) - np.log1p(d[f2] / wm2)) / d["Delta_years"], "WMH/WM", "log"

def annual_change_log1p_abs_hsn(d, f2, f3):
    v2 = d[f2] * d["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"]
    v3 = d[f3] * d["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_3"]
    return (np.log1p(v3) - np.log1p(v2)) / d["Delta_years"], "HSN", "log"

# ---------------- Regression and Diagnostics ----------------
def fit_cluster(formula, data, groups):
    """
    OLS with cluster-robust SEs by `groups` if possible; otherwise HC3.
    """
    if (groups in data.columns) and (data[groups].dropna().nunique() > 1):
        return smf.ols(formula, data=data).fit(
            cov_type="cluster",
            cov_kwds={"groups": data[groups]}
        )
    return smf.ols(formula, data=data).fit(cov_type="HC3")

def diagnostic_compare(d, outcome_raw, outcome_log, rhs_terms, categorical_in):
    """
    Compare raw vs log1p models via:
      - Breusch–Pagan (heteroscedasticity)
      - Jarque–Bera (normality)
      - AIC
    Recommend log1p if ≥2 criteria favor log1p.
    """
    keep_cols = [outcome_raw, outcome_log, "group", "match_id"] + rhs_terms
    tmp = d[keep_cols].dropna().copy()
    if tmp.empty:
        return None

    rhs = " + ".join(
        [f"C({v})" if v in categorical_in else v for v in rhs_terms]
    ) if rhs_terms else ""
    f_raw = f"{outcome_raw} ~ C(group)" + (f" + {rhs}" if rhs else "")
    f_log = f"{outcome_log} ~ C(group)" + (f" + {rhs}" if rhs else "")

    m_raw = fit_cluster(f_raw, tmp, "match_id")
    m_log = fit_cluster(f_log, tmp, "match_id")

    bp_raw = het_breuschpagan(m_raw.resid, m_raw.model.exog)[1]
    bp_log = het_breuschpagan(m_log.resid, m_log.model.exog)[1]
    jb_raw = jarque_bera(m_raw.resid)[1]
    jb_log = jarque_bera(m_log.resid)[1]
    aic_raw, aic_log = m_raw.aic, m_log.aic

    rec = (
        (bp_log > bp_raw)
        + (jb_log > jb_raw)
        + (aic_log < aic_raw - 2)
    ) >= 2

    return {
        "n": len(tmp),
        "BP_p_raw": float(bp_raw),
        "BP_p_log": float(bp_log),
        "JB_p_raw": float(jb_raw),
        "JB_p_log": float(jb_log),
        "AIC_raw": float(aic_raw),
        "AIC_log": float(aic_log),
        "Recommend_log1p": bool(rec),
    }

def run_models(df, covars_in, categorical_in, change_kind="raw_abs",
               baseline_adjust=False, return_models=False):
    """
    Run OLS models for all WMH regions.

    change_kind: 'raw_abs', 'raw_ratio', 'log_abs', 'log_ratio'
    baseline_adjust: if True, include appropriate baseline at Instance 2.
    """
    rows = []
    fits = {}

    for region, (f2, f3) in wmh_fields.items():
        d = df.copy()

        # outcome
        if change_kind == "raw_ratio":
            y, scale, scl = annual_change_raw_ratio(d, f2, f3)
        elif change_kind == "raw_abs":
            y, scale, scl = annual_change_raw_abs_hsn(d, f2, f3)
        elif change_kind == "log_ratio":
            y, scale, scl = annual_change_log1p_ratio(d, f2, f3)
        else:  # "log_abs"
            y, scale, scl = annual_change_log1p_abs_hsn(d, f2, f3)

        d = d.assign(Annual_Change=y).dropna(
            subset=["Annual_Change", "match_id", "group"]
        )

        rhs_terms = list(covars_in)

        if baseline_adjust:
            d["Baseline_cov"] = _baseline_column(d, change_kind, f2)
            rhs_terms = ["Baseline_cov"] + rhs_terms

        rhs = " + ".join(
            [f"C({c})" if c in categorical_in else c for c in rhs_terms]
        ) if rhs_terms else ""
        formula = "Annual_Change ~ C(group)" + (f" + {rhs}" if rhs else "")

        fit = fit_cluster(formula, d, "match_id")
        if return_models:
            fits[region] = fit

        key = "C(group)[T.Study]"
        beta = ci_l = ci_u = pval = np.nan
        if key in fit.params.index:
            beta = float(fit.params[key])
            ci_l, ci_u = [float(x) for x in fit.conf_int().loc[key]]
            pval = float(fit.pvalues[key])

        pct = pct_l = pct_u = np.nan
        if scl == "log" and np.isfinite(beta):
            pct = 100 * (np.exp(beta) - 1)
            pct_l = 100 * (np.exp(ci_l) - 1)
            pct_u = 100 * (np.exp(ci_u) - 1)

        outcome = f"{region} · {scale}" + ("" if scl == "raw" else " (%/yr)")
        rows.append({
            "Outcome": outcome,
            "β": beta,
            "95% CI Low": ci_l,
            "95% CI High": ci_u,
            "p": pval,
            "%/year": pct,
            "%/year CI Low": pct_l,
            "%/year CI High": pct_u,
            "N Study": int(d.loc[d.group == 'Study', 'Participant_ID'].nunique())
                       if "Participant_ID" in d.columns else int((d.group == 'Study').sum()),
            "N Control": int(d.loc[d.group == 'Control', 'Participant_ID'].nunique())
                         if "Participant_ID" in d.columns else int((d.group == 'Control').sum()),
        })

    res = pd.DataFrame(rows)
    if return_models:
        return res, fits
    return res

def inc_dec_tests(df, use_ratio=True):
    """
    Proportion of participants with increased vs decreased WMH (Study vs Control).
    Returns counts, percentages, difference, CIs, and p-values.
    """
    rows = []
    for region, (f2, f3) in wmh_fields.items():
        d = df.copy()
        if use_ratio:
            y, _, _ = annual_change_raw_ratio(d, f2, f3)
        else:
            y, _, _ = annual_change_raw_abs_hsn(d, f2, f3)
        d = d.assign(Annual_Change=y).dropna(subset=["Annual_Change", "group"])
        d["Increased"] = (d["Annual_Change"] > 0).astype(int)

        inc_s = int(d.loc[d.group == "Study", "Increased"].sum())
        inc_c = int(d.loc[d.group == "Control", "Increased"].sum())
        n_s = int((d.group == "Study").sum())
        n_c = int((d.group == "Control").sum())
        dec_s, dec_c = n_s - inc_s, n_c - inc_c

        table = np.array([[inc_s, dec_s], [inc_c, dec_c]], dtype=int)

        if (table < 5).any():
            test = "Fisher exact"
            _, p_primary = stats.fisher_exact(table)
        else:
            test = "Chi-square (Yates)"
            _, p_primary, _, _ = stats.chi2_contingency(table, correction=True)

        prop_s = inc_s / n_s if n_s > 0 else np.nan
        prop_c = inc_c / n_c if n_c > 0 else np.nan
        ci_l, ci_h = confint_proportions_2indep(
            inc_s, n_s, inc_c, n_c, method="newcomb"
        )
        _, p_prop = proportions_ztest([inc_s, inc_c], [n_s, n_c])

        rows.append({
            "Outcome": region,
            "Study Increase": inc_s,
            "Study Decrease": dec_s,
            "Study N": n_s,
            "Control Increase": inc_c,
            "Control Decrease": dec_c,
            "Control N": n_c,
            "Increase % (Study)": 100 * prop_s if n_s > 0 else np.nan,
            "Increase % (Control)": 100 * prop_c if n_c > 0 else np.nan,
            "Prop Diff (S-C)": (prop_s - prop_c) if (n_s > 0 and n_c > 0) else np.nan,
            "95% CI Low": ci_l,
            "95% CI High": ci_h,
            "Test": test,
            "p_primary": p_primary,
            "p_prop_z": float(p_prop),
        })

    return pd.DataFrame(rows)

# ---------------- Formatting helpers ----------------
def _format_pq(p, q=None):
    def fnum(x, prefix):
        if x is None or (isinstance(x, float) and np.isnan(x)):
            return None
        return f"{prefix}<0.001" if x < 0.001 else f"{prefix}={x:.3f}"
    p_txt = fnum(p, "p")
    q_txt = fnum(q, "q") if q is not None else None
    if p_txt and q_txt:
        return f"{p_txt}, {q_txt}"
    if p_txt:
        return p_txt
    if q_txt:
        return q_txt
    return None

def _stars_by_p(p):
    try:
        p = float(p)
    except Exception:
        return ""
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return ""

def _add_pq_bracket(ax, x_left, x_right, y_base,
                    label_txt, barw=0.52, fontsize=8, dy=0.8):
    if not label_txt:
        return
    dx = barw * 0.60
    y_top = y_base + 1.0
    ax.plot(
        [x_left - dx, x_left - dx, x_right + dx, x_right + dx],
        [y_base, y_top, y_top, y_base],
        color="black", lw=0.9, zorder=5,
    )
    ax.text(
        (x_left + x_right) / 2.0, y_top + dy,
        label_txt,
        ha="center", va="bottom",
        fontsize=fontsize, fontstyle="italic",
        color="black", zorder=6,
    )

# ---------------- Baseline covariate helper ----------------
def _baseline_column(df, change_kind, f2):
    """
    Baseline covariate (Instance 2) matching change_kind.
    """
    if change_kind in ("raw_abs", "log_abs"):
        base = df[f2] * df["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"]
        if change_kind == "log_abs":
            base = np.log1p(base)
        return pd.to_numeric(base, errors="coerce")
    else:
        wm2 = df["Volume_of_white_matter_Instance_2"].replace(0, np.nan)
        base = pd.to_numeric(df[f2], errors="coerce") / wm2
        if change_kind == "log_ratio":
            base = np.log1p(base)
        return pd.to_numeric(base, errors="coerce")

# ---------------- Diagnostics plots ----------------
def plot_model_diagnostics(fit, out_base):
    """
    Save residual diagnostics:
      - Residuals vs fitted
      - Normal Q–Q
      - Histogram of standardized residuals
    """
    if fit is None:
        return

    # Residuals vs Fitted
    fig, ax = plt.subplots(figsize=(6.4, 4.0))
    ax.scatter(fit.fittedvalues, fit.resid, s=14, alpha=0.75, edgecolors="none")
    ax.axhline(0, color="grey", ls="--", lw=1)
    ax.set_xlabel("Fitted values")
    ax.set_ylabel("Residuals")
    ax.set_title("Residuals vs Fitted")
    ax.grid(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(out_base + "_resid_vs_fitted.png", **PNG_KW)
    plt.savefig(out_base + "_resid_vs_fitted.pdf", **PDF_KW)
    plt.savefig(out_base + "_resid_vs_fitted.svg", **SVG_KW)
    plt.savefig(out_base + "_resid_vs_fitted.tiff", **TIFF_KW)
    plt.close()

    # Q–Q plot
    fig = plt.figure(figsize=(6.0, 6.0))
    ax = fig.add_subplot(111)
    sm.ProbPlot(fit.resid).qqplot(line="s", ax=ax)
    ax.set_title("Normal Q–Q")
    ax.grid(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(out_base + "_qq.png", **PNG_KW)
    plt.savefig(out_base + "_qq.pdf", **PDF_KW)
    plt.savefig(out_base + "_qq.svg", **SVG_KW)
    plt.savefig(out_base + "_qq.tiff", **TIFF_KW)
    plt.close()

    # Histogram of standardized residuals
    fig, ax = plt.subplots(figsize=(6.4, 4.0))
    try:
        infl = fit.get_influence()
        zresid = infl.resid_studentized_internal
    except Exception:
        zresid = (fit.resid - np.nanmean(fit.resid)) / np.nanstd(fit.resid)
    ax.hist(zresid, bins=30, edgecolor="white")
    ax.set_xlabel("Standardized residuals")
    ax.set_ylabel("Count")
    ax.set_title("Histogram of standardized residuals")
    ax.grid(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(out_base + "_hist_stdresid.png", **PNG_KW)
    plt.savefig(out_base + "_hist_stdresid.pdf", **PDF_KW)
    plt.savefig(out_base + "_hist_stdresid.svg", **SVG_KW)
    plt.savefig(out_base + "_hist_stdresid.tiff", **TIFF_KW)
    plt.close()

# ---------------- Methods record ----------------
def write_methods_record(outdir):
    txt = []
    txt.append("Longitudinal WMH change (Instance 3 - Instance 2)")
    txt.append("")
    txt.append("Outcomes:")
    txt.append("- HSN absolute change per year (raw and log1p back-transform to %/year).")
    txt.append("- WMH/WM ratio change per year (raw and log1p back-transform to %/year).")
    txt.append("")
    txt.append("Primary model: OLS with Study vs Control indicator; covariates:")
    txt.append("- Age at Instance 2, Sex, BMI at Instance 0, Genetic ethnic grouping, Smoking (ever),")
    txt.append("  Alcohol intake frequency (ordinal), Townsend deprivation index at recruitment.")
    txt.append("SE: Cluster-robust by matched set (match_id) if available; otherwise HC3.")
    txt.append("")
    txt.append("Sensitivity model: Adds baseline value at Instance 2 matching the outcome scale.")
    txt.append("")
    txt.append("Multiple testing:")
    txt.append("- For HSN log1p models (Periventricular, Deep), BH-FDR applied; Total excluded (q=NaN).")
    txt.append("- Inc/Dec tests: only Periventricular and Deep included in FDR; Total excluded.")
    txt.append("")
    txt.append("Figures:")
    txt.append("- Forest plots match main analysis style; PNG/PDF/SVG and 600 dpi TIFF with LZW compression.")
    path = os.path.join(outdir, "METHODS_Longitudinal.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(txt))
    return path

# ---------------- Forest Plotting ----------------
def forest(df_plot, ttl, outpath, logscale=False):
    """
    Publication-styled forest plot:
      - Deep-blue markers & CIs
      - Outcome order: Total → Periventricular → Deep
      - Y-axis labels: Total WMH / PWMH / DWMH
      - logscale=True uses %/year; else raw β.
    """
    if df_plot is None or df_plot.empty:
        return

    plt.rcParams.update(PUB_RC)

    dfp = df_plot.copy()
    region = dfp["Outcome"].astype(str).str.split("·", n=1, expand=True)[0].str.strip()
    dfp["Region"] = region

    dfp = dfp[dfp["Region"].isin(PLOT_ORDER)].copy()
    if dfp.empty:
        return

    dfp["Region"] = pd.Categorical(dfp["Region"], categories=PLOT_ORDER, ordered=True)
    dfp = dfp.sort_values("Region").reset_index(drop=True)

    if logscale:
        x = pd.to_numeric(dfp["%/year"], errors="coerce").to_numpy()
        xl = pd.to_numeric(dfp["%/year CI Low"], errors="coerce").to_numpy()
        xu = pd.to_numeric(dfp["%/year CI High"], errors="coerce").to_numpy()
        xlabel = "Annualized % change (SA - Control)"
    else:
        x = pd.to_numeric(dfp["β"], errors="coerce").to_numpy()
        xl = pd.to_numeric(dfp["95% CI Low"], errors="coerce").to_numpy()
        xu = pd.to_numeric(dfp["95% CI High"], errors="coerce").to_numpy()
        xlabel = "Annualized change (β ± 95% CI)"

    n = len(dfp)
    y = np.arange(n)[::-1]
    fig_h = max(3.6, 0.85 * n)

    fig = plt.figure(figsize=(6.8, fig_h))
    ax = plt.gca()

    ax.errorbar(
        x, y,
        xerr=[x - xl, xu - x],
        fmt="o",
        ms=MS,
        mfc=COLOR_PRIMARY,
        mec=COLOR_PRIMARY,
        ecolor=COLOR_PRIMARY,
        elinewidth=ELW,
        capsize=CAPSIZE,
        linestyle="none",
        zorder=3,
    )

    ax.axvline(0, color="grey", linestyle="--", linewidth=1)
    ax.grid(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    ylab = [DISPLAY_LABEL.get(r, r) for r in dfp["Region"].astype(str)]
    ax.set_yticks(y)
    ax.set_yticklabels(ylab)

    xmin = np.nanmin(xl) if np.isfinite(np.nanmin(xl)) else -1.0
    xmax = np.nanmax(xu) if np.isfinite(np.nanmax(xu)) else 1.0
    span = xmax - xmin if (xmax - xmin) > 0 else 1.0
    pad_left, pad_right = 0.08 * span, 0.30 * span
    ax.set_xlim(xmin - pad_left, xmax + pad_right)

    ax.set_xlabel(xlabel)
    ax.set_title(ttl)
    plt.tight_layout()

    os.makedirs(os.path.dirname(outpath), exist_ok=True)
    base, _ = os.path.splitext(outpath)
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".svg", **SVG_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)

# ---------------- Single-region Inc/Dec plots ----------------
def plot_inc_dec_paper(tests_table, ttl, outdir):
    """
    Single-region stacked bar plots for Increase vs Decrease.
    - q-values only shown for Periventricular & Deep (if available).
    - Total WMH: p only (as primary); q suppressed.
    """
    if tests_table.empty:
        return

    plt.rcParams.update(PUB_RC)
    os.makedirs(outdir, exist_ok=True)

    OI_INCREASE = "#D55E00"
    OI_DECREASE = "#009E73"

    for region in tests_table["Outcome"].unique():
        t = tests_table[tests_table["Outcome"] == region].iloc[0]
        inc_c, dec_c, n_c = float(t["Control Increase"]), float(t["Control Decrease"]), float(t["Control N"])
        inc_s, dec_s, n_s = float(t["Study Increase"]),   float(t["Study Decrease"]),   float(t["Study N"])

        inc_pct_c = 100 * inc_c / n_c if n_c > 0 else np.nan
        inc_pct_s = 100 * inc_s / n_s if n_s > 0 else np.nan
        dec_pct_c = 100 - inc_pct_c if not np.isnan(inc_pct_c) else np.nan
        dec_pct_s = 100 - inc_pct_s if not np.isnan(inc_pct_s) else np.nan

        p_main = pd.to_numeric(t["p_primary"], errors="coerce")
        p_main = float(p_main) if pd.notna(p_main) else np.nan

        # q 仅 Periventricular/Deep 显示
        q_main = None
        if region in ("Periventricular", "Deep"):
            if "q_FDR_incdec_combined" in t.index:
                q_try = pd.to_numeric(t["q_FDR_incdec_combined"], errors="coerce")
                q_main = float(q_try) if pd.notna(q_try) else None

        pq_txt_base = _format_pq(p_main, q_main)

        fig, ax = plt.subplots(figsize=(6.2, 4.3))
        x = np.array([0, 1], dtype=float)
        barw = 0.52

        ax.bar(x, [inc_pct_c, inc_pct_s], width=barw,
               color=OI_INCREASE, edgecolor="black", linewidth=0.8,
               label="Increase", zorder=3)
        ax.bar(x, [dec_pct_c, dec_pct_s], width=barw,
               bottom=[inc_pct_c, inc_pct_s],
               color=OI_DECREASE, edgecolor="black", linewidth=0.8,
               label="Decrease", zorder=3)

        ax.set_xticks(x)
        ax.set_xticklabels(["Control", "Study"])
        ax.set_ylabel("Participants (%)")
        ax.set_ylim(0, 104)
        ax.set_yticks([0, 25, 50, 75, 100])
        ax.set_title(f"{region} {ttl} — Increase vs Decrease by Group", pad=16)

        ax.grid(axis="y", linestyle=":", alpha=0.45)
        ax.set_axisbelow(True)
        ax.margins(x=0.18)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        # % labels
        for i, ip in enumerate([inc_pct_c, inc_pct_s]):
            if not np.isnan(ip):
                ax.text(x[i], ip * 0.50, f"{ip:.1f}%",
                        ha="center", va="center", fontsize=9, color="white")
        for i, (ip, dp) in enumerate([(inc_pct_c, dec_pct_c),
                                      (inc_pct_s, dec_pct_s)]):
            if not np.isnan(dp):
                ax.text(x[i], ip + dp * 0.50, f"{dp:.1f}%",
                        ha="center", va="center", fontsize=9, color="white")

        # p / q label (Peri/Deep may show q)
        pq_txt_base = _format_pq(p_main, q_main)
        stars = _stars_by_p(p_main)
        label_txt = f"{stars} {pq_txt_base}" if stars and pq_txt_base else (stars or pq_txt_base)

        if label_txt:
            _add_pq_bracket(
                ax,
                x_left=x[0],
                x_right=x[1],
                y_base=101.0,
                label_txt=label_txt,
                barw=barw,
                fontsize=9,
                dy=0.4,
            )

        ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), frameon=False)
        plt.subplots_adjust(left=0.10, right=0.76, bottom=0.14, top=0.88)

        fname = f"{safe_name(region)}_IncDec_{safe_name(ttl)}_Paper"
        base = os.path.join(outdir, fname)
        plt.savefig(base + ".png", **PNG_KW)
        plt.savefig(base + ".pdf", **PDF_KW)
        plt.savefig(base + ".svg", **SVG_KW)
        plt.savefig(base + ".tiff", **TIFF_KW)
        plt.close()

# ---------------- 2-panel Figure X (HSN) ----------------
def plot_figureX_hsn(la_table: pd.DataFrame,
                     incdec_hsn: pd.DataFrame,
                     outdir: str,
                     filename: str = "FigureX_Longitudinal_WMH.png"):
    """
    Two-panel figure:
      Panel A: HSN log1p annualized % change (SA - Control), Total/PWMH/DWMH.
               Significance by FDR for PWMH/DWMH; p-only for Total.
      Panel B: Inc/Dec stacked bars; q only for PWMH/DWMH.
    """
    if la_table is None or la_table.empty or incdec_hsn is None or incdec_hsn.empty:
        return

    plt.rcParams.update(PUB_RC)
    os.makedirs(outdir, exist_ok=True)

    # --- Panel A prep ---
    region_order_raw = ["Total", "Periventricular", "Deep"]
    wanted = [f"{lab} · HSN (%/yr)" for lab in region_order_raw]

    la = la_table.copy()
    la = la[la["Outcome"].isin(wanted)].copy()
    if la.empty:
        return
    la["Outcome"] = pd.Categorical(la["Outcome"], categories=wanted, ordered=True)
    la = la.sort_values("Outcome").reset_index(drop=True)
    la["Region"] = la["Outcome"].str.split("·", n=1, expand=True)[0].str.strip()

    region_map = {"Total": "Total WMH", "Periventricular": "PWMH", "Deep": "DWMH"}
    y_labels_left = [region_map.get(r, r) for r in la["Region"]]

    x = pd.to_numeric(la["%/year"], errors="coerce").to_numpy()
    xl = pd.to_numeric(la["%/year CI Low"], errors="coerce").to_numpy()
    xu = pd.to_numeric(la["%/year CI High"], errors="coerce").to_numpy()

    def _pick_q(row):
        if "q_FDR_primary" in row and pd.notna(row["q_FDR_primary"]):
            return row["q_FDR_primary"]
        if "q_FDR_sensitivity" in row and pd.notna(row["q_FDR_sensitivity"]):
            return row["q_FDR_sensitivity"]
        return np.nan

    qvals = pd.to_numeric(la.apply(_pick_q, axis=1), errors="coerce").to_numpy()
    pvals = pd.to_numeric(la.get("p", np.nan), errors="coerce").to_numpy()


    use_q = np.isfinite(qvals)
    sig_mask = np.where(use_q, (qvals <= 0.05), (pvals <= 0.05))


    COLOR_NONSIG = "#7f7f7f"
    point_colors = np.where(sig_mask, COLOR_PRIMARY, COLOR_NONSIG)

    # --- Panel B prep (Inc/Dec HSN) ---
    t = incdec_hsn.copy()
    t = t[t["Outcome"].isin(region_order_raw)].copy()

    rows = []
    for rg in region_order_raw:
        r = t[t["Outcome"] == rg]
        if r.empty:
            continue
        r = r.iloc[0]
        n_c = float(r["Control N"])
        n_s = float(r["Study N"])
        inc_c = 100.0 * float(r["Control Increase"]) / n_c if n_c > 0 else np.nan
        inc_s = 100.0 * float(r["Study Increase"]) / n_s if n_s > 0 else np.nan
        dec_c = 100.0 - inc_c if not np.isnan(inc_c) else np.nan
        dec_s = 100.0 - inc_s if not np.isnan(inc_s) else np.nan

        q_show = None
        if (rg in ("Periventricular", "Deep") and
            ("q_FDR_incdec_combined" in r.index) and
            pd.notna(r["q_FDR_incdec_combined"])):
            q_show = float(r["q_FDR_incdec_combined"])

        rows.append({
            "Region": region_map.get(rg, rg),
            "inc_c": inc_c, "dec_c": dec_c,
            "inc_s": inc_s, "dec_s": dec_s,
            "p": float(r["p_primary"]) if pd.notna(r["p_primary"]) else np.nan,
            "q": q_show,
        })

    if not rows:
        return
    dfB = pd.DataFrame(rows)

    # --- Figure layout ---
    fig = plt.figure(figsize=(11.6, 4.0))
    gs = fig.add_gridspec(1, 2, width_ratios=[1.5, 1.45], wspace=0.30)

    # Panel A
    axA = fig.add_subplot(gs[0, 0])
    y = np.linspace(0, len(la) - 1, len(la)) * 0.7
    y = y[::-1]
    ypad = 0.35
    axA.set_ylim(-ypad, y.max() + ypad)

    span_ref = np.nanmax(xu) - np.nanmin(xl) if np.isfinite(np.nanmax(xu) - np.nanmin(xl)) else 1.0

    for i in range(len(y)):
        if any(np.isnan([x[i], xl[i], xu[i]])):
            continue
        axA.errorbar(
            x[i], y[i],
            xerr=[[x[i] - xl[i]], [xu[i] - x[i]]],
            fmt="o", ms=MS,
            mfc=point_colors[i], mec=point_colors[i],
            ecolor=point_colors[i], elinewidth=ELW,
            capsize=CAPSIZE, linestyle="none", zorder=3,
        )

    axA.axvline(0, color="grey", linestyle="--", linewidth=1)
    axA.grid(False)
    axA.spines["top"].set_visible(False)
    axA.spines["right"].set_visible(False)
    axA.set_yticks(y)
    axA.set_yticklabels(y_labels_left)

    xmin = np.nanmin(xl) if np.isfinite(np.nanmin(xl)) else -1.0
    xmax = np.nanmax(xu) if np.isfinite(np.nanmax(xu)) else 1.0
    span = xmax - xmin if (xmax - xmin) > 0 else 1.0
    axA.set_xlim(xmin - 0.08 * span, xmax + 0.15 * span)

    axA.set_xlabel("Annualized % change (SA - Control)")

    # Panel B
    axB = fig.add_subplot(gs[0, 1])
    COL_INC = COLOR_PRIMARY
    COL_DEC = "#9c9c9c"
    edge_lw = 0.8

    centers = np.arange(len(dfB)) * 2.0
    barw = 0.48

    for i, r in dfB.iterrows():
        xc, xs = centers[i] - 0.36, centers[i] + 0.36
        axB.bar(xc, r["inc_c"], width=barw, color=COL_INC,
                edgecolor=COL_INC, linewidth=edge_lw, zorder=3)
        axB.bar(xc, r["dec_c"], width=barw, bottom=r["inc_c"],
                color=COL_DEC, edgecolor=COL_DEC, linewidth=edge_lw, zorder=3)
        axB.bar(xs, r["inc_s"], width=barw, color=COL_INC,
                edgecolor=COL_INC, linewidth=edge_lw, zorder=3)
        axB.bar(xs, r["dec_s"], width=barw, bottom=r["inc_s"],
                color=COL_DEC, edgecolor=COL_DEC, linewidth=edge_lw, zorder=3)

        stars = _stars_by_p(r["p"])
        label_txt = stars
        if label_txt:
            from numpy import nansum
            y_base = max(nansum([r["inc_c"], r["dec_c"]]),
                         nansum([r["inc_s"], r["dec_s"]])) + 2.0
            _add_pq_bracket(
                axB,
                x_left=xc,
                x_right=xs,
                y_base=y_base,
                label_txt=label_txt,
                barw=barw,
                fontsize=11,
                dy=0.6,
            )

    tick_positions = []
    tick_labels = []
    for i, region_name in enumerate(dfB["Region"]):
        xc, xs = centers[i] - 0.36, centers[i] + 0.36
        tick_positions.extend([xc, xs])
        tick_labels.extend(["CO", "SA"])
        axB.text(centers[i], -10.5, region_name,
                 ha="center", va="top", fontsize=11)

    axB.set_xticks(tick_positions)
    axB.set_xticklabels(tick_labels, fontsize=9)
    plt.subplots_adjust(bottom=0.18)

    axB.set_ylim(0, 115)
    axB.set_ylabel("Participants (%)")
    axB.grid(axis="y", linestyle=":", alpha=0.45)
    axB.set_axisbelow(True)
    axB.spines["top"].set_visible(False)
    axB.spines["right"].set_visible(False)

    legend_items = [
        mpl.patches.Patch(facecolor=COL_INC, edgecolor=COL_INC, label="Increase"),
        mpl.patches.Patch(facecolor=COL_DEC, edgecolor=COL_DEC, label="Decrease"),
    ]
    axB.legend(handles=legend_items,
               loc="center left", bbox_to_anchor=(1.02, 0.5),
               frameon=False, handlelength=1.4, ncol=1)

    axB.set_xlim(min(tick_positions) - 1.0, max(tick_positions) + 1.0)

    boxA = axA.get_position()
    fig.text(boxA.x0 - 0.07, boxA.y1 + 0.005, "A",
             fontsize=18, fontweight="bold")
    boxB = axB.get_position()
    fig.text(boxB.x0 - 0.09, boxB.y1 + 0.005, "B",
             fontsize=18, fontweight="bold")

    plt.subplots_adjust(left=0.18, right=0.80, bottom=0.12, top=0.94)

    base = os.path.join(outdir, os.path.splitext(filename)[0])
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".svg", **SVG_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)

# ---------------- Follow-up duration summary ----------------
def _summarize_series_years(s: pd.Series):
    s = pd.to_numeric(s, errors="coerce").dropna()
    if s.empty:
        return dict(N=0, Mean=np.nan, SD=np.nan,
                    Median=np.nan, Q1=np.nan, Q3=np.nan,
                    IQR=np.nan, Min=np.nan, Max=np.nan)
    mean = float(s.mean())
    sd = float(s.std(ddof=1))
    q1 = float(s.quantile(0.25))
    q3 = float(s.quantile(0.75))
    return dict(
        N=int(s.size),
        Mean=mean,
        SD=sd,
        Median=float(s.median()),
        Q1=q1,
        Q3=q3,
        IQR=q3 - q1,
        Min=float(s.min()),
        Max=float(s.max()),
    )

def write_followup_summary(df: pd.DataFrame, label: str, subdir: str):
    """Summarize follow-up duration (Instance 2 → Instance 3) by group and print summary."""
    os.makedirs(subdir, exist_ok=True)
    rows = []
    parts = [
        ("Overall", df),
        ("Control", df[df["group"] == "Control"]),
        ("Study", df[df["group"] == "Study"]),
    ]
    for gname, gdf in parts:
        stats_row = _summarize_series_years(gdf["Delta_years"])
        rows.append({"Cohort": label, "Group": gname, **stats_row})
    tab = pd.DataFrame(rows)

    # Save summary table
    out_csv = os.path.join(subdir, "FollowUp_Duration.csv")
    tab.to_csv(out_csv, index=False, encoding="utf-8-sig")

    def fmt(x):
        return "NA" if (x is None or (isinstance(x, float) and np.isnan(x))) else f"{x:.2f}"

    def sentence(row):
        n = row["N"]
        mean, sd = fmt(row["Mean"]), fmt(row["SD"])
        med, q1, q3 = fmt(row["Median"]), fmt(row["Q1"]), fmt(row["Q3"])
        mn, mx = fmt(row["Min"]), fmt(row["Max"])
        grp = row["Group"]
        return (f"{label} — {grp}: follow-up between Instance 2 and 3 "
                f"was median {med} years (IQR {q1}–{q3}), mean {mean}±{sd} years; "
                f"range {mn}–{mx}; N={n}.")

    lines = [sentence(r) for _, r in tab.iterrows()]
    out_txt = os.path.join(subdir, "FollowUp_Duration_narrative.txt")
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    # Print concise console summary
    print(f"[{label}] Follow-up duration summary (years):")
    for _, r in tab.iterrows():
        med, q1, q3 = fmt(r["Median"]), fmt(r["Q1"]), fmt(r["Q3"])
        print(f"  {r['Group']}: median {med} (IQR {q1}–{q3}); N={r['N']}")



# ---------------- Pipeline for one cohort ----------------
def run_for_one_cohort(label, filepath):
    """
    Full longitudinal pipeline (I3–I2) for one cohort.
    Uses core covariates only (no CMC, no e4, no education).
    """
    print(f"\n===== {label} =====")
    subdir = os.path.join(OUTDIR, label)
    os.makedirs(subdir, exist_ok=True)

    # Input & preprocessing
    df = clean_cols(pd.read_csv(filepath))

    df["Delta_years"] = (
        pd.to_datetime(df["Date_of_attending_assessment_centre_Instance_3"])
        - pd.to_datetime(df["Date_of_attending_assessment_centre_Instance_2"])
    ).dt.days / 365.25

    df = df[(df["Delta_years"] > 0) & (df["Delta_years"] < 8)].copy()

    needed = [c for p in wmh_fields.values() for c in p] + [
        "Volume_of_white_matter_Instance_2",
        "Volume_of_white_matter_Instance_3",
        "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2",
        "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_3",
        "Participant_ID",
        "match_id",
        "group",
        "Delta_years",
    ]
    df = df.dropna(subset=[c for c in needed if c in df.columns]).copy()
    df["group"] = pd.Categorical(df["group"], ["Control", "Study"])

    # Covariates (no CMC, no e4, no education)
    covars = [
        "Age_at_Instance_2",
        "Sex",
        "Body_mass_index_BMI_Instance_0",
        "Genetic_ethnic_grouping",
        "Smoking_Ever",
        "Alcohol_intake_frequency_ordinal",
        "Townsend_deprivation_index_at_recruitment",
    ]
    categorical_all = {"Sex", "Genetic_ethnic_grouping", "Smoking_Ever"}

    covars_in = [v for v in covars if v in df.columns]
    categorical_in = [v for v in covars_in if v in categorical_all]

    # Simple imputation
    for v in covars_in:
        if pd.api.types.is_numeric_dtype(df[v]):
            df[v] = pd.to_numeric(df[v], errors="coerce")
            df[v] = df[v].fillna(df[v].median())
        else:
            m = df[v].mode(dropna=True)
            df[v] = df[v].fillna(m.iloc[0] if not m.empty else df[v])

    # Diagnostics: raw vs log1p
    diag = []
    for region, (f2, f3) in wmh_fields.items():
        # HSN
        d_abs = df.copy()
        y1, _, _ = annual_change_raw_abs_hsn(d_abs, f2, f3)
        y2, _, _ = annual_change_log1p_abs_hsn(d_abs, f2, f3)
        d_abs["Y_raw_abs"] = y1
        d_abs["Y_log_abs"] = y2
        r = diagnostic_compare(d_abs, "Y_raw_abs", "Y_log_abs",
                               covars_in, categorical_in)
        if r:
            diag.append({"Outcome": f"{region} · HSN", **r})
        # Ratio
        d_rat = df.copy()
        y1, _, _ = annual_change_raw_ratio(d_rat, f2, f3)
        y2, _, _ = annual_change_log1p_ratio(d_rat, f2, f3)
        d_rat["Y_raw_ratio"] = y1
        d_rat["Y_log_ratio"] = y2
        r = diagnostic_compare(d_rat, "Y_raw_ratio", "Y_log_ratio",
                               covars_in, categorical_in)
        if r:
            diag.append({"Outcome": f"{region} · WMH/WM", **r})

    if diag:
        pd.DataFrame(diag).to_csv(
            os.path.join(subdir, "Diagnostics_raw_vs_log1p.csv"),
            index=False,
            encoding="utf-8-sig",
        )

    # Primary models
    rr = run_models(df, covars_in, categorical_in, "raw_ratio")
    ra = run_models(df, covars_in, categorical_in, "raw_abs")
    lr = run_models(df, covars_in, categorical_in, "log_ratio")
    la = run_models(df, covars_in, categorical_in, "log_abs")  # HSN log1p (%/yr)

    # FDR for rr/ra/lr (simple, all regions)
    for tab in [rr, ra, lr]:
        if not tab.empty:
            tab["FDR q"] = multipletests(tab["p"], method="fdr_bh")[1]

    # FDR for la: only Periventricular & Deep within cohort; Total excluded
    la = la.copy()
    la["q_FDR_primary"] = np.nan
    la["q_FDR_sensitivity"] = np.nan
    if not la.empty:
        # 仅对 PWMH / DWMH 做 FDR；若只有 1 个有效 p，则 q = p（确保每个区域都有 q）
        region_name = la["Outcome"].astype(str).str.split("·", n=1, expand=True)[0].str.strip()
        mask_pd = region_name.isin(["Periventricular", "Deep"])
        idx = la.index[mask_pd & pd.to_numeric(la["p"], errors="coerce").notna()]

        if len(idx) == 1:
            qvals = pd.to_numeric(la.loc[idx, "p"], errors="coerce").values  # q = p
        elif len(idx) >= 2:
            p_clean = pd.to_numeric(la.loc[idx, "p"], errors="coerce").values
            qvals = multipletests(p_clean, method="fdr_bh", alpha=0.05)[1]
        else:
            qvals = None

        if qvals is not None:
            if label == "Primary":
                la.loc[idx, "q_FDR_primary"] = qvals
            elif label == "Sensitivity":
                la.loc[idx, "q_FDR_sensitivity"] = qvals



    # Baseline-adjusted HSN log1p models
    def run_models_log_abs_with_baseline(df_in, covars_in, categorical_in):
        rows = []
        for region, (f2, f3) in wmh_fields.items():
            d = df_in.copy()
            y, _, _ = annual_change_log1p_abs_hsn(d, f2, f3)
            v2 = d[f2] * d["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"]
            d = d.assign(
                Annual_Change=y,
                Baseline_log1p=np.log1p(v2),
            ).dropna(subset=["Annual_Change", "Baseline_log1p",
                             "match_id", "group"])
            rhs_terms = [f"C({c})" if c in categorical_in else c
                         for c in covars_in] + ["Baseline_log1p"]
            rhs = " + ".join(rhs_terms)
            formula = "Annual_Change ~ C(group)" + (f" + {rhs}" if rhs else "")
            fit = fit_cluster(formula, d, "match_id")
            key = "C(group)[T.Study]"
            beta = ci_l = ci_u = pval = np.nan
            if key in fit.params.index:
                beta = float(fit.params[key])
                ci_l, ci_u = [float(x) for x in fit.conf_int().loc[key]]
                pval = float(fit.pvalues[key])
            pct = pct_l = pct_u = np.nan
            if np.isfinite(beta):
                pct = 100 * (np.exp(beta) - 1)
                pct_l = 100 * (np.exp(ci_l) - 1)
                pct_u = 100 * (np.exp(ci_u) - 1)
            outcome = f"{region} · HSN (%/yr) — baseline-adjusted"
            rows.append({
                "Outcome": outcome,
                "β": beta,
                "95% CI Low": ci_l,
                "95% CI High": ci_u,
                "p": pval,
                "%/year": pct,
                "%/year CI Low": pct_l,
                "%/year CI High": pct_u,
                "N Study": int(d.loc[d.group == 'Study', 'Participant_ID'].nunique())
                           if "Participant_ID" in d.columns else int((d.group == 'Study').sum()),
                "N Control": int(d.loc[d.group == 'Control', 'Participant_ID'].nunique())
                             if "Participant_ID" in d.columns else int((d.group == 'Control').sum()),
            })
        return pd.DataFrame(rows)

    la_baseline = run_models_log_abs_with_baseline(df, covars_in, categorical_in)

    # --- FDR: cohort-wide across PWMH + DWMH (Total excluded) BEFORE saving CSV ---
    def _apply_fdr_pwmh_dwmh_in_cohort(tbl: pd.DataFrame, cohort_label: str) -> pd.DataFrame:
        """
        Apply BH-FDR across all Periventricular + Deep rows within a cohort.
        - Only numeric p-values are included.
        - If only one valid test exists, set q = p.
        - Writes q into q_FDR_primary or q_FDR_sensitivity according to the cohort.
        """
        if tbl is None or tbl.empty:
            return tbl

        out = tbl.copy()
        # Ensure q columns exist
        if "q_FDR_primary" not in out.columns:
            out["q_FDR_primary"] = np.nan
        if "q_FDR_sensitivity" not in out.columns:
            out["q_FDR_sensitivity"] = np.nan

        # Identify PWMH/DWMH rows by region (left of "·")
        region = out["Outcome"].astype(str).str.split("·", n=1, expand=True)[0].str.strip()
        mask_pd = region.isin(["Periventricular", "Deep"])

        # Select valid indices with numeric p
        idx = out.index[mask_pd & pd.to_numeric(out["p"], errors="coerce").notna()]
        if len(idx) == 0:
            return out

        # Compute q: one test -> q=p; otherwise BH-FDR
        if len(idx) == 1:
            qvals = pd.to_numeric(out.loc[idx, "p"], errors="coerce").values
        else:
            pvals = pd.to_numeric(out.loc[idx, "p"], errors="coerce").values
            qvals = multipletests(pvals, method="fdr_bh", alpha=0.05)[1]

        if cohort_label == "Primary":
            out.loc[idx, "q_FDR_primary"] = qvals
        elif cohort_label == "Sensitivity":
            out.loc[idx, "q_FDR_sensitivity"] = qvals

        return out

    # Apply FDR to HSN log1p (unadjusted) and baseline-adjusted tables
    la = _apply_fdr_pwmh_dwmh_in_cohort(la, label)
    la_baseline = _apply_fdr_pwmh_dwmh_in_cohort(la_baseline, label)

    # --- Save regression tables AFTER FDR so q columns are in the CSVs ---
    rr.to_csv(os.path.join(subdir, "WMH_WM_ratio_raw_change.csv"),
              index=False, encoding="utf-8-sig")
    ra.to_csv(os.path.join(subdir, "WMH_HSN_abs_raw_change.csv"),
              index=False, encoding="utf-8-sig")
    lr.to_csv(os.path.join(subdir, "WMH_WM_ratio_log1p_pct_change.csv"),
              index=False, encoding="utf-8-sig")
    la.to_csv(os.path.join(subdir, "WMH_HSN_abs_log1p_pct_change.csv"),
              index=False, encoding="utf-8-sig")
    if la_baseline is not None and not la_baseline.empty:
        la_baseline.to_csv(
            os.path.join(subdir, "WMH_HSN_abs_log1p_pct_change_BaselineAdjusted.csv"),
            index=False, encoding="utf-8-sig"
        )


    # Forest plots
    forest(rr, "WMH/WM ratio (raw annualized change)",
           os.path.join(subdir, "Forest_Ratio_Raw.png"), logscale=False)
    forest(ra, "HSN-WMH (raw annualized change)",
           os.path.join(subdir, "Forest_HSN_Raw.png"), logscale=False)
    forest(lr, "WMH/WM ratio (annualized % change)",
           os.path.join(subdir, "Forest_Ratio_LogPct.png"), logscale=True)
    forest(la, "HSN-WMH (annualized % change, log1p)",
           os.path.join(subdir, "Forest_HSN_LogPct.png"), logscale=True)
    if not la_baseline.empty:
        forest(la_baseline,
               "HSN-WMH (annualized % change, log1p) — baseline-adjusted",
               os.path.join(subdir, "Forest_HSN_LogPct_BaselineAdj.png"),
               logscale=True)

    # Increase/Decrease tests
    tr = inc_dec_tests(df, use_ratio=True)
    ta = inc_dec_tests(df, use_ratio=False)


    # FDR only for Peri & Deep (Total excluded) — Ratio (tr)
    if not tr.empty:
        tr["FDR q (primary)"] = np.nan
        m = tr["Outcome"].isin(["Periventricular", "Deep"])
        idx = tr.index[m & pd.to_numeric(tr["p_primary"], errors="coerce").notna()]
        if len(idx) >= 1:
            p_clean = pd.to_numeric(tr.loc[idx, "p_primary"], errors="coerce").values
            tr.loc[idx, "FDR q (primary)"] = multipletests(p_clean, method="fdr_bh")[1]

    if not ta.empty:
        ta["FDR q (primary)"] = np.nan
        m = ta["Outcome"].isin(["Periventricular", "Deep"])
        idx = ta.index[m & pd.to_numeric(ta["p_primary"], errors="coerce").notna()]
        if len(idx) >= 1:
            p_clean = pd.to_numeric(ta.loc[idx, "p_primary"], errors="coerce").values
            ta.loc[idx, "FDR q (primary)"] = multipletests(p_clean, method="fdr_bh")[1]

    # Combined FDR across HSN+Ratio (Peri/Deep only), drop NaN and allow size=1
    comb = []
    for df_src, typ in [(tr, "Ratio"), (ta, "HSN")]:
        if not df_src.empty:
            sub = df_src[df_src["Outcome"].isin(["Periventricular", "Deep"])].copy()
            sub["p_primary"] = pd.to_numeric(sub["p_primary"], errors="coerce")
            for _, r in sub.dropna(subset=["p_primary"]).iterrows():
                comb.append((typ, r["Outcome"], float(r["p_primary"])))

    if comb:
        pvals = np.array([x[2] for x in comb], dtype=float)
        qvals = multipletests(pvals, method="fdr_bh")[1]
        if not tr.empty:
            tr["q_FDR_incdec_combined"] = np.nan
        if not ta.empty:
            ta["q_FDR_incdec_combined"] = np.nan
        for (typ, region, _), qv in zip(comb, qvals):
            if typ == "Ratio" and not tr.empty:
                tr.loc[tr["Outcome"] == region, "q_FDR_incdec_combined"] = qv
            if typ == "HSN" and not ta.empty:
                ta.loc[ta["Outcome"] == region, "q_FDR_incdec_combined"] = qv


    # Export Inc/Dec tables
    tr.to_csv(os.path.join(subdir, "IncreaseDecrease_Ratio.csv"),
              index=False, encoding="utf-8-sig")
    ta.to_csv(os.path.join(subdir, "IncreaseDecrease_HSN.csv"),
              index=False, encoding="utf-8-sig")

    # Plots
    plot_inc_dec_paper(tr, "WMH/WM", subdir)
    plot_inc_dec_paper(ta, "HSN", subdir)

    # Within-group Wilcoxon (I2 vs I3 HSN volumes)
    rows = []
    for region, (f2, f3) in wmh_fields.items():
        for grp in ["Control", "Study"]:
            dgrp = df[df["group"] == grp].copy()
            y2 = dgrp[f2] * dgrp["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"]
            y3 = dgrp[f3] * dgrp["Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_3"]
            dgrp = dgrp.assign(y2=y2, y3=y3).dropna(subset=["y2", "y3"])
            if len(dgrp) == 0:
                stat, p = np.nan, np.nan
            else:
                try:
                    stat, p = stats.wilcoxon(dgrp["y2"], dgrp["y3"])
                except ValueError:
                    stat, p = np.nan, np.nan
            rows.append({
                "Outcome": region,
                "Group": grp,
                "N": len(dgrp),
                "I2 mean": np.mean(dgrp["y2"]) if len(dgrp) > 0 else np.nan,
                "I3 mean": np.mean(dgrp["y3"]) if len(dgrp) > 0 else np.nan,
                "Mean change": np.mean(dgrp["y3"] - dgrp["y2"]) if len(dgrp) > 0 else np.nan,
                "Median change": np.median(dgrp["y3"] - dgrp["y2"]) if len(dgrp) > 0 else np.nan,
                "Wilcoxon p": p,
            })

    pd.DataFrame(rows).to_csv(
        os.path.join(subdir, "WithinGroup_Wilcoxon.csv"),
        index=False, encoding="utf-8-sig"
    )

    # Follow-up summary
    write_followup_summary(df, label, subdir)

    # Two-panel combined figure
    plot_figureX_hsn(la, ta, subdir,
                     filename="FigureX_Longitudinal_WMH.png")

    # Methods record
    write_methods_record(subdir)

# ---------------- Run pipeline for both cohorts ----------------
for lab, fp in COHORTS.items():
    run_for_one_cohort(lab, fp)

print(f"\nAll outputs saved to: {os.path.abspath(OUTDIR)}")

# =====================================================================
# Combine BOTH cohorts & ALL FOUR longitudinal analyses into ONE eTable
# =====================================================================

import os as _os, re as _re
from pathlib import Path

OUTDIR = "Annual_Change"

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches
except Exception:
    Document = None

def _safe(s: str) -> str:
    return _re.sub(r"[^0-9A-Za-z._-]+", "_", str(s))

def _set_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left', 'right', 'top', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = OxmlElement(f"w:{edge}")
            for key in ("val", "sz", "color", "space"):
                if key in edge_data:
                    tag.set(qn(f"w:{key}"), str(edge_data[key]))
            tcBorders.append(tag)

def _three_line(table, header_row_idx=0):
    for row in table.rows:
        for cell in row.cells:
            _set_border(cell,
                        left={"val": "nil"},
                        right={"val": "nil"},
                        top={"val": "nil"},
                        bottom={"val": "nil"})
    for cell in table.rows[header_row_idx].cells:
        _set_border(cell,
                    top={"val": "single", "sz": 8, "color": "000000"},
                    bottom={"val": "single", "sz": 8, "color": "000000"})
    for cell in table.rows[-1].cells:
        _set_border(cell,
                    bottom={"val": "single", "sz": 8, "color": "000000"})

def _right_align(table, idxs):
    for row in table.rows[1:]:
        for j in idxs:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _fmt_num(x, nd=1):
    return "" if pd.isna(x) else f"{float(x):.{nd}f}"

def _fmt_p(x):
    return "—" if pd.isna(x) else ("<0.001" if float(x) < 1e-3 else f"{float(x):.3f}")

def _fmt_q(x):
    return "—" if pd.isna(x) else ("<0.001" if float(x) < 1e-3 else f"{float(x):.3f}")

ANL = [
    dict(key="HSN_log1p",
         file="WMH_HSN_abs_log1p_pct_change.csv",
         metric="HSN (log1p %/yr)",
         eff="%/year", lo="%/year CI Low", hi="%/year CI High",
         pct=True, q_mode="by_cohort"),
    dict(key="WMHWM_log1p",
         file="WMH_WM_ratio_log1p_pct_change.csv",
         metric="WMH/WM (log1p %/yr)",
         eff="%/year", lo="%/year CI Low", hi="%/year CI High",
         pct=True, q_mode=None),
    dict(key="HSN_raw",
         file="WMH_HSN_abs_raw_change.csv",
         metric="HSN (β/yr)",
         eff="β", lo="95% CI Low", hi="95% CI High",
         pct=False, q_mode=None),
    dict(key="WMHWM_raw",
         file="WMH_WM_ratio_raw_change.csv",
         metric="WMH/WM (β/yr)",
         eff="β", lo="95% CI Low", hi="95% CI High",
         pct=False, q_mode=None),
]

REGION_MAP = {"Total": "Total WMH", "Periventricular": "PWMH", "Deep": "DWMH"}

def _pick_q_row(row, cohort_label: str):
    c = cohort_label.lower()
    if c.startswith("primary") and ("q_FDR_primary" in row.index):
        return row.get("q_FDR_primary", np.nan)
    if c.startswith("sensitivity") and ("q_FDR_sensitivity" in row.index):
        return row.get("q_FDR_sensitivity", np.nan)
    return np.nan

def _gather_all():
    rows = []
    order_region = ["Total", "Periventricular", "Deep"]
    for cohort_label in COHORTS.keys():
        subdir = _os.path.join(OUTDIR, cohort_label)
        for cfg in ANL:
            p = _os.path.join(subdir, cfg["file"])
            if not _os.path.exists(p):
                print(f"[Warn] Missing: {p}")
                continue
            df = pd.read_csv(p)
            region = df["Outcome"].astype(str).str.split("·", n=1, expand=True)[0].str.strip()
            df["Region"] = region
            df = df[df["Region"].isin(order_region)].copy()
            for _, r in df.iterrows():
                eff = pd.to_numeric(r.get(cfg["eff"]), errors="coerce")
                lo = pd.to_numeric(r.get(cfg["lo"]), errors="coerce")
                hi = pd.to_numeric(r.get(cfg["hi"]), errors="coerce")
                pval = pd.to_numeric(r.get("p"), errors="coerce")
                qval = _pick_q_row(r, cohort_label) if cfg["q_mode"] == "by_cohort" else np.nan
                rows.append({
                    "Cohort": cohort_label,
                    "Metric": cfg["metric"],
                    "Outcome": REGION_MAP.get(r["Region"], r["Region"]),
                    "Effect": eff,
                    "CI Low": lo,
                    "CI High": hi,
                    "p": pval,
                    "q": qval,
                    "is_pct": cfg["pct"],
                })
    return pd.DataFrame(rows)

def export_single_eTable(all_df: pd.DataFrame,
                         title="eTable 7. Longitudinal WMH changes across analytic cohorts and metrics",
                         legend=(
                             "Adjusted annualized between-group differences are shown for Primary and "
                             "Sensitivity cohorts. Metrics: HSN (log1p %/yr), WMH/WM (log1p %/yr), "
                             "HSN (β/yr), WMH/WM (β/yr). For HSN (log1p %/yr), BH-FDR was applied "
                             "within cohort to PWMH and DWMH only; Total WMH excluded. Models adjust "
                             "for age at Instance 2, sex, BMI at Instance 0, genetic ethnic grouping, "
                             "smoking (ever), alcohol-intake frequency, Townsend index; cluster-robust "
                             "SEs by match_id or HC3 where clustering not feasible. Two-sided α=0.05."
                         )):
    if all_df is None or all_df.empty:
        print("[Info] No results to export.")
        return None

    all_df = all_df.copy()
    all_df["Analytic cohort"] = all_df["Cohort"].replace({
        "Primary": "Primary PSM cohort",
        "Sensitivity": "PSM–Sensitivity cohort",
    })
    all_df = all_df.drop(columns=["Cohort"])

    order_metric = [a["metric"] for a in ANL]
    order_region = ["Total WMH", "PWMH", "DWMH"]

    all_df["Analytic cohort"] = pd.Categorical(
        all_df["Analytic cohort"],
        categories=["Primary PSM cohort", "PSM–Sensitivity cohort"],
        ordered=True,
    )
    all_df["Metric"] = pd.Categorical(
        all_df["Metric"], categories=order_metric, ordered=True
    )
    all_df["Outcome"] = pd.Categorical(
        all_df["Outcome"], categories=order_region, ordered=True
    )

    df = all_df.sort_values(
        ["Analytic cohort", "Metric", "Outcome"]
    ).reset_index(drop=True)

    csv_path = _os.path.join(OUTDIR,
                             "Longitudinal_AllAnalyses_AllCohorts.csv")
    df_out = df[[
        "Analytic cohort", "Metric", "Outcome",
        "Effect", "CI Low", "CI High", "p", "q"
    ]].copy()
    df_out.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"Combined CSV saved: {csv_path}")

    if Document is None:
        print("python-docx not available: Word eTable skipped.")
        return None

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True

    headers = [
        "Analytic cohort", "Metric", "Outcome",
        "Effect", "95% CI (low, high)", "p", "q",
    ]
    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, r in df.iterrows():
        cells = t.add_row().cells
        ci_txt = "" if any(pd.isna([r["CI Low"], r["CI High"]])) else \
                 f"({_fmt_num(r['CI Low'], 1 if r['is_pct'] else 3)}, {_fmt_num(r['CI High'], 1 if r['is_pct'] else 3)})"
        vals = [
            str(r["Analytic cohort"]),
            str(r["Metric"]),
            str(r["Outcome"]),
            _fmt_num(r["Effect"], 1 if r["is_pct"] else 3),
            ci_txt,
            _fmt_p(r["p"]),
            _fmt_q(r["q"]),
        ]
        for j, v in enumerate(vals):
            cells[j].text = v

    _right_align(t, idxs=[3, 4, 5, 6])
    _three_line(t, header_row_idx=0)

    doc.add_paragraph().add_run(legend).italic = True

    docx_path = _os.path.join(
        OUTDIR, "Longitudinal_AllAnalyses_AllCohorts.docx"
    )
    doc.save(docx_path)
    print(f"Word eTable saved: {docx_path}")
    return docx_path

_all = _gather_all()
export_single_eTable(_all)

# =====================================================================
# eTable: Increase vs Decrease (between-group)
# =====================================================================

def export_incdec_eTable(cohorts=COHORTS,
                         title="eTable 8. Increase vs Decrease tests across analytic cohorts and WMH regions",
                         legend=(
                             "Proportion of participants with increased WMH per year shown as n (%). "
                             "Tests: Fisher's exact or Chi-square. BH-FDR applied only to PWMH and DWMH; "
                             "Total WMH excluded."
                         )):
    rows = []
    for cohort_label in cohorts.keys():
        if cohort_label == "Primary":
            cohort_name = "Primary PSM cohort"
        elif cohort_label == "Sensitivity":
            cohort_name = "PSM–Sensitivity cohort"
        else:
            cohort_name = cohort_label

        subdir = _os.path.join(OUTDIR, cohort_label)
        for typ, fname in [("HSN", "IncreaseDecrease_HSN.csv"),
                           ("WMH/WM", "IncreaseDecrease_Ratio.csv")]:
            fpath = _os.path.join(subdir, fname)
            if not _os.path.exists(fpath):
                print(f"[Warn] Missing: {fpath}")
                continue
            df = pd.read_csv(fpath)

            for _, r in df.iterrows():
                n_s = int(r.get("Study N", np.nan)) if not pd.isna(r.get("Study N", np.nan)) else 0
                n_c = int(r.get("Control N", np.nan)) if not pd.isna(r.get("Control N", np.nan)) else 0
                inc_s = int(r.get("Study Increase", np.nan)) if not pd.isna(r.get("Study Increase", np.nan)) else 0
                inc_c = int(r.get("Control Increase", np.nan)) if not pd.isna(r.get("Control Increase", np.nan)) else 0
                pct_s = (100 * inc_s / n_s) if n_s > 0 else np.nan
                pct_c = (100 * inc_c / n_c) if n_c > 0 else np.nan

                study_str = f"{inc_s} ({pct_s:.1f}%)" if n_s > 0 else "—"
                ctrl_str = f"{inc_c} ({pct_c:.1f}%)" if n_c > 0 else "—"
                n_str = f"{n_s} / {n_c}" if (n_s > 0 and n_c > 0) else "—"

                rows.append({
                    "Analytic cohort": cohort_name,
                    "N (SA/Control)": n_str,
                    "Metric": typ,
                    "Outcome": REGION_MAP.get(r["Outcome"], r["Outcome"]),
                    "Study Increase n (%)": study_str,
                    "Control Increase n (%)": ctrl_str,
                    "Prop Diff (S-C)": _fmt_num(r.get("Prop Diff (S-C)"), 3),
                    "95% CI": ("" if any(pd.isna([r.get("95% CI Low"),
                                                  r.get("95% CI High")]))
                               else f"({_fmt_num(r['95% CI Low'], 3)}, "
                                    f"{_fmt_num(r['95% CI High'], 3)})"),
                    "p": _fmt_p(r.get("p_primary")),
                    "q": _fmt_q(r.get("q_FDR_incdec_combined", np.nan)),
                })

    if not rows:
        print("[Info] No Increase/Decrease results found.")
        return None

    df_all = pd.DataFrame(rows)

    metric_order = ["HSN", "WMH/WM"]
    outcome_order = ["Total WMH", "PWMH", "DWMH"]

    df_all["Metric"] = pd.Categorical(
        df_all["Metric"], categories=metric_order, ordered=True
    )
    df_all["Outcome"] = pd.Categorical(
        df_all["Outcome"], categories=outcome_order, ordered=True
    )
    df_all["Analytic cohort"] = pd.Categorical(
        df_all["Analytic cohort"],
        categories=["Primary PSM cohort", "PSM–Sensitivity cohort"],
        ordered=True,
    )

    df_all = df_all.sort_values(
        ["Analytic cohort", "Metric", "Outcome"]
    ).reset_index(drop=True)

    csv_path = _os.path.join(OUTDIR, "IncreaseDecrease_AllCohorts.csv")
    df_all.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"CSV saved: {csv_path}")

    if Document is None:
        print("python-docx not available: Word eTable skipped.")
        return None

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True

    headers = [
        "Analytic cohort", "N (SA/Control)", "Metric", "Outcome",
        "Study Increase n (%)", "Control Increase n (%)",
        "Prop Diff (S-C)", "95% CI", "p", "q",
    ]
    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, r in df_all.iterrows():
        cells = t.add_row().cells
        vals = [str(r[h]) for h in headers]
        for j, v in enumerate(vals):
            cells[j].text = v

    _right_align(t, idxs=[1, 4, 5, 6, 7, 8, 9])
    _three_line(t, header_row_idx=0)

    doc.add_paragraph().add_run(legend).italic = True

    docx_path = _os.path.join(OUTDIR, "IncreaseDecrease_AllCohorts.docx")
    doc.save(docx_path)
    print(f"Word eTable saved: {docx_path}")
    return docx_path

export_incdec_eTable()

# =====================================================================
# Within-group Increase vs Decrease (binomial vs 50%)
# =====================================================================

def export_within_group_incdec(cohorts=COHORTS,
                               title="eTable 11. Within-group Increase vs Decrease tests (SA and Control, by cohort)",
                               legend=(
                                   "Within-group exact binomial tests (two-sided, p0=0.5) assess whether the "
                                   "proportion with WMH increase vs decrease deviates from 50%, separately in "
                                   "SA and Control. BH-FDR applied only to PWMH and DWMH; Total WMH excluded."
                               )):
    rows = []

    for cohort_label in cohorts.keys():
        cohort_name = ("Primary PSM cohort" if cohort_label == "Primary"
                       else "PSM–Sensitivity cohort" if cohort_label == "Sensitivity"
                       else cohort_label)
        subdir = _os.path.join(OUTDIR, cohort_label)

        for typ, fname in [("HSN", "IncreaseDecrease_HSN.csv"),
                           ("WMH/WM", "IncreaseDecrease_Ratio.csv")]:
            fpath = _os.path.join(subdir, fname)
            if not _os.path.exists(fpath):
                print(f"[Warn] Missing: {fpath}")
                continue
            df = pd.read_csv(fpath)
            df = df.rename(columns=lambda x: x.replace("Study", "SA") if isinstance(x, str) else x)

            for _, r in df.iterrows():
                outcome = REGION_MAP.get(r["Outcome"], r["Outcome"])
                for grp in ["SA", "Control"]:
                    n = int(r.get(f"{grp} N", np.nan)) if not pd.isna(r.get(f"{grp} N", np.nan)) else 0
                    inc = int(r.get(f"{grp} Increase", np.nan)) if not pd.isna(r.get(f"{grp} Increase", np.nan)) else 0
                    dec = int(r.get(f"{grp} Decrease", np.nan)) if not pd.isna(r.get(f"{grp} Decrease", np.nan)) else 0
                    nn = inc + dec
                    if nn <= 0:
                        continue

                    bt = stats.binomtest(k=inc, n=nn, p=0.5, alternative="two-sided")
                    pval = float(bt.pvalue)

                    rows.append({
                        "Analytic cohort": cohort_name,
                        "Metric": typ,
                        "Outcome": outcome,
                        "Group": grp,
                        "N": nn,
                        "Increase n": inc,
                        "Decrease n": dec,
                        "Increase %": f"{100 * inc / nn:.1f}%",
                        "Decrease %": f"{100 * dec / nn:.1f}%",
                        "Test": "Exact binomial (p0=0.5)",
                        "p": pval,
                        "q": np.nan,
                    })

    if not rows:
        print("[Info] No within-group results found.")
        return None

    df_all = pd.DataFrame(rows)

    df_all["Analytic cohort"] = pd.Categorical(
        df_all["Analytic cohort"],
        categories=["Primary PSM cohort", "PSM–Sensitivity cohort"],
        ordered=True,
    )

    outcome_order = {
        "Total WMH": 0,
        "PWMH": 1, "Periventricular WMH": 1, "Periventricular": 1,
        "DWMH": 2, "Deep WMH": 2, "Deep": 2,
    }
    df_all["Outcome_order"] = df_all["Outcome"].map(outcome_order).fillna(99)

    df_all = df_all.sort_values(
        ["Analytic cohort", "Metric", "Group", "Outcome_order"]
    ).reset_index(drop=True)
    df_all = df_all.drop(columns=["Outcome_order"])

    # FDR within (cohort × metric × group) for PWMH/DWMH only
    df_all["q"] = np.nan
    for (cohort_name, typ, grp), sub in df_all.groupby(["Analytic cohort", "Metric", "Group"]):
        mask = sub["Outcome"].isin(["PWMH", "Periventricular WMH", "Periventricular",
                                    "DWMH", "Deep WMH", "Deep"])
        idx = sub.index[mask & pd.to_numeric(sub["p"], errors="coerce").notna()]
        if len(idx) >= 1:
            p_clean = pd.to_numeric(sub.loc[idx, "p"], errors="coerce").values
            qvals = multipletests(p_clean, method="fdr_bh")[1]
            df_all.loc[idx, "q"] = qvals


    csv_path = _os.path.join(OUTDIR, "WithinGroup_IncDec_AllCohorts.csv")
    df_out = df_all.copy()
    df_out["p"] = df_out["p"].apply(_fmt_p)
    df_out["q"] = df_out["q"].apply(_fmt_q)
    df_out.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"Within-group Inc/Dec CSV saved: {csv_path}")

    if Document is None:
        print("python-docx not available: Word eTable skipped.")
        return None

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True

    headers = [
        "Analytic cohort", "Metric", "Outcome", "Group", "N",
        "Increase n", "Decrease n",
        "Increase %", "Decrease %",
        "p", "q",
    ]
    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, r in df_out.iterrows():
        cells = t.add_row().cells
        vals = [str(r[h]) for h in headers]
        for j, v in enumerate(vals):
            cells[j].text = v

    _right_align(t, idxs=[4, 5, 6, 7, 8, 9, 10])
    _three_line(t, header_row_idx=0)

    doc.add_paragraph().add_run(legend).italic = True

    docx_path = _os.path.join(OUTDIR, "WithinGroup_IncDec_AllCohorts.docx")
    doc.save(docx_path)
    print(f"Word eTable saved: {docx_path}")
    return docx_path

export_within_group_incdec()

# =====================================================================
# SA I2 - Dx distribution (final; explicit Dx column)
# =====================================================================

SAVE_DIR = os.path.join(OUTDIR, "SA_Dx_Timing")

def _plot_sa_i2_minus_dx_for_cohort(cohort_label, csv_path, outdir=SAVE_DIR):
    """Plot SA (Study) distribution of (Instance 2 date – SA diagnosis date) in years, per cohort."""
    if not os.path.exists(csv_path):
        print(f"[Skip] {cohort_label}: file not found -> {csv_path}")
        return

    # Load and normalize column names
    df = clean_cols(pd.read_csv(csv_path))

    # Required columns (diagnosis column corrected)
    req = [
        "group",
        "Date_of_attending_assessment_centre_Instance_2",
        "Date_G47_first_reported_sleep_disorders",
    ]
    missing = [c for c in req if c not in df.columns]
    if missing:
        print(f"[Skip] {cohort_label}: missing columns -> {missing}")
        return

    # SA only
    d = df[df["group"].astype(str).str.strip().str.capitalize() == "Study"].copy()
    if d.empty:
        print(f"[Info] {cohort_label}: Study group is empty.")
        return

    # Compute (I2 – Dx) in years
    i2 = pd.to_datetime(d["Date_of_attending_assessment_centre_Instance_2"], errors="coerce")
    dx = pd.to_datetime(d["Date_G47_first_reported_sleep_disorders"], errors="coerce")
    mask = i2.notna() & dx.notna()
    if not mask.any():
        print(f"[Info] {cohort_label}: no valid I2 and Dx pairs.")
        return

    diff_years = (i2[mask] - dx[mask]).dt.days / 365.25
    diff_years = diff_years.replace([np.inf, -np.inf], np.nan).dropna()
    if diff_years.empty:
        print(f"[Info] {cohort_label}: I2–Dx data empty after cleaning.")
        return

    # Descriptives
    n = diff_years.size
    mu = diff_years.mean()
    sd = diff_years.std(ddof=1)
    q1, med, q3 = diff_years.quantile([0.25, 0.50, 0.75])
    mn, mx = diff_years.min(), diff_years.max()
    n_pos = int((diff_years > 0).sum())
    n_neg = int((diff_years < 0).sum())

    print(
        f"[{cohort_label}] SA I2–Dx (years): N={n} | mean={mu:.2f}, SD={sd:.2f} | "
        f"median={med:.2f}, IQR={q1:.2f}–{q3:.2f} | min={mn:.2f}, max={mx:.2f} | "
        f"positive={n_pos}, negative={n_neg}"
    )

    # Histogram
    lo, hi = int(np.floor(mn)), int(np.ceil(mx))
    bins = np.arange(lo, hi + 1, 1) if hi > lo else 15

    plt.rcParams.update(PUB_RC)
    fig, ax = plt.subplots(figsize=(6.8, 4.2))
    ax.hist(diff_years, bins=bins, edgecolor="white", label="SA (Study)")
    ax.axvline(0, color="grey", ls="--", lw=1)
    ax.set_title(f"{cohort_label} — SA: Timing of I2 relative to diagnosis (I2 – Dx)")
    ax.set_xlabel("Years (I2 – Dx) [positive = Dx before I2; negative = Dx after I2]")
    ax.set_ylabel("Participants")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    txt = (
        f"N={n}\nMean={mu:.2f}, SD={sd:.2f}\n"
        f"Median={med:.2f} (IQR {q1:.2f}–{q3:.2f})\n"
        f"Min={mn:.2f}, Max={mx:.2f}"
    )
    ax.text(
        0.98, 0.98, txt,
        transform=ax.transAxes,
        ha="right", va="top",
        fontsize=9,
        bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="#cccccc"),
    )

    ax.legend(loc="center left", bbox_to_anchor=(1.02, 0.5), frameon=False)
    plt.tight_layout()

    # Save
    out_subdir = os.path.join(outdir, cohort_label)
    os.makedirs(out_subdir, exist_ok=True)
    base = os.path.join(out_subdir, "SA_I2_minus_Dx_Distribution")
    fig.savefig(base + ".png", **PNG_KW)
    fig.savefig(base + ".pdf", **PDF_KW)
    fig.savefig(base + ".svg", **SVG_KW)
    fig.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)



for _lab, _fp in COHORTS.items():
    _plot_sa_i2_minus_dx_for_cohort(_lab, _fp, outdir=SAVE_DIR)

# =====================================================================
# SA-only — Dx duration (I2 - Dx) by Increase vs Decrease (Total HSN)
# =====================================================================

from scipy import stats as _stats

def _summarize_series_noclip(s: pd.Series):
    s = pd.to_numeric(s, errors="coerce").dropna()
    if s.empty:
        return dict(N=0, Mean=np.nan, SD=np.nan,
                    Median=np.nan, Q1=np.nan, Q3=np.nan,
                    IQR=np.nan, Min=np.nan, Max=np.nan)
    mean = float(s.mean())
    sd = float(s.std(ddof=1))
    q1 = float(s.quantile(0.25))
    q3 = float(s.quantile(0.75))
    return dict(
        N=int(s.size),
        Mean=mean,
        SD=sd,
        Median=float(s.median()),
        Q1=q1,
        Q3=q3,
        IQR=q3 - q1,
        Min=float(s.min()),
        Max=float(s.max()),
    )

def _analyze_sa_dx_duration_by_incdec(cohort_label, filepath, out_root):
    """
    SA-only, split by Total WMH HSN absolute change (Increase vs Decrease),
    outcome = (I2 - Dx) in years (no clipping). Mann–Whitney U test.
    """
    if not os.path.exists(filepath):
        print(f"[Skip] {cohort_label}: file not found -> {filepath}")
        return

    df = clean_cols(pd.read_csv(filepath))

    col_I2 = "Date_of_attending_assessment_centre_Instance_2"
    col_T2 = "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"
    col_T3 = "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_3"
    col_S2 = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
    col_S3 = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_3"
    col_DX = "Date_G47_first_reported_sleep_disorders"

    needed = [col_I2, col_T2, col_T3, col_S2, col_S3, col_DX, "group"]
    miss = [c for c in needed if c not in df.columns]
    if miss:
        print(f"[Skip] {cohort_label}: missing columns -> {miss}")
        return

    d = df[df["group"].astype(str).str.strip().str.capitalize() == "Study"].copy()
    if d.empty:
        print(f"[Info] {cohort_label}: SA group empty.")
        return

    d[col_I2] = pd.to_datetime(d[col_I2], errors="coerce")
    d[col_DX] = pd.to_datetime(d[col_DX], errors="coerce")
    d["Dx_years_I2_minus_Dx"] = (
        d[col_I2] - d[col_DX]
    ).dt.days / 365.25

    d = d.dropna(subset=[
        "Dx_years_I2_minus_Dx",
        col_T2, col_T3, col_S2, col_S3
    ]).copy()
    if d.empty:
        print(f"[Info] {cohort_label}: no valid I2-Dx values in SA group.")
        return

    v2 = pd.to_numeric(d[col_T2], errors="coerce") * pd.to_numeric(d[col_S2], errors="coerce")
    v3 = pd.to_numeric(d[col_T3], errors="coerce") * pd.to_numeric(d[col_S3], errors="coerce")
    d["IncDec"] = np.where(v3 > v2, "Increase", "Decrease")

    inc = d.loc[d["IncDec"] == "Increase", "Dx_years_I2_minus_Dx"]
    dec = d.loc[d["IncDec"] == "Decrease", "Dx_years_I2_minus_Dx"]
    all_sa = d["Dx_years_I2_minus_Dx"]

    sum_inc = _summarize_series_noclip(inc)
    sum_dec = _summarize_series_noclip(dec)
    sum_all = _summarize_series_noclip(all_sa)

    if sum_inc["N"] > 0 and sum_dec["N"] > 0:
        try:
            U, p_mwu = _stats.mannwhitneyu(inc, dec, alternative="two-sided")
        except Exception:
            U, p_mwu = np.nan, np.nan
    else:
        U, p_mwu = np.nan, np.nan

    outdir = os.path.join(out_root, "SA_DxDuration_ByIncDec", cohort_label)
    os.makedirs(outdir, exist_ok=True)

    def _fmt(v, nd=2):
        return "NA" if (v is None or (isinstance(v, float) and np.isnan(v))) else f"{float(v):.{nd}f}"

    tbl = pd.DataFrame([
        {"Cohort": cohort_label, "Subgroup": "Increase", **sum_inc},
        {"Cohort": cohort_label, "Subgroup": "Decrease", **sum_dec},
        {"Cohort": cohort_label, "Subgroup": "Overall SA", **sum_all},
    ])
    tbl["Mean±SD (years)"] = tbl.apply(
        lambda r: f"{_fmt(r['Mean'])}±{_fmt(r['SD'])}", axis=1
    )
    tbl["Median [IQR] (years)"] = tbl.apply(
        lambda r: f"{_fmt(r['Median'])} [{_fmt(r['Q1'])}, {_fmt(r['Q3'])}]", axis=1
    )
    tbl["Min–Max (years)"] = tbl.apply(
        lambda r: f"{_fmt(r['Min'])}–{_fmt(r['Max'])}", axis=1
    )
    keep = [
        "Cohort", "Subgroup", "N",
        "Mean±SD (years)",
        "Median [IQR] (years)",
        "Min–Max (years)",
    ]

    if pd.notna(p_mwu):
        p_txt = "<0.001" if p_mwu < 0.001 else f"{p_mwu:.3f}"
    else:
        p_txt = "NA"

    out_rows = tbl[keep].copy()
    out_rows.loc[len(out_rows)] = [
        cohort_label,
        "Mann–Whitney U (two-sided) p",
        "",
        "",
        "",
        p_txt,
    ]

    csv_path = os.path.join(outdir, f"{cohort_label}_SA_DxDuration_IncDec.csv")
    out_rows.to_csv(csv_path, index=False, encoding="utf-8-sig")
    print(f"[OK] SA Dx-duration (I2-Dx) table saved: {csv_path}")

    # Figure
    plt.rcParams.update(PUB_RC)
    fig, ax = plt.subplots(figsize=(6.8, 4.6))
    d["IncDec"] = pd.Categorical(
        d["IncDec"], categories=["Increase", "Decrease"], ordered=True
    )
    plot_data = [
        d.loc[d["IncDec"] == "Increase", "Dx_years_I2_minus_Dx"],
        d.loc[d["IncDec"] == "Decrease", "Dx_years_I2_minus_Dx"],
    ]

    bp = ax.boxplot(
        plot_data,
        positions=[0, 1],
        widths=0.55,
        vert=True,
        showfliers=True,
        patch_artist=True,
        medianprops=dict(linewidth=1.4),
        boxprops=dict(linewidth=1.0),
        whiskerprops=dict(linewidth=1.0),
        capprops=dict(linewidth=1.0),
    )

    col_inc = COLOR_PRIMARY
    col_dec = "#9c9c9c"
    bp["boxes"][0].set_facecolor(col_inc)
    bp["boxes"][1].set_facecolor(col_dec)

    rng = np.random.default_rng(42)
    for i, yv in enumerate(plot_data):
        x = np.full(len(yv), i, dtype=float) + rng.uniform(-0.10, 0.10, size=len(yv))
        ax.scatter(
            x,
            yv,
            s=18,
            alpha=0.70,
            edgecolors="white",
            linewidths=0.4,
            c=(col_inc if i == 0 else col_dec),
            label=("Increase" if i == 0 else "Decrease"),
            zorder=3,
        )

    ax.set_xticks([0, 1])
    ax.set_xticklabels(["Increase", "Decrease"])
    ax.set_ylabel("Diagnosis duration (years, I2 - Dx)")
    ax.set_title(
        f"{cohort_label} — SA: I2 - Dx duration by subgroup (Increase vs Decrease)",
        pad=12,
    )

    if pd.notna(p_mwu):
        ax.text(
            0.5,
            0.98,
            f"Mann–Whitney U: p={p_txt}",
            ha="center",
            va="top",
            transform=ax.transAxes,
            fontsize=10,
        )

    ax.grid(axis="y", linestyle=":", alpha=0.45)
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    ax.legend(
        loc="center left",
        bbox_to_anchor=(1.02, 0.5),
        frameon=False,
    )
    plt.subplots_adjust(right=0.76, left=0.12, bottom=0.15, top=0.88)

    base = os.path.join(outdir, f"{cohort_label}_SA_DxDuration_IncDec")
    fig.savefig(base + ".png", **PNG_KW)
    fig.savefig(base + ".pdf", **PDF_KW)
    fig.savefig(base + ".svg", **SVG_KW)
    fig.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)

for _lab, _fp in COHORTS.items():
    _analyze_sa_dx_duration_by_incdec(_lab, _fp, OUTDIR)

# ========================= End of longitudinal script =========================
