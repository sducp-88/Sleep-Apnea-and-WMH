# Extracted from WMH_Analysis.ipynb, code cell 12.
# Notebook heading: # Stratified OLS of HSN-WMH (log1p) with in-stratum BH-FDR
# Run this file from the repository root unless a local CONFIG section is edited.

# Stratified OLS of HSN-WMH (log1p) with in-stratum BH-FDR
# ----------------------------------------------------------------------------------
# Updates per request:
#   1) Keep the stratification formerly called "SA severity" but RENAME it to
#      "SA diagnosis method" using the column SA_ascertain_group with levels:
#         ["Self_Report_Only", "Hospital_Secondary", "Hospital_Primary"].
#      As with other strata, only matched sets where the Study arm belongs to
#      a given level are included for that level.
#   2) Do NOT add CMC as an adjustment covariate; legends and models have no CMC.
#
# Analysis notes:
#   • Within each stratum level, apply BH-FDR to {Periventricular WMH, Deep WMH};
#     Total WMH is excluded from FDR (q = NaN).
#   • Cohorts: primary_cohort_stratified.csv & sensitivity_cohort_stratified.csv.
#   • Outputs under Stratify/<Cohort>; a combined CSV & Word table under Stratify/.
# ----------------------------------------------------------------------------------

import os, re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import statsmodels.api as sm
from patsy import dmatrices
from statsmodels.stats.multitest import multipletests

# ---------- Optional Word export ----------
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception:
    Document = None

# ---------------- Paths & cohorts ----------------
COHORTS = {
    "Primary":     "primary_cohort_stratified.csv",
    "Sensitivity": "sensitivity_cohort_stratified.csv",
}
OUTDIR = "Stratify"
os.makedirs(OUTDIR, exist_ok=True)

# ---- Export defaults ----
SAVEFIG_DPI = 600
PNG_KW  = {"dpi": SAVEFIG_DPI, "bbox_inches": "tight", "facecolor": "white"}
PDF_KW  = {"bbox_inches": "tight"}
TIFF_KW = {"format": "tiff", "dpi": SAVEFIG_DPI, "bbox_inches": "tight",
           "facecolor": "white", "pil_kwargs": {"compression": "tiff_lzw"}}

# ---------------- Columns & settings ----------------
GROUP_COL, MATCH_ID = "group", "match_id"
SCALE_COL = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
RAW_DEEP  = "Total_volume_of_deep_white_matter_hyperintensities_Instance_2"
RAW_PV    = "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"
RAW_TOT   = ("Total_volume_of_white_matter_hyperintensities_from_T1_"
             "and_T2_FLAIR_images_Instance_2")

OUTCOME_BUILD = [
    ("Log_HSNorm_Total_WMH",           RAW_TOT,  "Total WMH"),
    ("Log_HSNorm_PeriVentricular_WMH", RAW_PV,   "Periventricular WMH"),
    ("Log_HSNorm_Deep_WMH",            RAW_DEEP, "Deep WMH"),
]
OUTCOME_ORDER = ["Total WMH", "Periventricular WMH", "Deep WMH"]

# Prespecified covariates (NO CMC here)
ADJUST_ALL = [
    "Sex","Age_at_Instance_2","Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0","Genetic_ethnic_grouping",
    "Smoking_Ever","Alcohol_intake_frequency_ordinal",
]
CATEGORICAL_ALL = {"Sex","Genetic_ethnic_grouping","Smoking_Ever"}

# ---------------- Helpers ----------------
def safe_name(s: str) -> str:
    return re.sub(r"[^0-9A-Za-z._-]+", "_", str(s))

def clean_cols(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns=lambda c: pd.Series(c).str.replace(r"[^0-9a-zA-Z]+","_", regex=True).iloc[0])

def ensure_outcomes(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    if GROUP_COL in df.columns:
        df[GROUP_COL] = pd.Categorical(df[GROUP_COL], ["Control","Study"])
    scale = pd.to_numeric(df.get(SCALE_COL, 1.0), errors="coerce").fillna(1.0)
    for out_col, src, _lab in OUTCOME_BUILD:
        v = pd.to_numeric(df.get(src), errors="coerce")
        v = v * scale
        df[out_col] = np.log1p(v)
    return df

def build_formula(outcome: str, data: pd.DataFrame) -> str:
    rhs = ["C(group)"]
    for v in ADJUST_ALL:
        if v in data.columns:
            rhs.append(f"C({v})" if v in CATEGORICAL_ALL else v)
    return f"{outcome} ~ " + " + ".join(rhs)

def fit_cluster(formula: str, data: pd.DataFrame):
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    base = sm.OLS(y, X).fit(cov_type="HC3")
    if MATCH_ID not in data.columns:
        return base
    groups = data.loc[y.index, MATCH_ID].dropna()
    if groups.nunique() <= 1:
        return base
    return sm.OLS(y.loc[groups.index], X.loc[groups.index]).fit(
        cov_type="cluster", cov_kwds={"groups": groups}
    )

# ---------------- Forest plot ----------------
def forest_plot(df_plot: pd.DataFrame, title: str, out_png: str, use_abbrev=True):
    if df_plot is None or df_plot.empty: return
    ORDER = ["Deep WMH", "Periventricular WMH", "Total WMH"]
    ABBR  = {"Periventricular WMH": "PWMH", "Deep WMH": "DWMH"}

    dfp = df_plot.copy()
    dfp["Outcome"] = pd.Categorical(dfp["Outcome"], categories=ORDER, ordered=True)
    dfp = dfp.sort_values(["Outcome","Stratum"]).reset_index(drop=True)

    def row_color(row):
        qcols = [c for c in row.index if c.lower().startswith("q_fdr")]
        qvals = [row[c] for c in qcols if pd.notna(row[c])]
        if qvals and np.nanmin(qvals) <= 0.05: return "#1f3b4d"
        if pd.notna(row.get("p")) and float(row["p"]) <= 0.05: return "#1f3b4d"
        return "#7f7f7f"

    colors = [row_color(r) for _, r in dfp.iterrows()]
    x  = pd.to_numeric(dfp["% change"], errors="coerce").to_numpy()
    xl = pd.to_numeric(dfp["% CI Low"], errors="coerce").to_numpy()
    xu = pd.to_numeric(dfp["% CI High"], errors="coerce").to_numpy()
    n  = len(dfp); y = np.arange(n)

    plt.rcParams.update({"font.family":"Arial","font.size":12,
                         "axes.labelsize":12,"xtick.labelsize":11,"ytick.labelsize":11,
                         "savefig.dpi":SAVEFIG_DPI})
    fig = plt.figure(figsize=(6.2, max(3.6, 0.55*n))); ax = plt.gca()

    for i in range(n):
        if np.isnan(x[i]) or np.isnan(xl[i]) or np.isnan(xu[i]): continue
        ax.errorbar(x[i], y[i],
                    xerr=[[x[i]-xl[i]], [xu[i]-x[i]]],
                    fmt="o", mfc=colors[i], mec=colors[i],
                    ecolor=colors[i], elinewidth=1.8, capsize=4, markersize=6.0,
                    linestyle="none", zorder=3)

    xmin = np.nanmin(xl) if np.isfinite(np.nanmin(xl)) else -1.0
    xmax = np.nanmax(xu) if np.isfinite(np.nanmax(xu)) else  1.0
    pad  = max(0.12*(xmax-xmin if xmax>xmin else 1.0), 2.0)
    ax.set_xlim(xmin - pad, xmax + pad)

    def fmt_outcome(o): return ABBR.get(o, o) if use_abbrev else o
    ylabels = (dfp["Outcome"].map(fmt_outcome).astype(str) + " (" + dfp["Stratum"].astype(str) + ")").tolist()
    ax.set_yticks(y); ax.set_yticklabels(ylabels)

    ax.axvline(0, color="grey", linestyle="--", linewidth=1)
    ax.grid(False); ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.set_xlabel("% change (SA - Control)")
    plt.tight_layout()

    plt.savefig(out_png, **PNG_KW)
    plt.savefig(os.path.splitext(out_png)[0] + ".pdf", **PDF_KW)
    plt.savefig(os.path.splitext(out_png)[0] + ".tiff", **TIFF_KW)
    plt.close(fig)

# ---------------- Word helpers ----------------
def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge); tag = OxmlElement(f'w:{edge}')
            for key in ("val","sz","color","space"):
                if key in edge_data: tag.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(tag)

def _apply_three_line_table(table, header_row_idx=0):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, left={"val":"nil"}, right={"val":"nil"},
                                   top={"val":"nil"},  bottom={"val":"nil"})
    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell, top={"val":"single","sz":8,"color":"000000"},
                              bottom={"val":"single","sz":8,"color":"000000"})
    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom={"val":"single","sz":8,"color":"000000"})

def _right_align_numeric(table, numeric_col_idx):
    for row in table.rows[1:]:
        for j in numeric_col_idx:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def export_stratified_word(df_all: pd.DataFrame, outdir: str,
                           title: str = "eTable. Stratified analyses of WMH",
                           legend: str | None = None,
                           filename: str = "Stratified_AllCohorts_Supplement.docx"):
    if Document is None or df_all is None or df_all.empty:
        print("Word export skipped."); return None

    # --- 兼容列名：内部统一使用 df["Cohort"] 处理 ---
    df = df_all.copy()
    if "Cohort" not in df.columns and "Analytic Cohort" in df.columns:
        df = df.rename(columns={"Analytic Cohort": "Cohort"})

    def _fmt_num(x, nd=1):  return "" if pd.isna(x) else f"{float(x):.{nd}f}"
    def _fmt_p(x):          return "—" if pd.isna(x) else ("<0.001" if float(x) < 1e-3 else f"{float(x):.3f}")
    def _fmt_q(x):          return "—" if pd.isna(x) else ("<0.001" if float(x) < 1e-3 else f"{float(x):.3f}")

    # 选择 q 列（若存在分队列的 q）
    def _pick_q(row):
        c = str(row.get("Cohort","")).lower()
        if c.startswith("primary"):     return row.get("q_FDR_primary", np.nan)
        if c.startswith("psm–sensitivity") or c.startswith("sensitivity"):
            return row.get("q_FDR_sensitivity", np.nan)
        qa, qb = row.get("q_FDR_primary", np.nan), row.get("q_FDR_sensitivity", np.nan)
        return qa if pd.notna(qa) else qb

    df["q"] = df.apply(_pick_q, axis=1)
    df["95% CI"] = df.apply(
        lambda r: "" if any(pd.isna([r.get("% CI Low"), r.get("% CI High")]))
        else f"({_fmt_num(r['% CI Low'])}, {_fmt_num(r['% CI High'])})", axis=1
    )

    # 输出列：注意第一列显示为“Analytic Cohort”
    cols_out = ["Cohort","Stratification","Stratum","Outcome","N","% change","95% CI","p","q"]
    df_out = pd.DataFrame({
        "Cohort": df["Cohort"],  # 内部 Cohort，下面表头改成 Analytic Cohort
        "Stratification": df.get("Stratification", ""),
        "Stratum": df.get("Stratum", ""),
        "Outcome": df.get("Outcome", ""),
        "N": df.get("N", ""),
        "% change": df.get("% change", np.nan).apply(lambda v: _fmt_num(v, 1)),
        "95% CI": df["95% CI"],
        "p": df.get("p", np.nan).apply(_fmt_p),
        "q": df.get("q", np.nan).apply(_fmt_q),
    })[cols_out]

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph(); p.add_run(title).bold = True

    # 表头：第一列改为 Analytic Cohort
    headers = ["Analytic Cohort","Stratification","Stratum","Outcome",
            "N (stratum / total SA)","%Δ","95% CI (low, high)","p","q"]
    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, row in df_out.iterrows():
        cells = t.add_row().cells
        cells[0].text = "" if pd.isna(row["Cohort"]) else str(row["Cohort"])
        cells[1].text = "" if pd.isna(row["Stratification"]) else str(row["Stratification"])
        cells[2].text = "" if pd.isna(row["Stratum"]) else str(row["Stratum"])
        cells[3].text = "" if pd.isna(row["Outcome"]) else str(row["Outcome"])
        cells[4].text = "" if pd.isna(row["N"]) else str(row["N"])
        cells[5].text = "" if pd.isna(row["% change"]) else str(row["% change"])
        cells[6].text = "" if pd.isna(row["95% CI"]) else str(row["95% CI"])
        cells[7].text = "" if pd.isna(row["p"]) else str(row["p"])
        cells[8].text = "" if pd.isna(row["q"]) else str(row["q"])

    _right_align_numeric(t, numeric_col_idx=[5,6,7,8])
    _apply_three_line_table(t, header_row_idx=0)

    if legend is None:
        legend = ("Legend: Adjusted between-group differences in log1p-transformed WMH expressed as percent change "
                  "(Study vs Control) with 95% CIs. N is Study/Control per stratum. Benjamini–Hochberg FDR is applied "
                  "within each stratum to {Periventricular, Deep}; Total WMH excluded. Models adjust for prespecified "
                  "covariates; cluster-robust SEs by matched set (match_id) or HC3 when clustering is infeasible.")
    doc.add_paragraph().add_run(legend).italic = True

    path = os.path.join(outdir, filename)
    doc.save(path); print(f"Word table saved: {path}")
    return path


# ---------------- One stratification block ----------------
def run_one_block(df: pd.DataFrame,
                  strat_col: str,
                  levels: list,
                  label_map: dict,
                  title: str,
                  cohort_label: str) -> pd.DataFrame:
    """
    Run a full stratified analysis for a single stratification variable.

    Key changes:
      • N is reported as "stratum SA / total SA participating in this stratification".
        - Numerator: number of SA (group == 'Study') whose value == current stratum level.
        - Denominator: number of SA (group == 'Study') that have a non-missing value in this
          stratification variable and whose value is one of `levels`.
      • Matched sets are restricted so that the Study arm belongs to the current stratum level.
      • Within each stratum, BH-FDR is applied to {Periventricular WMH, Deep WMH}; Total WMH excluded.
    """
    # Skip if the stratification column is missing or entirely NA
    if strat_col not in df.columns or df[strat_col].dropna().empty:
        print(f"[Skip] {strat_col} missing → {title}")
        return pd.DataFrame()

    # Denominator for N: all SA (Study) that are eligible for THIS stratification (non-missing value in `levels`)
    sa_total_in_strat = int(((df[GROUP_COL] == "Study") & df[strat_col].isin(levels)).sum())

    out_rows = []

    for lvl in levels:
        lvl_label = label_map.get(lvl, str(lvl))

        # Keep matched sets where the Study arm belongs to this level
        study_ids_in_level = df[(df[strat_col] == lvl) & (df[GROUP_COL] == "Study")][MATCH_ID].dropna().unique()
        if study_ids_in_level.size == 0:
            continue
        sub = df[df[MATCH_ID].isin(study_ids_in_level)].copy()

        # Minimal pragmatic imputation for prespecified covariates
        for v in ADJUST_ALL:
            if v not in sub.columns:
                continue
            if pd.api.types.is_numeric_dtype(sub[v]):
                sub[v] = pd.to_numeric(sub[v], errors="coerce").fillna(sub[v].median())
            else:
                m = sub[v].mode(dropna=True)
                sub[v] = sub[v].fillna(m.iloc[0] if not m.empty else sub[v])

        # N formatting: "stratum SA / total SA participating in this stratification"
        sa_in_level = int(((df[GROUP_COL] == "Study") & (df[strat_col] == lvl)).sum())
        n_label = f"{sa_in_level}/{sa_total_in_strat}" if sa_total_in_strat > 0 else f"{sa_in_level}/0"

        # Fit three WMH outcomes (log1p head-size normalized already prepared upstream)
        for out_col, out_lab in {
            "Log_HSNorm_Total_WMH": "Total WMH",
            "Log_HSNorm_PeriVentricular_WMH": "Periventricular WMH",
            "Log_HSNorm_Deep_WMH": "Deep WMH",
        }.items():
            if out_col not in sub.columns:
                continue

            formula = build_formula(out_col, sub)
            model = fit_cluster(formula, sub)

            coef_name = "C(group)[T.Study]"
            if coef_name not in model.params.index:
                continue

            beta = float(model.params[coef_name])
            ci_l, ci_u = [float(x) for x in model.conf_int().loc[coef_name]]
            pval = float(model.pvalues[coef_name])

            # Back-transform β to percent change on the original (non-log) scale
            pct   = 100.0 * (np.exp(beta) - 1.0)
            pct_l = 100.0 * (np.exp(ci_l) - 1.0)
            pct_u = 100.0 * (np.exp(ci_u) - 1.0)

            out_rows.append({
                "Stratification": title,
                "Stratum": lvl_label,
                "Outcome": out_lab,
                "Beta": beta,
                "CI Lower": ci_l,
                "CI Upper": ci_u,
                "p": pval,
                "% change": pct,
                "% CI Low": pct_l,
                "% CI High": pct_u,
                "q_FDR_primary": np.nan,
                "q_FDR_sensitivity": np.nan,
                "N": n_label,
            })

        # Within-level BH-FDR for {Periventricular, Deep}; Total excluded
        if out_rows:
            idx_this_level = [i for i, r in enumerate(out_rows) if r["Stratum"] == lvl_label]
            idx_fdr = [i for i in idx_this_level if out_rows[i]["Outcome"] in ("Periventricular WMH", "Deep WMH")]
            if idx_fdr:
                pvals = np.array([out_rows[i]["p"] for i in idx_fdr], dtype=float)
                qvals = multipletests(pvals, method="fdr_bh")[1]
                for i_row, qv in zip(idx_fdr, qvals):
                    if cohort_label == "Primary":
                        out_rows[i_row]["q_FDR_primary"] = float(qv)
                    elif cohort_label == "Sensitivity":
                        out_rows[i_row]["q_FDR_sensitivity"] = float(qv)

    return pd.DataFrame(out_rows)


# ---------------- Build STRATA dynamically (includes SA_ascertain_group) ----------------
def build_strata(df_example: pd.DataFrame):
    """
    Create the STRATA list with:
      • Sex (adaptive mapping)
      • Sleepy_SA_Stratum (1/0)
      • Snoring_Stratum (1/0)
      • SA_MRI_interval_stratum (LE10_years vs GT10_years)
      • SA_ascertain_group (Self_Report_Only / Hospital_Secondary / Hospital_Primary)
    """
    strata = []

    # Sex (adaptive to "Male/Female" or 1/0)
    if "Sex" in df_example.columns:
        uniq = set(str(x) for x in pd.Series(df_example["Sex"]).dropna().unique())
        if {"Male","Female"} & uniq:
            levels = ["Male","Female"]; lmap = {"Male":"Male","Female":"Female"}
        else:
            levels = [1,0];            lmap = {1:"Male", 0:"Female"}
        strata.append(("Sex", levels, lmap, "Sex", "Sex"))

    if "Sleepy_SA_Stratum" in df_example.columns:
        strata.append(("Sleepy_SA_Stratum", [1,0], {1:"Sleepy", 0:"Not Sleepy"}, "Daytime Sleepiness", "Sleepiness"))
    if "Snoring_Stratum" in df_example.columns:
        strata.append(("Snoring_Stratum", [1,0], {1:"Snoring", 0:"No Snore"}, "Snoring Status", "Snoring"))
    if "SA_MRI_interval_stratum" in df_example.columns:
        strata.append(("SA_MRI_interval_stratum",
                       ["LE10_years","GT10_years"],
                       {"LE10_years":"≤10 y", "GT10_years":">10 y"},
                       "Diagnosis–MRI Interval", "Dx_MRI_Interval"))

    # NEW: SA diagnosis method stratification (renamed from “severity”)
    if "SA_ascertain_group" in df_example.columns:
        strata.append(("SA_ascertain_group",
                       ["Self_Report_Only","Hospital_Secondary","Hospital_Primary"],
                       {"Self_Report_Only":"Self-Report",
                        "Hospital_Secondary":"Hosp-Secondary",
                        "Hospital_Primary":"Hosp-Primary"},
                       "SA Diagnosis Method", "SA_Diagnosis"))

    return strata

# ---------------- One cohort end-to-end ----------------
def run_for_cohort(label: str, csv_path: str) -> pd.DataFrame:
    print(f"\n===== {label} ({csv_path}) =====")
    subdir = os.path.join(OUTDIR, label); os.makedirs(subdir, exist_ok=True)

    df_raw = pd.read_csv(csv_path)
    df = ensure_outcomes(clean_cols(df_raw))

    STRATA = build_strata(df)

    all_blocks = []
    for strat_col, levels, lmap, panel_title, safe_key in STRATA:
        block_df = run_one_block(df, strat_col, levels, lmap, panel_title, label)
        if block_df is None or block_df.empty: continue

        all_blocks.append(block_df)

        block_csv = os.path.join(subdir, f"Stratified_{safe_key}.csv")
        block_df.to_csv(block_csv, index=False, encoding="utf-8-sig")

        out_png = os.path.join(subdir, f"Forest_{safe_key}.png")
        forest_plot(block_df, panel_title, out_png)

        # Console preview
        show = block_df.copy()
        show["Outcome"] = pd.Categorical(show["Outcome"], OUTCOME_ORDER, ordered=True)
        show = show.sort_values(["Outcome","Stratum"])
        cols = ["Stratification","Stratum","Outcome","% change","% CI Low","% CI High","p",
                "q_FDR_primary","q_FDR_sensitivity","N"]
        print("\n--", panel_title, "--")
        print(show[cols].round(4).to_string(index=False))

    if all_blocks:
        res_all = pd.concat(all_blocks, ignore_index=True)
        res_all = res_all[[
            "Stratification","Stratum","Outcome","N",
            "Beta","CI Lower","CI Upper","p",
            "q_FDR_primary","q_FDR_sensitivity",
            "% change","% CI Low","% CI High"
        ]]
        res_all.insert(0, "Analytic Cohort", label)
        res_all["Analytic Cohort"] = res_all["Analytic Cohort"].replace({
            "Primary": "Primary PSM cohort",
            "Sensitivity": "PSM–sensitivity cohort"
        })
        res_csv = os.path.join(subdir, "Stratified_AllResults.csv")
        res_all.to_csv(res_csv, index=False, encoding="utf-8-sig")
    else:
        res_all = pd.DataFrame()

    # README
    note = [
        "Head-size–normalized volumes with log1p transform.",
        "Covariates: " + ", ".join(ADJUST_ALL) + " (no CMC).",
        "Cluster-robust SEs by match_id where feasible; otherwise HC3.",
        "Within each stratum: BH-FDR to {Periventricular, Deep}; Total excluded.",
        "Forest figure outcome order: Total → Periventricular → Deep."
    ]
    with open(os.path.join(subdir, "README_Strata.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(note))

    return res_all

# ---------------- Run both cohorts & export combined table ----------------
all_res = []
for lab, fp in COHORTS.items():
    if not os.path.exists(fp):
        print(f"[Warning] File not found: {fp} → skip {lab}")
        continue
    res = run_for_cohort(lab, fp)
    if res is not None and not res.empty:
        all_res.append(res)

if all_res:
    comb = pd.concat(all_res, ignore_index=True)

    # 如果是旧列名，统一成“Analytic Cohort”
    if "Cohort" in comb.columns and "Analytic Cohort" not in comb.columns:
        comb = comb.rename(columns={"Cohort": "Analytic Cohort"})

    # 统一显示文本为论文用语
    comb["Analytic Cohort"] = comb["Analytic Cohort"].replace({
        "Primary": "Primary PSM cohort",
        "Sensitivity": "PSM–sensitivity cohort",
        "primary": "Primary PSM cohort",
        "sensitivity": "PSM–sensitivity cohort",
    })

    # 保存 CSV
    out_csv = os.path.join(OUTDIR, "Stratified_AllCohorts_Combined.csv")
    comb.to_csv(out_csv, index=False, encoding="utf-8-sig")
    print(f"Combined CSV saved: {out_csv}")

    # Word：函数内部会兼容列名并把表头写为“Analytic Cohort”
    export_stratified_word(comb, OUTDIR, filename="Stratified_AllCohorts_Supplement.docx")


print(f"\nAll stratified outputs saved under: {os.path.abspath(OUTDIR)}")
