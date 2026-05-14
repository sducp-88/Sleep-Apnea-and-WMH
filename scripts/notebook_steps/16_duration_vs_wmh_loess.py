# Extracted from WMH_Analysis.ipynb, code cell 15.
# Notebook heading: # Individual scatter + LOESS analysis: SA diagnosis duration vs WMH differences
# Run this file from the repository root unless a local CONFIG section is edited.

# Individual scatter + LOESS analysis: SA diagnosis duration vs WMH differences
# -----------------------------------------------------------------------------
# Plots: unchanged (publication-grade, PNG/PDF/SVG/TIFF)
# Tables:
#   1) Publishable Word three-line table of piecewise slopes (every 5 years)
#   2) Publishable Word three-line eTable: RCS (restricted cubic splines)
#      omnibus nonlinearity tests across cohorts/regions, on Percent & Log scales
#      (FDR only for PWMH & DWMH, per main analysis)
# -----------------------------------------------------------------------------

import os, re, warnings
warnings.filterwarnings("ignore", category=RuntimeWarning)

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

# statsmodels
import statsmodels.formula.api as smf
from statsmodels.nonparametric.smoothers_lowess import lowess
from statsmodels.stats.multitest import multipletests

# RCS splines
from patsy import cr

# Word export
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------- CONFIG ----------------
COHORTS = {
    "Primary":     "primary_cohort.csv",
    "Sensitivity": "sensitivity_cohort.csv",
}
OUTDIR = "Duration_I2_Individual"
os.makedirs(OUTDIR, exist_ok=True)

SUBDIR = {
    "plots_log": os.path.join(OUTDIR, "Plots_log"),
    "plots_log_sens": os.path.join(OUTDIR, "Plots_log_Sensitivity"),
    "plots_pct": os.path.join(OUTDIR, "Plots_pct"),
    "plots_pct_sens": os.path.join(OUTDIR, "Plots_pct_Sensitivity"),
    "plots_pct_sens_panel": os.path.join(OUTDIR, "Plots_pct_Sens_Panel"),
    "tables": os.path.join(OUTDIR, "Tables"),
}
for _p in SUBDIR.values():
    os.makedirs(_p, exist_ok=True)

# Collector for all per-point rows across cohorts/regions
POINT_ROWS_ALL = []

# Column names (Instance 2 date, first SA diagnosis date)
COL_I2_DATE = "Date_of_attending_assessment_centre_Instance_2"
COL_DX_DATE = "Date_G47_first_reported_sleep_disorders"

# WMH columns and scaling (Instance 2)
WMH_I2_COLS = {
    "Total":          "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2",
    "Periventricular":"Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2",
    "Deep":           "Total_volume_of_deep_white_matter_hyperintensities_Instance_2",
}
SCALE_I2 = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"

# Display labels and plotting order
DISPLAY_LABEL = {"Total":"Total WMH","Periventricular":"PWMH","Deep":"DWMH"}
PLOT_ORDER = ["Total","Periventricular","Deep"]
FDR_REGIONS = {"PWMH", "DWMH"}  # only these two get FDR q-values

# Plot style
plt.rcParams.update({
    "font.family": "Arial","figure.dpi": 120,"savefig.dpi": 600,
    "font.size": 12,"axes.labelsize": 12,
    "xtick.labelsize": 11,"ytick.labelsize": 11,
    "figure.autolayout": True
})
EXPORT_KW = {
    "png": {"dpi":600,"bbox_inches":"tight","facecolor":"white"},
    "pdf": {"bbox_inches":"tight"},
    "svg": {"bbox_inches":"tight"},
    "tiff": {"dpi":600,"bbox_inches":"tight"},
}

# ---------------- UTILITIES ----------------
def clean_cols(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns=lambda x: re.sub(r"[^0-9a-zA-Z]+","_",str(x)))

def ensure_datetime(s: pd.Series) -> pd.Series:
    return pd.to_datetime(s, errors="coerce")

def ensure_numeric(s: pd.Series) -> pd.Series:
    return pd.to_numeric(s, errors="coerce")

def safe_name(s: str) -> str:
    return re.sub(r"[^0-9A-Za-z._-]+","_",str(s)).strip("_")

def prepare_dataframe(fp: str) -> pd.DataFrame:
    """Compute diagnosis duration and log1p head-size–normalized WMH at Instance 2."""
    df = clean_cols(pd.read_csv(fp, low_memory=False))

    d_i2 = ensure_datetime(df.get(COL_I2_DATE))
    d_dx = ensure_datetime(df.get(COL_DX_DATE))
    df["Diagnosis_duration_years"] = (d_i2 - d_dx).dt.days / 365.25
    df.loc[df["Diagnosis_duration_years"] < 0, "Diagnosis_duration_years"] = np.nan
    df.loc[df["Diagnosis_duration_years"] > 60, "Diagnosis_duration_years"] = np.nan

    scale = ensure_numeric(df.get(SCALE_I2)).fillna(1.0)
    for reg, col in WMH_I2_COLS.items():
        v = ensure_numeric(df.get(col)) * scale
        df[f"Log_HSN_{reg}_I2"] = np.log1p(v)

    return df

# ---------------- INDIVIDUAL DIFFERENCES ----------------
def compute_individual_diff(df: pd.DataFrame, region: str, label: str) -> pd.DataFrame:
    """Compute SA–Control differences within matched sets (Δ log1p and %)."""
    y_col = f"Log_HSN_{region}_I2"

    if "group" not in df.columns:
        if "treatment_var" in df.columns:
            df["group"] = df["treatment_var"].map({0:"Control", 1:"Study"})
        elif "SA" in df.columns:
            df["group"] = df["SA"].map({0:"Control", 1:"Study"})
        else:
            raise ValueError(f"{label}: 'group'/'treatment_var'/'SA' column not found.")

    sa   = df[df["group"] == "Study"].copy()
    ctrl = df[df["group"] == "Control"].copy()

    ctrl_mean = ctrl.groupby("match_id")[y_col].mean().rename("ctrl_mean")
    sa = sa.merge(ctrl_mean, on="match_id", how="left")

    sa["Diff_log"] = sa[y_col] - sa["ctrl_mean"]        # Δ log1p
    sa["Diff_pct"] = np.expm1(sa["Diff_log"]) * 100.0   # % difference

    return sa[["match_id","Diagnosis_duration_years","Diff_log","Diff_pct"]]

# ---------------- PLOTTING (unchanged) ----------------
def plot_scatter_loess_log(dat: pd.DataFrame, region: str, label: str):
    plt.figure(figsize=(6,4))
    sns.scatterplot(data=dat, x="Diagnosis_duration_years", y="Diff_log",
                    alpha=0.5, s=28, color="#1f3b4d")
    sns.regplot(data=dat, x="Diagnosis_duration_years", y="Diff_log",
                lowess=True, scatter=False, color="red", line_kws={"lw":1.8})
    plt.axhline(0, color="grey", ls="--")
    plt.xlabel("Diagnosis duration (years)")
    plt.ylabel("Δ log1p WMH")
    plt.title(f"{label} – {DISPLAY_LABEL.get(region,region)} (log1p)")
    base = os.path.join(SUBDIR["plots_log"], f"IndivDiffLog_{safe_name(label)}_{region}")
    for ext, kw in EXPORT_KW.items():
        plt.savefig(base+f".{ext}", **kw)
    plt.close()

def plot_scatter_loess_pct(dat: pd.DataFrame, region: str, label: str):
    plt.figure(figsize=(6,4))
    sns.scatterplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                    alpha=0.5, s=28, color="#1f3b4d")
    sns.regplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                lowess=True, scatter=False, color="red", line_kws={"lw":1.8})
    plt.axhline(0, color="grey", ls="--")
    plt.xlabel("Diagnosis duration (years)")
    plt.ylabel("Δ WMH vs Controls (%)")
    plt.title(f"{label} – {DISPLAY_LABEL.get(region,region)} (%)")
    base = os.path.join(SUBDIR["plots_pct"], f"IndivDiffPct_{safe_name(label)}_{region}")
    for ext, kw in EXPORT_KW.items():
        plt.savefig(base+f".{ext}", **kw)
    plt.close()

def plot_scatter_loess_log_sensitivity(dat: pd.DataFrame, region: str, label: str,
                                       max_years: int = 30, right_limit: float = 0.4):
    dat = dat[dat["Diagnosis_duration_years"] <= max_years].copy()
    if dat.empty: return
    fig, ax1 = plt.subplots(figsize=(6,4))
    sns.scatterplot(data=dat, x="Diagnosis_duration_years", y="Diff_log",
                    alpha=0.4, s=28, color="#1f3b4d", ax=ax1)
    ax1.axhline(0, color="grey", ls="--")
    ax1.set_xlabel("Diagnosis duration (years)")
    ax1.set_ylabel("Δ log1p WMH – raw (all points)", color="#1f3b4d")
    ax1.tick_params(axis="y", labelcolor="#1f3b4d")

    ax2 = ax1.twinx()
    ax2.set_ylim(-right_limit, right_limit)
    ax2.set_ylabel(f"Δ log1p WMH – LOESS (±{right_limit})", color="red")
    ax2.tick_params(axis="y", labelcolor="red")

    smoothed = lowess(dat["Diff_log"], dat["Diagnosis_duration_years"], frac=0.3, return_sorted=True)
    ax2.plot(smoothed[:,0], smoothed[:,1], color="red", lw=1.8)

    plt.title(f"{label} – {DISPLAY_LABEL.get(region,region)} (log1p, ≤{max_years}y)")
    base = os.path.join(SUBDIR["plots_log_sens"],
                        f"IndivDiffLogDualRight_{safe_name(label)}_{region}_le{max_years}y_pm{right_limit}")
    for ext, kw in EXPORT_KW.items():
        plt.savefig(base+f".{ext}", **kw)
    plt.close()

def plot_scatter_loess_pct_sensitivity(dat: pd.DataFrame, region: str, label: str,
                                       max_years: int = 30, ylim: tuple = (-100, 100)):
    dat = dat[dat["Diagnosis_duration_years"] <= max_years].copy()
    if dat.empty: return
    plt.figure(figsize=(6,4))
    sns.scatterplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                    alpha=0.5, s=28, color="#1f3b4d")
    sns.regplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                lowess=True, scatter=False, color="red", line_kws={"lw":1.8})
    plt.axhline(0, color="grey", ls="--")
    plt.xlabel("Diagnosis duration (years)")
    plt.ylabel("Δ WMH vs Controls (%)")
    plt.title(f"{label} – {DISPLAY_LABEL.get(region,region)} (%, ≤{max_years}y)")
    plt.ylim(*ylim)
    base = os.path.join(SUBDIR["plots_pct_sens"],
                        f"IndivDiffPct_{safe_name(label)}_{region}_le{max_years}y")
    for ext, kw in EXPORT_KW.items():
        plt.savefig(base+f".{ext}", **kw)
    plt.close()

def plot_scatter_loess_pct_dualaxis(dat: pd.DataFrame, region: str, label: str,
                                    max_years: int = 30, right_ylim: tuple = (0, 40)):
    dat = dat[dat["Diagnosis_duration_years"] <= max_years].copy()
    if dat.empty: return

    left_min = float(np.nanmin(dat["Diff_pct"]))
    left_max = float(np.nanmax(dat["Diff_pct"]))

    fig, ax1 = plt.subplots(figsize=(6,4))
    sns.scatterplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                    alpha=0.4, s=28, color="#1f3b4d", ax=ax1)
    ax1.axhline(0, color="grey", ls="--")
    ax1.set_xlabel("Diagnosis duration (years)")
    ax1.set_ylabel("Δ WMH vs Controls (%) – raw (all points)", color="#1f3b4d")
    ax1.tick_params(axis="y", labelcolor="#1f3b4d")
    ax1.set_ylim(left_min, left_max)

    ax2 = ax1.twinx()
    sns.regplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                lowess=True, scatter=False, color="red", line_kws={"lw":1.8}, ax=ax2)

    left_span = left_max - left_min
    rel0 = (0.0 - left_min) / left_span if left_span > 0 else 0.5
    right_span = right_ylim[1] - right_ylim[0]
    right_min_aligned = right_ylim[0] - rel0 * right_span
    right_max_aligned = right_min_aligned + right_span
    ax2.set_ylim(right_min_aligned, right_max_aligned)

    ax2.set_ylabel("Δ WMH vs Controls (%) – LOESS (0–50%)", color="red")
    ax2.tick_params(axis="y", labelcolor="red")

    plt.title(f"{label} – {DISPLAY_LABEL.get(region,region)} (%, ≤{max_years}y)")
    base = os.path.join(SUBDIR["plots_pct_sens"],
                        f"IndivDiffPctDualFixed_{safe_name(label)}_{region}_le{max_years}y")
    for ext, kw in EXPORT_KW.items():
        plt.savefig(base+f".{ext}", **kw)
    plt.close()

def plot_pct_threepanel_dualaxis(results_by_region: dict, label: str,
                                 max_years: int = 30, right_ylim: tuple = (0, 40)):
    outdir = os.path.join(SUBDIR["plots_pct_sens_panel"], safe_name(label))
    os.makedirs(outdir, exist_ok=True)
    fig, axes = plt.subplots(1, 3, figsize=(14, 3.5), sharex=True)
    panel_name = {"Total": "Total WMH", "Periventricular": "PWMH", "Deep": "DWMH"}
    from matplotlib.ticker import MaxNLocator

    for i, region in enumerate(PLOT_ORDER):
        dat = results_by_region.get(region)
        if dat is None or dat.empty: continue
        dat = dat[dat["Diagnosis_duration_years"] <= max_years].copy()
        if dat.empty: continue

        ax1 = axes[i]
        data_min = float(np.nanmin(dat["Diff_pct"]))
        left_min = min(0.0, data_min)
        left_max = 2000.0 if region in ("Total", "Periventricular") else 7000.0

        sns.scatterplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                        alpha=0.3, s=20, color="#4a6c8c", ax=ax1, zorder=1)
        ax1.axhline(0, color="grey", ls="--", lw=1, zorder=2)
        ax1.set_xlabel("Diagnosis duration (years)")
        ax1.set_ylabel("% change (SA – Control)" if i == 0 else "", color="black")
        ax1.tick_params(axis="y", labelcolor="black")
        ax1.set_ylim(left_min, left_max)
        ax1.yaxis.set_major_locator(MaxNLocator(nbins=8, prune='upper' if region=="Deep" else None))

        ax2 = ax1.twinx()
        sns.regplot(data=dat, x="Diagnosis_duration_years", y="Diff_pct",
                    lowess=True, scatter=False, color="#e76f51",
                    line_kws={"lw": 2.0, "zorder": 3}, ax=ax2)

        left_span = left_max - left_min
        rel0 = (0.0 - left_min) / left_span if left_span > 0 else 0.5
        right_span = right_ylim[1] - right_ylim[0]
        right_min_aligned = right_ylim[0] - rel0 * right_span
        right_max_aligned = right_min_aligned + right_span
        ax2.set_ylim(right_min_aligned, right_max_aligned)
        if i == 2:
            ax2.set_ylabel("LOESS %", color="#e76f51")
        else:
            ax2.set_ylabel("")

        ax2.tick_params(axis="y", labelcolor="#e76f51")
        panel_tag = chr(65 + i)
        ax1.text(-0.10, 1.11, panel_tag, transform=ax1.transAxes,
                 fontsize=18, fontweight="bold", va="top", ha="right")
        ax1.text(0.50, 1.095, panel_name.get(region, region),
                 transform=ax1.transAxes, ha="center", va="top",
                 fontsize=14, fontweight="normal", color="black", clip_on=False)

    plt.tight_layout()
    base = os.path.join(outdir, f"Figure_Percent_LT{max_years}")
    for ext, kw in EXPORT_KW.items():
        plt.savefig(base + f".{ext}", **kw)
    plt.close()

# ---------------- SEGMENT SLOPES: every 5 years ----------------
def make_5y_bins(max_years: int = 30):
    """Return bins and labels for [0,5], (5,10], ..., (25,30] years."""
    edges = list(range(0, max_years + 5, 5))  # 0,5,10,...,30
    labels = [f"{edges[i]}–{edges[i+1]}y" for i in range(len(edges)-1)]
    return edges, labels

def compute_segment_slopes_5y(dat: pd.DataFrame, region: str, label: str,
                              max_years: int = 30) -> pd.DataFrame | None:
    """Piecewise linear slopes for Δ% within 5-year segments up to max_years (HC3 SE)."""
    dat = dat[dat["Diagnosis_duration_years"] <= max_years].copy()
    if dat.empty:
        return None

    edges, labels = make_5y_bins(max_years)
    dat["Segment"] = pd.cut(
        dat["Diagnosis_duration_years"], bins=edges, labels=labels, include_lowest=True, right=True
    )

    rows = []
    for seg, sub in dat.groupby("Segment", dropna=True):
        if len(sub) < 10:  # small-segment guard
            continue
        m = smf.ols("Diff_pct ~ Diagnosis_duration_years", data=sub).fit(cov_type="HC3")
        slope = m.params["Diagnosis_duration_years"]
        ci_low, ci_high = m.conf_int().loc["Diagnosis_duration_years"]
        pval = m.pvalues["Diagnosis_duration_years"]
        rows.append({
            "Analytic cohort": label,
            "WMH region": DISPLAY_LABEL.get(region, region),
            "Duration segment (years)": seg,
            "N": int(len(sub)),
            "β per year": float(slope),
            "CI low": float(ci_low),
            "CI high": float(ci_high),
            "p": float(pval),
        })

    df = pd.DataFrame(rows)
    # NOTE: FDR will be applied later across PWMH & DWMH only (per main analysis)
    return df

# ---------------- WORD EXPORT: helpers ----------------
def _set_doc_normal_tnr10(doc: Document):
    """Force Normal style to Times New Roman 10pt for full document."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)

def _apply_table_tnr10(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

# ---------------- WORD EXPORT: publishable tables ----------------
PRETTY_COLS = [
    "Analytic cohort", "WMH region", "Duration segment (years)", "N",
    "β per year", "CI low", "CI high", "p", "q (FDR)"
]

def save_publishable_segment_table(df: pd.DataFrame, outpath: str,
                                   title: str, legend_text: str):
    """Save a publishable Word three-line table (Times New Roman 10pt) with Title+Legend."""
    os.makedirs(os.path.dirname(outpath), exist_ok=True)
    df = df[PRETTY_COLS].copy()

    # Sort: cohort, region (Total→PWMH→DWMH), then segment in chronological order
    def _seg_start(s):
        try:
            return int(str(s).split("–")[0])
        except Exception:
            return 9999
    region_order = {"Total WMH": 0, "PWMH": 1, "DWMH": 2}
    df["_seg_order"] = df["Duration segment (years)"].map(_seg_start)
    df["_reg_order"] = df["WMH region"].map(lambda x: region_order.get(str(x), 9))
    df = (
        df.sort_values(["Analytic cohort", "_reg_order", "_seg_order"])
          .drop(columns=["_seg_order", "_reg_order"])
    )

    # Format numbers
    df_fmt = df.copy()
    for c in ["β per year","CI low","CI high"]:
        df_fmt[c] = pd.to_numeric(df_fmt[c], errors="coerce").map(lambda x: f"{x:.3f}" if pd.notna(x) else "")
    for c in ["p","q (FDR)"]:
        df_fmt[c] = pd.to_numeric(df_fmt[c], errors="coerce").map(lambda x: f"{x:.3g}" if pd.notna(x) else "")

    doc = Document()
    _set_doc_normal_tnr10(doc)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    # Table
    t = doc.add_table(rows=1, cols=len(df_fmt.columns))
    hdr = t.rows[0].cells
    for j, c in enumerate(df_fmt.columns):
        hdr[j].text = str(c)

    for _, row in df_fmt.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df_fmt.columns):
            cells[j].text = "" if pd.isna(row[c]) else str(row[c])

    # Three-line borders
    tbl = t._tbl
    for i, tr in enumerate(tbl.tr_lst):
        for tc in tr.tc_lst:
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            if i == 0:
                top = OxmlElement("w:top"); top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "16")
                bot = OxmlElement("w:bottom"); bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "16")
                tcBorders.append(top); tcBorders.append(bot)
            elif i == len(tbl.tr_lst) - 1:
                bot = OxmlElement("w:bottom"); bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "16")
                tcBorders.append(bot)
            tcPr.append(tcBorders)

    # Font unify
    _apply_table_tnr10(t)

    # Legend
    p_leg = doc.add_paragraph()
    r_leg = p_leg.add_run(legend_text)
    r_leg.italic = True
    r_leg.font.name = "Times New Roman"
    r_leg.font.size = Pt(10)

    doc.save(outpath)
    print(f"[OK] Publishable segment slope table saved: {outpath}")

# ===== RCS omnibus nonlinearity test & eTable (supplement) =====================
RCS_ROWS = []  # collect per cohort × region × scale

def rcs_global_test(dat: pd.DataFrame, region: str, label: str,
                    x_col: str = "Diagnosis_duration_years",
                    y_col: str = "Diff_pct",
                    df_spline: int = 4,
                    max_years: int = 30):
    """
    Omnibus nonlinearity test: linear vs restricted cubic spline (RCS).
    Full model:   y ~ cr(x, df=df_spline)
    Restricted:   y ~ x
    Test:         nested F-test via results.compare_f_test()
    """
    d = dat[[x_col, y_col]].dropna().copy()
    d = d[d[x_col] <= max_years]
    if len(d) < (df_spline + 2):
        return None  # insufficient data

    # Restricted linear model
    m_lin = smf.ols(f"{y_col} ~ {x_col}", data=d).fit()

    # Full RCS model (natural cubic spline)
    m_rcs = smf.ols(f"{y_col} ~ cr({x_col}, df={df_spline})", data=d).fit()

    F, p, df_diff = m_rcs.compare_f_test(m_lin)  # omnibus nonlinearity
    row = {
        "Analytic cohort": label,
        "WMH region": DISPLAY_LABEL.get(region, region),
        "Scale": "Percent" if y_col == "Diff_pct" else "Log (Δ log1p)",
        "N": int(len(d)),
        "df_spline": int(df_spline),
        "F (nonlinearity)": float(F) if F is not None else np.nan,
        "p (nonlinearity)": float(p) if p is not None else np.nan,
        "df_diff": int(df_diff) if df_diff is not None else np.nan,
        # context: overall model fit
        "F (full model)": float(m_rcs.fvalue) if m_rcs.fvalue is not None else np.nan,
        "p (full model)": float(m_rcs.f_pvalue) if m_rcs.f_pvalue is not None else np.nan,
    }
    RCS_ROWS.append(row)
    print(f"[RCS] {label} – {region} – {row['Scale']}: "
          f"F={row['F (nonlinearity)']:.3f}, p={row['p (nonlinearity)']:.3g}, "
          f"df_diff={row['df_diff']}")
    return row

RCS_PRETTY_COLS = [
    "Analytic cohort", "WMH region", "Scale", "N", "df_spline",
    "F (nonlinearity)", "p (nonlinearity)", "q (FDR, nonlinearity)",
    "df_diff", "F (full model)", "p (full model)"
]

def save_publishable_rcs_table(df: pd.DataFrame, outpath: str,
                               title: str, legend_text: str):
    """Save a publishable Word three-line eTable for RCS omnibus tests."""
    os.makedirs(os.path.dirname(outpath), exist_ok=True)

    # Order by cohort, region (Total→PWMH→DWMH), then Scale (Log→Percent)
    region_order = {"Total WMH":0, "PWMH":1, "DWMH":2}
    scale_order  = {"Log (Δ log1p)":0, "Percent":1}
    df["_rord"] = df["WMH region"].map(lambda x: region_order.get(x, 9))
    df["_sord"] = df["Scale"].map(lambda x: scale_order.get(x, 9))
    df = df.sort_values(["Analytic cohort", "_rord", "_sord"]).drop(columns=["_rord","_sord"])

    # Format numbers
    df_fmt = df.copy()
    for c in ["F (nonlinearity)", "F (full model)"]:
        df_fmt[c] = pd.to_numeric(df_fmt[c], errors="coerce").map(lambda x: f"{x:.3f}" if pd.notna(x) else "")
    for c in ["p (nonlinearity)", "q (FDR, nonlinearity)", "p (full model)"]:
        df_fmt[c] = pd.to_numeric(df_fmt[c], errors="coerce").map(lambda x: f"{x:.3g}" if pd.notna(x) else "")

    df_fmt = df_fmt[RCS_PRETTY_COLS]

    doc = Document()
    _set_doc_normal_tnr10(doc)

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    t = doc.add_table(rows=1, cols=len(df_fmt.columns))
    hdr = t.rows[0].cells
    for j, c in enumerate(df_fmt.columns):
        hdr[j].text = str(c)

    for _, row in df_fmt.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df_fmt.columns):
            cells[j].text = "" if pd.isna(row[c]) else str(row[c])

    # Three-line borders
    tbl = t._tbl
    for i, tr in enumerate(tbl.tr_lst):
        for tc in tr.tc_lst:
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            if i == 0:
                top = OxmlElement("w:top"); top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "16")
                bot = OxmlElement("w:bottom"); bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "16")
                tcBorders.append(top); tcBorders.append(bot)
            elif i == len(tbl.tr_lst) - 1:
                bot = OxmlElement("w:bottom"); bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "16")
                tcBorders.append(bot)
            tcPr.append(tcBorders)

    _apply_table_tnr10(t)

    p_leg = doc.add_paragraph()
    r_leg = p_leg.add_run(legend_text)
    r_leg.italic = True
    r_leg.font.name = "Times New Roman"
    r_leg.font.size = Pt(10)

    doc.save(outpath)
    print(f"[OK] Publishable RCS omnibus eTable saved: {outpath}")
# ============================================================================

# ---------------- MAIN ----------------
def run_one_cohort(label: str, fp: str) -> dict:
    """Compute diffs (for plots), make all plots; return per-region data."""
    df = prepare_dataframe(fp)
    if "group" not in df.columns:
        if "treatment_var" in df.columns:
            df["group"] = df["treatment_var"].map({0:"Control", 1:"Study"})
        elif "SA" in df.columns:
            df["group"] = df["SA"].map({0:"Control", 1:"Study"})
        else:
            raise ValueError(f"{label}: no 'group'/'treatment_var'/'SA' column.")

    results_by_region = {}
    for reg in PLOT_ORDER:
        dat = compute_individual_diff(df, reg, label)
        if dat is None or dat.empty:
            continue
        results_by_region[reg] = dat

        # --- Save per-cohort×region per-point CSV ---
        per_csv = os.path.join(SUBDIR["tables"], f"IndivDiff_{safe_name(label)}_{reg}.csv")
        dat.to_csv(per_csv, index=False, encoding="utf-8-sig")
        print(f"[OK] CSV saved: {per_csv}")

        # --- Append rows to the global collector for the combined CSV ---
        dat2 = dat.copy()
        dat2.insert(0, "WMH region", DISPLAY_LABEL.get(reg, reg))
        dat2.insert(0, "Analytic cohort", label)
        POINT_ROWS_ALL.append(dat2)

        # Plots (unchanged)
        plot_scatter_loess_log(dat, reg, label)
        plot_scatter_loess_pct(dat, reg, label)
        plot_scatter_loess_log_sensitivity(dat, reg, label, max_years=30)
        plot_scatter_loess_pct_sensitivity(dat, reg, label, max_years=30)
        plot_scatter_loess_pct_dualaxis(dat, reg, label, max_years=30)

        # --- RCS omnibus tests on Percent and Log scales ---
        rcs_global_test(dat, reg, label, y_col="Diff_pct", df_spline=4, max_years=30)
        rcs_global_test(dat, reg, label, y_col="Diff_log", df_spline=4, max_years=30)

    # Three-panel figure using accumulated regions
    if results_by_region:
        plot_pct_threepanel_dualaxis(results_by_region, label, max_years=30)

    return results_by_region

if __name__ == "__main__":
    # Collect all 5-year segment slopes across cohorts/regions
    seg_results_all = []

    for lab, fp in COHORTS.items():
        if not os.path.exists(fp):
            print(f"[Skip] {lab}: file not found → {fp}")
            continue

        res = run_one_cohort(lab, fp)
        for reg, dat in res.items():
            seg_tbl = compute_segment_slopes_5y(dat, reg, lab, max_years=30)
            if seg_tbl is not None and not seg_tbl.empty:
                seg_results_all.append(seg_tbl)

    # Merge and export the single publishable 5y-segment table
    if seg_results_all:
        df_all = pd.concat(seg_results_all, ignore_index=True)

        # === FDR only for PWMH & DWMH (not for Total WMH) ===
        df_all["q (FDR)"] = np.nan
        m = df_all["WMH region"].isin(FDR_REGIONS) & df_all["p"].notna()
        if m.any():
            df_all.loc[m, "q (FDR)"] = multipletests(df_all.loc[m, "p"], method="fdr_bh")[1]

        title = "eTable Y. Piecewise 5-year slopes of SA–control differences in white matter hyperintensities by diagnosis duration"
        legend = (
            "Legend. Slopes (β per year) are estimated within 5-year segments using OLS with HC3 standard errors. "
            "Outcome is Δ WMH vs controls (%) derived from expm1 of log-differences. "
            "FDR is applied to PWMH and DWMH only, consistent with the main analysis. "
            "Segments with N<10 are omitted."
        )

        docx_path = os.path.join(SUBDIR["tables"], "SegmentSlopes_5y_Publishable.docx")
        save_publishable_segment_table(df_all, docx_path, title, legend)

        # Also export CSV for reproducibility
        csv_path = os.path.join(SUBDIR["tables"], "SegmentSlopes_5y_Publishable.csv")
        df_all.to_csv(csv_path, index=False, encoding="utf-8-sig")
        print(f"[OK] CSV saved: {csv_path}")

    # === Export combined per-point CSV ===
    if POINT_ROWS_ALL:
        df_points_all = pd.concat(POINT_ROWS_ALL, ignore_index=True)
        ordered_cols = [
            "Analytic cohort", "WMH region",
            "match_id", "Diagnosis_duration_years",
            "Diff_log", "Diff_pct"
        ]
        df_points_all = df_points_all[ordered_cols]
        out_points_csv = os.path.join(SUBDIR["tables"], "IndivDiff_AllCohorts_AllRegions.csv")
        df_points_all.to_csv(out_points_csv, index=False, encoding="utf-8-sig")
        print(f"[OK] Combined per-point CSV saved: {out_points_csv}")

    # === RCS omnibus eTable (Supplement; FDR only for PWMH & DWMH) ============
    if RCS_ROWS:
        df_rcs = pd.DataFrame(RCS_ROWS)

        # FDR for nonlinearity p-values — PWMH & DWMH only
        df_rcs["q (FDR, nonlinearity)"] = np.nan
        m = df_rcs["WMH region"].isin(FDR_REGIONS) & df_rcs["p (nonlinearity)"].notna()
        if m.any():
            df_rcs.loc[m, "q (FDR, nonlinearity)"] = multipletests(
                df_rcs.loc[m, "p (nonlinearity)"], method="fdr_bh"
            )[1]

        # Save CSV
        out_rcs_csv = os.path.join(SUBDIR["tables"], "RCS_Omnibus_Nonlinearity.csv")
        df_rcs.to_csv(out_rcs_csv, index=False, encoding="utf-8-sig")
        print(f"[OK] RCS omnibus tests saved: {out_rcs_csv}")

        # Save Word eTable
        rcs_title = "eTable X. Omnibus restricted cubic spline tests for nonlinearity of SA–control WMH differences by diagnosis duration"
        rcs_legend = (
            "Legend. Omnibus tests compare a restricted cubic spline model (df=4) with a linear model for diagnosis duration "
            "using a nested F-test. Outcome is Δ WMH vs controls (%) and Δ log1p WMH (log scale). "
            "FDR is applied to PWMH and DWMH only, consistent with the main analysis. "
            "N denotes observations within ≤30 years."
        )
        rcs_docx = os.path.join(SUBDIR["tables"], "eTable_RCS_Omnibus_Nonlinearity.docx")
        save_publishable_rcs_table(df_rcs, rcs_docx, rcs_title, rcs_legend)

    print("\nAll cohorts processed. Plots saved; publishable tables exported.")

    # === How many had duration >30y (count once using Total WMH) ===
try:
    df_pts = pd.read_csv(os.path.join(SUBDIR["tables"], "IndivDiff_AllCohorts_AllRegions.csv"))
    d = df_pts[df_pts["WMH region"] == "Total WMH"]

    by_cohort = (
        d.groupby("Analytic cohort")["Diagnosis_duration_years"]
         .agg(N_total=lambda s: int(s.notna().sum()),
              N_le30=lambda s: int((s <= 30).sum()),
              N_gt30=lambda s: int((s > 30).sum()))
         .reset_index()
    )

    overall_gt30 = int((d["Diagnosis_duration_years"] > 30).sum())

    txt = ["Counts for duration filter (use Total WMH to avoid triple-counting):"]
    for _, r in by_cohort.iterrows():
        txt.append(f"- {r['Analytic cohort']}: N_total={r['N_total']}, N_<=30y={r['N_le30']}, N_>30y={r['N_gt30']}")
    txt.append(f"Overall N_>30y removed: {overall_gt30}")

    out_txt = os.path.join(SUBDIR["tables"], "Counts_Duration_gt30_excluded.txt")
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(txt))
    print("\n".join(txt))
    print(f"[OK] Saved: {out_txt}")
except Exception as e:
    print("[WARN] Could not compute >30y counts:", e)
