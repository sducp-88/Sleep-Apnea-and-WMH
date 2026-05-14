# Extracted from WMH_Analysis.ipynb, code cell 7.
# Notebook heading: # SA -> WMH -> Cognitive mediation analysis (No CMC)
# Run this file from the repository root unless a local CONFIG section is edited.

# SA -> WMH -> Cognitive mediation analysis (No CMC)
# =============================================================================
# Fixes:
#   1) Sobel test normal CDF no longer uses sm.stats.norm (not available).
#      - Prefer scipy.stats.norm if available.
#      - Fallback to math.erf-based normal CDF.
#   2) Add progress bars for bootstrap with tqdm if available,
#      else a lightweight text progress indicator.
#
# Style aligned with user's cognitive secondary outcomes script.
#
# X (Exposure): group (Study vs Control)
# M (Mediators): WMH (head-size normalized + log1p)
# Y (Outcomes): 4 cognitive endpoints (Instance 2; standardized; higher=better)
#
# Cohorts:
#   - primary_cohort.csv      -> Primary PSM cohort
#   - sensitivity_cohort.csv  -> PSM–sensitivity cohort
#
# Covariates (Base; same as cognitive script):
#   Sex, Age_at_Instance_2, Townsend_deprivation_index_at_recruitment,
#   Body_mass_index_BMI_Instance_0, Genetic_ethnic_grouping,
#   Smoking_Ever, Alcohol_intake_frequency_ordinal,
#   has_degree, e4_count
#
# SE:
#   - Prefer cluster-robust by match_id where feasible, else HC3.
#
# Inference:
#   - Sobel test for indirect effect.
#   - 2000x bootstrap for indirect effect and proportion mediated (PM).
#     * Cluster bootstrap by match_id when available.
#
# Outputs:
#   - Mediation/<Cohort>/Mediation_Results.csv
#   - Mediation/Mediation_Results_AllCohorts_Combined.csv
#   - Mediation/Mediation_Results_AllCohorts.docx (publication-ready)
# =============================================================================

import os
import re
import math
import numpy as np
import pandas as pd
import statsmodels.api as sm
from patsy import dmatrices

# Optional normal CDF from scipy
try:
    from scipy.stats import norm as _scipy_norm
except Exception:
    _scipy_norm = None

# Optional progress bar
try:
    from tqdm.auto import tqdm
except Exception:
    tqdm = None

# Optional Word export
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception:
    Document = None


# ======================== CONFIG ========================

OUTROOT = "Mediation"
os.makedirs(OUTROOT, exist_ok=True)

COHORT_FILES = {
    "Primary PSM cohort":     "primary_cohort.csv",
    "PSM–sensitivity cohort": "sensitivity_cohort.csv",
}

GROUP_COL = "group"
CLUSTER_VAR = "match_id"
GROUP_LEVELS = ["Control", "Study"]

# WMH raw + scaling
SCALE_COL = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
RAW_TOT = "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"
RAW_PV  = "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"
RAW_DW  = "Total_volume_of_deep_white_matter_hyperintensities_Instance_2"

# Cognitive raw variables
RT_COL   = "Mean_time_to_correctly_identify_matches_Instance_2"
TM_COL   = "Duration_to_complete_alphanumeric_path_trail_2_Instance_2"
FI_COL   = "Fluid_intelligence_score_Instance_2"
MEM_COL  = "Maximum_digits_remembered_correctly_Instance_2"

# Endpoint labels/order (RT, TMT-B first)
COG_ENDPOINTS = [
    ("RT_z",  "Reaction time"),
    ("TM_z",  "Trail making test-B"),
    ("FI_z",  "Fluid intelligence"),
    ("MEM_z", "Digit span"),
]

# Mediators order/labels
WMH_MEDIATORS = [
    ("Log_HSNorm_Total_WMH_T1_T2",     "Total WMH"),
    ("Log_HSNorm_PeriVentricular_WMH", "PWMH"),
    ("Log_HSNorm_Deep_WMH",            "DWMH"),
]

# Base covariates
BASE_COVARIATES = [
    "Sex",
    "Age_at_Instance_2",
    "Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "Alcohol_intake_frequency_ordinal",
    "has_degree",
    "e4_count",
]

CATEGORICAL_VARS = {
    "Sex",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "has_degree",
}

BOOTSTRAP_N = 2000
RANDOM_SEED = 20251209


# ======================== PROGRESS ========================

def _iter_progress(iterable, total=None, desc=None):
    """
    Unified progress iterator.
    - Uses tqdm if available.
    - Else prints lightweight progress updates.
    """
    if tqdm is not None:
        return tqdm(iterable, total=total, desc=desc, leave=False)

    # Fallback minimal progress (prints at ~10% intervals)
    if total is None:
        total = len(iterable) if hasattr(iterable, "__len__") else None
    if total is None:
        return iterable

    step = max(1, total // 10)
    def gen():
        for i, x in enumerate(iterable, 1):
            if i == 1 or i % step == 0 or i == total:
                print(f"{desc or 'Progress'}: {i}/{total}")
            yield x
    return gen()


# ======================== HELPERS ========================

def clean_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = df.columns.str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
    return df

def ensure_group(df: pd.DataFrame) -> pd.DataFrame:
    if GROUP_COL not in df.columns:
        raise ValueError("Missing 'group' column.")
    df[GROUP_COL] = pd.Categorical(df[GROUP_COL], GROUP_LEVELS)
    return df

def build_wmh_vars(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    if SCALE_COL not in df.columns:
        raise ValueError("Missing scaling factor column for WMH.")
    scale = pd.to_numeric(df[SCALE_COL], errors="coerce")

    df["HSNorm_Total_WMH_T1_T2"] = pd.to_numeric(df.get(RAW_TOT), errors="coerce") * scale
    df["HSNorm_PeriVentricular_WMH"] = pd.to_numeric(df.get(RAW_PV), errors="coerce") * scale
    df["HSNorm_Deep_WMH"] = pd.to_numeric(df.get(RAW_DW), errors="coerce") * scale

    df["Log_HSNorm_Total_WMH_T1_T2"] = np.log1p(df["HSNorm_Total_WMH_T1_T2"])
    df["Log_HSNorm_PeriVentricular_WMH"] = np.log1p(df["HSNorm_PeriVentricular_WMH"])
    df["Log_HSNorm_Deep_WMH"] = np.log1p(df["HSNorm_Deep_WMH"])
    return df

def standardize(series: pd.Series) -> pd.Series:
    s = pd.to_numeric(series, errors="coerce")
    mu = s.mean()
    sd = s.std(ddof=1)
    return (s - mu) / sd if sd > 0 else pd.Series(np.nan, index=s.index)

def prepare_cognitive_scores(df: pd.DataFrame, cohort_label: str) -> pd.DataFrame:
    df = df.copy()

    # Reaction time
    rt = pd.to_numeric(df.get(RT_COL), errors="coerce")
    n_total = rt.notna().sum()
    rt_z = standardize(np.log(rt))
    rt_z = -rt_z
    n_used = rt_z.notna().sum()
    print(f"[{cohort_label}] RT: removed {n_total - n_used}, retained {n_used}")
    df["RT_z"] = rt_z

    # TMT-B
    tm = pd.to_numeric(df.get(TM_COL), errors="coerce")
    invalid = tm.isna() | (tm == 0)
    n_total = (~tm.isna()).sum()
    tm_clean = tm.mask(invalid)
    tm_z = standardize(np.log(tm_clean))
    tm_z = -tm_z
    n_used = tm_z.notna().sum()
    print(f"[{cohort_label}] TMT-B: removed {n_total - n_used} (incl. zeros), retained {n_used}")
    df["TM_z"] = tm_z

    # FI
    fi = pd.to_numeric(df.get(FI_COL), errors="coerce")
    invalid = fi.isna() | (fi == -1)
    n_total = (~fi.isna()).sum()
    fi_clean = fi.mask(invalid)
    fi_z = standardize(fi_clean)
    n_used = fi_z.notna().sum()
    print(f"[{cohort_label}] FI: removed {n_total - n_used} (incl. -1), retained {n_used}")
    df["FI_z"] = fi_z

    # Digit span
    mem = pd.to_numeric(df.get(MEM_COL), errors="coerce")
    invalid = mem.isna() | (mem == -1)
    n_total = (~mem.isna()).sum()
    mem_clean = mem.mask(invalid)
    mem_z = standardize(mem_clean)
    n_used = mem_z.notna().sum()
    print(f"[{cohort_label}] Digit span: removed {n_total - n_used} (incl. -1), retained {n_used}")
    df["MEM_z"] = mem_z

    return df

def impute_covariates(df: pd.DataFrame, covariates: list, categoricals: set) -> pd.DataFrame:
    df = df.copy()
    for v in covariates:
        if v not in df.columns:
            continue
        if v in categoricals:
            mv = df[v].mode(dropna=True)
            if not mv.empty:
                df[v] = df[v].fillna(mv.iloc[0])
        else:
            df[v] = pd.to_numeric(df[v], errors="coerce")
            df[v] = df[v].fillna(df[v].median())
    return df

def build_formula(outcome: str, rhs_vars: list, categoricals: set) -> str:
    terms = []
    for v in rhs_vars:
        if v in categoricals:
            terms.append(f"C({v})")
        else:
            terms.append(v)
    return f"{outcome} ~ " + " + ".join(terms)

def find_group_coef(params_index):
    for name in params_index:
        if name.startswith("C(group)[T.") and name.endswith("Study]"):
            return name
    for name in params_index:
        if name.startswith("C(group)[T."):
            return name
    return None

def fit_ols_robust(formula: str, data: pd.DataFrame, cluster_var: str):
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    if y.empty:
        return None, "NA", 0, 0

    base = sm.OLS(y, X).fit(cov_type="HC3")

    if cluster_var not in data.columns:
        return base, "HC3 (no cluster var)", 0, int(len(y))

    groups = data.loc[y.index, cluster_var]
    if groups.notna().sum() < 2 or groups.nunique() < 2:
        return base, "HC3 (<=1 cluster)", int(groups.nunique()), int(len(y))

    ok = groups.notna()
    y2, X2, g2 = y.loc[ok], X.loc[ok], groups.loc[ok]
    model = sm.OLS(y2, X2).fit(cov_type="cluster", cov_kwds={"groups": g2})
    return model, f"cluster (n_clusters={g2.nunique()})", int(g2.nunique()), int(len(y2))


# -------- Normal CDF (robust) --------

def _normal_cdf(x: float) -> float:
    if _scipy_norm is not None:
        return float(_scipy_norm.cdf(x))
    # fallback using erf
    return 0.5 * (1.0 + math.erf(x / math.sqrt(2.0)))


def sobel_p(a, sa, b, sb):
    """
    Sobel test for indirect effect.
    Returns z, p (two-sided).
    """
    denom = math.sqrt((b*b*sa*sa) + (a*a*sb*sb))
    if denom == 0 or np.isnan(denom):
        return np.nan, np.nan
    z = (a*b) / denom
    p = 2 * (1 - _normal_cdf(abs(z)))
    return z, p


# -------- Bootstrap sampling --------

def cluster_bootstrap_sample(df: pd.DataFrame, cluster_var: str, rng: np.random.Generator):
    if cluster_var in df.columns:
        clusters = df[cluster_var].dropna().unique()
        if len(clusters) >= 2:
            sampled = rng.choice(clusters, size=len(clusters), replace=True)
            parts = [df[df[cluster_var] == cid] for cid in sampled]
            return pd.concat(parts, ignore_index=True)
    idx = rng.integers(0, len(df), size=len(df))
    return df.iloc[idx].reset_index(drop=True)


# -------- Path computation --------

def compute_paths(df: pd.DataFrame, mediator_col: str, outcome_col: str):
    covs = [v for v in BASE_COVARIATES if v in df.columns]
    cats = set(CATEGORICAL_VARS) & set(covs)

    # a: M ~ group + covs
    rhs_a = ["group"] + covs
    form_a = build_formula(mediator_col, rhs_a, cats | {"group"})
    d_a = impute_covariates(df, covs, cats)

    m_a, se_info_a, _, _ = fit_ols_robust(form_a, d_a, CLUSTER_VAR)
    if m_a is None:
        return None

    gcoef_a = find_group_coef(m_a.params.index)
    if gcoef_a is None:
        return None

    a = float(m_a.params[gcoef_a])
    a_se = float(m_a.bse[gcoef_a])
    a_p = float(m_a.pvalues[gcoef_a])

    # c total: Y ~ group + covs
    rhs_c = ["group"] + covs
    form_c = build_formula(outcome_col, rhs_c, cats | {"group"})
    d_c = impute_covariates(df, covs, cats)

    m_c, se_info_c, _, nobs_c = fit_ols_robust(form_c, d_c, CLUSTER_VAR)
    if m_c is None:
        return None

    gcoef_c = find_group_coef(m_c.params.index)
    if gcoef_c is None:
        return None

    ctotal = float(m_c.params[gcoef_c])
    ctotal_se = float(m_c.bse[gcoef_c])
    ctotal_p = float(m_c.pvalues[gcoef_c])

    # b & c': Y ~ group + M + covs
    rhs_b = ["group", mediator_col] + covs
    form_b = build_formula(outcome_col, rhs_b, cats | {"group"})
    d_b = impute_covariates(df, covs, cats)

    m_b, se_info_b, _, nobs_b = fit_ols_robust(form_b, d_b, CLUSTER_VAR)
    if m_b is None:
        return None

    gcoef_b = find_group_coef(m_b.params.index)
    if gcoef_b is None or mediator_col not in m_b.params.index:
        return None

    b = float(m_b.params[mediator_col])
    b_se = float(m_b.bse[mediator_col])
    b_p = float(m_b.pvalues[mediator_col])

    cprime = float(m_b.params[gcoef_b])
    cprime_se = float(m_b.bse[gcoef_b])
    cprime_p = float(m_b.pvalues[gcoef_b])

    return {
        "a": a, "a_se": a_se, "a_p": a_p, "se_info_a": se_info_a,
        "b": b, "b_se": b_se, "b_p": b_p, "se_info_b": se_info_b,
        "cprime": cprime, "cprime_se": cprime_se, "cprime_p": cprime_p,
        "ctotal": ctotal, "ctotal_se": ctotal_se, "ctotal_p": ctotal_p,
        "se_info_c": se_info_c,
        "n_obs_c": nobs_c, "n_obs_b": nobs_b,
    }


def bootstrap_indirect(df: pd.DataFrame, mediator_col: str, outcome_col: str,
                       n_boot=2000, seed=0, desc=None):
    rng = np.random.default_rng(seed)
    ab_list = []
    pm_list = []

    iterator = _iter_progress(range(n_boot), total=n_boot, desc=desc or "Bootstrap")

    for _ in iterator:
        db = cluster_bootstrap_sample(df, CLUSTER_VAR, rng)
        paths = compute_paths(db, mediator_col, outcome_col)
        if paths is None:
            continue

        a = paths["a"]
        b = paths["b"]
        ctotal = paths["ctotal"]

        ab = a * b
        ab_list.append(ab)

        if pd.notna(ctotal) and ctotal != 0:
            pm_list.append(ab / ctotal)

    if len(ab_list) < 50:
        return {
            "ab_boot_mean": np.nan,
            "ab_ci_low": np.nan,
            "ab_ci_high": np.nan,
            "pm_boot_mean": np.nan,
            "pm_ci_low": np.nan,
            "pm_ci_high": np.nan,
            "n_boot_ok": len(ab_list)
        }

    ab_arr = np.array(ab_list, dtype=float)
    pm_arr = np.array(pm_list, dtype=float) if pm_list else np.array([], dtype=float)

    ab_ci = np.nanpercentile(ab_arr, [2.5, 97.5])

    if pm_arr.size >= 50:
        pm_ci = np.nanpercentile(pm_arr, [2.5, 97.5])
        pm_mean = float(np.nanmean(pm_arr))
    else:
        pm_ci = [np.nan, np.nan]
        pm_mean = np.nan

    return {
        "ab_boot_mean": float(np.nanmean(ab_arr)),
        "ab_ci_low": float(ab_ci[0]),
        "ab_ci_high": float(ab_ci[1]),
        "pm_boot_mean": pm_mean,
        "pm_ci_low": float(pm_ci[0]),
        "pm_ci_high": float(pm_ci[1]),
        "n_boot_ok": int(len(ab_arr))
    }


# ======================== WORD TABLE HELPERS ========================

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = OxmlElement(f'w:{edge}')
            for key in ("val","sz","color","space"):
                if key in edge_data:
                    tag.set(qn(f"w:{key}"), str(edge_data[key]))
            tcBorders.append(tag)

def _apply_three_line_table(table, header_row_idx=0):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                             left={"val":"nil"}, right={"val":"nil"},
                             top={"val":"nil"}, bottom={"val":"nil"})
    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell,
                         top={"val":"single","sz":8,"color":"000000"},
                         bottom={"val":"single","sz":8,"color":"000000"})
    for cell in table.rows[-1].cells:
        _set_cell_border(cell,
                         bottom={"val":"single","sz":8,"color":"000000"})

def _right_align_numeric(table, idx_list):
    for row in table.rows[1:]:
        for j in idx_list:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def export_mediation_word(df_all: pd.DataFrame, out_path: str,
                          title: str, legend: str):
    if Document is None or df_all is None or df_all.empty:
        print("Word export skipped (python-docx unavailable or empty results).")
        return None

    df = df_all.copy()

    cohort_order = ["Primary PSM cohort", "PSM–sensitivity cohort"]
    outcome_order = [lab for _, lab in COG_ENDPOINTS]
    med_order = [lab for _, lab in WMH_MEDIATORS]

    df["Cohort"] = pd.Categorical(df["Cohort"], categories=cohort_order, ordered=True)
    df["Outcome"] = pd.Categorical(df["Outcome"], categories=outcome_order, ordered=True)
    df["Mediator"] = pd.Categorical(df["Mediator"], categories=med_order, ordered=True)

    df = df.sort_values(["Cohort", "Outcome", "Mediator"]).reset_index(drop=True)

    def _fmt(x, nd=3):
        return "—" if pd.isna(x) else f"{float(x):.{nd}f}"

    def _fmt_p(x):
        if pd.isna(x): return "—"
        x = float(x)
        return "<0.001" if x < 1e-3 else f"{x:.3f}"

    def _fmt_ci(lo, hi, nd=3):
        if any(pd.isna([lo, hi])): return "—"
        return f"({_fmt(lo, nd)}, {_fmt(hi, nd)})"

    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run(title).bold = True

    headers = [
        "Analytic Cohort", "Outcome", "Mediator", "N (SA/Control)",
        "a (SE)", "p_a",
        "b (SE)", "p_b",
        "Indirect a×b", "Sobel p",
        "Bootstrap 95% CI (a×b)",
        "Direct c′ (SE)", "p_c′",
        "Total c (SE)", "p_c",
        "PM", "Bootstrap 95% CI (PM)"
    ]

    t = doc.add_table(rows=1, cols=len(headers))
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h

    for _, r in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(r.get("Cohort", ""))
        cells[1].text = str(r.get("Outcome", ""))
        cells[2].text = str(r.get("Mediator", ""))
        cells[3].text = str(r.get("N", ""))

        cells[4].text = f"{_fmt(r.get('a'))} ({_fmt(r.get('a_se'))})"
        cells[5].text = _fmt_p(r.get("a_p"))

        cells[6].text = f"{_fmt(r.get('b'))} ({_fmt(r.get('b_se'))})"
        cells[7].text = _fmt_p(r.get("b_p"))

        cells[8].text = _fmt(r.get("ab"))
        cells[9].text = _fmt_p(r.get("sobel_p"))
        cells[10].text = _fmt_ci(r.get("ab_ci_low"), r.get("ab_ci_high"))

        cells[11].text = f"{_fmt(r.get('cprime'))} ({_fmt(r.get('cprime_se'))})"
        cells[12].text = _fmt_p(r.get("cprime_p"))

        cells[13].text = f"{_fmt(r.get('ctotal'))} ({_fmt(r.get('ctotal_se'))})"
        cells[14].text = _fmt_p(r.get("ctotal_p"))

        cells[15].text = _fmt(r.get("pm"))
        cells[16].text = _fmt_ci(r.get("pm_ci_low"), r.get("pm_ci_high"))

    _right_align_numeric(t, [3,4,5,6,7,8,9,10,11,12,13,14,15,16])
    _apply_three_line_table(t, header_row_idx=0)

    leg = doc.add_paragraph()
    leg.add_run(legend).italic = True

    doc.save(out_path)
    print(f"Word table saved: {out_path}")
    return out_path


# ======================== MAIN ========================

all_rows = []
np.random.seed(RANDOM_SEED)

legend_text = (
    "Legend: Mediation analyses evaluated whether white matter hyperintensity (WMH) burden "
    "mediates the association between sleep apnea (SA) and cognitive outcomes. "
    "SA was defined by the group variable (Study vs Control). "
    "Three WMH mediators (Total WMH, periventricular WMH [PWMH], deep WMH [DWMH]) "
    "were head-size–normalized and log1p-transformed. "
    "Cognitive outcomes were standardized (z-scores) with higher values indicating better performance; "
    "reaction time and Trail making test-B were log-transformed and multiplied by −1 prior to standardization. "
    "All models adjusted for prespecified covariates (age at Instance 2, sex, Townsend deprivation index, "
    "baseline BMI, genetic ethnic grouping, smoking status, alcohol intake frequency, educational attainment, "
    "and APOE ε4 allele count). Cardiovascular–metabolic condition (CMC) variables were not included. "
    "Standard errors were cluster-robust by matched set (match_id) where feasible; otherwise HC3 robust SEs were used. "
    "Indirect effects (a×b) were evaluated using Sobel tests and 2000 bootstrap resamples "
    "(cluster bootstrap by match_id when available)."
)

for cohort_label, fp in COHORT_FILES.items():
    if not os.path.exists(fp):
        print(f"[Warning] File not found: {fp} (skip {cohort_label})")
        continue

    print(f"\n=== {cohort_label} ===")
    subdir = os.path.join(OUTROOT, re.sub(r'[^0-9A-Za-z]+', '_', cohort_label).strip("_"))
    os.makedirs(subdir, exist_ok=True)

    df = pd.read_csv(fp)
    df = clean_columns(df)
    df = ensure_group(df)
    df = build_wmh_vars(df)
    df = prepare_cognitive_scores(df, cohort_label)

    # Keep only rows with group defined
    df = df.dropna(subset=[GROUP_COL])

    cohort_rows = []

    for m_col, m_label in WMH_MEDIATORS:
        if m_col not in df.columns:
            continue

        for y_col, y_label in COG_ENDPOINTS:
            if y_col not in df.columns:
                continue

            # Outcome-specific exclusion
            d = df.copy()
            d = d[~d[y_col].isna()]
            d = d[~d[m_col].isna()]
            d = d.dropna(subset=[GROUP_COL])

            if d.empty:
                continue

            # N as SA/Control after endpoint & mediator filtering
            n_sa = int((d[GROUP_COL] == "Study").sum())
            n_ctrl = int((d[GROUP_COL] == "Control").sum())
            n_str = f"{n_sa}/{n_ctrl}"

            # Compute analytic paths
            paths = compute_paths(d, m_col, y_col)
            if paths is None:
                continue

            a, a_se, a_p = paths["a"], paths["a_se"], paths["a_p"]
            b, b_se, b_p = paths["b"], paths["b_se"], paths["b_p"]
            cprime, cprime_se, cprime_p = paths["cprime"], paths["cprime_se"], paths["cprime_p"]
            ctotal, ctotal_se, ctotal_p = paths["ctotal"], paths["ctotal_se"], paths["ctotal_p"]

            ab = a * b

            # Sobel
            z_sobel, p_sobel = sobel_p(a, a_se, b, b_se)

            # Proportion mediated (robust check)
            pm = (ab / ctotal) if (pd.notna(ctotal) and ctotal != 0) else np.nan

            # Bootstrap with progress
            desc = f"{cohort_label} | {m_label} | {y_label}"
            boot = bootstrap_indirect(
                d, m_col, y_col,
                n_boot=BOOTSTRAP_N,
                seed=RANDOM_SEED + (hash((cohort_label, m_label, y_label)) % 100000),
                desc=desc
            )

            row = {
                "Cohort": cohort_label,
                "Outcome": y_label,
                "Outcome_col": y_col,
                "Mediator": m_label,
                "Mediator_col": m_col,
                "N": n_str,

                "a": a, "a_se": a_se, "a_p": a_p,
                "b": b, "b_se": b_se, "b_p": b_p,
                "ab": ab,
                "sobel_z": z_sobel, "sobel_p": p_sobel,

                "ab_boot_mean": boot["ab_boot_mean"],
                "ab_ci_low": boot["ab_ci_low"],
                "ab_ci_high": boot["ab_ci_high"],

                "cprime": cprime, "cprime_se": cprime_se, "cprime_p": cprime_p,
                "ctotal": ctotal, "ctotal_se": ctotal_se, "ctotal_p": ctotal_p,

                "pm": pm,
                "pm_boot_mean": boot["pm_boot_mean"],
                "pm_ci_low": boot["pm_ci_low"],
                "pm_ci_high": boot["pm_ci_high"],

                "n_boot_ok": boot["n_boot_ok"],
                "SE_a": paths["se_info_a"],
                "SE_b": paths["se_info_b"],
                "SE_c": paths["se_info_c"],
            }

            cohort_rows.append(row)
            all_rows.append(row)

    # Save per-cohort CSV
    cohort_df = pd.DataFrame(cohort_rows)
    if not cohort_df.empty:
        out_csv = os.path.join(subdir, "Mediation_Results.csv")
        cohort_df.to_csv(out_csv, index=False, encoding="utf-8-sig")
        print(f"[{cohort_label}] CSV saved: {out_csv}")


# Combined outputs
if all_rows:
    comb = pd.DataFrame(all_rows)

    out_comb_csv = os.path.join(OUTROOT, "Mediation_Results_AllCohorts_Combined.csv")
    comb.to_csv(out_comb_csv, index=False, encoding="utf-8-sig")
    print(f"Combined CSV saved: {out_comb_csv}")

    out_docx = os.path.join(OUTROOT, "Mediation_Results_AllCohorts.docx")
    export_mediation_word(
        comb,
        out_docx,
        title="Table. Mediation of the association between sleep apnea and cognitive outcomes by WMH burden (No CMC)",
        legend=legend_text
    )

print("\nSA → WMH → Cognitive mediation analysis completed.")
