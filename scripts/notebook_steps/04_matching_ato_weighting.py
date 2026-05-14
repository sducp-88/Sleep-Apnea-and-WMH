# Extracted from WMH_Analysis.ipynb, code cell 3.
# Notebook heading: # Matching and ATO Weighting Pipeline 
# Run this file from the repository root unless a local CONFIG section is edited.

# Matching and ATO Weighting Pipeline 
"""
UK Biobank Propensity Score Matching and ATO Weighting Pipeline
=======================================================================================
This script implements a reproducible pipeline for propensity score (PS)–based cohort
construction and overlap weighting, designed for publication-quality epidemiologic
analyses of UK Biobank neuroimaging data.

Workflow
--------
1. Neuro Exclusion (Primary Cohort)
   • Exclude Study participants with any recorded neurologic disease prior to the
     sleep apnea (SA) index date.
   • Control participants are not excluded at this stage.

2. Propensity Score Estimation
   • Logistic regression using prespecified baseline covariates:
     age at MRI, sex, baseline BMI, Townsend deprivation index, genetic ethnic
     grouping, smoking status, and alcohol intake frequency.
   • Continuous covariates are median-imputed; categorical covariates are mode-imputed
     and label-encoded.
   • Primary engine: scikit-learn (lbfgs); fallback: NumPy IRLS.

3. 1:10 Nearest-Neighbor Matching (PS logit scale)
   • Matching performed without replacement and without caliper.
   • Primary cohort: post-matching risk-set pruning removes controls with any
     neurologic disease preceding the assigned SA index date.
   • Sensitivity cohort: same matching procedure but without pruning.

4. Balance Diagnostics
   • Standardized mean differences (SMDs) calculated for each covariate before and
     after matching.
   • Results exported as CSV for transparency and reproducibility.

5. Supplementary eMethod (Word Document)
   • Generates a publication-ready description of the study population, covariates,
     PS estimation, matching algorithm, and balance assessment.

6. ATO (Average Treatment effect in the Overlap population) Weighting
   • Primary ATO: MRI-anchored symmetric neuro exclusion at Instance 2, followed by
     PS estimation and overlap weights.
   • Sensitivity ATO: pseudo-index construction stratified by PS quantiles, symmetric
     neuro exclusion, PS re-estimation, and overlap weights.
   • Outputs include weighted balance diagnostics and effective sample size.

Outputs
-------
• Primary and sensitivity matched cohorts (CSV)
• Balance diagnostics before and after matching
• Neurology-style eMethod (Word, publication-ready)
• ATO cohorts (primary and sensitivity) with weights and diagnostics

Dependencies
------------
• pandas, numpy, python-docx
• scikit-learn (preferred; automatic fallback to NumPy IRLS if unavailable)
"""

from pathlib import Path
from textwrap import dedent
import pandas as pd
import numpy as np

# sklearn logistic regression (preferred)
USE_SKLEARN = True
try:
    from sklearn.linear_model import LogisticRegression
except Exception:
    USE_SKLEARN = False

# Word export
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =========================
# Configuration
# =========================
INPUT_FILE = "data_processed.csv"

PRIMARY_INPUT_AFTER_NEURO = "data_processed_exclude_neuro.csv"
PRIMARY_OUT     = "primary_cohort.csv"
SENSITIVITY_OUT = "sensitivity_cohort.csv"

MATCH_DIR         = Path("Matching")
MATCH_DIR.mkdir(exist_ok=True)
MATCH_DIR_PRIMARY = MATCH_DIR / "Primary"
MATCH_DIR_SENS    = MATCH_DIR / "Sensitivity"
for p in [MATCH_DIR_PRIMARY, MATCH_DIR_SENS]:
    p.mkdir(parents=True, exist_ok=True)

# ID & core fields
COL_ID_CANDIDATES = ["Participant_ID", "eid"]  # fallback: first column
COL_TREAT         = "treatment_var"            # 1=Study, 0=Control

# ---- Canonical names used INSIDE code ----
# We'll automatically map real CSV columns (with or without trailing underscore) to these canonical names.
COL_SA_DATE_CANON = "Date_G47_first_reported_sleep_disorders"  # canonical
NEURO_DATE_COLS_CANON = [
    "Date_I63_first_reported_cerebral_infarction",
    "Date_I64_first_reported_stroke_not_specified_as_haemorrhage_or_infarction",
    "Date_G45_first_reported_transient_cerebral_ischaemic_attacks_and_related_syndromes",
    "Date_G35_first_reported_multiple_sclerosis",
    "Date_I67_first_reported_other_cerebrovascular_diseases",
]

# Prespecified covariates (use-if-present; others skipped with a warning)
BASE_COVARIATES = [
    "Sex",
    "Age_at_Instance_2",
    "Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "Alcohol_intake_frequency_ordinal",
]

# Matching parameters
RATIO = 10
WITH_REPLACEMENT = False


# =========================
# Column aliasing (handles trailing underscores removed by your CSV)
# =========================
# Map from our canonical name -> list of acceptable variants in the CSV (priority order)
COLUMN_ALIASES = {
    # SA index date
    COL_SA_DATE_CANON: [
        "Date_G47_first_reported_sleep_disorders",          # preferred
        "Date_G47_first_reported_sleep_disorders"           # without trailing underscore
    ],
    # Neuro dates
    "Date_I63_first_reported_cerebral_infarction": [
        "Date_I63_first_reported_cerebral_infarction",
        "Date_I63_first_reported_cerebral_infarction"
    ],
    "Date_I64_first_reported_stroke_not_specified_as_haemorrhage_or_infarction": [
        "Date_I64_first_reported_stroke_not_specified_as_haemorrhage_or_infarction",
        "Date_I64_first_reported_stroke_not_specified_as_haemorrhage_or_infarction"
    ],
    "Date_G45_first_reported_transient_cerebral_ischaemic_attacks_and_related_syndromes": [
        "Date_G45_first_reported_transient_cerebral_ischaemic_attacks_and_related_syndromes",
        "Date_G45_first_reported_transient_cerebral_ischaemic_attacks_and_related_syndromes"
    ],
    "Date_G35_first_reported_multiple_sclerosis": [
        "Date_G35_first_reported_multiple_sclerosis",
        "Date_G35_first_reported_multiple_sclerosis"
    ],
    "Date_I67_first_reported_other_cerebrovascular_diseases": [
        "Date_I67_first_reported_other_cerebrovascular_diseases",
        "Date_I67_first_reported_other_cerebrovascular_diseases"
    ],
}

def apply_column_aliases(df: pd.DataFrame) -> pd.DataFrame:
    """
    Ensure that the DataFrame has our canonical column names by renaming
    any existing variant (with/without trailing underscore) to the canonical.
    """
    df = df.copy()
    current = set(df.columns)
    rename_map = {}
    for canon, variants in COLUMN_ALIASES.items():
        for v in variants:
            if v in current:
                rename_map[v] = canon
                break
    if rename_map:
        df = df.rename(columns=rename_map)
    return df


# =========================
# Utilities
# =========================
def pick_id_column(df: pd.DataFrame) -> str:
    for c in COL_ID_CANDIDATES:
        if c in df.columns:
            return c
    return df.columns[0]

def ensure_datetime(df: pd.DataFrame, cols: list) -> pd.DataFrame:
    out = df.copy()
    for c in cols:
        if c in out.columns:
            out[c] = pd.to_datetime(out[c], errors="coerce")
    return out

def ps_logit(p: np.ndarray) -> np.ndarray:
    p = np.clip(p, 1e-6, 1 - 1e-6)
    return np.log(p / (1 - p))

def fill_and_encode_inplace(df: pd.DataFrame, cols: list) -> tuple[pd.DataFrame, list]:
    """
    Prepare covariates for logistic regression:
      - numeric: median imputation
      - non-numeric: mode imputation -> categorical codes
    Returns (prepared_df, actually_used_covs)
    """
    used = []
    out = df.copy()
    for col in cols:
        if col not in out.columns:
            print(f"[WARN] Missing covariate dropped: {col}")
            continue
        s = out[col]
        if pd.api.types.is_numeric_dtype(s):
            out[col] = s.fillna(s.median()) if s.notna().any() else s.fillna(0)
        else:
            if s.isna().all():
                out[col] = "Unknown"
            else:
                mode_val = s.mode(dropna=True)
                mode_val = mode_val.iloc[0] if not mode_val.empty else "Unknown"
                out[col] = s.fillna(mode_val)
            out[col] = out[col].astype("category").cat.codes
        used.append(col)
    return out, used

def smd_numeric(x: pd.Series, g: pd.Series) -> float:
    """Standardized mean difference for numeric variable by binary group g (1=Study, 0=Control)."""
    x1 = x[g == 1].astype(float)
    x0 = x[g == 0].astype(float)
    if len(x1) < 2 or len(x0) < 2:
        return np.nan
    m1, m0 = x1.mean(), x0.mean()
    s1, s0 = x1.std(ddof=1), x0.std(ddof=1)
    sp = np.sqrt(((len(x1)-1)*s1**2 + (len(x0)-1)*s0**2) / max(len(x1)+len(x0)-2, 1))
    return float((m1 - m0) / sp) if sp > 0 else np.nan

def balance_table(df: pd.DataFrame, covs: list, treat_col: str, label: str) -> pd.DataFrame:
    rows = []
    for c in covs:
        if c not in df.columns:
            continue
        s = df[c]
        if not pd.api.types.is_numeric_dtype(s):
            s = s.astype("category").cat.codes
        smd = smd_numeric(s, df[treat_col])
        rows.append({"label": label, "covariate": c, "SMD": smd})
    return pd.DataFrame(rows)

def write_text(path: Path, text: str):
    path.write_text(text.strip() + "\n", encoding="utf-8")

# Fallback: tiny NumPy IRLS logistic regression (only used if sklearn unavailable)
def _irls_logistic_predict_proba(X: np.ndarray, y: np.ndarray, max_iter: int = 1000, tol: float = 1e-6) -> np.ndarray:
    X = np.c_[np.ones((X.shape[0], 1)), X]  # add intercept
    beta = np.zeros(X.shape[1])
    for _ in range(max_iter):
        z = X @ beta
        p = 1.0 / (1.0 + np.exp(-np.clip(z, -50, 50)))
        W = p * (1 - p) + 1e-8
        z_tilde = z + (y - p) / W
        XT_W = X.T * W
        H = XT_W @ X
        g = XT_W @ z_tilde
        try:
            beta_new = np.linalg.solve(H, g)
        except np.linalg.LinAlgError:
            beta_new = np.linalg.pinv(H) @ g
        if np.max(np.abs(beta_new - beta)) < tol:
            beta = beta_new
            break
        beta = beta_new
    p_final = 1.0 / (1.0 + np.exp(-np.clip(X @ beta, -50, 50)))
    return np.clip(p_final, 1e-9, 1-1e-9)


# =========================
# Step 1: Neuro exclusion (Study only -> primary input)
# =========================
def neuro_exclusion_study_only(input_file: str, output_file: str):
    df = pd.read_csv(input_file)
    # --- apply alias mapping so we can use canonical names below ---
    df = apply_column_aliases(df)

    need_cols = [COL_TREAT, COL_SA_DATE_CANON, *NEURO_DATE_COLS_CANON]
    miss = [c for c in need_cols if c not in df.columns]
    if miss:
        raise KeyError(f"Missing required columns for neuro exclusion (after aliasing): {miss}")

    df = ensure_datetime(df, [COL_SA_DATE_CANON] + NEURO_DATE_COLS_CANON)
    study_mask = df[COL_TREAT] == 1
    lt_matrix = pd.DataFrame({c: df[c].lt(df[COL_SA_DATE_CANON]) for c in NEURO_DATE_COLS_CANON})
    exclude_mask = study_mask & lt_matrix.any(axis=1)

    kept = df[~exclude_mask].copy()
    excluded = df[exclude_mask].copy()
    kept.to_csv(output_file, index=False)

    summ = dedent(f"""
    ===== Neuro Exclusion (Study only; ANY neuro date < SA index_date) =====
    Input : {input_file}
    Output: {output_file}

    Original: total={len(df)}, Study={(df[COL_TREAT]==1).sum()}, Control={(df[COL_TREAT]==0).sum()}
    Excluded: total={len(excluded)}, Study={(excluded[COL_TREAT]==1).sum()}, Control={(excluded[COL_TREAT]==0).sum()}
    Final   : total={len(kept)}, Study={(kept[COL_TREAT]==1).sum()}, Control={(kept[COL_TREAT]==0).sum()}

    Note: Column aliasing enabled (with/without trailing underscores).
    """).strip()
    write_text(MATCH_DIR / "Neuro_Exclusion_Summary.txt", summ)
    print(summ)


# =========================
# Step 2A: Propensity score estimation
# =========================
def estimate_propensity(df_in: pd.DataFrame, covariates: list) -> tuple[pd.DataFrame, list, str]:
    """
    Preferred: sklearn LogisticRegression(lbfgs, max_iter=1000).
    Fallback : pure-NumPy IRLS with identical input/output interface.
    Returns: (df_with_ps, used_covariates, engine_str)
    """
    if COL_TREAT not in df_in.columns:
        raise KeyError(f"'{COL_TREAT}' not found.")

    df = df_in.copy()
    # also ensure aliasing here in case user modified columns between steps
    df = apply_column_aliases(df)

    df[COL_TREAT] = pd.to_numeric(df[COL_TREAT], errors="coerce")
    df = df.dropna(subset=[COL_TREAT])

    df, used_covs = fill_and_encode_inplace(df, covariates)
    if not used_covs:
        raise ValueError("No usable covariates for propensity estimation.")

    X = df[used_covs].to_numpy()
    y = df[COL_TREAT].astype(int).to_numpy()

    if USE_SKLEARN:
        model = LogisticRegression(max_iter=1000, solver="lbfgs")
        model.fit(X, y)
        df["propensity_score"] = np.clip(model.predict_proba(X)[:, 1], 1e-9, 1-1e-9)
        engine = "sklearn-LogisticRegression (lbfgs)"
    else:
        df["propensity_score"] = _irls_logistic_predict_proba(X, y, max_iter=1000, tol=1e-6)
        engine = "NumPy-IRLS (fallback)"

    return df, used_covs, engine


# =========================
# Step 2B: Matching (Primary with pruning; Sensitivity without)
# =========================
def match_primary(df_ps: pd.DataFrame, ratio: int, out_csv: str, diag_dir: Path):
    # Ensure aliasing (for dates) before referencing canonical column names
    df_ps = apply_column_aliases(df_ps)
    id_col = pick_id_column(df_ps)

    need = [id_col, COL_TREAT, "propensity_score", COL_SA_DATE_CANON, "hospitalization_exposure"] + NEURO_DATE_COLS_CANON
    miss = [c for c in need if c not in df_ps.columns]
    if miss:
        raise KeyError(f"Missing required columns for primary matching (after aliasing): {miss}")

    df = ensure_datetime(df_ps, [COL_SA_DATE_CANON] + NEURO_DATE_COLS_CANON).copy()
    df["_orig_index_"] = np.arange(len(df))

    # ensure exposure is numeric/binary
    df["hospitalization_exposure"] = pd.to_numeric(df["hospitalization_exposure"], errors="coerce").fillna(0).astype(int)

    df["_logit_ps"] = ps_logit(df["propensity_score"].to_numpy())
    treated  = df[df[COL_TREAT] == 1].copy().sort_values("_logit_ps").reset_index(drop=True)
    controls = df[df[COL_TREAT] == 0].copy().sort_values("_logit_ps").reset_index(drop=True)
    controls["_available"] = True

    matched_rows, treated_used = [], []
    drop_no_avail = 0
    drop_no_same_exp = 0

    def pick_k_neighbors_same_exposure(trow, controls_df, k):
        # exact-match pool: same hospitalization_exposure AND still available
        pool = controls_df[(controls_df["_available"]) &
                           (controls_df["hospitalization_exposure"] == trow["hospitalization_exposure"])]
        if pool.empty:
            return [], True  # no same-exposure controls
        diffs = np.abs(pool["_logit_ps"].values - trow["_logit_ps"])
        order = np.argsort(diffs)
        ordered = pool.iloc[order]
        return ordered.index[:min(k, len(ordered))].tolist(), False

    for _, t in treated.iterrows():
        if not controls["_available"].any():
            drop_no_avail += 1
            continue
        idxs, no_same = pick_k_neighbors_same_exposure(t, controls, ratio)
        if len(idxs) == 0:
            if no_same:
                drop_no_same_exp += 1
            else:
                drop_no_avail += 1
            continue
        treated_used.append(t[id_col])
        for ci in idxs:
            crow = controls.loc[ci].copy()
            if not WITH_REPLACEMENT:
                controls.at[ci, "_available"] = False
            row = crow.to_dict()
            row["match_id_orig"] = t[id_col]
            matched_rows.append(row)

    matched_treated_df = df[(df[COL_TREAT]==1) & (df[id_col].isin(treated_used))].copy()
    matched_treated_df["group"] = "Study"
    matched_treated_df["match_id_orig"] = matched_treated_df[id_col].values
    matched_treated_df["index_date"] = matched_treated_df[COL_SA_DATE_CANON]

    matched_controls_df = pd.DataFrame(matched_rows)
    if matched_controls_df.empty:
        raise ValueError("No controls matched in primary stream.")
    matched_controls_df["group"] = "Control"
    study_idx_map = matched_treated_df.set_index("match_id_orig")[COL_SA_DATE_CANON].to_dict()
    matched_controls_df["index_date"] = matched_controls_df["match_id_orig"].map(study_idx_map)

    matched_all = pd.concat([matched_treated_df, matched_controls_df], ignore_index=True)

    # Risk-set pruning on controls: drop if ANY neuro date < index_date
    control_mask = (matched_all["group"] == "Control")
    lt_matrix = pd.DataFrame({
        c: matched_all.loc[control_mask, c].lt(matched_all.loc[control_mask, "index_date"])
        for c in NEURO_DATE_COLS_CANON
    })
    control_bad = lt_matrix.any(axis=1)
    bad_idx = matched_all.loc[control_mask].index[control_bad.values]
    n_ctrl_dropped = len(bad_idx)

    matched_all_pruned = matched_all.drop(index=bad_idx).copy()

    # Remove SA cases without remaining controls
    post_counts = (matched_all_pruned[matched_all_pruned["group"]=="Control"]
                   .groupby("match_id_orig")[id_col].count())
    studies_no_ctrl = (matched_all_pruned[matched_all_pruned["group"]=="Study"]
                       .loc[lambda d: ~d["match_id_orig"].isin(post_counts.index), ["match_id_orig"]])
    n_study_dropped = len(studies_no_ctrl)
    if n_study_dropped > 0:
        matched_all_pruned = matched_all_pruned[
            ~matched_all_pruned["match_id_orig"].isin(studies_no_ctrl["match_id_orig"])
        ].copy()

    # Recode match_id
    study_order = (matched_all_pruned[matched_all_pruned["group"]=="Study"]
                   .sort_values("_orig_index_")["match_id_orig"].drop_duplicates().tolist())
    id_map = {orig_id: i for i, orig_id in enumerate(study_order)}
    matched_all_pruned["match_id"] = matched_all_pruned["match_id_orig"].map(id_map).astype(int)

    # Sort & save
    matched_all_pruned["group"] = pd.Categorical(matched_all_pruned["group"],
                                                 categories=["Control","Study"], ordered=True)
    matched_all_pruned = matched_all_pruned.sort_values(by=["group", "match_id", id_col]).reset_index(drop=True)
    matched_all_pruned = matched_all_pruned.drop(
        columns=[c for c in ["_available","_logit_ps","_orig_index_","match_id_orig"] if c in matched_all_pruned.columns]
    )
    matched_all_pruned.to_csv(out_csv, index=False, encoding="utf-8-sig")

    # Diagnostics
    n_study_after = (matched_all_pruned["group"]=="Study").sum()
    n_ctrl_after  = (matched_all_pruned["group"]=="Control").sum()
    avg_ctrls     = (n_ctrl_after / n_study_after) if n_study_after else 0.0

    bal_before = balance_table(df_ps.assign(treatment_var=df_ps[COL_TREAT].astype(int)),
                                BASE_COVARIATES, "treatment_var", "before")
    bal_after  = balance_table(matched_all_pruned.assign(treatment_var=(matched_all_pruned["group"]=="Study").astype(int)),
                               BASE_COVARIATES, "treatment_var", "after")
    pd.concat([bal_before, bal_after], ignore_index=True)\
      .to_csv(diag_dir / "Balance_Before_After.csv", index=False, encoding="utf-8")

    summ = dedent(f"""
    ===== Primary Matching (1:{RATIO}, exact-match on hospitalization_exposure, no caliper, no replacement) =====
    Unmatched SA (no controls available): {drop_no_avail}
    Unmatched SA (no same-exposure controls): {drop_no_same_exp}
    Risk-set pruning:
      Controls dropped (ANY neuro < index_date) : {n_ctrl_dropped}
      Studies dropped (no controls left)        : {n_study_dropped}

    Final matched sizes:
      Study (after)   : {n_study_after}
      Control (after) : {n_ctrl_after}
      Avg controls per Study: {avg_ctrls:.2f} (target={RATIO})

    Note: Column aliasing enabled; exposure exact-matching enforced.
    """).strip()
    write_text(diag_dir / "PSM_Summary.txt", summ)
    print(summ)


def match_sensitivity(df_ps: pd.DataFrame, ratio: int, out_csv: str, diag_dir: Path):
    df_ps = apply_column_aliases(df_ps)
    id_col = pick_id_column(df_ps)

    need = [id_col, COL_TREAT, "propensity_score", COL_SA_DATE_CANON, "hospitalization_exposure"]
    miss = [c for c in need if c not in df_ps.columns]
    if miss:
        raise KeyError(f"Missing required columns for sensitivity matching (after aliasing): {miss}")

    df = ensure_datetime(df_ps, [COL_SA_DATE_CANON]).copy()
    df["hospitalization_exposure"] = pd.to_numeric(df["hospitalization_exposure"], errors="coerce").fillna(0).astype(int)

    treated = df[df[COL_TREAT] == 1].reset_index(drop=True)
    control = df[df[COL_TREAT] == 0].reset_index(drop=True)

    t_scores = ps_logit(treated["propensity_score"].to_numpy())
    c_scores = ps_logit(control["propensity_score"].to_numpy())

    matched_flags = np.zeros(len(control), dtype=bool)
    match_results, drop_no_avail, drop_no_same_exp = [], 0, 0

    for i, ps in enumerate(t_scores):
        # available controls, SAME exposure only
        same_exp_mask = (control["hospitalization_exposure"].values == treated.loc[i, "hospitalization_exposure"])
        avail = np.where((~matched_flags) & same_exp_mask)[0]
        if avail.size == 0:
            # no same-exposure controls left
            drop_no_same_exp += 1
            continue
        diffs = np.abs(c_scores[avail] - ps)
        k = min(ratio, len(avail))
        best = avail[np.argpartition(diffs, k-1)[:k]]
        matched_flags[best] = True
        match_results.append((i, best.tolist()))

    # Assemble
    matched_controls = []
    for t_idx, c_idxs in match_results:
        for c_idx in c_idxs:
            r = control.iloc[c_idx].copy()
            r["match_id_tmp"] = t_idx
            matched_controls.append(r)
    matched_controls = pd.DataFrame(matched_controls)

    matched_treated_ids = [t_idx for t_idx, _ in match_results]
    treated_sub = treated.loc[matched_treated_ids].copy()
    treated_sub["old_idx"] = treated_sub.index

    id_map = {old: new for new, old in enumerate(matched_treated_ids)}
    treated_sub["match_id"] = treated_sub["old_idx"].map(id_map)
    if not matched_controls.empty:
        matched_controls["match_id"] = matched_controls["match_id_tmp"].map(id_map)

    treated_sub["group"] = "Study"
    if not matched_controls.empty:
        matched_controls["group"] = "Control"

    # index_date propagation
    study_index_map = treated_sub.set_index("match_id")[COL_SA_DATE_CANON].to_dict()
    treated_sub["index_date"] = treated_sub[COL_SA_DATE_CANON]
    if not matched_controls.empty:
        matched_controls["index_date"] = matched_controls["match_id"].map(study_index_map)

    matched_all = pd.concat([treated_sub, matched_controls], ignore_index=True)

    # Cleanup & sort
    for tmp in ["match_id_tmp", "old_idx"]:
        if tmp in matched_all.columns:
            matched_all = matched_all.drop(columns=[tmp])

    matched_all["group"] = pd.Categorical(matched_all["group"], categories=["Control","Study"], ordered=True)
    matched_all = matched_all.sort_values(by=["group","match_id", id_col]).reset_index(drop=True)

    matched_all.to_csv(out_csv, index=False, encoding="utf-8-sig")

    # Diagnostics
    n_study = (matched_all["group"]=="Study").sum()
    n_ctrl  = (matched_all["group"]=="Control").sum()
    matched_counts = [len(c_idxs) for _, c_idxs in match_results]
    dist_text = ""
    if matched_counts:
        vc = pd.Series(matched_counts).value_counts().sort_index()
        dist_text = "\n".join([f"  {k} controls: {v}" for k, v in vc.items()])

    bal_before = balance_table(df_ps, BASE_COVARIATES, COL_TREAT, "before")
    bal_after  = balance_table(matched_all.assign(treatment_var=(matched_all["group"]=="Study").astype(int)),
                               BASE_COVARIATES, "treatment_var", "after")
    pd.concat([bal_before, bal_after], ignore_index=True)\
      .to_csv(diag_dir / "Balance_Before_After.csv", index=False, encoding="utf-8")

    summ = dedent(f"""
    ===== Sensitivity Matching (1:{RATIO}, exact-match on hospitalization_exposure, no caliper, no replacement, no pruning) =====
    SA with no same-exposure controls: {drop_no_same_exp}
    Final matched sizes:
      Study   : {n_study}
      Control : {n_ctrl}
    """).strip()
    write_text(diag_dir / "PSM_Summary.txt", summ)
    print(summ)



# =========================
# eMethod (Neurology-style Word)
# =========================
def _doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

def export_emethod_matching(primary_used_covs: list, sens_used_covs: list, engine_primary: str, engine_sens: str):
    doc = Document()
    _doc_style(doc)
    doc.add_heading("Supplementary eMethod: Propensity Score Estimation and Matching", level=1)

    # Study population
    doc.add_heading("Study population", level=2)
    para = ("UK Biobank imaging participants with usable T1- and T2-FLAIR MRI at Instance 2 were included. "
            "Participants with missing WMH outcomes were excluded upstream. The primary cohort additionally "
            "excluded Study participants with any neurologic diagnosis recorded prior to the sleep apnea (SA) "
            "index date; the sensitivity cohort applied no such exclusion.")
    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Variables
    doc.add_heading("Variables", level=2)
    para = ("Prespecified covariates were age at MRI, sex, baseline BMI (Instance 0), Townsend deprivation index, "
            "genetic ethnic grouping (binary), smoking (ever vs never), and alcohol intake frequency (ordinal). "
            "Missing data were handled by median imputation for continuous covariates and mode imputation with "
            "label encoding for categorical covariates.")
    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Propensity score estimation
    doc.add_heading("Propensity score estimation", level=2)
    para = (f"Propensity scores were estimated via logistic regression separately for the two cohorts. "
            f"For the primary cohort, the engine was {engine_primary}; covariates actually used: "
            f"{', '.join(primary_used_covs) if primary_used_covs else 'None'}. "
            f"For the sensitivity cohort, the engine was {engine_sens}; covariates actually used: "
            f"{', '.join(sens_used_covs) if sens_used_covs else 'None'}. "
            f"Column aliasing was implemented to harmonize date-field names with/without trailing underscores.")
    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Matching procedure
    doc.add_heading("Matching procedure", level=2)
    para = (f"Each SA case was matched to up to {RATIO} controls using nearest-neighbor matching on the logit of the "
            "propensity score, without replacement and without a caliper.")
    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sensitivity analysis
    doc.add_heading("Sensitivity analysis", level=2)
    para = ("In the primary cohort, post-matching risk-set pruning excluded controls with any neurologic diagnosis "
            "preceding the matched SA index date; SA cases without remaining controls were removed. "
            "The sensitivity cohort conducted the same matching without pruning.")
    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Balance assessment
    doc.add_heading("Balance assessment", level=2)
    para = ("Covariate balance was evaluated using standardized mean differences (SMDs) before and after matching. "
            "Diagnostics are provided in the Matching directory.")
    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    out_path = MATCH_DIR / "eMethod_Matching.docx"
    doc.save(out_path)
    print(f"[OK] Neurology-style eMethod saved: {out_path}")


# =========================
# Main
# =========================
def main():
    # Step 1: Neuro exclusion -> primary input
    neuro_exclusion_study_only(INPUT_FILE, PRIMARY_INPUT_AFTER_NEURO)

    # Step 2: Primary stream
    df_primary_base = pd.read_csv(PRIMARY_INPUT_AFTER_NEURO)
    df_primary_base = apply_column_aliases(df_primary_base)
    df_primary_ps, primary_used_covs, engine_primary = estimate_propensity(df_primary_base, BASE_COVARIATES)
    match_primary(df_primary_ps, RATIO, PRIMARY_OUT, MATCH_DIR_PRIMARY)

    # Step 2: Sensitivity stream
    df_sens_base = pd.read_csv(INPUT_FILE)
    df_sens_base = apply_column_aliases(df_sens_base)
    df_sens_ps, sens_used_covs, engine_sens = estimate_propensity(df_sens_base, BASE_COVARIATES)
    match_sensitivity(df_sens_ps, RATIO, SENSITIVITY_OUT, MATCH_DIR_SENS)

    # eMethod (Neurology style)
    export_emethod_matching(primary_used_covs, sens_used_covs, engine_primary, engine_sens)

    print("\n[DONE] Outputs:")
    print(f"  • {PRIMARY_INPUT_AFTER_NEURO}")
    print(f"  • {PRIMARY_OUT}")
    print(f"  • {SENSITIVITY_OUT}")
    print(f"  • Matching diagnostics under: {MATCH_DIR}/")
    if not USE_SKLEARN:
        print("  [Note] sklearn not available in this environment; used NumPy IRLS fallback for PS.")

if __name__ == "__main__":
    main()

# =========================
# ATO (Overlap Weighting) cohorts
# =========================

# --- New outputs & dirs ---
ATO_DIR = Path("ATO")
ATO_DIR_SYM = ATO_DIR / "SymmetricExclusion"
ATO_DIR_NOEX = ATO_DIR / "NoExclusion"
for p in [ATO_DIR, ATO_DIR_SYM, ATO_DIR_NOEX]:
    p.mkdir(parents=True, exist_ok=True)

ATO_SYM_OUT    = "ato_sensitivity_sym.csv"
ATO_NOEX_OUT   = "ato_sensitivity_noexclusion.csv"

# --- Add a canonical I2 MRI date column & aliases (best-effort guesses; extend if needed) ---
COL_I2_DATE_CANON = "Date_of_attending_assessment_centre_Instance_2"
COLUMN_ALIASES[COL_I2_DATE_CANON] = [
    "Date_of_attending_assessment_centre_Instance_2",
    "Date_of_attending_assessment_centre_instance_2",
    "Date_of_attending_assessment_centre_Instance_2_",
    "Date_of_attending_assessment_centre_instance_2_",
]

def detect_i2_date_column(df: pd.DataFrame) -> str | None:
    """
    Heuristic: try to find an I2 date column if aliases failed.
    Looks for case-insensitive tokens around 'Instance_2' and any of {'MRI','Imaging','assessment','centre','center','date'}.
    """
    cand = []
    toks_any = ("mri", "imaging", "assessment", "centre", "center", "date")
    for c in df.columns:
        low = c.lower()
        if "instance_2" in low and any(t in low for t in toks_any):
            cand.append(c)
    if not cand:
        for c in df.columns:
            low = c.lower()
            if "instance 2" in low and any(t in low for t in toks_any):
                cand.append(c)
    return cand[0] if cand else None

def ensure_i2_date_present(df: pd.DataFrame) -> pd.DataFrame:
    """Ensure df has COL_I2_DATE_CANON (rename or detect)."""
    df2 = apply_column_aliases(df)
    if COL_I2_DATE_CANON not in df2.columns:
        guess = detect_i2_date_column(df2)
        if guess is None:
            raise KeyError(
                f"Could not find I2 MRI date column. "
                f"Add an alias for '{COL_I2_DATE_CANON}' to COLUMN_ALIASES or ensure a suitable column exists."
            )
        df2 = df2.rename(columns={guess: COL_I2_DATE_CANON})
    return ensure_datetime(df2, [COL_I2_DATE_CANON])

# ---------- Weighted balance diagnostics ----------
def wmean(x: np.ndarray, w: np.ndarray) -> float:
    w = np.asarray(w, float)
    x = np.asarray(x, float)
    s = w.sum()
    return float((w * x).sum() / s) if s > 0 else np.nan

def wvar(x: np.ndarray, w: np.ndarray) -> float:
    mu = wmean(x, w)
    s = w.sum()
    if s <= 0:
        return np.nan
    return float((w * (x - mu) ** 2).sum() / s)

def wstd(x: np.ndarray, w: np.ndarray) -> float:
    v = wvar(x, w)
    return float(np.sqrt(v)) if v >= 0 else np.nan

def weighted_smd(x: pd.Series, treat: pd.Series, w: pd.Series) -> float:
    """Weighted SMD using pooled weighted SD."""
    x = x.copy()
    if not pd.api.types.is_numeric_dtype(x):
        x = x.astype("category").cat.codes
    g1 = (treat == 1).values
    g0 = (treat == 0).values
    w = w.values.astype(float)
    xv = x.values.astype(float)

    mu1 = wmean(xv[g1], w[g1]); mu0 = wmean(xv[g0], w[g0])
    s1  = wstd(xv[g1], w[g1]);  s0  = wstd(xv[g0], w[g0])
    denom = np.sqrt((s1**2 + s0**2) / 2.0) if np.isfinite(s1) and np.isfinite(s0) else np.nan
    if denom is None or not np.isfinite(denom) or denom == 0:
        return np.nan
    return float((mu1 - mu0) / denom)

def weighted_balance_table(df: pd.DataFrame, covs: list, treat_col: str, wcol: str, label: str) -> pd.DataFrame:
    rows = []
    for c in covs:
        if c not in df.columns:
            continue
        smd = weighted_smd(df[c], df[treat_col], df[wcol])
        rows.append({"label": label, "covariate": c, "Weighted_SMD": smd})
    return pd.DataFrame(rows)

def effective_sample_size(w: np.ndarray) -> float:
    w = np.asarray(w, float)
    num = (w.sum()) ** 2
    den = (w ** 2).sum()
    return float(num / den) if den > 0 else np.nan

# ---------- ATO weights ----------
def add_ato_weights(df: pd.DataFrame, ps_col: str = "propensity_score", treat_col: str = COL_TREAT,
                    normalize: bool = True, add_group_norm: bool = True) -> pd.DataFrame:
    out = df.copy()
    if ps_col not in out.columns:
        raise KeyError(f"Propensity score column '{ps_col}' not found.")
    if treat_col not in out.columns:
        raise KeyError(f"Treatment column '{treat_col}' not found.")
    p = np.clip(out[ps_col].astype(float).values, 1e-6, 1 - 1e-6)
    z = out[treat_col].astype(int).values
    w = np.where(z == 1, 1.0 - p, p)  # ATO weights
    out["ato_weight"] = w

    if normalize:
        out["ato_weight_norm"] = out["ato_weight"] * (len(out) / out["ato_weight"].sum())

    if add_group_norm:
        out["ato_weight_gnorm"] = out["ato_weight"]
        for g in [0, 1]:
            mask = (out[treat_col] == g)
            s = out.loc[mask, "ato_weight"].sum()
            if s > 0:
                out.loc[mask, "ato_weight_gnorm"] = out.loc[mask, "ato_weight"] * (mask.sum() / s)
    return out

# ---------- ATO 1: MRI-anchored symmetric exclusion ----------
def build_ato_sensitivity_sym(input_csv: str, covariates: list) -> pd.DataFrame:
    df = pd.read_csv(input_csv)
    df = ensure_i2_date_present(df)
    df = apply_column_aliases(df)
    need_cols = [COL_TREAT, COL_I2_DATE_CANON] + NEURO_DATE_COLS_CANON
    miss = [c for c in need_cols if c not in df.columns]
    if miss:
        raise KeyError(f"[ATO-Sym] Missing required columns: {miss}")

    df = ensure_datetime(df, [COL_I2_DATE_CANON, COL_SA_DATE_CANON] + NEURO_DATE_COLS_CANON)

    # Exclude ANY neuro date < I2 date (both Study and Control)
    neuro_before_i2 = pd.DataFrame({c: df[c].le(df[COL_I2_DATE_CANON]) for c in NEURO_DATE_COLS_CANON})
    df = df.loc[~neuro_before_i2.any(axis=1)].copy()

    df_ps, used_covs, engine = estimate_propensity(df, covariates)
    df_ps = add_ato_weights(df_ps, ps_col="propensity_score", treat_col=COL_TREAT,
                            normalize=True, add_group_norm=True)

    diag = ATO_DIR_SYM
    wbal = weighted_balance_table(df_ps, used_covs, COL_TREAT, "ato_weight_norm", "ATO_sensitivity_sym")
    wbal.to_csv(diag / "Weighted_Balance.csv", index=False, encoding="utf-8-sig")

    ess_all  = effective_sample_size(df_ps["ato_weight_norm"].values)
    ess_treat = effective_sample_size(df_ps.loc[df_ps[COL_TREAT]==1, "ato_weight_norm"].values)
    ess_ctrl  = effective_sample_size(df_ps.loc[df_ps[COL_TREAT]==0, "ato_weight_norm"].values)

    summ = dedent(f"""
    ===== ATO Sensitivity (MRI-anchored symmetric neuro exclusion at I2) =====
    Input : {input_csv}
    Rows after exclusion: {len(df_ps)}

    Propensity engine: {engine}
    Covariates used  : {', '.join(used_covs) if used_covs else 'None'}

    Effective Sample Size (normalized weights):
      ESS (all)   : {ess_all:.1f}
      ESS (Study) : {ess_treat:.1f}
      ESS (Control): {ess_ctrl:.1f}
    """).strip()
    write_text(diag / "ATO_Summary.txt", summ)
    print(summ)

    out_cols_first = [pick_id_column(df_ps), COL_TREAT, "propensity_score",
                      "ato_weight", "ato_weight_norm", "ato_weight_gnorm"]
    out_cols = out_cols_first + [c for c in df_ps.columns if c not in out_cols_first]
    df_ps[out_cols].to_csv(ATO_SYM_OUT, index=False, encoding="utf-8-sig")
    return df_ps

# ---------- ATO 2: No exclusion ----------
def build_ato_sensitivity_noex(input_csv: str, covariates: list) -> pd.DataFrame:
    df = pd.read_csv(input_csv)
    df = ensure_i2_date_present(df)
    df = apply_column_aliases(df)
    df = ensure_datetime(df, [COL_I2_DATE_CANON, COL_SA_DATE_CANON] + NEURO_DATE_COLS_CANON)

    
    df_ps, used_covs, engine = estimate_propensity(df, covariates)
    df_ps = add_ato_weights(df_ps, ps_col="propensity_score", treat_col=COL_TREAT,
                            normalize=True, add_group_norm=True)

    diag = ATO_DIR_NOEX
    wbal = weighted_balance_table(df_ps, used_covs, COL_TREAT, "ato_weight_norm", "ATO_sensitivity_noex")
    wbal.to_csv(diag / "Weighted_Balance.csv", index=False, encoding="utf-8-sig")

    ess_all  = effective_sample_size(df_ps["ato_weight_norm"].values)
    ess_treat = effective_sample_size(df_ps.loc[df_ps[COL_TREAT]==1, "ato_weight_norm"].values)
    ess_ctrl  = effective_sample_size(df_ps.loc[df_ps[COL_TREAT]==0, "ato_weight_norm"].values)

    summ = dedent(f"""
    ===== ATO Sensitivity (No neuro exclusion) =====
    Input : {input_csv}
    Rows after exclusion: {len(df_ps)} (no exclusion applied)

    Propensity engine: {engine}
    Covariates used  : {', '.join(used_covs) if used_covs else 'None'}

    Effective Sample Size (normalized weights):
      ESS (all)   : {ess_all:.1f}
      ESS (Study) : {ess_treat:.1f}
      ESS (Control): {ess_ctrl:.1f}
    """).strip()
    write_text(diag / "ATO_Summary.txt", summ)
    print(summ)

    out_cols_first = [pick_id_column(df_ps), COL_TREAT, "propensity_score",
                      "ato_weight", "ato_weight_norm", "ato_weight_gnorm"]
    out_cols = out_cols_first + [c for c in df_ps.columns if c not in out_cols_first]
    df_ps[out_cols].to_csv(ATO_NOEX_OUT, index=False, encoding="utf-8-sig")
    return df_ps

# ---------- Convenient entrypoint ----------
def main_ato():
    """
    Build two ATO sensitivity cohorts:
      - Symmetric exclusion at I2
      - No exclusion
    """
    df_sym = build_ato_sensitivity_sym(INPUT_FILE, BASE_COVARIATES)
    df_noex = build_ato_sensitivity_noex(INPUT_FILE, BASE_COVARIATES)

    print("\n[DONE] ATO outputs (sensitivity cohorts):")
    print(f"  • {ATO_SYM_OUT}    (diagnostics in {ATO_DIR_SYM}/)")
    print(f"  • {ATO_NOEX_OUT}   (diagnostics in {ATO_DIR_NOEX}/)")
    if not USE_SKLEARN:
        print("  [Note] sklearn not available; NumPy IRLS fallback was used where needed.")


# Optionally run ATO when this file is executed directly
if __name__ == "__main__":
    
    main_ato()
