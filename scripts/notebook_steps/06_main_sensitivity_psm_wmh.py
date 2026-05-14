# Extracted from WMH_Analysis.ipynb, code cell 5.
# Notebook heading: # Main + Sensitiviy PSM Cohort WMH pipeline
# Run this file from the repository root unless a local CONFIG section is edited.

# Main + Sensitiviy PSM Cohort WMH pipeline
"""
WMH pipeline
- Head-size normalized log1p OLS
- Forest plots (single + overlay)
- Table 1 (p-values) CSV + Word (three-line table)
- Supplement SMD tables + Love plot
- Manuscript OLS table (Primary & Sensitivity together) CSV + Word (three-line)
- Figure legends (TXT + Word)

"""

import os, re
import pandas as pd
import numpy as np
import statsmodels.api as sm
import matplotlib as mpl
import matplotlib.pyplot as plt
from patsy import dmatrices
from statsmodels.stats.diagnostic import het_breuschpagan
from statsmodels.stats.multitest import multipletests
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# optional deps
try:
    import scipy.stats as st
except Exception:
    st = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

# ---------------- CONFIG ----------------
CLUSTER_VAR = "match_id"
GROUP_BASELINE = ["Control", "Study"]

MAIN_OUT_DIR = "Main_Outcome"
os.makedirs(MAIN_OUT_DIR, exist_ok=True)

file_sets = {
    "Primary 1:10 (NoNeuro)":       "primary_cohort.csv",
    "Sensitivity 1:10 (WithNeuro)": "sensitivity_cohort.csv"
}

SHORT = {
    "Log_HSNorm_Total_WMH_T1_T2":     "Total WMH",
    "Log_HSNorm_PeriVentricular_WMH": "Periventricular WMH",
    "Log_HSNorm_Deep_WMH":            "Deep WMH",
}

# plot order (and table order)
PLOT_ORDER = ["Total WMH", "Periventricular WMH", "Deep WMH"]

# covariates
BASE_ADJ = [
    "Sex",
    "Age_at_Instance_2",
    "Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "Alcohol_intake_frequency_ordinal",
]
CATEGORICAL = {"Sex", "Smoking_Ever", "Genetic_ethnic_grouping"}

CATEGORICAL.update({"CMC_score_cat"})

# threshold for gatekeeping
PRIMARY_P_ALPHA = 0.05

# Table 1 settings
TABLE1_USE_PVALUE = True      # show p-values
TABLE1_INCLUDE_SMD = False    # add an SMD column in Table 1 (usually False; SMD in supplement)

# Word export appearance
THREELINE_TABLES = True       # export Word tables as three-line tables
RIGHT_ALIGN_NUMERIC = True    # right-align numeric columns

# labels shown on figures (tables keep the full names)
DISPLAY_LABEL = {
    "Total WMH":            "Total WMH",
    "Periventricular WMH":  "PWMH",
    "Deep WMH":             "DWMH",
}
# global style
mpl.rcParams.update({
    "font.family": "Arial",
    "figure.dpi": 120, "savefig.dpi": 600,
    "font.size": 12, "axes.labelsize": 12,
    "xtick.labelsize": 11, "ytick.labelsize": 11,
    "figure.autolayout": True
})

# --- High-quality TIFF export defaults ---
SAVEFIG_DPI = 600
PNG_KW  = {"dpi": SAVEFIG_DPI, "bbox_inches": "tight", "facecolor": "white"}
PDF_KW  = {"bbox_inches": "tight"}
SVG_KW  = {"bbox_inches": "tight"}
TIFF_KW = {
    "dpi": SAVEFIG_DPI,
    "bbox_inches": "tight",
    "facecolor": "white",
    "format": "tiff",
    # Pillow compression (will be ignored if backend not Pillow)
    "pil_kwargs": {"compression": "tiff_lzw"}
}
# ---- PLOT STYLE TOGGLES ----
SHOW_X_GRID = False            # TOGGLE: draw vertical dotted grid on x-axis (Neurology style: False)
STRIP_TOP_RIGHT_SPINES = True  # TOGGLE: remove top/right panel spines for a cleaner look

# ---------------- UTILITIES ----------------
def build_formula(outcome, adjust_vars, categorical_vars, group_var="group"):
    terms = [f"C({group_var})"]
    for v in adjust_vars:
        if v == group_var:
            continue
        terms.append(f"C({v})" if v in categorical_vars else v)
    return f"{outcome} ~ " + " + ".join(terms)

def find_group_term(params_index, group_var="group", target_level="Study"):
    cands = [p for p in params_index if p.startswith(f"C({group_var})[T.")]
    if not cands: return None
    for p in cands:
        if p.endswith(f"[T.{target_level}]"): return p
    return cands[0]

def fit_with_cluster(formula, data, cluster_var):
    """return: model, SE_info, y_used, X_used, n_clusters, n_obs"""
    y, X = dmatrices(formula, data, return_type="dataframe", NA_action="drop")
    n_obs = len(y)
    base = sm.OLS(y, X).fit(cov_type="HC3")
    if cluster_var not in data.columns:
        return base, "HC3 (no cluster var)", y, X, 0, n_obs
    groups = data.loc[y.index, cluster_var]
    if groups.notna().sum() == 0:
        return base, "HC3 (all cluster NaN)", y, X, 0, n_obs
    if groups.isna().any():
        m = groups.notna(); y, X, groups = y.loc[m], X.loc[m], groups.loc[m]; n_obs = len(y)
    n_clusters = int(groups.nunique())
    if n_clusters > 1:
        mod = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": groups.to_numpy()})
        used = f"cluster (groups={cluster_var}, n_clusters={n_clusters}, n={n_obs})"
        return mod, used, y, X, n_clusters, n_obs
    return base, "HC3 (<=1 cluster)", y, X, n_clusters, n_obs

def count_group_n(df_like):
    if "Participant_ID" in df_like.columns:
        n_c = int(df_like[df_like["group"]=="Control"]["Participant_ID"].nunique())
        n_s = int(df_like[df_like["group"]=="Study"]["Participant_ID"].nunique())
    else:
        n_c = int((df_like["group"]=="Control").sum())
        n_s = int((df_like["group"]=="Study").sum())
    return n_c, n_s, n_c + n_s

# ---------------- PLOTTING ----------------
def plot_single_forest(res_log: pd.DataFrame, label: str, outdir: str):
    """Horizontal forest plot, compact layout, legend omitted (single cohort)."""
    if res_log is None or res_log.empty:
        return

    # keep your order logic
    order_map = {v: i for i, v in enumerate(PLOT_ORDER)}
    dfp = res_log.sort_values(by="Outcome", key=lambda s: s.map(order_map)).copy()

    # --- NEW: ensure numeric + valid CI and pre-compute xlim that includes 0 ---
    for col in ["% Change", "% CI Lower", "% CI Upper"]:
        dfp[col] = pd.to_numeric(dfp[col], errors="coerce")

    # fix accidental swapped CI (rare but defensive)
    bad = dfp["% CI Lower"] > dfp["% CI Upper"]
    if bad.any():
        lo = dfp.loc[bad, "% CI Upper"].values
        hi = dfp.loc[bad, "% CI Lower"].values
        dfp.loc[bad, "% CI Lower"] = lo
        dfp.loc[bad, "% CI Upper"] = hi

    # compute axis limits from CI and FORCE include 0
    xmin = float(np.nanmin(dfp["% CI Lower"]))
    xmax = float(np.nanmax(dfp["% CI Upper"]))
    xmin = min(xmin, 0.0)
    xmax = max(xmax, 0.0)
    # add 8% padding; also guard degenerate range
    span = xmax - xmin
    if not np.isfinite(span) or span <= 0:
        span = 1.0
    pad = 0.08 * span

    fig = plt.figure(figsize=(6.8, 3.8))
    ax = plt.gca()
    y = np.arange(len(dfp))[::-1]
    pad_y = 0.45

    # build symmetric xerr safely (non-negative widths; NaN -> 0)
    lo = (dfp["% Change"] - dfp["% CI Lower"]).astype(float).to_numpy()
    hi = (dfp["% CI Upper"] - dfp["% Change"]).astype(float).to_numpy()
    lo = np.where(np.isfinite(lo) & (lo >= 0), lo, 0.0)
    hi = np.where(np.isfinite(hi) & (hi >= 0), hi, 0.0)

    ax.errorbar(
        dfp["% Change"], y,
        xerr=[lo, hi],
        fmt='o', ms=5.5, mfc="#1f3b4d", mec="#1f3b4d",
        ecolor="#1f3b4d", elinewidth=1.6, capsize=4,
        linestyle="none", color="#1f3b4d"
    )

    ax.axvline(0, color='grey', linestyle='--', linewidth=1)  # zero reference line
    ax.set_yticks(y, [DISPLAY_LABEL.get(v, v) for v in dfp["Outcome"]])
    ax.set_ylim(y.min()-pad_y, y.max()+pad_y)
    ax.set_xlabel("% change (Study - Control)")

    # --- NEW: apply the xlim after plotting so it always shows 0 ---
    ax.set_xlim(xmin - pad, xmax + pad)

    if SHOW_X_GRID:
        ax.grid(axis='x', linestyle=':', linewidth=0.7, alpha=0.8)
    else:
        ax.grid(False)
    if STRIP_TOP_RIGHT_SPINES:
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    plt.tight_layout()
    safe = re.sub(r"[^0-9a-zA-Z]+", "_", label).strip("_")
    base = os.path.join(outdir, f"{safe}_LOG1P_pct_CI")
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".svg", **SVG_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)


def plot_overlay_forest(res_primary: pd.DataFrame,
                        res_sensitivity: pd.DataFrame,
                        outdir: str,
                        basename: str):
    """Overlay: Primary vs Sensitivity. Legend outside (right-centered), no box."""
    if res_primary is None or res_primary.empty or res_sensitivity is None or res_sensitivity.empty:
        return

    order_map = {v: i for i, v in enumerate(PLOT_ORDER)}
    L = res_primary.sort_values("Outcome", key=lambda s: s.map(order_map)).reset_index(drop=True).copy()
    R = res_sensitivity.sort_values("Outcome", key=lambda s: s.map(order_map)).reset_index(drop=True).copy()

    if not L["Outcome"].equals(R["Outcome"]):
        common = pd.Index(L["Outcome"]).intersection(R["Outcome"])
        L = L[L["Outcome"].isin(common)].sort_values("Outcome", key=lambda s: s.map(order_map))
        R = R[R["Outcome"].isin(common)].sort_values("Outcome", key=lambda s: s.map(order_map))

    # --- NEW: force include 0 in combined xlim ---
    for df in (L, R):
        for col in ["% Change", "% CI Lower", "% CI Upper"]:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    xmin = float(np.nanmin([L["% CI Lower"].min(), R["% CI Lower"].min()]))
    xmax = float(np.nanmax([L["% CI Upper"].max(), R["% CI Upper"].max()]))
    xmin = min(xmin, 0.0)
    xmax = max(xmax, 0.0)
    span = xmax - xmin
    if not np.isfinite(span) or span <= 0:
        span = 1.0
    pad_x = 0.07 * span
    xmin -= pad_x
    xmax += pad_x

    c_primary, c_sens = "#1f3b4d", "#6e6e6e"
    m_primary, m_sens = "o", "s"
    lw, cap, jitter, pad_y = 1.6, 4, 0.12, 0.45

    fig, ax = plt.subplots(figsize=(6.9, 3.8))
    y = np.arange(len(L))[::-1]

    # build xerr safely
    L_lo = (L["% Change"] - L["% CI Lower"]).astype(float).to_numpy()
    L_hi = (L["% CI Upper"] - L["% Change"]).astype(float).to_numpy()
    L_lo = np.where(np.isfinite(L_lo) & (L_lo >= 0), L_lo, 0.0)
    L_hi = np.where(np.isfinite(L_hi) & (L_hi >= 0), L_hi, 0.0)

    R_lo = (R["% Change"] - R["% CI Lower"]).astype(float).to_numpy()
    R_hi = (R["% CI Upper"] - R["% Change"]).astype(float).to_numpy()
    R_lo = np.where(np.isfinite(R_lo) & (R_lo >= 0), R_lo, 0.0)
    R_hi = np.where(np.isfinite(R_hi) & (R_hi >= 0), R_hi, 0.0)

    # Primary
    ax.errorbar(
        L["% Change"], y + jitter,
        xerr=[L_lo, L_hi],
        fmt=m_primary, ms=5.5, mfc=c_primary, mec=c_primary,
        ecolor=c_primary, elinewidth=lw, capsize=cap,
        linestyle="none", label="Primary PSM cohort"
    )
    # Sensitivity
    ax.errorbar(
        R["% Change"], y - jitter,
        xerr=[R_lo, R_hi],
        fmt=m_sens, ms=5.2, mfc="white", mec=c_sens,
        ecolor=c_sens, elinewidth=lw, capsize=cap,
        linestyle="none", label="PSM-sensitivity cohort"
    )

    ax.axvline(0, color="grey", linestyle="--", linewidth=1)
    ax.set_yticks(y, [DISPLAY_LABEL.get(v, v) for v in L["Outcome"]])
    ax.set_ylim(y.min() - pad_y, y.max() + pad_y)
    ax.set_xlabel("% change (SA - Control)")
    ax.set_xlim(xmin, xmax)  # <-- ensure 0 is visible

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    ax.legend(
        frameon=False, loc="center left",
        bbox_to_anchor=(1.02, 0.5),
        handlelength=1.8, fontsize=9,
        labelspacing=0.6,
        markerscale=0.9
    )

    plt.subplots_adjust(right=0.78)

    safe = re.sub(r"[^0-9a-zA-Z]+", "_", basename).strip("_")
    base = os.path.join(outdir, safe)
    plt.savefig(base + ".png", **PNG_KW)
    plt.savefig(base + ".pdf", **PDF_KW)
    plt.savefig(base + ".svg", **SVG_KW)
    plt.savefig(base + ".tiff", **TIFF_KW)
    plt.close(fig)


# ---------------- Table 1 helpers ----------------
def _as_num(s): return pd.to_numeric(s, errors="coerce")
def _format_mean_sd(x): x=_as_num(x); return f"{x.mean():.2f} ({x.std(ddof=1):.2f})"
def _format_median_iqr(x):
    x=_as_num(x); q1,q3=x.quantile([0.25,0.75]); return f"{x.median():.2f} [{q1:.2f}, {q3:.2f}]"
def _is_true_series(s, true=("yes","true","ever","y","1")):
    if pd.api.types.is_numeric_dtype(s): return s.astype(float)==1.0
    ss=s.astype(str).str.strip().str.lower(); return ss.isin(true)
def _is_female_series(s):
    if pd.api.types.is_numeric_dtype(s): return s.astype(float)==0.0
    ss=s.astype(str).str.strip().str.lower(); return ss.str.startswith("f")|(ss=="female")
def _is_white_series(s):
    if pd.api.types.is_numeric_dtype(s): return s.astype(float)==1.0
    ss=s.astype(str).str.strip().str.lower()
    return ss.str.contains("white")|ss.str.contains("british")|ss.str.contains("cauc")
def _pct(b): b=b.fillna(False); return 100.0*b.mean()
def _smd_cont(x1,x0):
    x1=_as_num(x1); x0=_as_num(x0); m1,m0=x1.mean(),x0.mean(); s1,s0=x1.std(ddof=1),x0.std(ddof=1)
    sp=np.sqrt((s1**2+s0**2)/2.0); return (m1-m0)/sp if sp>0 else np.nan
def _smd_binary(b1,b0):
    p1,p0=b1.mean(),b0.mean(); p=(p1+p0)/2.0; denom=np.sqrt(p*(1-p)); return (p1-p0)/denom if denom>0 else np.nan
def _p_ttest_welch(x1,x0):
    if st is None: return np.nan
    x1=_as_num(x1).dropna(); x0=_as_num(x0).dropna()
    if len(x1)<2 or len(x0)<2: return np.nan
    return float(st.ttest_ind(x1,x0,equal_var=False).pvalue)
def _p_mannwhitney(x1,x0):
    if st is None: return np.nan
    x1=_as_num(x1).dropna(); x0=_as_num(x0).dropna()
    if len(x1)<1 or len(x0)<1: return np.nan
    try: return float(st.mannwhitneyu(x1,x0,alternative="two-sided").pvalue)
    except Exception: return np.nan
def _p_cat_2x2(b1,b0):
    if st is None: return np.nan
    a=int(b1.sum()); b=int((~b1).sum()); c=int(b0.sum()); d=int((~b0).sum())
    table=np.array([[a,b],[c,d]],dtype=float)
    try:
        chi2,p_chi,_,exp = st.chi2_contingency(table, correction=False)
        if (exp<5).any():
            _,p = st.fisher_exact(table, alternative="two-sided"); return float(p)
        return float(p_chi)
    except Exception:
        return np.nan

def make_table1(dmod: pd.DataFrame, analysis_label: str, id_col="Participant_ID") -> pd.DataFrame:
    dfb = dmod.copy()
    g1 = dfb["group"] == "Study"
    g0 = dfb["group"] == "Control"

    if id_col in dfb.columns:
        n1 = dfb.loc[g1, id_col].nunique()
        n0 = dfb.loc[g0, id_col].nunique()
    else:
        n1 = int(g1.sum())
        n0 = int(g0.sum())

    rows = []

    def add(name, s_val, c_val, p=None, smd=None):
        row = {"Characteristic": name, "Study": s_val, "Control": c_val}
        if TABLE1_USE_PVALUE:
            row["p_value"] = p
        if TABLE1_INCLUDE_SMD:
            row["SMD"] = smd
        rows.append(row)

    # ---- Formatters (1 decimal place throughout) ----
    def _as_num(s): 
        return pd.to_numeric(s, errors="coerce")

    def _format_mean_sd(x):
        x = _as_num(x)
        return f"{x.mean():.1f} ({x.std(ddof=1):.1f})"

    def _format_median_iqr(x):
        x = _as_num(x)
        q1, q3 = x.quantile([0.25, 0.75])
        return f"{x.median():.1f} [{q1:.1f}, {q3:.1f}]"

    def _pct(b):
        b = b.fillna(False)
        return 100.0 * b.mean()

    # ---- Variables ----
    if "Age_at_Instance_2" in dfb.columns:
        s = dfb.loc[g1, "Age_at_Instance_2"]
        c = dfb.loc[g0, "Age_at_Instance_2"]
        add("Age at Instance 2, mean (SD), y",
            _format_mean_sd(s), _format_mean_sd(c),
            p=_p_ttest_welch(s, c),
            smd=round(_smd_cont(s, c), 3))

    if "Sex" in dfb.columns:
        sF = _is_female_series(dfb.loc[g1, "Sex"])
        cF = _is_female_series(dfb.loc[g0, "Sex"])
        add("Female sex, %",
            f"{_pct(sF):.1f}", f"{_pct(cF):.1f}",
            p=_p_cat_2x2(sF.astype(bool), cF.astype(bool)),
            smd=round(_smd_binary(sF.astype(float), cF.astype(float)), 3))

    if "Body_mass_index_BMI_Instance_0" in dfb.columns:
        s = dfb.loc[g1, "Body_mass_index_BMI_Instance_0"]
        c = dfb.loc[g0, "Body_mass_index_BMI_Instance_0"]
        add("BMI at Instance 0, mean (SD), kg/m²",
            _format_mean_sd(s), _format_mean_sd(c),
            p=_p_ttest_welch(s, c),
            smd=round(_smd_cont(s, c), 3))

    if "Genetic_ethnic_grouping" in dfb.columns:
        sW = _is_white_series(dfb.loc[g1, "Genetic_ethnic_grouping"])
        cW = _is_white_series(dfb.loc[g0, "Genetic_ethnic_grouping"])
        add("Genetic ethnic grouping, White %",
            f"{_pct(sW):.1f}", f"{_pct(cW):.1f}",
            p=_p_cat_2x2(sW.astype(bool), cW.astype(bool)),
            smd=round(_smd_binary(sW.astype(float), cW.astype(float)), 3))

    if "Townsend_deprivation_index_at_recruitment" in dfb.columns:
        s = dfb.loc[g1, "Townsend_deprivation_index_at_recruitment"]
        c = dfb.loc[g0, "Townsend_deprivation_index_at_recruitment"]
        add("Townsend deprivation index, mean (SD)",
            _format_mean_sd(s), _format_mean_sd(c),
            p=_p_ttest_welch(s, c),
            smd=round(_smd_cont(s, c), 3))

    if "Smoking_Ever" in dfb.columns:
        sE = _is_true_series(dfb.loc[g1, "Smoking_Ever"])
        cE = _is_true_series(dfb.loc[g0, "Smoking_Ever"])
        add("Smoking, ever, %",
            f"{_pct(sE):.1f}", f"{_pct(cE):.1f}",
            p=_p_cat_2x2(sE.astype(bool), cE.astype(bool)),
            smd=round(_smd_binary(sE.astype(float), cE.astype(float)), 3))

    if "Alcohol_intake_frequency_ordinal" in dfb.columns:
        s = dfb.loc[g1, "Alcohol_intake_frequency_ordinal"]
        c = dfb.loc[g0, "Alcohol_intake_frequency_ordinal"]
        add("Alcohol intake frequency, median [IQR]",
            _format_median_iqr(s), _format_median_iqr(c),
            p=_p_mannwhitney(s, c),
            smd=round(_smd_cont(_as_num(s), _as_num(c)), 3))
        

    # ---- CMC (comorbidity count): show only ≥1, n (%) WITHOUT p-value/SMD ----
    _cmc_candidates = ["CMC_score_raw", "CMC_score_cat", "CMC", "Comorbidity_count"]
    _cmc_found = [c for c in _cmc_candidates if c in dfb.columns]

    if _cmc_found:
        cmc_col = _cmc_found[0]
        cmc_s = pd.to_numeric(dfb.loc[g1, cmc_col], errors="coerce")
        cmc_c = pd.to_numeric(dfb.loc[g0, cmc_col], errors="coerce")

        car_s = cmc_s >= 1
        car_c = cmc_c >= 1

        n_s = int(car_s.sum(skipna=True))
        n_c = int(car_c.sum(skipna=True))

        den_s = int(cmc_s.notna().sum())
        den_c = int(cmc_c.notna().sum())

        pct_s = (100.0 * n_s / den_s) if den_s > 0 else float("nan")
        pct_c = (100.0 * n_c / den_c) if den_c > 0 else float("nan")

        add("CMC ≥1, n (%)",
            f"{n_s} ({pct_s:.1f}%)",
            f"{n_c} ({pct_c:.1f}%)",
            p=None)

    # ---- APOE ε4: show only carrier (≥1), n (%) WITHOUT p-value/SMD ----
    _apoe_count_candidates = [
        "APOE_e4_count", "APOE_e4count", "e4_count", "APOE_e4",
        "APOE_e4_cnt", "APOE_e4_num"
    ]
    _apoe_carrier_candidates = ["APOE_e4_carrier", "APOE_e4_any", "APOE_e4_has"]

    _ap_cnt = [c for c in _apoe_count_candidates if c in dfb.columns]
    _ap_car = [c for c in _apoe_carrier_candidates if c in dfb.columns]

    if _ap_cnt or _ap_car:
        if _ap_cnt:
            ap_col = _ap_cnt[0]
            ap_s = pd.to_numeric(dfb.loc[g1, ap_col], errors="coerce")
            ap_c = pd.to_numeric(dfb.loc[g0, ap_col], errors="coerce")
            car_s = ap_s >= 1
            car_c = ap_c >= 1
            den_s = int(ap_s.notna().sum())
            den_c = int(ap_c.notna().sum())
        else:
            ap_col = _ap_car[0]
            # accept 1/0, True/False, or strings "1"/"0"
            ap_s_raw = dfb.loc[g1, ap_col]
            ap_c_raw = dfb.loc[g0, ap_col]
            # normalize to boolean with NaN respected
            def _to_bool(s):
                if pd.api.types.is_bool_dtype(s):
                    return s
                s2 = pd.to_numeric(s, errors="coerce")
                if s2.notna().any():
                    return s2 == 1
                ss = s.astype(str).str.strip().str.lower()
                return ss.isin(["1", "true", "yes", "y"])
            car_s = _to_bool(ap_s_raw)
            car_c = _to_bool(ap_c_raw)
            den_s = int(pd.notna(ap_s_raw).sum())
            den_c = int(pd.notna(ap_c_raw).sum())

        n_s = int(pd.Series(car_s).sum(skipna=True))
        n_c = int(pd.Series(car_c).sum(skipna=True))

        pct_s = (100.0 * n_s / den_s) if den_s > 0 else float("nan")
        pct_c = (100.0 * n_c / den_c) if den_c > 0 else float("nan")

        add("APOE ε4 carrier (≥1), n (%)",
            f"{n_s} ({pct_s:.1f}%)",
            f"{n_c} ({pct_c:.1f}%)",
            p=None)
            # ---- Has Degree: show only Yes n (%) WITHOUT p-value/SMD ----
    _degree_candidates = [
        "Has_Degree", "has_degree", "degree", "Qualifications_has_degree"
    ]
    _degree_found = [c for c in _degree_candidates if c in dfb.columns]

    if _degree_found:
        deg_col = _degree_found[0]
        deg_s = dfb.loc[g1, deg_col]
        deg_c = dfb.loc[g0, deg_col]

        # Normalize to boolean (1/0, True/False, string)
        def _to_bool(s):
            if pd.api.types.is_bool_dtype(s):
                return s
            s2 = pd.to_numeric(s, errors="coerce")
            if s2.notna().any():
                return s2 == 1
            ss = s.astype(str).str.strip().str.lower()
            return ss.isin(["1", "true", "yes", "y"])

        car_s = _to_bool(deg_s)
        car_c = _to_bool(deg_c)

        n_s = int(pd.Series(car_s).sum(skipna=True))
        n_c = int(pd.Series(car_c).sum(skipna=True))
        den_s = int(pd.notna(deg_s).sum())
        den_c = int(pd.notna(deg_c).sum())

        pct_s = (100.0 * n_s / den_s) if den_s > 0 else float("nan")
        pct_c = (100.0 * n_c / den_c) if den_c > 0 else float("nan")

        add("Has degree, n (%)",
            f"{n_s} ({pct_s:.1f}%)",
            f"{n_c} ({pct_c:.1f}%)",
            p=None)




    # ---- Build Table ----
    tbl = pd.DataFrame(rows)
    top = {"Characteristic": "Sample size (unique participants), n",
           "Study": n1, "Control": n0}
    if TABLE1_USE_PVALUE:
        top["p_value"] = np.nan
    if TABLE1_INCLUDE_SMD:
        top["SMD"] = np.nan
    tbl = pd.concat([pd.DataFrame([top]), tbl], ignore_index=True)
    tbl.insert(0, "Analysis", analysis_label)

    return tbl


# ---------------- WORD helpers (three-line tables) ----------------
def _set_cell_border(cell, **kwargs):
    """
    Set cell borders: _set_cell_border(cell, top={"val":"single","sz":8,"color":"000000"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left', 'right', 'top', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = OxmlElement(f'w:{edge}')
            for key in ["val","sz","color","space"]:
                if key in edge_data:
                    tag.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(tag)

def apply_three_line_table(table, header_row_idx=0):
    """
    Convert a Word table to classic 'three-line' format:
    - top border on the header row (table top)
    - bottom border on the header row (header underline)
    - bottom border on the last row (table bottom)
    - remove vertical borders
    """
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, left={"val":"nil"}, right={"val":"nil"}, top={"val":"nil"}, bottom={"val":"nil"})

    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell, top={"val":"single","sz":8,"color":"000000"})

    for cell in table.rows[header_row_idx].cells:
        _set_cell_border(cell, bottom={"val":"single","sz":8,"color":"000000"})

    for cell in table.rows[-1].cells:
        _set_cell_border(cell, bottom={"val":"single","sz":8,"color":"000000"})

def set_table_font(table, font_name="Times New Roman", font_size=10):
    """Apply uniform font and size to all table content (Neurology style)."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)

def _right_align_numeric(table, numeric_col_idx):
    if not RIGHT_ALIGN_NUMERIC: return
    for row in table.rows[1:]:
        for j in numeric_col_idx:
            if j < len(row.cells):
                for p in row.cells[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ---------------- Word exporters ----------------
def export_table1_word(table_df: pd.DataFrame, analysis_label: str, outdir: str, timestamp: str):
    if Document is None:
        print("python-docx not available: Word export skipped."); return None
    doc = Document()
    title = doc.add_paragraph()
    r = title.add_run(f"Table 1. Baseline characteristics — {analysis_label}")
    r.bold = True; title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cols = ["Characteristic","Study","Control"]
    if TABLE1_USE_PVALUE: cols.append("p_value")
    if TABLE1_INCLUDE_SMD: cols.append("SMD")

    t = doc.add_table(rows=1, cols=len(cols))
    hdr = t.rows[0].cells
    for i,c in enumerate(cols): hdr[i].text = c

    for _, row in table_df[["Characteristic","Study","Control"] + (["p_value"] if TABLE1_USE_PVALUE else []) + (["SMD"] if TABLE1_INCLUDE_SMD else [])].iterrows():
        cells = t.add_row().cells
        for j,c in enumerate(cols):
            val = row[c] if c in row else ""
            if c in ("p_value","SMD") and isinstance(val,float) and not np.isnan(val):
                cells[j].text = f"{val:.3g}"
            else:
                cells[j].text = "" if pd.isna(val) else str(val)

    num_idx = [1,2] + ([3] if TABLE1_USE_PVALUE else []) + ([4] if TABLE1_INCLUDE_SMD else [])
    _right_align_numeric(t, num_idx)

    if THREELINE_TABLES:
        apply_three_line_table(t, header_row_idx=0)

    legend = doc.add_paragraph()
    legend.add_run(
        "Legend: Values are mean (SD), median [IQR], or n (%). "
        "p-values from Welch’s t test (continuous), Mann–Whitney U for the ordinal alcohol-frequency variable, "
        "and χ² or Fisher’s exact test for categorical variables, as appropriate. "
        "Groups are Study (sleep apnea) and Control at Instance 2."
    ).italic = True

    safe = re.sub(r"[^0-9a-zA-Z]+","_", analysis_label).strip("_")
    path = os.path.join(outdir, f"{safe}_Table1_Baseline.docx")
    doc.save(path)
    return path

def export_smd_word(df_combined: pd.DataFrame, outdir: str, timestamp: str):
    if Document is None:
        print("python-docx not available: SMD Word export skipped."); return None
    doc = Document()
    title = doc.add_paragraph()
    r = title.add_run("Supplementary Table S1. Covariate balance after matching (SMD)")
    r.bold = True; title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cols = ["Characteristic","SMD (Primary)","|SMD| (Primary)","SMD (Sensitivity)","|SMD| (Sensitivity)"]
    t = doc.add_table(rows=1, cols=len(cols)); hdr = t.rows[0].cells
    for i,c in enumerate(cols): hdr[i].text = c

    for _, row in df_combined.iterrows():
        cells = t.add_row().cells
        for j,c in enumerate(cols):
            val = row.get(c, "")
            cells[j].text = f"{val:.3f}" if (isinstance(val,float) and not np.isnan(val)) else ("" if pd.isna(val) else str(val))

    _right_align_numeric(t, [1,2,3,4])
    if THREELINE_TABLES:
        apply_three_line_table(t, header_row_idx=0)

    foot = doc.add_paragraph()
    foot.add_run(
        "Legend: Standardized mean differences (SMD) summarize post-matching covariate balance. "
        "Values <0.10 indicate acceptable balance. Continuous SMDs use pooled SD; "
        "binary SMDs use Cohen’s h. Alcohol intake frequency is treated as an ordinal continuous variable."
    ).italic = True

    path = os.path.join(outdir, f"Supplement_SMD_Balance.docx")
    doc.save(path)
    return path

# -------- OLS results table (manuscript, adjusted) --------
def _fmt_pct_ci(pct, lo, hi):
    if any(pd.isna(x) for x in [pct, lo, hi]): return "—"
    return f"{pct:.1f}% ({lo:.1f}, {hi:.1f})"
def _fmt_p(p):  return "—" if pd.isna(p) else ("<0.001" if p < 1e-3 else f"{p:.3f}")
def _fmt_q(q):  return "—" if pd.isna(q) else ("<0.001" if q < 1e-3 else f"{q:.3f}")
def _fmt_logci(beta, lo, hi):
    """Format raw log1p coefficient and CI as: β (lo, hi)."""
    if any(pd.isna(x) for x in [beta, lo, hi]):
        return "—"
    return f"{beta:.3f} ({lo:.3f}, {hi:.3f})"

def make_ols_results_table(res_primary: pd.DataFrame,
                           res_sens: pd.DataFrame,
                           meta_primary: pd.DataFrame|None,
                           meta_sens: pd.DataFrame|None) -> pd.DataFrame:
    order_map = {v:i for i,v in enumerate(PLOT_ORDER)}
    rp = res_primary.sort_values("Outcome", key=lambda s: s.map(order_map))
    rs = res_sens.sort_values("Outcome", key=lambda s: s.map(order_map))

    def _get_ns(meta_df):
        try:
            m=meta_df.iloc[0]; return int(m["N_study(used)"]), int(m["N_control(used)"])
        except Exception: return (np.nan, np.nan)

    n_sp,n_cp=_get_ns(meta_primary); n_ss,n_cs=_get_ns(meta_sens)
    rows=[]
    for oc in PLOT_ORDER:
        r1=rp[rp["Outcome"]==oc]; r2=rs[rs["Outcome"]==oc]
        if r1.empty or r2.empty: 
            continue
        r1=r1.iloc[0]; r2=r2.iloc[0]
        rows.append({
            "Outcome": oc,
            "Primary %Δ (95% CI)": _fmt_pct_ci(r1["% Change"], r1["% CI Lower"], r1["% CI Upper"]),
            "Primary log1p β (95% CI)": _fmt_logci(r1["Log β"], r1["Log CI Lower"], r1["Log CI Upper"]),
            "Primary p": _fmt_p(r1["p_value"]),
            "Primary q": _fmt_q(r1.get("q_value", np.nan)),

            "Sensitivity %Δ (95% CI)": _fmt_pct_ci(r2["% Change"], r2["% CI Lower"], r2["% CI Upper"]),
            "Sensitivity log1p β (95% CI)": _fmt_logci(r2["Log β"], r2["Log CI Lower"], r2["Log CI Upper"]),
            "Sensitivity p": _fmt_p(r2["p_value"]),
            "Sensitivity q": _fmt_q(r2.get("q_value", np.nan)),
        })
    df=pd.DataFrame(rows)
    df.attrs["N_primary"]=(n_sp,n_cp)
    df.attrs["N_sensitivity"]=(n_ss,n_cs)
    return df

def export_ols_table_word(df: pd.DataFrame, outdir: str, timestamp: str):
    """
    Export OLS results to a Word table in vertical format (first column = Analytic cohort).
    Adjusted model version.
    """
    if Document is None:
        print("python-docx not available: OLS Word export skipped.")
        return None

    (nsp, ncp) = df.attrs.get("N_primary", (np.nan, np.nan))
    (nss, ncs) = df.attrs.get("N_sensitivity", (np.nan, np.nan))

    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Table 2. Adjusted differences in WMH (linear models)")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cols = ["Analytic cohort", "Outcome", "%Δ (95% CI)", "log1p β (95% CI)", "p", "q"]
    t = doc.add_table(rows=1, cols=len(cols))
    hdr = t.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c

    for _, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = "Primary PSM cohort"
        cells[1].text = row["Outcome"]
        cells[2].text = row["Primary %Δ (95% CI)"]
        cells[3].text = row["Primary log1p β (95% CI)"]
        cells[4].text = row["Primary p"]
        cells[5].text = row["Primary q"]

    for _, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = "PSM–sensitivity cohort"
        cells[1].text = row["Outcome"]
        cells[2].text = row["Sensitivity %Δ (95% CI)"]
        cells[3].text = row["Sensitivity log1p β (95% CI)"]
        cells[4].text = row["Sensitivity p"]
        cells[5].text = row["Sensitivity q"]

    _right_align_numeric(t, [2, 3, 4, 5])

    if THREELINE_TABLES:
        apply_three_line_table(t, header_row_idx=0)
    set_table_font(t, font_name="Times New Roman", font_size=10)

    leg = doc.add_paragraph()
    leg.add_run(
        "Values are adjusted between-group differences in log1p WMH expressed as percent change "
        "(Study vs Control) with 95% confidence intervals; the raw log1p coefficient and its 95% CI "
        "are provided for reference. p values are from the coefficient of the Study indicator. "
        "q values are Benjamini–Hochberg FDR–adjusted within the secondary outcome family "
        "(Periventricular, Deep); Total is the primary endpoint and not FDR-adjusted. "
        "Models adjusted for age at Instance 2, sex, BMI at Instance 0, smoking (ever), alcohol-intake "
        "frequency (ordinal), Townsend deprivation index, and genetic ethnic grouping; cluster-robust "
        "standard errors by matched set (match_id). Gatekeeping: secondary outcomes are interpreted "
        "only if Total WMH is significant at α=0.05."
    ).italic = True

    ninfo = doc.add_paragraph()
    ninfo.add_run(
        f"N (Study / Control): Primary = {nsp} / {ncp}; Sensitivity = {nss} / {ncs}."
    ).italic = True

    path = os.path.join(outdir, f"OLS_Results_Manuscript_Table.docx")
    doc.save(path)
    print(f"OLS results Word table saved: {path}")
    return path

# -------- NEW: Unadjusted OLS results Word table --------
def export_unadjusted_table_word(df: pd.DataFrame, outdir: str, timestamp: str):
    """
    Export UNADJUSTED OLS results (only C(group)) to a Word table,
    mirroring the adjusted table's layout and three-line styling.
    """
    if Document is None:
        print("python-docx not available: Unadjusted Word export skipped.")
        return None

    (nsp, ncp) = df.attrs.get("N_primary", (np.nan, np.nan))
    (nss, ncs) = df.attrs.get("N_sensitivity", (np.nan, np.nan))

    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Supplementary Table U1. Unadjusted differences in WMH (linear models)")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    cols = ["Analytic cohort", "Outcome", "%Δ (95% CI)", "log1p β (95% CI)", "p", "q"]
    t = doc.add_table(rows=1, cols=len(cols))
    hdr = t.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c

    for _, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = "Primary PSM cohort"
        cells[1].text = row["Outcome"]
        cells[2].text = row["Primary %Δ (95% CI)"]
        cells[3].text = row["Primary log1p β (95% CI)"]
        cells[4].text = row["Primary p"]
        cells[5].text = row["Primary q"]

    for _, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = "PSM–sensitivity cohort"
        cells[1].text = row["Outcome"]
        cells[2].text = row["Sensitivity %Δ (95% CI)"]
        cells[3].text = row["Sensitivity log1p β (95% CI)"]
        cells[4].text = row["Sensitivity p"]
        cells[5].text = row["Sensitivity q"]

    _right_align_numeric(t, [2, 3, 4, 5])

    if THREELINE_TABLES:
        apply_three_line_table(t, header_row_idx=0)
    set_table_font(t, font_name="Times New Roman", font_size=10)

    # Legend tailored for unadjusted models
    leg = doc.add_paragraph()
    leg.add_run(
        "Values are UNADJUSTED between-group differences in log1p WMH expressed as percent change "
        "(Study vs Control) with 95% confidence intervals; the raw log1p coefficient and its 95% CI "
        "are provided for reference. p values are from the coefficient of the Study indicator. "
        "q values are Benjamini–Hochberg FDR–adjusted within the secondary outcome family "
        "(Periventricular, Deep); Total is the primary endpoint and not FDR-adjusted. "
        "Models include only the group indicator with cluster-robust standard errors by matched set (match_id)."
    ).italic = True

    ninfo = doc.add_paragraph()
    ninfo.add_run(
        f"N (Study / Control): Primary = {nsp} / {ncp}; Sensitivity = {nss} / {ncs}."
    ).italic = True

    path = os.path.join(outdir, f"UNADJ_Results_Manuscript_Table.docx")
    doc.save(path)
    print(f"Unadjusted results Word table saved: {path}")
    return path

# ---------------- MAIN ----------------
timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
all_log, all_meta, all_tbl1 = [], [], []
res_by_label, meta_by_label, smd_by_label = {}, {}, {}

# NEW containers for UNADJUSTED results
res_unadj_by_label, meta_unadj_by_label = {}, {}

for label, file in file_sets.items():
    print(f"\n=== {label} (log1p main analysis) ===")
    df = pd.read_csv(file)
    df.columns = df.columns.str.replace(r"[^0-9a-zA-Z]+","_", regex=True)

    if "group" not in df.columns or df["group"].dropna().nunique()<2:
        print("Skipping: invalid 'group'."); continue
    df["group"]=pd.Categorical(df["group"], categories=GROUP_BASELINE, ordered=False)

    sc="Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
    if sc not in df.columns:
        print("Skipping: missing scaling factor."); continue
    scale=pd.to_numeric(df[sc], errors="coerce")

    # head-size normalized WMH + log1p
    df["HSNorm_Deep_WMH"]            = pd.to_numeric(df["Total_volume_of_deep_white_matter_hyperintensities_Instance_2"], errors="coerce") * scale
    df["HSNorm_PeriVentricular_WMH"] = pd.to_numeric(df["Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"], errors="coerce") * scale
    df["HSNorm_Total_WMH_T1_T2"]     = pd.to_numeric(df["Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"], errors="coerce") * scale
    df["Log_HSNorm_Deep_WMH"]            = np.log1p(df["HSNorm_Deep_WMH"])
    df["Log_HSNorm_PeriVentricular_WMH"] = np.log1p(df["HSNorm_PeriVentricular_WMH"])
    df["Log_HSNorm_Total_WMH_T1_T2"]     = np.log1p(df["HSNorm_Total_WMH_T1_T2"])

    present_adj=[v for v in BASE_ADJ if v in df.columns]
    categorical_in=set([v for v in present_adj if v in CATEGORICAL])

    # impute covariates
    dmod=df.copy()
    for v in present_adj:
        if pd.api.types.is_numeric_dtype(dmod[v]):
            dmod[v]=pd.to_numeric(dmod[v], errors="coerce").fillna(dmod[v].median())
        else:
            mode_vals=dmod[v].mode(dropna=True)
            dmod[v]=dmod[v].fillna(mode_vals.iloc[0] if not mode_vals.empty else dmod[v])
    dmod=dmod.dropna(subset=["group"])

    outcomes=[
        "Log_HSNorm_Total_WMH_T1_T2",
        "Log_HSNorm_PeriVentricular_WMH",
        "Log_HSNorm_Deep_WMH"
    ]

    # -------- Adjusted models --------
    rows, meta = [], []
    for oc in outcomes:
        if oc not in dmod.columns: print(f"Missing {oc}; skip."); continue
        fml = build_formula(oc, present_adj, categorical_in)
        model, seinfo, y_used, X_used, n_clu, n_obs = fit_with_cluster(fml, dmod, CLUSTER_VAR)
        ols_plain = sm.OLS(y_used, X_used).fit()
        try: bp_p = float(het_breuschpagan(ols_plain.resid, ols_plain.model.exog)[1])
        except Exception: bp_p = np.nan
        used_df = dmod.loc[y_used.index]
        n_ctrl, n_study, n_tot = count_group_n(used_df)
        key = find_group_term(model.params.index, "group", "Study")
        if key is None: print("Group contrast not found; skip."); continue
        beta=float(model.params[key]); ci_l,ci_u=[float(x) for x in model.conf_int().loc[key]]
        pval=float(model.pvalues[key]); pct, pct_l, pct_u = 100*(np.exp(beta)-1), 100*(np.exp(ci_l)-1), 100*(np.exp(ci_u)-1)
        rows.append({"Analysis":label,"Outcome":SHORT[oc],"% Change":pct,"% CI Lower":pct_l,"% CI Upper":pct_u,
                     "Log β":beta,"Log CI Lower":ci_l,"Log CI Upper":ci_u,"p_value":pval,"SE_type":seinfo})
        meta.append({"Analysis":label,"Outcome":SHORT[oc],"Model formula":fml,"SE_type":seinfo,"BP_p":bp_p,
                     "N_total(used)":n_tot,"N_control(used)":n_ctrl,"N_study(used)":n_study,
                     "n_clusters":n_clu,"n_obs":n_obs,"Covariates_used":", ".join(present_adj),
                     "Back-transform":"%Δ = (e^β - 1) × 100%"})

    res=pd.DataFrame(rows); meta_df=pd.DataFrame(meta)

    # ---------- FDR + gatekeeping (adjusted) ----------
    if not res.empty:
        p_total = res.loc[res["Outcome"]=="Total WMH","p_value"]
        total_sig = (len(p_total)>0) and (float(p_total.iloc[0]) < PRIMARY_P_ALPHA)

        mask_sec = res["Outcome"].isin(["Periventricular WMH","Deep WMH"])
        if mask_sec.any():
            rej, qvals, _, _ = multipletests(res.loc[mask_sec,"p_value"].values, method="fdr_bh", alpha=0.05)
            res.loc[mask_sec,"q_value"]=qvals; res.loc[mask_sec,"sig_FDR(q≤0.05)"]=rej

        res.loc[res["Outcome"]=="Total WMH",["q_value","sig_FDR(q≤0.05)"]]=np.nan
        res["gatekeep_total_sig"]=total_sig
        res["secondary_interpretable"]=np.where(res["Outcome"].isin(["Periventricular WMH","Deep WMH"]), bool(total_sig), np.nan)

    # cache (adjusted)
    res_by_label[label]=res.copy(); meta_by_label[label]=meta_df.copy()

    if not res.empty:
        order_map={v:i for i,v in enumerate(PLOT_ORDER)}
        print(res.sort_values("Outcome", key=lambda s: s.map(order_map)).round(6).to_string(index=False))
        safe=re.sub(r"[^0-9a-zA-Z]+","_", label).strip("_")
        res.to_csv(os.path.join(MAIN_OUT_DIR, f"{safe}_LOG1P_results_gatekept.csv"), index=False, encoding="utf-8-sig")
        meta_df.to_csv(os.path.join(MAIN_OUT_DIR, f"{safe}_LOG1P_metadata.csv"), index=False, encoding="utf-8-sig")
        plot_single_forest(res, label, MAIN_OUT_DIR)

    # --- NEW: Unadjusted models (only C(group)) ---
    rows_u, meta_u = [], []
    for oc in outcomes:
        if oc not in dmod.columns: print(f"Missing {oc}; skip (unadjusted)."); continue
        fml_u = build_formula(oc, [], set())  # only C(group)
        model_u, seinfo_u, y_u, X_u, n_clu_u, n_obs_u = fit_with_cluster(fml_u, dmod, CLUSTER_VAR)
        ols_plain_u = sm.OLS(y_u, X_u).fit()
        try: bp_p_u = float(het_breuschpagan(ols_plain_u.resid, ols_plain_u.model.exog)[1])
        except Exception: bp_p_u = np.nan
        used_u = dmod.loc[y_u.index]
        n_ctrl_u, n_study_u, n_tot_u = count_group_n(used_u)
        key_u = find_group_term(model_u.params.index, "group", "Study")
        if key_u is None: print("Group contrast not found (unadjusted); skip."); continue
        beta_u=float(model_u.params[key_u]); ci_lu,ci_uu=[float(x) for x in model_u.conf_int().loc[key_u]]
        pval_u=float(model_u.pvalues[key_u]); pct_u, pct_lu, pct_uu = 100*(np.exp(beta_u)-1), 100*(np.exp(ci_lu)-1), 100*(np.exp(ci_uu)-1)
        rows_u.append({"Analysis":label,"Outcome":SHORT[oc],"% Change":pct_u,"% CI Lower":pct_lu,"% CI Upper":pct_uu,
                       "Log β":beta_u,"Log CI Lower":ci_lu,"Log CI Upper":ci_uu,"p_value":pval_u,"SE_type":seinfo_u})
        meta_u.append({"Analysis":label,"Outcome":SHORT[oc],"Model formula":fml_u,"SE_type":seinfo_u,"BP_p":bp_p_u,
                       "N_total(used)":n_tot_u,"N_control(used)":n_ctrl_u,"N_study(used)":n_study_u,
                       "n_clusters":n_clu_u,"n_obs":n_obs_u,"Covariates_used":"(none; unadjusted)",
                       "Back-transform":"%Δ = (e^β - 1) × 100%"} )

    res_u = pd.DataFrame(rows_u); meta_u_df = pd.DataFrame(meta_u)

    # FDR for unadjusted (secondary outcomes only); no gatekeeping flag needed for export
    if not res_u.empty:
        mask_sec_u = res_u["Outcome"].isin(["Periventricular WMH","Deep WMH"])
        if mask_sec_u.any():
            rej_u, qvals_u, _, _ = multipletests(res_u.loc[mask_sec_u,"p_value"].values, method="fdr_bh", alpha=0.05)
            res_u.loc[mask_sec_u,"q_value"]=qvals_u; res_u.loc[mask_sec_u,"sig_FDR(q≤0.05)"]=rej_u
        res_u.loc[res_u["Outcome"]=="Total WMH",["q_value","sig_FDR(q≤0.05)"]]=np.nan

        # cache + CSV
        res_unadj_by_label[label] = res_u.copy()
        meta_unadj_by_label[label] = meta_u_df.copy()
        safe = re.sub(r"[^0-9a-zA-Z]+","_", label).strip("_")
        res_u.to_csv(os.path.join(MAIN_OUT_DIR, f"{safe}_UNADJ_results.csv"), index=False, encoding="utf-8-sig")
        meta_u_df.to_csv(os.path.join(MAIN_OUT_DIR, f"{safe}_UNADJ_metadata.csv"), index=False, encoding="utf-8-sig")
        print(f"Unadjusted results saved for {label}.")

    # --- Table 1 ---
    try:
        tbl1 = make_table1(dmod, label, id_col="Participant_ID")
        safe=re.sub(r"[^0-9a-zA-Z]+","_", label).strip("_")
        csv_path=os.path.join(MAIN_OUT_DIR, f"{safe}_Table1_Baseline.csv")
        tbl1.to_csv(csv_path, index=False, encoding="utf-8-sig"); print(f"Table 1 CSV saved: {csv_path}")
        docx_path = export_table1_word(tbl1, label, MAIN_OUT_DIR, timestamp)
        if docx_path: print(f"Table 1 Word saved: {docx_path}")
        elif Document is None: print("Install 'python-docx' for Word export:  pip install python-docx")
        all_tbl1.append(tbl1)
    except Exception as e:
        print(f"Table 1 failed for {label}: {e}")

    # --- SMD for supplement ---
    try:
        smd_tbl = pd.DataFrame({
            "Characteristic":[
                "Age at Instance 2, y",
                "Female sex",
                "BMI at Instance 0, kg/m²",
                "Genetic ethnic grouping, White",
                "Townsend deprivation index",
                "Smoking, ever",
                "Alcohol intake frequency (ordinal)"
            ],
            "SMD":[
                _smd_cont(dmod.loc[dmod["group"]=="Study","Age_at_Instance_2"],
                          dmod.loc[dmod["group"]=="Control","Age_at_Instance_2"]) if "Age_at_Instance_2" in dmod.columns else np.nan,
                _smd_binary(_is_female_series(dmod.loc[dmod["group"]=="Study","Sex"]).astype(float),
                            _is_female_series(dmod.loc[dmod["group"]=="Control","Sex"]).astype(float)) if "Sex" in dmod.columns else np.nan,
                _smd_cont(dmod.loc[dmod["group"]=="Study","Body_mass_index_BMI_Instance_0"],
                          dmod.loc[dmod["group"]=="Control","Body_mass_index_BMI_Instance_0"]) if "Body_mass_index_BMI_Instance_0" in dmod.columns else np.nan,
                _smd_binary(_is_white_series(dmod.loc[dmod["group"]=="Study","Genetic_ethnic_grouping"]).astype(float),
                            _is_white_series(dmod.loc[dmod["group"]=="Control","Genetic_ethnic_grouping"]).astype(float)) if "Genetic_ethnic_grouping" in dmod.columns else np.nan,
                _smd_cont(dmod.loc[dmod["group"]=="Study","Townsend_deprivation_index_at_recruitment"],
                          dmod.loc[dmod["group"]=="Control","Townsend_deprivation_index_at_recruitment"]) if "Townsend_deprivation_index_at_recruitment" in dmod.columns else np.nan,
                _smd_binary(_is_true_series(dmod.loc[dmod["group"]=="Study","Smoking_Ever"]).astype(float),
                            _is_true_series(dmod.loc[dmod["group"]=="Control","Smoking_Ever"]).astype(float)) if "Smoking_Ever" in dmod.columns else np.nan,
                _smd_cont(_as_num(dmod.loc[dmod["group"]=="Study","Alcohol_intake_frequency_ordinal"]),
                          _as_num(dmod.loc[dmod["group"]=="Control","Alcohol_intake_frequency_ordinal"])) if "Alcohol_intake_frequency_ordinal" in dmod.columns else np.nan
            ]
        })
        smd_tbl["|SMD|"]=smd_tbl["SMD"].abs()
        safe=re.sub(r"[^0-9a-zA-Z]+","_", label).strip("_")
        smd_csv=os.path.join(MAIN_OUT_DIR, f"{safe}_SMD.csv")
        smd_tbl.to_csv(smd_csv, index=False, encoding="utf-8-sig"); print(f"SMD table saved: {smd_csv}")
        smd_by_label[label]=smd_tbl.rename(columns={"SMD":"SMD","|SMD|":"|SMD|"})
    except Exception as e:
        print(f"SMD computation failed for {label}: {e}")

    all_log.append(res); all_meta.append(meta_df)

# Overlay forest figure
LEFT_KEY="Primary 1:10 (NoNeuro)"; RIGHT_KEY="Sensitivity 1:10 (WithNeuro)"
if LEFT_KEY in res_by_label and RIGHT_KEY in res_by_label and \
   (not res_by_label[LEFT_KEY].empty) and (not res_by_label[RIGHT_KEY].empty):
    plot_overlay_forest(
        res_primary=res_by_label[LEFT_KEY],
        res_sensitivity=res_by_label[RIGHT_KEY],
        outdir=MAIN_OUT_DIR,
        basename=f"Overlay_Forest_Primary_vs_Sensitivity"
    )
    print("Overlay figure saved (Primary vs Sensitivity).")
else:
    print("Overlay figure not created: missing one of the cohorts.")

# SMD supplement (combined table + Love plot)
def combine_smd_tables(tbl_primary: pd.DataFrame, tbl_sens: pd.DataFrame) -> pd.DataFrame:
    p = tbl_primary.set_index("Characteristic")[["SMD","|SMD|"]].rename(
        columns={"SMD":"SMD (Primary)","|SMD|":"|SMD| (Primary)"})
    s = tbl_sens.set_index("Characteristic")[["SMD","|SMD|"]].rename(
        columns={"SMD":"SMD (Sensitivity)","|SMD|":"|SMD| (Sensitivity)"})
    combo = p.join(s, how="outer").reset_index()
    combo["Balanced Primary (<0.1)"] = combo["|SMD| (Primary)"] < 0.1
    combo["Balanced Sensitivity (<0.1)"] = combo["|SMD| (Sensitivity)"] < 0.1
    order = ["Age at Instance 2, y","Female sex","BMI at Instance 0, kg/m²",
             "Genetic ethnic grouping, White","Townsend deprivation index",
             "Smoking, ever","Alcohol intake frequency (ordinal)"]
    combo = combo.sort_values("Characteristic", key=lambda s: s.map({v:i for i,v in enumerate(order)}))
    return combo

def plot_love_overlay(tbl_primary: pd.DataFrame, tbl_sens: pd.DataFrame, outdir: str, basename: str):
    order=["Age at Instance 2, y","Female sex","BMI at Instance 0, kg/m²",
           "Genetic ethnic grouping, White","Townsend deprivation index",
           "Smoking, ever","Alcohol intake frequency (ordinal)"]
    p=tbl_primary.set_index("Characteristic").reindex(order)
    s=tbl_sens.set_index("Characteristic").reindex(order)
    y=np.arange(len(order))[::-1]; jitter=0.10; pad_y=0.45
    fig=plt.figure(figsize=(6.8,3.6)); ax=plt.gca()
    ax.errorbar(p["|SMD|"], y+jitter, fmt="o", ms=5.5, mfc="#1f3b4d", mec="#1f3b4d",
                linestyle="none", color="#1f3b4d", label="Primary")
    ax.errorbar(s["|SMD|"], y-jitter, fmt="s", ms=5.2, mfc="white", mec="#6e6e6e",
                linestyle="none", color="#6e6e6e", label="Sensitivity")
    ax.axvline(0.1, color="grey", linestyle="--", linewidth=1)
    ax.axvline(0.2, color="grey", linestyle=":",  linewidth=1)
    ax.set_yticks(y, order); ax.set_xlabel("|SMD|")
    xmax=np.nanmax([p["|SMD|"], s["|SMD|"]]); ax.set_xlim(left=0, right=float(xmax)*1.2+0.02)
    ax.set_ylim(y.min()-pad_y, y.max()+pad_y); ax.grid(axis="x", linestyle=":", linewidth=0.7, alpha=0.8)
    leg=ax.legend(frameon=False, loc="upper right", handlelength=1.6)
    for attr in ("legendHandles","legend_handles"):
        if hasattr(leg,attr):
            try:
                for lh in getattr(leg,attr):
                    if hasattr(lh,"set_linewidth"): lh.set_linewidth(1.6)
            except Exception: pass; break
    plt.tight_layout(); safe=re.sub(r"[^0-9a-zA-Z]+","_", basename).strip("_"); base=os.path.join(outdir, safe)
    plt.savefig(base+".png", dpi=600); plt.savefig(base+".pdf"); plt.savefig(base+".svg"); plt.close(fig)

if LEFT_KEY in smd_by_label and RIGHT_KEY in smd_by_label:
    smd_combo = combine_smd_tables(smd_by_label[LEFT_KEY], smd_by_label[RIGHT_KEY])
    combo_csv = os.path.join(MAIN_OUT_DIR, f"Supplement_SMD_Balance.csv")
    smd_combo.to_csv(combo_csv, index=False, encoding="utf-8-sig"); print(f"Supplementary SMD CSV saved: {combo_csv}")
    if Document is not None:
        path_docx = export_smd_word(smd_combo, MAIN_OUT_DIR, timestamp)
        print(f"Supplementary SMD Word saved: {path_docx}")
    else:
        print("Install 'python-docx' for SMD Word:  pip install python-docx")
    plot_love_overlay(smd_by_label[LEFT_KEY], smd_by_label[RIGHT_KEY],
                      outdir=MAIN_OUT_DIR,
                      basename=f"Love_Plot_Primary_vs_Sensitivity")
    print("Love plot (overlay) saved.")
else:
    print("SMD supplement not created: missing SMD tables.")

# Figure legends
figure_legends = f"""Figure X. Adjusted differences in WMH (overlay forest plot)
Primary (deep blue, filled circles) and Sensitivity (dark gray, open squares) cohorts are overlaid in a single panel without connecting lines. Points show adjusted between-group differences in log1p-transformed WMH volumes expressed as percent change (Study vs Control); error bars indicate 95% CIs. Models adjusted for age at Instance 2, sex, BMI at Instance 0, smoking (ever), alcohol-intake frequency (ordinal), Townsend deprivation index, and genetic ethnic grouping; cluster-robust standard errors by matched set (match_id). Total WMH is the primary endpoint; Periventricular and Deep WMH are secondary endpoints with BH-FDR applied within the secondary family. The Sensitivity cohort repeats matching without pre-index or index-date exclusions.

Figure Y-1. Primary cohort forest plot
Single-cohort version of Figure X for the primary cohort (excluding pre-index neurological/cerebrovascular diagnoses; index-date assignment and risk-set pruning applied in controls).

Figure Y-2. Sensitivity cohort forest plot
Single-cohort version of Figure X for the sensitivity cohort (no pre-index or index-date exclusions)."""
leg_txt=os.path.join(MAIN_OUT_DIR, f"Figure_legends.txt")
with open(leg_txt,"w",encoding="utf-8") as f: f.write(figure_legends)
print(f"Figure legends (TXT) saved: {leg_txt}")
if Document is not None:
    doc=Document(); p=doc.add_paragraph(); r=p.add_run("Figure Legends"); r.bold=True
    for block in figure_legends.strip().split("\n\n"): doc.add_paragraph(block)
    leg_docx=os.path.join(MAIN_OUT_DIR, f"Figure_legends.docx")
    doc.save(leg_docx); print(f"Figure legends (Word) saved: {leg_docx}")
else:
    print("Install 'python-docx' for Word legends:  pip install python-docx")

# OLS manuscript table (Primary & Sensitivity together; with q) -- ADJUSTED
if LEFT_KEY in res_by_label and RIGHT_KEY in res_by_label and \
   (not res_by_label[LEFT_KEY].empty) and (not res_by_label[RIGHT_KEY].empty):
    ols_tbl = make_ols_results_table(
        res_by_label[LEFT_KEY], res_by_label[RIGHT_KEY],
        meta_by_label.get(LEFT_KEY), meta_by_label.get(RIGHT_KEY)
    )
    csv_ols = os.path.join(MAIN_OUT_DIR, f"OLS_Results_Manuscript_Table.csv")
    ols_tbl.to_csv(csv_ols, index=False, encoding="utf-8-sig"); print(f"OLS results table CSV saved: {csv_ols}")
    if Document is not None:
        docx_ols = export_ols_table_word(ols_tbl, MAIN_OUT_DIR, timestamp)
        print(f"OLS results table Word saved: {docx_ols}")
    else:
        print("Install 'python-docx' for OLS Word export:  pip install python-docx")
else:
    print("OLS results table not created: missing one of the cohorts.")

# NEW: Unadjusted manuscript-like table (Primary & Sensitivity together)
if LEFT_KEY in res_unadj_by_label and RIGHT_KEY in res_unadj_by_label and \
   (not res_unadj_by_label[LEFT_KEY].empty) and (not res_unadj_by_label[RIGHT_KEY].empty):
    unadj_tbl = make_ols_results_table(
        res_unadj_by_label[LEFT_KEY], res_unadj_by_label[RIGHT_KEY],
        meta_unadj_by_label.get(LEFT_KEY), meta_unadj_by_label.get(RIGHT_KEY)
    )
    csv_unadj = os.path.join(MAIN_OUT_DIR, f"UNADJ_Results_Manuscript_Table.csv")
    unadj_tbl.to_csv(csv_unadj, index=False, encoding="utf-8-sig")
    print(f"Unadjusted results table CSV saved: {csv_unadj}")
    if Document is not None:
        docx_unadj = export_unadjusted_table_word(unadj_tbl, MAIN_OUT_DIR, timestamp)
        print(f"Unadjusted results table Word saved: {docx_unadj}")
    else:
        print("Install 'python-docx' for unadjusted Word export:  pip install python-docx")
else:
    print("Unadjusted results table not created: missing one of the cohorts or results.")

# combine CSV exports for record
final_log  = pd.concat(all_log,  ignore_index=True) if all_log  else pd.DataFrame()
final_meta = pd.concat(all_meta, ignore_index=True) if all_meta else pd.DataFrame()
if not final_log.empty:
    final_log.to_csv(os.path.join(MAIN_OUT_DIR, f"Main_all_LOG1P_results_gatekept.csv"),
                     index=False, encoding="utf-8-sig")
if not final_meta.empty:
    final_meta.to_csv(os.path.join(MAIN_OUT_DIR, f"Main_all_LOG1P_metadata.csv"),
                      index=False, encoding="utf-8-sig")
if all_tbl1:
    all_tbl1_df = pd.concat(all_tbl1, ignore_index=True)
    all_tbl1_df.to_csv(os.path.join(MAIN_OUT_DIR, f"All_Table1_Baseline.csv"),
                       index=False, encoding="utf-8-sig")

    # Export combined Word Table 1 (Primary + Sensitivity with block format)
    if Document is not None:
        doc = Document()
        r = doc.add_paragraph().add_run("Table 1. Baseline characteristics — Primary PSM and PSM–Sensitivity cohort")
        r.bold = True
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        cols = ["Characteristic", "SA", "Control"]
        if TABLE1_USE_PVALUE: cols.append("p_value")

        t = doc.add_table(rows=1, cols=len(cols))
        hdr = t.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = c

        for cohort in ["Primary 1:10 (NoNeuro)", "Sensitivity 1:10 (WithNeuro)"]:
            sub = all_tbl1_df[all_tbl1_df["Analysis"] == cohort]

            row = t.add_row().cells
            row[0].text = "Primary PSM cohort" if "Primary" in cohort else "PSM–sensitivity cohort"
            for j in range(1, len(cols)):
                row[j].text = ""

            for _, r0 in sub.iterrows():
                cells = t.add_row().cells
                cells[0].text = str(r0["Characteristic"])
                cells[1].text = str(r0["Study"])
                cells[2].text = str(r0["Control"])
                if TABLE1_USE_PVALUE:
                    val = r0.get("p_value", "")
                    if isinstance(val, float) and not np.isnan(val):
                        cells[3].text = "<0.001" if val < 0.001 else f"{val:.3f}"
                    else:
                        cells[3].text = str(val) if val not in (None, np.nan) else ""

        num_idx = [1, 2] + ([3] if TABLE1_USE_PVALUE else [])
        _right_align_numeric(t, num_idx)

        if THREELINE_TABLES:
            apply_three_line_table(t, header_row_idx=0)
        set_table_font(t, font_name="Times New Roman", font_size=10)

        legend = doc.add_paragraph()
        legend.add_run(
            "Legend: Values are mean (SD), median [IQR], or n (%). "
            "p-values from Welch’s t test (continuous), Mann–Whitney U (alcohol intake), "
            "and χ² or Fisher’s exact test (categorical). "
            "SA = sleep apnea; Control = matched controls."
        ).italic = True

        doc.save(os.path.join(MAIN_OUT_DIR, f"Table1_Baseline_Primary_Sensitivity.docx"))
        print("Combined Table 1 Word (block format) saved.")

print(f'\nAll outputs saved in "{MAIN_OUT_DIR}/" (timestamp={timestamp}).')
print("Tables exported as three-line tables in Word. Figures saved as PNG/PDF/SVG.")


# ================== CMC-adjusted sensitivity analysis (+CMC adjustment explicitly labeled) ==================
# Goal:
#   Conduct a sensitivity analysis in both cohorts by adding CMC (comorbidity count)
#   as an additional covariate to the main fully adjusted models.
#   This block generates eFigures (with “+CMC adjustment” label) and eTables (CSV + Word),
#   in the same format as the main analysis.

SUPP_DIR = os.path.join(MAIN_OUT_DIR, "eSupp_CMC")
os.makedirs(SUPP_DIR, exist_ok=True)

res_cmc_by_label, meta_cmc_by_label = {}, {}

def _choose_cmc_var(df):
    """Choose which CMC variable to use (prefer categorical)."""
    if "CMC_score_cat" in df.columns:
        return "CMC_score_cat", True
    if "CMC_score_raw" in df.columns:
        return "CMC_score_raw", False
    return None, None

def _attach_cmc_and_impute(dmod0: pd.DataFrame, base_adj: list, categorical_set: set):
    """
    Add the chosen CMC variable to the adjustment set and perform imputation:
      - numeric → median
      - categorical → mode
    """
    dmod = dmod0.copy()
    cmc_var, is_cat = _choose_cmc_var(dmod)
    if cmc_var is None:
        return dmod, base_adj, categorical_set, None

    adj_vars = base_adj + [cmc_var]
    cat_vars = set(categorical_set)
    if is_cat:
        cat_vars.add(cmc_var)

    for v in adj_vars:
        if pd.api.types.is_numeric_dtype(dmod[v]):
            dmod[v] = pd.to_numeric(dmod[v], errors="coerce").fillna(dmod[v].median())
        else:
            mode_vals = dmod[v].mode(dropna=True)
            dmod[v] = dmod[v].fillna(mode_vals.iloc[0] if not mode_vals.empty else dmod[v])

    return dmod, adj_vars, cat_vars, cmc_var

# Iterate through both cohorts
for label, file in file_sets.items():
    print(f"\n=== {label} (CMC-adjusted sensitivity analysis) ===")
    if label not in res_by_label:
        print("Skipping CMC sensitivity: main analysis results not found for this cohort.")
        continue

    df0 = pd.read_csv(file)
    df0.columns = df0.columns.str.replace(r"[^0-9a-zA-Z]+", "_", regex=True)
    if "group" not in df0.columns or df0["group"].dropna().nunique() < 2:
        print("Skipping: invalid or missing 'group' variable.")
        continue
    df0["group"] = pd.Categorical(df0["group"], categories=GROUP_BASELINE, ordered=False)

    # Ensure log-transformed outcomes exist
    need_build = not all(c in df0.columns for c in [
        "Log_HSNorm_Total_WMH_T1_T2",
        "Log_HSNorm_PeriVentricular_WMH",
        "Log_HSNorm_Deep_WMH"
    ])
    if need_build:
        sc = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"
        if sc not in df0.columns:
            print("Skipping: missing head-size scaling factor.")
            continue
        scale = pd.to_numeric(df0[sc], errors="coerce")
        df0["HSNorm_Deep_WMH"] = pd.to_numeric(df0["Total_volume_of_deep_white_matter_hyperintensities_Instance_2"], errors="coerce") * scale
        df0["HSNorm_PeriVentricular_WMH"] = pd.to_numeric(df0["Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2"], errors="coerce") * scale
        df0["HSNorm_Total_WMH_T1_T2"] = pd.to_numeric(df0["Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2"], errors="coerce") * scale
        df0["Log_HSNorm_Deep_WMH"] = np.log1p(df0["HSNorm_Deep_WMH"])
        df0["Log_HSNorm_PeriVentricular_WMH"] = np.log1p(df0["HSNorm_PeriVentricular_WMH"])
        df0["Log_HSNorm_Total_WMH_T1_T2"] = np.log1p(df0["HSNorm_Total_WMH_T1_T2"])

    # Prepare covariates
    present_adj = [v for v in BASE_ADJ if v in df0.columns]
    categorical_in = set([v for v in present_adj if v in CATEGORICAL])

    # Impute base covariates first
    dmod_base = df0.copy()
    for v in present_adj:
        if pd.api.types.is_numeric_dtype(dmod_base[v]):
            dmod_base[v] = pd.to_numeric(dmod_base[v], errors="coerce").fillna(dmod_base[v].median())
        else:
            mode_vals = dmod_base[v].mode(dropna=True)
            dmod_base[v] = dmod_base[v].fillna(mode_vals.iloc[0] if not mode_vals.empty else dmod_base[v])

    # Add and impute CMC
    dmod, adj_vars, cat_vars, cmc_used = _attach_cmc_and_impute(dmod_base, present_adj, categorical_in)
    if cmc_used is None:
        print("Skipping: no CMC variable found.")
        continue
    dmod = dmod.dropna(subset=["group"])

    outcomes = [
        "Log_HSNorm_Total_WMH_T1_T2",
        "Log_HSNorm_PeriVentricular_WMH",
        "Log_HSNorm_Deep_WMH"
    ]

    rows, meta = [], []
    for oc in outcomes:
        if oc not in dmod.columns:
            print(f"Outcome {oc} missing; skipped.")
            continue

        fml = build_formula(oc, adj_vars, cat_vars)
        model, seinfo, y_used, X_used, n_clu, n_obs = fit_with_cluster(fml, dmod, CLUSTER_VAR)

        ols_plain = sm.OLS(y_used, X_used).fit()
        try:
            bp_p = float(het_breuschpagan(ols_plain.resid, ols_plain.model.exog)[1])
        except Exception:
            bp_p = np.nan

        used_df = dmod.loc[y_used.index]
        n_ctrl, n_study, n_tot = count_group_n(used_df)

        key = find_group_term(model.params.index, "group", "Study")
        if key is None:
            print("Group contrast not found; skipped.")
            continue

        beta = float(model.params[key])
        ci_l, ci_u = [float(x) for x in model.conf_int().loc[key]]
        pval = float(model.pvalues[key])
        pct, pct_l, pct_u = 100*(np.exp(beta)-1), 100*(np.exp(ci_l)-1), 100*(np.exp(ci_u)-1)

        rows.append({
            "Analysis": label,
            "Outcome": SHORT[oc],
            "% Change": pct, "% CI Lower": pct_l, "% CI Upper": pct_u,
            "Log β": beta, "Log CI Lower": ci_l, "Log CI Upper": ci_u,
            "p_value": pval, "SE_type": seinfo
        })
        meta.append({
            "Analysis": label, "Outcome": SHORT[oc],
            "Model formula": fml, "SE_type": seinfo, "BP_p": bp_p,
            "N_total(used)": n_tot, "N_control(used)": n_ctrl, "N_study(used)": n_study,
            "n_clusters": n_clu, "n_obs": n_obs,
            "Covariates_used": ", ".join(adj_vars),
            "CMC_var": cmc_used,
            "Back-transform": "%Δ = (e^β − 1) × 100%"
        })

    res_cmc = pd.DataFrame(rows)
    meta_cmc = pd.DataFrame(meta)

    # FDR + gatekeeping (same as main analysis)
    if not res_cmc.empty:
        p_total = res_cmc.loc[res_cmc["Outcome"] == "Total WMH", "p_value"]
        total_sig = (len(p_total) > 0) and (float(p_total.iloc[0]) < PRIMARY_P_ALPHA)

        mask_sec = res_cmc["Outcome"].isin(["Periventricular WMH", "Deep WMH"])
        if mask_sec.any():
            rej, qvals, _, _ = multipletests(res_cmc.loc[mask_sec, "p_value"].values,
                                             method="fdr_bh", alpha=0.05)
            res_cmc.loc[mask_sec, "q_value"] = qvals
            res_cmc.loc[mask_sec, "sig_FDR(q≤0.05)"] = rej

        res_cmc.loc[res_cmc["Outcome"] == "Total WMH", ["q_value", "sig_FDR(q≤0.05)"]] = np.nan
        res_cmc["gatekeep_total_sig"] = total_sig
        res_cmc["secondary_interpretable"] = np.where(
            res_cmc["Outcome"].isin(["Periventricular WMH", "Deep WMH"]),
            bool(total_sig), np.nan
        )

    res_cmc_by_label[label] = res_cmc.copy()
    meta_cmc_by_label[label] = meta_cmc.copy()
    safe = re.sub(r"[^0-9a-zA-Z]+", "_", label).strip("_")

    res_cmc.to_csv(os.path.join(SUPP_DIR, f"{safe}_CMCadj_results.csv"), index=False, encoding="utf-8-sig")
    meta_cmc.to_csv(os.path.join(SUPP_DIR, f"{safe}_CMCadj_metadata.csv"), index=False, encoding="utf-8-sig")

    # Single-cohort eFigure (explicitly labeled as “+CMC adjustment”)
    if not res_cmc.empty:
        plot_single_forest(res_cmc, f"eFigure (+CMC adjustment) — {label}", SUPP_DIR)

# Overlay eFigure (Primary vs Sensitivity)
if (LEFT_KEY in res_cmc_by_label and RIGHT_KEY in res_cmc_by_label and
    not res_cmc_by_label[LEFT_KEY].empty and not res_cmc_by_label[RIGHT_KEY].empty):
    plot_overlay_forest(
        res_primary=res_cmc_by_label[LEFT_KEY],
        res_sensitivity=res_cmc_by_label[RIGHT_KEY],
        outdir=SUPP_DIR,
        basename="eFigure_CMCadj_Overlay_Primary_vs_Sensitivity (+CMC adjustment)"
    )
    print("Overlay eFigure (+CMC adjustment) saved.")
else:
    print("Overlay eFigure (+CMC adjustment) not created: one or both cohorts missing.")

# eTable: merged results (both cohorts)
if (LEFT_KEY in res_cmc_by_label and RIGHT_KEY in res_cmc_by_label and
    not res_cmc_by_label[LEFT_KEY].empty and not res_cmc_by_label[RIGHT_KEY].empty):
    etbl = make_ols_results_table(
        res_cmc_by_label[LEFT_KEY], res_cmc_by_label[RIGHT_KEY],
        meta_cmc_by_label.get(LEFT_KEY), meta_cmc_by_label.get(RIGHT_KEY)
    )
    etbl_csv = os.path.join(SUPP_DIR, "eTable_CMCadj_Results.csv")
    etbl.to_csv(etbl_csv, index=False, encoding="utf-8-sig")
    print(f"eTable (+CMC adjustment) CSV saved: {etbl_csv}")

    if Document is not None:
        path_docx = export_ols_table_word(etbl, SUPP_DIR, timestamp)
        if path_docx:
            new_path = os.path.join(SUPP_DIR, "eTable_CMCadj_Results.docx")
            try:
                os.replace(path_docx, new_path)
            except Exception:
                pass
            print(f"eTable (+CMC adjustment) Word saved: {new_path}")
    else:
        print("Install 'python-docx' to enable Word export.")
else:
    print("eTable (+CMC adjustment) not created: one or both cohorts missing.")
# ================== END CMC-adjusted sensitivity analysis ==================
