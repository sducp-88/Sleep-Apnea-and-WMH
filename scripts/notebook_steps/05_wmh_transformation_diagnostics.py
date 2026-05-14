# Extracted from WMH_Analysis.ipynb, code cell 4.
# Notebook heading: # WMH Transformation Diagnostics 
# Run this file from the repository root unless a local CONFIG section is edited.

# WMH Transformation Diagnostics 
"""
WMH Transformation Diagnostics (publication-ready, robust to missing statsmodels)
================================================================================
This script evaluates whether head-size–normalized WMH outcomes should be modeled
on the raw scale or on the log1p scale. It:

  • Loads two matched cohorts:
      - primary_cohort.csv
      - sensitivity_cohort.csv
  • Computes head-size–normalized outcomes: raw × volumetric scaling factor (T1)
  • Fits OLS models (raw vs log1p) with prespecified covariates
  • Runs diagnostics: Breusch–Pagan (heteroscedasticity), Jarque–Bera (normality), AIC
  • Exports histograms, QQ plots, residual–vs–fitted plots (both PNG and PDF)
  • Writes a publication-grade summary CSV and a Word eMethod

Robustness:
  - If statsmodels imports successfully, it is used.
  - If not, the script falls back to pure NumPy OLS and approximate p-values
    (Wilson–Hilferty for χ²), and implements QQ plots without SciPy.

Outputs
-------
  • Diagnostics/wmh_transform_diagnostics_summary.csv
  • Diagnostics/figures/*_hist.(png|pdf), *_qq_raw.(png|pdf), *_qq_log.(png|pdf), *_rvf.(png|pdf)
  • Diagnostics/eMethod_WMH_Transform_Diagnostics.docx

Dependencies: pandas, numpy, matplotlib, patsy, python-docx
(Optional)   : statsmodels  (if available, used for exact tests/QQ plots)
"""

import re
from pathlib import Path
import math

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# Try statsmodels; if not available or broken, use fallback
HAVE_SM = True
try:
    import statsmodels.api as sm
    from statsmodels.stats.diagnostic import het_breuschpagan
    from statsmodels.stats.stattools import jarque_bera
except Exception:
    HAVE_SM = False

from patsy import dmatrices

# Word export
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =========================
# File sets (two cohorts)
# =========================
FILE_SETS = {
    "Primary": "primary_cohort.csv",
    "Sensitivity": "sensitivity_cohort.csv",
}

# =========================
# Output folders
# =========================
DIAG_DIR = Path("Diagnostics")
FIG_DIR = DIAG_DIR / "figures"
DIAG_DIR.mkdir(exist_ok=True)
FIG_DIR.mkdir(parents=True, exist_ok=True)

# =========================
# Column names (robust resolution; we DO NOT mutate your column names)
# =========================
SCALE_FIELD_EXPECTED = "Volumetric_scaling_from_T1_head_image_to_standard_space_Instance_2"

OUTCOME_DEFS = {
    # new_col : (raw_WMH_col_in_input, pretty_label)
    "HSNorm_Deep_WMH": (
        "Total_volume_of_deep_white_matter_hyperintensities_Instance_2",
        "Head-size Normalized Deep WMH"
    ),
    "HSNorm_PeriVentricular_WMH": (
        "Total_volume_of_peri_ventricular_white_matter_hyperintensities_Instance_2",
        "Head-size Normalized Periventricular WMH"
    ),
    "HSNorm_Total_WMH_T1_T2": (
        "Total_volume_of_white_matter_hyperintensities_from_T1_and_T2_FLAIR_images_Instance_2",
        "Head-size Normalized Total WMH"
    ),
}

# Prespecified covariates (only those present will be used)
BASE_ADJUST_VARS = [
    "Sex",
    "Age_at_Instance_2",
    "Townsend_deprivation_index_at_recruitment",
    "Body_mass_index_BMI_Instance_0",
    "Genetic_ethnic_grouping",
    "Smoking_Ever",
    "Alcohol_intake_frequency_ordinal",
]

CATEGORICAL_VARS_ALL = {"Sex", "Genetic_ethnic_grouping", "Smoking_Ever", "group"}
GROUP_COL = "group"  # expected: "Control"/"Study"


# =========================
# Helpers: robust column resolution
# =========================
def normalize_name(name: str) -> str:
    """Lowercase; replace non-alnum with '_'; strip underscores."""
    return re.sub(r"[^0-9a-zA-Z]+", "_", name).strip("_").lower()

def resolve_column(df: pd.DataFrame, desired: str) -> str:
    """
    Return the actual column name in df corresponding to `desired`,
    allowing for case/spacing/underscore differences. Raises KeyError if not found.
    """
    if desired in df.columns:
        return desired
    norm_map = {normalize_name(c): c for c in df.columns}
    key = normalize_name(desired)
    if key in norm_map:
        return norm_map[key]
    # Relaxed endswith match
    candidates = [c for k, c in norm_map.items() if k.endswith(key)]
    if len(candidates) == 1:
        return candidates[0]
    raise KeyError(f"Column not found (wanted='{desired}'). Sample columns: {list(df.columns)[:10]} ...")


# =========================
# Modeling utilities (statsmodels OR NumPy fallback)
# =========================
def build_formula(outcome: str, present_adjust_vars: list, categorical_vars: set, group_var: str = GROUP_COL) -> str:
    """OLS formula: outcome ~ C(group) + covariates (categoricals wrapped in C())."""
    terms = [f"C({group_var})"]
    for v in present_adjust_vars:
        if v == group_var:
            continue
        terms.append(f"C({v})" if v in categorical_vars else v)
    return f"{outcome} ~ " + " + ".join(terms)

def fit_ols_sm(formula: str, data: pd.DataFrame):
    """OLS via statsmodels; returns (model, y_df, X_df)."""
    y, X = dmatrices(formula, data, return_type='dataframe', NA_action='drop')
    model = sm.OLS(y, X).fit()
    return model, y, X

# --- Fallback math helpers (no SciPy) ---
def _norm_ppf(p: np.ndarray) -> np.ndarray:
    """
    Acklam's approximation for the inverse standard normal CDF (ppf).
    Valid for p in (0,1). Vectorized.
    """
    p = np.asarray(p, dtype=float)
    # Coefficients
    a = [-3.969683028665376e+01,  2.209460984245205e+02,
         -2.759285104469687e+02,  1.383577518672690e+02,
         -3.066479806614716e+01,  2.506628277459239e+00]
    b = [-5.447609879822406e+01,  1.615858368580409e+02,
         -1.556989798598866e+02,  6.680131188771972e+01,
         -1.328068155288572e+01]
    c = [-7.784894002430293e-03, -3.223964580411365e-01,
         -2.400758277161838e+00, -2.549732539343734e+00,
          4.374664141464968e+00,  2.938163982698783e+00]
    d = [7.784695709041462e-03,  3.224671290700398e-01,
         2.445134137142996e+00,  3.754408661907416e+00]

    plow = 0.02425
    phigh = 1 - plow
    q = np.zeros_like(p)
    # lower region
    mask = p < plow
    if mask.any():
        pp = p[mask]
        ql = np.sqrt(-2*np.log(pp))
        q[mask] = (((((c[0]*ql + c[1])*ql + c[2])*ql + c[3])*ql + c[4])*ql + c[5]) / \
                   ((((d[0]*ql + d[1])*ql + d[2])*ql + d[3])*ql + 1)
    # central
    mask = (p >= plow) & (p <= phigh)
    if mask.any():
        pp = p[mask] - 0.5
        r = pp*pp
        q[mask] = (((((a[0]*r + a[1])*r + a[2])*r + a[3])*r + a[4])*r + a[5])*pp / \
                   (((((b[0]*r + b[1])*r + b[2])*r + b[3])*r + b[4])*r + 1)
    # upper
    mask = p > phigh
    if mask.any():
        pp = 1 - p[mask]
        ql = np.sqrt(-2*np.log(pp))
        q[mask] = -(((((c[0]*ql + c[1])*ql + c[2])*ql + c[3])*ql + c[4])*ql + c[5]) / \
                    ((((d[0]*ql + d[1])*ql + d[2])*ql + d[3])*ql + 1)
    return q

def _chi2_sf_wh(x: float, df: int) -> float:
    """
    Wilson–Hilferty approximation of χ² survival function P[Chi2_df >= x].
    Accurate enough for reporting in fallback mode.
    """
    if x < 0 or df <= 0:
        return float("nan")
    z = ((x/df)**(1/3) - (1 - 2/(9*df))) / math.sqrt(2/(9*df))
    # standard normal SF
    return 0.5 * math.erfc(z / math.sqrt(2))

def fit_ols_np(formula: str, data: pd.DataFrame):
    """
    OLS via NumPy with patsy matrices; returns a lightweight model-like dict.
    Provides fields: resid, fittedvalues, params, aic, model_exog, nobs, k_params
    """
    y, X = dmatrices(formula, data, return_type='dataframe', NA_action='drop')
    yv = y.values.ravel()
    Xv = X.values
    # Solve (X'X)β = X'y
    XtX = Xv.T @ Xv
    Xty = Xv.T @ yv
    try:
        beta = np.linalg.solve(XtX, Xty)
    except np.linalg.LinAlgError:
        beta = np.linalg.pinv(XtX) @ Xty
    fitted = Xv @ beta
    resid = yv - fitted
    n = yv.size
    k = Xv.shape[1]
    rss = float(np.dot(resid, resid))
    # Gaussian AIC up to an additive constant: n*ln(RSS/n) + 2k
    aic = n * math.log(rss / max(n, 1)) + 2 * k

    # Store what we need
    model = {
        "resid": resid,
        "fittedvalues": fitted,
        "params": beta,
        "aic": aic,
        "model_exog": Xv,  # for BP
        "nobs": n,
        "k_params": k,
        "y_df": y,
        "X_df": X,
    }
    return model

def bp_lm_np(model) -> float:
    """
    Breusch–Pagan LM statistic using auxiliary regression of squared residuals on exog.
    Returns LM = n * R^2. DoF approx = k - 1 (excluding intercept).
    """
    u2 = model["resid"]**2
    Z = model["model_exog"]  # includes intercept
    # OLS of u2 ~ Z
    beta = np.linalg.pinv(Z) @ u2
    u2_hat = Z @ beta
    # R^2
    ss_tot = float(((u2 - u2.mean())**2).sum())
    ss_res = float(((u2 - u2_hat)**2).sum())
    r2 = 1 - ss_res / ss_tot if ss_tot > 0 else 0.0
    LM = model["nobs"] * r2
    dof = max(model["k_params"] - 1, 1)
    p = _chi2_sf_wh(LM, dof)
    return LM, p

def jb_np(model) -> tuple[float, float, float, float]:
    """
    Jarque–Bera statistic and approximated p-value (χ²_2), along with skew/kurtosis.
    """
    r = model["resid"]
    n = r.size
    if n == 0:
        return float("nan"), float("nan"), float("nan"), float("nan")
    m = r.mean()
    s2 = np.mean((r - m)**2)
    if s2 <= 0:
        return float("nan"), float("nan"), float("nan"), float("nan")
    s = np.mean(((r - m)/math.sqrt(s2))**3)
    k = np.mean(((r - m)/math.sqrt(s2))**4)
    JB = n/6 * (s**2 + (k - 3)**2 / 4)
    p = _chi2_sf_wh(JB, 2)
    return JB, p, s, k

def qqplot_fallback(resid: np.ndarray, title: str, out_base: Path):
    """
    Normal QQ plot without SciPy/statsmodels:
      - theoretical quantiles via Acklam ppf
      - sample residuals sorted
    """
    r = np.asarray(resid, dtype=float)
    r = r[np.isfinite(r)]
    if r.size < 3:
        return
    r_sorted = np.sort(r)
    n = r_sorted.size
    probs = (np.arange(1, n+1) - 0.5) / n
    theo = _norm_ppf(probs)

    fig = plt.figure(figsize=(5.5, 4.5))
    ax = fig.add_subplot(1, 1, 1)
    ax.scatter(theo, r_sorted, s=8)
    # 45-degree fit line
    xline = np.linspace(theo.min(), theo.max(), 100)
    # Fit slope/intercept by least squares
    A = np.c_[theo, np.ones_like(theo)]
    b = np.linalg.lstsq(A, r_sorted, rcond=None)[0]
    ax.plot(xline, b[0]*xline + b[1], linestyle="--", color="gray")
    ax.set_title(title)
    ax.set_xlabel("Theoretical quantiles (N(0,1))")
    ax.set_ylabel("Sample residual quantiles")
    fig.tight_layout()
    # Save PNG+PDF
    fig.savefig(out_base.with_suffix(".png"), dpi=300, bbox_inches="tight")
    fig.savefig(out_base.with_suffix(".pdf"), bbox_inches="tight")
    plt.close(fig)


# =========================
# Plot saving helper (PNG + PDF)
# =========================
def _save_fig_both(fig: plt.Figure, base_path: Path):
    fig.savefig(base_path.with_suffix(".png"), dpi=300, bbox_inches="tight")
    fig.savefig(base_path.with_suffix(".pdf"), bbox_inches="tight")


# =========================
# Main analysis
# =========================
all_rows = []

for label, csv_path in FILE_SETS.items():
    print(f"\n=== Dataset: {label} ===")
    df = pd.read_csv(csv_path)

    # group checks
    if GROUP_COL not in df.columns or df[GROUP_COL].dropna().nunique() < 2:
        print("  [Skip] Invalid group column.")
        continue
    df[GROUP_COL] = pd.Categorical(df[GROUP_COL], categories=["Control", "Study"], ordered=False)

    # resolve scaling factor
    try:
        scale_col = resolve_column(df, SCALE_FIELD_EXPECTED)
    except KeyError as e:
        print(f"  [Skip] Scaling factor missing: {e}")
        continue
    scale = pd.to_numeric(df[scale_col], errors="coerce")

    # derive outcomes and log1p
    for new_out, (raw_expected, pretty) in OUTCOME_DEFS.items():
        try:
            raw_col = resolve_column(df, raw_expected)
        except KeyError as e:
            print(f"  [Skip] Outcome missing: {e}")
            continue
        df[new_out] = pd.to_numeric(df[raw_col], errors="coerce") * scale
        df[f"log_{new_out}"] = np.log1p(df[new_out])

    # covariates present
    present_adjust = [v for v in BASE_ADJUST_VARS if v in df.columns]
    categorical_vars = {v for v in present_adjust if v in CATEGORICAL_VARS_ALL} | {"group"}

    # gentle imputation
    df_model = df.copy()
    for v in present_adjust:
        if pd.api.types.is_numeric_dtype(df_model[v]):
            df_model[v] = pd.to_numeric(df_model[v], errors="coerce")
            df_model[v] = df_model[v].fillna(df_model[v].median())
        else:
            mode_val = df_model[v].mode(dropna=True)
            mode_val = mode_val.iloc[0] if not mode_val.empty else "Unknown"
            df_model[v] = df_model[v].fillna(mode_val)

    # run per-outcome diagnostics
    for new_out, (_, pretty) in OUTCOME_DEFS.items():
        raw_var = new_out
        log_var = f"log_{new_out}"
        if raw_var not in df_model.columns or log_var not in df_model.columns:
            continue

        # build formulas
        f_raw = build_formula(raw_var, present_adjust, categorical_vars, group_var=GROUP_COL)
        f_log = build_formula(log_var, present_adjust, categorical_vars, group_var=GROUP_COL)

        if HAVE_SM:
            # statsmodels path
            m_raw, y_raw, X_raw = fit_ols_sm(f_raw, df_model)
            m_log, y_log, X_log = fit_ols_sm(f_log, df_model)
            # BP & JB via statsmodels
            lm, lmp, fval, fp = het_breuschpagan(m_raw.resid, m_raw.model.exog)
            bp_lmp_raw = float(lmp)
            lm, lmp, fval, fp = het_breuschpagan(m_log.resid, m_log.model.exog)
            bp_lmp_log = float(lmp)
            jb_stat, jb_p_raw, _, _ = jarque_bera(m_raw.resid)
            jb_stat, jb_p_log, _, _ = jarque_bera(m_log.resid)
            aic_raw = float(m_raw.aic)
            aic_log = float(m_log.aic)
            # y series
            y_raw_vals = y_raw.iloc[:, 0].values
            y_log_vals = y_log.iloc[:, 0].values
        else:
            # NumPy fallback
            m_raw = fit_ols_np(f_raw, df_model)
            m_log = fit_ols_np(f_log, df_model)
            # BP (LM + approx p)
            LM_raw, bp_p_raw = bp_lm_np(m_raw)
            LM_log, bp_p_log  = bp_lm_np(m_log)
            bp_lmp_raw, bp_lmp_log = float(bp_p_raw), float(bp_p_log)
            # JB (stat + approx p); also get skew/kurt if needed
            jb_stat, jb_p_raw, _, _ = jb_np(m_raw)
            jb_stat, jb_p_log,  _, _ = jb_np(m_log)
            aic_raw = float(m_raw["aic"])
            aic_log = float(m_log["aic"])
            y_raw_vals = m_raw["y_df"].iloc[:, 0].values
            y_log_vals = m_log["y_df"].iloc[:, 0].values

        # distribution descriptors (on the modeling sample)
        def _desc(yv):
            yv = pd.Series(yv).dropna()
            return {
                "prop_zeros": float((yv == 0).mean()) if len(yv) else np.nan,
                "skew": float(yv.skew()) if len(yv) else np.nan,
                "kurt": float(yv.kurtosis()) if len(yv) else np.nan,
            }
        desc_raw = _desc(y_raw_vals)
        desc_log = _desc(y_log_vals)

        # record row
        row = {
            "Dataset": label,
            "Outcome": pretty,
            "n_raw": int(len(y_raw_vals)),
            "n_log": int(len(y_log_vals)),
            "Raw_prop_zeros": desc_raw["prop_zeros"],
            "Raw_skew": desc_raw["skew"],
            "Raw_kurt_excess": desc_raw["kurt"],
            "Log_skew": desc_log["skew"],
            "Log_kurt_excess": desc_log["kurt"],
            "BP_p_raw": bp_lmp_raw,
            "BP_p_log": bp_lmp_log,
            "JB_p_raw": float(jb_p_raw),
            "JB_p_log": float(jb_p_log),
            "AIC_raw": aic_raw,
            "AIC_log": aic_log,
            "Improved_BP": float(bp_lmp_log > bp_lmp_raw),
            "Improved_JB": float(jb_p_log > jb_p_raw),
            "Improved_AIC": float(aic_log < aic_raw - 2),
        }
        all_rows.append(row)

        # plots (PNG + PDF)
        save_prefix = FIG_DIR / f"{label}_{raw_var}"
        title_prefix = f"{label} — {pretty}"

        # 1) Histograms
        fig = plt.figure(figsize=(10, 4))
        ax1 = fig.add_subplot(1, 2, 1)
        ax1.hist(y_raw_vals, bins=40)
        ax1.set_title(f"{title_prefix} — Raw")
        ax1.set_xlabel("Outcome"); ax1.set_ylabel("Count")
        ax2 = fig.add_subplot(1, 2, 2)
        ax2.hist(y_log_vals, bins=40)
        ax2.set_title(f"{title_prefix} — log1p")
        ax2.set_xlabel("log1p(Outcome)")
        fig.tight_layout()
        _save_fig_both(fig, save_prefix.with_name(save_prefix.name + "_hist"))
        plt.close(fig)

        # 2) QQ plots
        if HAVE_SM:
            fig_raw = sm.qqplot(m_raw.resid, line='45')
            fig_raw.suptitle(f"{title_prefix} — QQ (Raw)")
            _save_fig_both(fig_raw, save_prefix.with_name(save_prefix.name + "_qq_raw"))
            plt.close(fig_raw)

            fig_log = sm.qqplot(m_log.resid, line='45')
            fig_log.suptitle(f"{title_prefix} — QQ (log1p)")
            _save_fig_both(fig_log, save_prefix.with_name(save_prefix.name + "_qq_log"))
            plt.close(fig_log)
        else:
            qqplot_fallback(m_raw["resid"], f"{title_prefix} — QQ (Raw)",
                            save_prefix.with_name(save_prefix.name + "_qq_raw"))
            qqplot_fallback(m_log["resid"], f"{title_prefix} — QQ (log1p)",
                            save_prefix.with_name(save_prefix.name + "_qq_log"))

        # 3) Residual vs Fitted
        if HAVE_SM:
            r_raw = m_raw.resid
            f_raw = m_raw.fittedvalues
            r_log = m_log.resid
            f_log = m_log.fittedvalues
        else:
            r_raw = m_raw["resid"]
            f_raw = m_raw["fittedvalues"]
            r_log = m_log["resid"]
            f_log = m_log["fittedvalues"]

        fig = plt.figure(figsize=(10, 4))
        ax1 = fig.add_subplot(1, 2, 1)
        ax1.scatter(f_raw, r_raw, s=8)
        ax1.axhline(0, color='gray', linestyle='--')
        ax1.set_title(f"{title_prefix} — Residual vs Fitted (Raw)")
        ax1.set_xlabel("Fitted"); ax1.set_ylabel("Residuals")
        ax2 = fig.add_subplot(1, 2, 2)
        ax2.scatter(f_log, r_log, s=8)
        ax2.axhline(0, color='gray', linestyle='--')
        ax2.set_title(f"{title_prefix} — Residual vs Fitted (log1p)")
        ax2.set_xlabel("Fitted"); ax2.set_ylabel("Residuals")
        fig.tight_layout()
        _save_fig_both(fig, save_prefix.with_name(save_prefix.name + "_rvf"))
        plt.close(fig)

# Export summary CSV
diag_df = pd.DataFrame(all_rows)
summary_csv = DIAG_DIR / "wmh_transform_diagnostics_summary.csv"
diag_df.to_csv(summary_csv, index=False, encoding="utf-8-sig")
print(f"\nSaved summary: {summary_csv}")
print(f"Figures saved under: {FIG_DIR}  (each as .png and .pdf)")

# Console view with overall recommendation
if not diag_df.empty:
    view = diag_df[[
        "Dataset", "Outcome",
        "Raw_prop_zeros", "Raw_skew", "Raw_kurt_excess",
        "Log_skew", "Log_kurt_excess",
        "BP_p_raw", "BP_p_log", "JB_p_raw", "JB_p_log",
        "AIC_raw", "AIC_log",
        "Improved_BP", "Improved_JB", "Improved_AIC"
    ]].copy()
    view["Recommend_log1p"] = (view[["Improved_BP", "Improved_JB", "Improved_AIC"]].sum(axis=1) >= 2)
    print("\n=== Diagnostic summary (key fields) ===")
    print(view.to_string(index=False))


# =========================
# Word report (Methods + results + figure thumbnails)
# =========================
def word_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

def add_picture(doc: Document, path: Path, caption: str):
    if path.exists():
        doc.add_picture(str(path), width=Inches(5.8))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_word_report(diag_df: pd.DataFrame):
    doc = Document()
    word_style(doc)
    doc.add_heading("Supplementary eMethod: WMH Transformation Diagnostics", level=1)

    # Methods
    doc.add_heading("Objective", level=2)
    doc.add_paragraph(
        "We assessed whether head-size–normalized white matter hyperintensity (WMH) volumes "
        "should be modeled on the raw scale or on the log1p scale for regression analyses."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("Datasets", level=2)
    doc.add_paragraph(
        "Two analytic cohorts were evaluated: (i) the primary matched cohort "
        "(`primary_cohort.csv`), and (ii) the sensitivity matched cohort "
        "(`sensitivity_cohort.csv`)."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("Outcomes", level=2)
    doc.add_paragraph(
        "For each cohort, we computed head-size–normalized outcomes as the raw WMH volume "
        "multiplied by the volumetric scaling factor from T1 (registration to standard space). "
        "Outcomes considered: total WMH (T1+T2 FLAIR), periventricular WMH, and deep WMH. "
        "For each, a log1p-transformed variant was created."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("Modeling and diagnostics", level=2)
    sm_note = ("Exact BP/JB p-values and QQ plots were computed using statsmodels."
               if HAVE_SM else
               "OLS and test statistics were computed via NumPy; p-values were approximated "
               "using Wilson–Hilferty for χ² and QQ plots were generated without SciPy. "
               "This approximation is suitable for diagnostic purposes.")
    doc.add_paragraph(
        "Ordinary least squares (OLS) models were fit using prespecified covariates "
        "(age at MRI, sex, baseline BMI, Townsend deprivation index, genetic ethnic grouping, "
        "smoking ever, and alcohol intake frequency). Categorical covariates were entered as "
        "indicator variables. Missing data were handled by median imputation for continuous "
        "covariates and mode imputation for categorical covariates. For each outcome, raw and "
        "log1p models were compared using Breusch–Pagan (heteroscedasticity), Jarque–Bera "
        "(normality of residuals), and AIC. " + sm_note
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("Decision rule", level=2)
    doc.add_paragraph(
        "As a heuristic, we recommend the log1p transformation if at least two of the three "
        "indicators improved (higher BP p-value, higher JB p-value, and lower AIC by ≥2)."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Results table
    doc.add_heading("Summary results", level=2)
    if not diag_df.empty:
        cols = [
            "Dataset", "Outcome",
            "Raw_prop_zeros", "Raw_skew", "Raw_kurt_excess",
            "Log_skew", "Log_kurt_excess",
            "BP_p_raw", "BP_p_log", "JB_p_raw", "JB_p_log",
            "AIC_raw", "AIC_log", "Improved_BP", "Improved_JB", "Improved_AIC"
        ]
        tbl = doc.add_table(rows=1, cols=len(cols))
        hdr = tbl.rows[0].cells
        for j, c in enumerate(cols):
            hdr[j].text = c
        for _, r in diag_df[cols].iterrows():
            row_cells = tbl.add_row().cells
            for j, c in enumerate(cols):
                row_cells[j].text = str(r[c])
        doc.add_paragraph(
            "Note: 'Improved_*' flags indicate better fit on log1p vs raw."
        ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Figure thumbnails (embed PNG; PDFs are saved alongside)
    doc.add_heading("Figure thumbnails", level=2)
    for label, _ in FILE_SETS.items():
        for new_out, (_, pretty) in OUTCOME_DEFS.items():
            base = FIG_DIR / f"{label}_{new_out}"
            add_picture(doc, base.with_name(base.name + "_hist.png"),
                        f"{label} — {pretty}: Histograms (raw vs log1p)")
            add_picture(doc, base.with_name(base.name + "_qq_raw.png"),
                        f"{label} — {pretty}: QQ plot (raw residuals)")
            add_picture(doc, base.with_name(base.name + "_qq_log.png"),
                        f"{label} — {pretty}: QQ plot (log1p residuals)")
            add_picture(doc, base.with_name(base.name + "_rvf.png"),
                        f"{label} — {pretty}: Residual vs Fitted (raw vs log1p)")

    out_doc = DIAG_DIR / "eMethod_WMH_Transform_Diagnostics.docx"
    doc.save(out_doc)
    print(f"Saved Word report: {out_doc}")

# Build Word report
build_word_report(diag_df)
